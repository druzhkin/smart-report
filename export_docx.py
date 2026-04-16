"""McKinsey-level DOCX export. Title page, TOC, Executive Summary,
matrix, priority blocks with coloured bars, connections in framed boxes,
appendices. Embeds 4 infographics from infographics.py.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

from infographics import render_all
from models import Block, BlockHeader, Connection, Report

# ---------- palette ----------

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
NAVY_HEX = "1B3A5C"
GREY_BG_HEX = "F5F5F5"
WHITE_HEX = "FFFFFF"
RED_HEX = "C0392B"
YELLOW_HEX = "F1C40F"
GREEN_HEX = "27AE60"
GREY_TXT = RGBColor(0x55, 0x55, 0x55)

PRIORITY_HEX = {"high": RED_HEX, "medium": YELLOW_HEX, "low": GREEN_HEX}
PRIORITY_LABEL = {"high": "🔴 Высокий", "medium": "🟡 Средний", "low": "🟢 Низкий"}
PRIORITY_ORDER = {"high": 0, "medium": 1, "low": 2}

FONT = "Arial"


# ---------- xml helpers ----------

def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, hex_color: str = "BBBBBB", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), hex_color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_paragraph_left_border(paragraph, hex_color: str, size_pt: int = 24) -> None:
    """Thick coloured left border on a paragraph — visual priority marker."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pbdr.append(left)
    p_pr.append(pbdr)


def _set_paragraph_all_borders(paragraph, hex_color: str = "1B3A5C", size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), hex_color)
        pbdr.append(b)
    p_pr.append(pbdr)


def _add_page_number_footer(doc) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = GREY_TXT

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_toc(doc) -> None:
    """Insert Word TOC field. User must press F9 in Word to populate."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Оглавление — обновите через F9 в Word"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


# ---------- style setup ----------

def _setup_document(doc: Document) -> None:
    # margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # base Normal style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    # heading styles
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            hs = doc.styles[style_name]
            hs.font.name = FONT
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = NAVY


def _heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = NAVY


def _para(doc, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _sorted_blocks(report: Report) -> list[Block]:
    headers = {h.cell: h for h in report.block_headers}

    def _key(b: Block):
        h = headers.get(b.cell)
        prio = PRIORITY_ORDER.get(h.priority if h else "", 3)
        score = -(h.score_novelty + h.score_concreteness + h.score_applicability) if h else 0
        return (prio, score, b.cell)

    return sorted(report.blocks, key=_key)


# ---------- sections ----------

def _title_page(doc, report: Report) -> None:
    # Top spacer
    for _ in range(3):
        doc.add_paragraph()

    # Logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ ЛОГОТИП ]")
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.color.rgb = GREY_TXT

    for _ in range(2):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("АНАЛИТИЧЕСКИЙ ОТЧЁТ")
    run.font.name = FONT
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Subtitle = goal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report.goal)
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    for _ in range(4):
        doc.add_paragraph()

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—" * 20)
    run.font.color.rgb = NAVY

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%d.%m.%Y"))
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.color.rgb = GREY_TXT

    # Stats
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats = (
        f"{len(report.matrix.domains)} доменов · "
        f"{len(report.blocks)} блоков · "
        f"{len(report.connections)} кросс-доменных связей"
    )
    run = p.add_run(stats)
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = GREY_TXT

    doc.add_page_break()


def _toc_page(doc) -> None:
    _heading(doc, "Оглавление", level=1)
    _add_toc(doc)
    doc.add_page_break()


def _executive_summary(doc, report: Report, metrics_img: Path) -> None:
    _heading(doc, "Executive Summary", level=1)
    es = report.exec_summary
    if es is None:
        _para(doc, "Executive Summary не был сгенерирован.")
        doc.add_page_break()
        return

    # Goal restatement (highlighted)
    p = doc.add_paragraph()
    run = p.add_run("Цель: ")
    run.font.name = FONT
    run.font.bold = True
    run.font.color.rgb = NAVY
    run = p.add_run(es.goal_restate)
    run.font.name = FONT
    _set_paragraph_left_border(p, NAVY_HEX, size_pt=24)

    # Matrix mini-table
    _heading(doc, "Матрица доменов", level=2)
    _render_matrix_table(doc, report)

    # Top findings
    _heading(doc, "Топ-5 находок", level=2)
    for i, f in enumerate(es.top_findings, 1):
        _numbered_finding(doc, i, f.block_cell, _strip_markdown(f.headline))

    # Top connections
    _heading(doc, "Топ-3 кросс-доменных связи", level=2)
    for c in es.top_connections:
        doms = " ↔ ".join(c.domains) if c.domains else ""
        _numbered_finding(doc, None, doms, _strip_markdown(c.headline))

    # Critical gaps
    _heading(doc, "Критические пробелы", level=2)
    for g in es.key_gaps:
        _numbered_finding(doc, None, "пробел", _strip_markdown(g), bar_hex=RED_HEX)

    doc.add_page_break()


def _render_matrix_table(doc, report: Report) -> None:
    """Domain × description table with shaded header."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    hdr[0].text = "Домен"
    hdr[1].text = "Суть / слои"
    for c in hdr:
        _shade_cell(c, NAVY_HEX)
        _set_cell_borders(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = FONT
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(11)

    for d in report.matrix.domains:
        row = table.add_row().cells
        row[0].text = d.name
        layers = "; ".join(f"{l.name}" for l in d.layers)
        row[1].text = f"{d.rationale}\nСлои: {layers}"
        for c in row:
            _set_cell_borders(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10)
        _shade_cell(row[0], GREY_BG_HEX)


def _insert_image(doc, img_path: Path, caption: str, width_cm: float = 16.0) -> None:
    if not img_path or not img_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(img_path), width=Cm(width_cm))
    except Exception:
        return
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = FONT
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY_TXT


def _block_header_card(doc, block: Block, header: BlockHeader | None) -> None:
    """Header box with coloured priority bar."""
    prio = header.priority if header else "low"
    color = PRIORITY_HEX.get(prio, "CCCCCC")

    # Cell title with coloured left bar
    p = doc.add_paragraph()
    run = p.add_run(block.cell)
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    _set_paragraph_left_border(p, color, size_pt=36)

    if header is None:
        return

    # Priority label
    p = doc.add_paragraph()
    run = p.add_run(f"Приоритет: {PRIORITY_LABEL.get(prio, prio)}")
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.bold = True

    # Card lines
    for label, value in [
        ("Главный вывод", header.one_liner),
        ("Сильнейшая цифра", header.strongest_number),
        ("Главный пробел", header.main_gap),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        r.font.name = FONT
        r.font.bold = True
        r.font.color.rgb = NAVY
        r = p.add_run(value)
        r.font.name = FONT
        # blue highlight for numbers
        if label == "Сильнейшая цифра":
            r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
            r.font.bold = True

    p = doc.add_paragraph()
    r = p.add_run(
        f"Оценка: новизна {header.score_novelty}/3 · "
        f"конкретность {header.score_concreteness}/3 · "
        f"применимость {header.score_applicability}/3"
    )
    r.font.name = FONT
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY_TXT


_MD_URL_RE = __import__("re").compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)")
_BARE_URL_RE = __import__("re").compile(r"(?:https?://|www\.)\S+")


def _strip_markdown(text: str) -> str:
    text = _MD_URL_RE.sub(r"\1", text)
    text = _BARE_URL_RE.sub("", text)
    text = text.replace("**", "").replace("__", "")
    return text.strip()


def _render_md_table(doc, lines: list[str]) -> None:
    rows = [
        [c.strip() for c in ln.strip().strip("|").split("|")]
        for ln in lines
        if ln.strip().startswith("|")
    ]
    rows = [r for r in rows if not all(set(c) <= {"-", ":", " "} for c in r)]
    if len(rows) < 2:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = table.rows[ri].cells
        for ci in range(cols):
            val = row[ci] if ci < len(row) else ""
            cells[ci].text = _strip_markdown(val)
            _set_cell_borders(cells[ci])
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10)
                    if ri == 0:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if ri == 0:
                _shade_cell(cells[ci], NAVY_HEX)


def _render_summary(doc, summary: str) -> None:
    lines = summary.split("\n")
    buf: list[str] = []
    tbl: list[str] = []
    in_table = False

    def flush_text():
        text = "\n".join(buf).strip()
        buf.clear()
        if not text:
            return
        for para in text.split("\n\n"):
            s = _strip_markdown(para)
            if not s:
                continue
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.font.name = FONT
            r.font.size = Pt(11)

    def flush_table():
        if tbl:
            _render_md_table(doc, tbl)
            tbl.clear()

    for ln in lines:
        if ln.strip().startswith("|") and ln.strip().endswith("|"):
            if not in_table:
                flush_text()
                in_table = True
            tbl.append(ln)
        else:
            if in_table:
                flush_table()
                in_table = False
            buf.append(ln)
    if in_table:
        flush_table()
    else:
        flush_text()


def _numbered_finding(doc, num: int | None, tag: str, text: str, bar_hex: str = NAVY_HEX) -> None:
    """Render a finding as a bordered card: [N] TAG — headline."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.2)
    table.columns[1].width = Cm(14.5)

    lcell = table.rows[0].cells[0]
    rcell = table.rows[0].cells[1]
    lcell.width = Cm(1.2)
    rcell.width = Cm(14.5)

    _shade_cell(lcell, bar_hex)
    lp = lcell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(str(num) if num is not None else "•")
    lr.font.name = FONT
    lr.font.bold = True
    lr.font.size = Pt(12)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _shade_cell(rcell, GREY_BG_HEX)
    rp = rcell.paragraphs[0]
    if tag:
        tr = rp.add_run(f"{tag}  ")
        tr.font.name = FONT
        tr.font.size = Pt(9)
        tr.font.bold = True
        tr.font.color.rgb = NAVY
    tr = rp.add_run(text)
    tr.font.name = FONT
    tr.font.size = Pt(11)
    rcell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # spacing after card
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _blocks_section(doc, report: Report) -> None:
    _heading(doc, "Блоки (по приоритету)", level=1)
    headers = {h.cell: h for h in report.block_headers}
    for b in _sorted_blocks(report):
        h = headers.get(b.cell)
        _block_header_card(doc, b, h)

        # body
        _render_summary(doc, b.summary or "")

        if b.findings:
            _heading(doc, "Источники", level=3)
            for f in b.findings:
                p = doc.add_paragraph(style="List Bullet")
                prefix = "📊 " if f.has_numbers else ""
                r = p.add_run(f"{prefix}{f.claim}")
                r.font.name = FONT
                r.font.size = Pt(10)
                # blue highlight if numbers
                if f.has_numbers:
                    r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
                r = p.add_run(f"  [{f.source_type}] ")
                r.font.name = FONT
                r.font.size = Pt(9)
                r.font.italic = True
                r.font.color.rgb = GREY_TXT
                r = p.add_run(f.source)
                r.font.name = FONT
                r.font.size = Pt(9)
                r.font.color.rgb = GREY_TXT

        if b.assumptions:
            _heading(doc, "Ключевые допущения", level=3)
            for a in b.assumptions:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(a)
                r.font.name = FONT
                r.font.size = Pt(10)

        if b.gaps:
            _heading(doc, "Пробелы / куда копать", level=3)
            for g in b.gaps:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(g)
                r.font.name = FONT
                r.font.size = Pt(10)

        analogies_list = getattr(b, "analogies", None) or []
        if analogies_list:
            _heading(doc, "Исторические аналогии", level=3)
            for a in analogies_list:
                loc = getattr(a, "location", "") or ""
                matched = getattr(a, "matched", None) or []
                differed = getattr(a, "differed", None) or []
                confidence = getattr(a, "confidence", "moderate") or "moderate"
                # heading: location — situation
                p = doc.add_paragraph()
                heading_text = f"{loc} — {a.situation}" if loc else a.situation
                r = p.add_run(heading_text)
                r.font.name = FONT
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = NAVY
                _set_paragraph_left_border(p, "1B3A5C", size_pt=16)
                # matched
                if matched:
                    pm = doc.add_paragraph()
                    r = pm.add_run("Что совпадает: ")
                    r.font.name = FONT; r.font.size = Pt(9); r.font.bold = True
                    r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
                    r2 = pm.add_run("; ".join(matched))
                    r2.font.name = FONT; r2.font.size = Pt(9)
                # differed
                if differed:
                    pd = doc.add_paragraph()
                    r = pd.add_run("Что отличается: ")
                    r.font.name = FONT; r.font.size = Pt(9); r.font.bold = True
                    r.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
                    r2 = pd.add_run("; ".join(differed))
                    r2.font.name = FONT; r2.font.size = Pt(9)
                # lesson
                pl = doc.add_paragraph()
                r = pl.add_run("Урок: ")
                r.font.name = FONT; r.font.size = Pt(9); r.font.bold = True
                r2 = pl.add_run(a.lesson)
                r2.font.name = FONT; r2.font.size = Pt(9)
                # confidence chip if not high
                if confidence != "high":
                    pc = doc.add_paragraph()
                    r = pc.add_run(f"[{confidence}]")
                    r.font.name = FONT; r.font.size = Pt(8)
                    r.font.color.rgb = GREY_TXT; r.font.italic = True

        # Quadrant Crunch — assumption inversions
        inv_blocks = getattr(report, "assumption_inversions", None) or []
        block_inv = next((bi for bi in inv_blocks if bi.block_cell == b.cell), None)
        if block_inv and block_inv.inversions:
            _heading(doc, "Проверка допущений", level=3)
            if block_inv.unfalsifiable_flag:
                pw = doc.add_paragraph()
                r = pw.add_run("Внимание: вывод может быть нефальсифицируем.")
                r.font.name = FONT; r.font.size = Pt(9); r.font.italic = True
                r.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
            for inv in block_inv.inversions:
                dep_color = (
                    RGBColor(0xDC, 0x26, 0x26) if inv.dependency == "critical"
                    else RGBColor(0xD9, 0x77, 0x06) if inv.dependency == "important"
                    else GREY_TXT
                )
                p = doc.add_paragraph()
                r = p.add_run(f"[{inv.dependency}] ")
                r.font.name = FONT; r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = dep_color
                r2 = p.add_run(inv.assumption)
                r2.font.name = FONT; r2.font.size = Pt(10); r2.font.bold = True
                _set_paragraph_left_border(p, "1B3A5C" if inv.dependency == "minor" else
                                           ("DC2626" if inv.dependency == "critical" else "D97706"), size_pt=16)
                pi = doc.add_paragraph()
                r = pi.add_run("А если нет? ")
                r.font.name = FONT; r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = dep_color
                r2 = pi.add_run(inv.inversion)
                r2.font.name = FONT; r2.font.size = Pt(10)
                pc = doc.add_paragraph()
                r = pc.add_run("Последствие: ")
                r.font.name = FONT; r.font.size = Pt(9); r.font.bold = True
                r2 = pc.add_run(inv.consequence)
                r2.font.name = FONT; r2.font.size = Pt(9)
                pp = doc.add_paragraph()
                r = pp.add_run(f"Вероятность: {inv.probability}   Ранний сигнал: ")
                r.font.name = FONT; r.font.size = Pt(9); r.font.italic = True
                r.font.color.rgb = GREY_TXT
                r2 = pp.add_run(inv.early_signal)
                r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = GREY_TXT

        doc.add_paragraph()


NATURE_ICONS = {
    "paradox": "⚡ Парадокс",
    "causal_chain": "🔗 Причинная цепочка",
    "unexpected_confirmation": "✓ Подтверждение",
    "shared_variable": "◇ Общая переменная",
}


def _connections_section(doc, report: Report) -> None:
    if not report.connections:
        return
    _heading(doc, "Кросс-доменные связи (бисоциация)", level=1)
    for c in report.connections:
        # framed box
        p = doc.add_paragraph()
        icon = NATURE_ICONS.get(c.nature, c.nature)
        r = p.add_run(f"{icon}  ·  {' ↔ '.join(c.domains)}  ·  сила: {c.strength}")
        r.font.name = FONT
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = NAVY
        _set_paragraph_all_borders(p, NAVY_HEX, size=8)

        p = doc.add_paragraph()
        r = p.add_run("Общая переменная: ")
        r.font.name = FONT
        r.font.bold = True
        r = p.add_run(c.shared_entity)
        r.font.name = FONT

        p = doc.add_paragraph()
        r = p.add_run(c.description)
        r.font.name = FONT
        r.font.size = Pt(11)

        if c.anchors:
            p = doc.add_paragraph()
            r = p.add_run("Опоры в блоках:")
            r.font.name = FONT
            r.font.bold = True
            r.font.size = Pt(10)
            for a in c.anchors:
                pp = doc.add_paragraph(style="List Bullet")
                r = pp.add_run(a)
                r.font.name = FONT
                r.font.size = Pt(10)

        if c.novelty:
            p = doc.add_paragraph()
            r = p.add_run("Что нового: ")
            r.font.name = FONT
            r.font.bold = True
            r.font.color.rgb = NAVY
            r = p.add_run(c.novelty)
            r.font.name = FONT
            r.font.italic = True

        doc.add_paragraph()


def _scenarios_section(doc, report: Report) -> None:
    """Dedicated «Сценарии» section for predictive questions."""
    cone = getattr(report, "scenario_cone", None)
    if not cone:
        return
    horizon = getattr(cone, "question_horizon", "12-24 месяцев") or "12-24 месяцев"
    _heading(doc, f"Сценарии · горизонт {horizon}", level=1)

    verdict = getattr(cone, "conditional_verdict", "") or ""
    if verdict:
        p = doc.add_paragraph()
        _set_paragraph_left_border(p, "2563EB", size_pt=20)
        r = p.add_run(verdict)
        r.font.name = FONT
        r.font.size = Pt(11)
        r.font.italic = True
        r.font.color.rgb = NAVY

    uncertainties = getattr(cone, "key_uncertainties", []) or []
    if uncertainties:
        p = doc.add_paragraph()
        r = p.add_run("Ключевые неопределённости: ")
        r.font.name = FONT
        r.font.bold = True
        r = p.add_run("; ".join(uncertainties))
        r.font.name = FONT
        r.font.size = Pt(10)

    scenario_colors = {"base": "2563EB", "optim": "27AE60", "pessim": "C0392B"}

    for s in (getattr(cone, "scenarios", []) or []):
        name = getattr(s, "name", "") or ""
        prob = getattr(s, "probability", "") or ""
        desc = getattr(s, "description", "") or ""
        key_driver = getattr(s, "key_driver", "") or ""
        implications = getattr(s, "implications", []) or []
        indicators = getattr(s, "indicators", []) or []

        name_low = name.lower()
        if "optim" in name_low or "оптим" in name_low:
            color_hex = scenario_colors["optim"]
        elif "pessim" in name_low or "пессим" in name_low:
            color_hex = scenario_colors["pessim"]
        else:
            color_hex = scenario_colors["base"]

        _heading(doc, f"{name} ({prob})", level=2)
        p = doc.add_paragraph()
        _set_paragraph_left_border(p, color_hex, size_pt=24)
        r = p.add_run(desc)
        r.font.name = FONT
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        r = p.add_run("Ключевой драйвер: ")
        r.font.name = FONT
        r.font.bold = True
        r.font.color.rgb = NAVY
        r = p.add_run(key_driver)
        r.font.name = FONT

        if implications:
            _heading(doc, "Следствия для клиента", level=3)
            for imp in implications:
                pp = doc.add_paragraph(style="List Bullet")
                r = pp.add_run(imp)
                r.font.name = FONT
                r.font.size = Pt(10)

        if indicators:
            _heading(doc, "Индикаторы (I&W)", level=3)
            for ind in indicators:
                pp = doc.add_paragraph(style="List Bullet")
                r = pp.add_run(ind)
                r.font.name = FONT
                r.font.size = Pt(10)

        doc.add_paragraph()

    wild_card = getattr(cone, "wild_card", None)
    if wild_card:
        _heading(doc, "Wild Card", level=2)
        wc_prob = getattr(wild_card, "probability", "") or ""
        wc_desc = getattr(wild_card, "description", "") or ""
        wc_impact = getattr(wild_card, "impact", "") or ""
        p = doc.add_paragraph()
        r = p.add_run(f"P={wc_prob}: {wc_desc}")
        r.font.name = FONT
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        if wc_impact:
            p = doc.add_paragraph()
            r = p.add_run("Эффект: ")
            r.font.name = FONT
            r.font.bold = True
            r = p.add_run(wc_impact)
            r.font.name = FONT
            r.font.size = Pt(10)


def _appendix_sources(doc, report: Report) -> None:
    _heading(doc, "Приложение A. Список источников", level=1)
    # group by type
    buckets: dict[str, list[tuple[str, str, str]]] = {}
    for b in report.blocks:
        for f in b.findings:
            buckets.setdefault(f.source_type or "unknown", []).append(
                (b.cell, f.claim, f.source)
            )
    for stype in ["primary", "secondary", "opinion", "unknown"]:
        if stype not in buckets:
            continue
        _heading(doc, f"{stype.title()} ({len(buckets[stype])})", level=2)
        for cell, claim, src in buckets[stype]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"[{cell}] ")
            r.font.name = FONT
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = NAVY
            r = p.add_run(f"{claim}  —  {src}")
            r.font.name = FONT
            r.font.size = Pt(9)


def _appendix_gaps(doc, report: Report) -> None:
    _heading(doc, "Приложение B. Пробелы и куда копать", level=1)
    for b in report.blocks:
        if not b.gaps:
            continue
        _heading(doc, b.cell, level=3)
        for g in b.gaps:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(g)
            r.font.name = FONT
            r.font.size = Pt(10)


# ---------- public entry ----------

def export_mckinsey_docx(report: Report, path: Path, images: dict[str, Path] | None = None) -> Path:
    """Generate a McKinsey-style .docx at `path`. Returns path."""
    doc = Document()
    _setup_document(doc)
    _add_page_number_footer(doc)

    # Generate infographics once
    if images is None:
        images = render_all(report)

    # 1. Title
    _title_page(doc, report)

    # 2. TOC
    _toc_page(doc)

    # 3. Executive Summary
    _executive_summary(doc, report, images.get("metrics"))

    # 4. Matrix map image (page 2 after ES)
    _heading(doc, "Карта матрицы доменов", level=1)
    _insert_image(doc, images.get("matrix"), "Рис. 1. Визуальная схема доменов и слоёв (цвет = приоритет).")
    doc.add_page_break()

    # 5. Key Metrics Dashboard
    _heading(doc, "Ключевые метрики", level=1)
    _insert_image(doc, images.get("metrics"), "Рис. 2. Ключевые цифры отчёта с источниками.")
    doc.add_page_break()

    # 6. Priority heatmap
    _heading(doc, "Тепловая карта приоритетов", level=1)
    _insert_image(doc, images.get("heatmap"), "Рис. 3. Где золото, где пусто: приоритет по ячейкам матрицы.")
    doc.add_page_break()

    # 7. Blocks (sorted by priority)
    _blocks_section(doc, report)

    # 8. Connections graph + section
    if report.connections:
        doc.add_page_break()
        _heading(doc, "Граф кросс-доменных связей", level=1)
        _insert_image(doc, images.get("graph"), "Рис. 4. Связи между доменами: цвет = тип, толщина = сила.")
        _connections_section(doc, report)

    # 9. Scenarios (predictive questions only)
    if getattr(report, "scenario_cone", None):
        doc.add_page_break()
        _scenarios_section(doc, report)

    # 10. Appendices
    doc.add_page_break()
    _appendix_sources(doc, report)
    doc.add_page_break()
    _appendix_gaps(doc, report)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
