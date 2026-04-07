from __future__ import annotations

from difflib import SequenceMatcher
import json
import os
import re
from pathlib import Path

import httpx
from langsmith import traceable
from loguru import logger


class RendererError(Exception):
    pass

from backend.config import settings
from backend.pipeline.model_router import AgentTask, estimate_cost, get_model
from backend.pipeline.state import AgentState
from backend.schemas.report_schema import ReportOutput, ReportSection
from backend.utils.json_parse import parse_llm_json, supports_json_mode
from backend.utils.retry import llm_retry

_SYSTEM_PROMPT_TEMPLATE = """You are a McKinsey-grade report writer using data from professional research.
Generate a comprehensive structured report with title, executive_summary, and sections array.
Each section has: title, content (full markdown with data, analysis, citations), order, sources (list of URLs).

Requirements:
- CRITICAL: Write the ENTIRE report (title, executive_summary, all section titles and content) in {language_name}. Do NOT use any other language.
- Use the provided evidence graph as the primary truth source. Do not invent claims that are absent from the evidence.
- If evidence is partial or conflicting, state uncertainty explicitly instead of smoothing it away.
- Executive summary: 300-400 words, actionable insights, key metrics with specific numbers.
- Generate AT LEAST 6-8 detailed sections covering all major aspects of the topic.
- Each section: minimum 400 words, thorough analysis with specific data points, statistics, and trends.
- TABLES: Always add a blank line before and after markdown tables. Format:

  Some text here.

  | Column 1 | Column 2 | Column 3 |
  |----------|----------|----------|
  | Value 1  | Value 2  | Value 3  |

  Continue text after table.

- Use markdown formatting: bold **key metrics**, bullet lists, tables for comparisons.
- Cite sources inline as [Source](url).
- Professional, authoritative tone. No filler. Dense with data.

Return valid JSON:
{{
  "title": "...",
  "executive_summary": "...",
  "sections": [
    {{"title": "...", "content": "...", "order": 1, "sources": ["url1", "url2"]}}
  ]
}}"""

LANG_NAMES = {"ru": "Russian", "en": "English", "de": "German", "fr": "French", "es": "Spanish", "zh": "Chinese", "ja": "Japanese", "ko": "Korean"}

_INTERNAL_META_PHRASES = (
    "i appreciate the detailed query",
    "what the search results cover",
    "limitations of the provided search results",
    "citation verification failed",
    "evidence quality:",
    "hallucinated",
    "open questions:",
)

_RECOMMENDATION_HINTS = (
    "recommendation",
    "рекомендац",
    "top-",
    "top ",
    "итоговый рейтинг",
    "best model",
    "best models",
    "лучшие модел",
    "рейтинг",
)


def _get_renderer_prompt(lang: str = "en") -> str:
    return _SYSTEM_PROMPT_TEMPLATE.format(language_name=LANG_NAMES.get(lang, lang))


def _looks_internal_or_meta(text: str) -> bool:
    lowered = (text or "").lower()
    return any(phrase in lowered for phrase in _INTERNAL_META_PHRASES)


def _sanitize_signal(text: str, *, max_len: int = 220) -> str | None:
    cleaned = " ".join((text or "").split()).strip()
    cleaned = re.sub(r"^[#\-*]+\s*", "", cleaned)
    if not cleaned:
        return None
    if _looks_internal_or_meta(cleaned):
        return None
    if len(cleaned) < 25:
        return None
    if cleaned.lower().startswith("sources"):
        return None
    return cleaned[:max_len].rstrip()


def _clean_section_content(content: str) -> str:
    lines = [line.rstrip() for line in (content or "").splitlines()]
    filtered: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            filtered.append("")
            continue
        if stripped in {"•", "вЂў"}:
            continue
        if _looks_internal_or_meta(stripped):
            continue
        filtered.append(line)
    normalized = "\n".join(filtered)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
    return normalized


def _is_source_only_content(content: str) -> bool:
    lines = [line.strip() for line in (content or "").splitlines() if line.strip()]
    if not lines:
        return True
    non_link_lines = [
        line
        for line in lines
        if not re.match(r"^[-*]?\s*https?://", line, flags=re.IGNORECASE)
        and line.lower() != "sources"
    ]
    return len(non_link_lines) < 2


def _dedupe_and_prune_sections(sections: list[ReportSection]) -> list[ReportSection]:
    deduped: list[ReportSection] = []
    fingerprints: list[str] = []

    for section in sorted(sections, key=lambda item: item.order):
        cleaned_content = _clean_section_content(section.content)
        if not cleaned_content:
            continue
        if _is_source_only_content(cleaned_content):
            continue
        fingerprint = re.sub(r"\W+", " ", cleaned_content.lower()).strip()
        if len(fingerprint) < 40:
            continue
        if any(
            fingerprint == seen
            or SequenceMatcher(None, fingerprint[:2500], seen[:2500]).ratio() >= 0.93
            for seen in fingerprints
        ):
            continue
        fingerprints.append(fingerprint)
        deduped.append(
            ReportSection(
                title=section.title,
                content=cleaned_content,
                order=len(deduped) + 1,
                sources=section.sources,
            )
        )

    if not deduped:
        return sections
    return deduped


def _is_recommendation_section(section: ReportSection) -> bool:
    title = (section.title or "").lower()
    content_head = "\n".join((section.content or "").splitlines()[:8]).lower()
    blob = f"{title}\n{content_head}"
    return any(marker in blob for marker in _RECOMMENDATION_HINTS)


def _apply_recommendation_policy(sections: list[ReportSection], allow_recommendations: bool) -> list[ReportSection]:
    if allow_recommendations:
        return sections

    filtered = [section for section in sections if not _is_recommendation_section(section)]
    if not filtered:
        return sections[:3]
    return [
        ReportSection(
            title=section.title,
            content=section.content,
            order=index + 1,
            sources=section.sources,
        )
        for index, section in enumerate(filtered)
    ]


def _tokenize(text: str) -> set[str]:
    return {token for token in re.findall(r"[a-zA-Zа-яА-Я0-9]+", (text or "").lower()) if len(token) > 2}


def _score_branch_match(section_title: str, section_description: str, branch: dict) -> int:
    section_tokens = _tokenize(f"{section_title} {section_description}")
    branch_tokens = _tokenize(
        " ".join(
            [
                branch.get("query", ""),
                " ".join(branch.get("findings", [])[:4]),
                " ".join(source.get("title", "") for source in branch.get("sources", [])),
            ]
        )
    )
    if not section_tokens or not branch_tokens:
        return 0
    return len(section_tokens & branch_tokens)


def _build_section_packets(
    section_targets: list[dict],
    branch_packets: list[dict],
    claim_table: list[dict],
) -> list[dict]:
    if not section_targets:
        return []

    packets: list[dict] = []
    for index, section in enumerate(section_targets, start=1):
        ranked = sorted(
            branch_packets,
            key=lambda branch: (
                _score_branch_match(section.get("title", ""), section.get("description", ""), branch),
                branch.get("confidence", 0.0),
            ),
            reverse=True,
        )
        matched = [
            branch
            for branch in ranked
            if _score_branch_match(section.get("title", ""), section.get("description", ""), branch) > 0
        ][:3]
        if not matched and ranked:
            matched = ranked[:1]

        usable_claims = [claim for claim in claim_table if claim.get("usable")]
        claim_ranked = sorted(
            usable_claims,
            key=lambda claim: (
                _score_branch_match(
                    section.get("title", ""),
                    section.get("description", ""),
                    {
                        "query": claim.get("section_hint", ""),
                        "findings": [claim.get("claim", "")],
                        "sources": [{"title": claim.get("source_title", "")}],
                    },
                ),
                float(claim.get("claim_score", 0.0)),
            ),
            reverse=True,
        )
        matched_claims = claim_ranked[:6]

        allowed_sources = list(
            dict.fromkeys(
                source.get("url", "")
                for branch in matched
                for source in branch.get("sources", [])
                if source.get("url")
            )
        )
        allowed_sources.extend(
            claim.get("source_url", "")
            for claim in matched_claims
            if claim.get("source_url")
        )
        allowed_sources = list(dict.fromkeys(allowed_sources))
        packets.append(
            {
                "title": section.get("title", f"Section {index}"),
                "description": section.get("description", ""),
                "required": section.get("required", True),
                "order": index,
                "matched_branches": matched,
                "matched_claims": matched_claims,
                "allowed_sources": allowed_sources,
                "subsections": section.get("subsections", []),
            }
        )
    return packets


def _build_executive_summary_packet(
    branch_packets: list[dict],
    claim_table: list[dict],
    evidence_items: list[dict],
    critique: dict | None,
    citation: dict | None,
    reflect: dict | None,
    synthesis_payload: dict | None,
) -> dict:
    top_signals: list[str] = []
    for claim in sorted(claim_table, key=lambda item: float(item.get("claim_score", 0.0)), reverse=True):
        if not claim.get("usable"):
            continue
        sanitized = _sanitize_signal(str(claim.get("claim", "")))
        if sanitized:
            top_signals.append(sanitized)
        if len(top_signals) >= 6:
            break

    for branch in sorted(branch_packets, key=lambda item: item.get("confidence", 0.0), reverse=True):
        for finding in branch.get("findings", [])[:3]:
            sanitized = _sanitize_signal(finding)
            if sanitized:
                top_signals.append(sanitized)
        if len(top_signals) >= 6:
            break
    if not top_signals:
        for item in evidence_items[:8]:
            sanitized = _sanitize_signal(item.get("claim", ""))
            if sanitized:
                top_signals.append(sanitized)
            if len(top_signals) >= 6:
                break

    open_questions: list[str] = []
    requires_verification_notice = False
    if critique:
        for issue in critique.get("blocking_issues", [])[:4]:
            issue_lower = (issue or "").lower()
            if "citation" in issue_lower or "halluc" in issue_lower:
                requires_verification_notice = True
                continue
            sanitized = _sanitize_signal(issue, max_len=180)
            if sanitized:
                open_questions.append(sanitized)
    if reflect:
        for gap in reflect.get("gaps", [])[:4]:
            sanitized = _sanitize_signal(gap, max_len=180)
            if sanitized:
                open_questions.append(sanitized)

    unique_domains = {
        source.get("domain", "")
        for branch in branch_packets
        for source in branch.get("sources", [])
        if source.get("domain")
    }
    if claim_table:
        unique_domains.update(
            claim.get("domain", "")
            for claim in claim_table
            if claim.get("domain")
        )
    return {
        "top_signals": list(dict.fromkeys(signal for signal in top_signals if signal))[:6],
        "open_questions": list(dict.fromkeys(question for question in open_questions if question))[:4],
        "branch_count": len(branch_packets),
        "evidence_count": len(evidence_items),
        "unique_domain_count": len(unique_domains),
        "citation_passed": citation.get("passed") if citation else None,
        "verified_count": citation.get("verified_count", 0) if citation else 0,
        "citation_total": citation.get("total", 0) if citation else 0,
        "requires_verification_notice": requires_verification_notice,
        "synthesis_ready": bool((synthesis_payload or {}).get("ready")),
        "allow_recommendations": bool((synthesis_payload or {}).get("allow_recommendations")),
        "synthesis_blockers": list((synthesis_payload or {}).get("blocking_reasons") or [])[:3],
    }


def _compose_executive_summary(packet: dict, llm_summary: str) -> str:
    llm_summary = (llm_summary or "").strip()
    if llm_summary and len(llm_summary) >= 120:
        base = _clean_section_content(llm_summary)
    else:
        top_signals = packet.get("top_signals", [])
        opening = top_signals[0] if top_signals else "The current evidence base supports a directional conclusion, but confidence varies by branch."
        base = (
            f"{opening} "
            f"The report draws on {packet.get('branch_count', 0)} research branches, "
            f"{packet.get('evidence_count', 0)} structured evidence items, and "
            f"{packet.get('unique_domain_count', 0)} unique source domains."
        )

    bullets = packet.get("top_signals", [])[:4]
    summary_parts = [base]
    if bullets:
        summary_parts.append("Key signals:\n" + "\n".join(f"- {item}" for item in bullets))

    citation_total = packet.get("citation_total", 0)
    if citation_total:
        if packet.get("citation_passed") is False:
            quality_line = (
                "Evidence quality is mixed; low-confidence claims were de-prioritized in this summary."
            )
        else:
            quality_line = (
                f"Evidence quality: {packet.get('verified_count', 0)}/{citation_total} citations verified."
            )
        summary_parts.append(quality_line)
    elif packet.get("requires_verification_notice"):
        summary_parts.append("Some claims require additional verification before operational decisions.")

    open_questions = packet.get("open_questions", [])[:3]
    if open_questions:
        summary_parts.append(
            "Open questions:\n" + "\n".join(f"- {question}" for question in open_questions)
        )

    if not packet.get("synthesis_ready", True):
        blockers = packet.get("synthesis_blockers", [])
        summary_parts.append(
            "Synthesis gate warning: evidence quality was insufficient for decision-grade recommendations."
        )
        if blockers:
            summary_parts.append("Current blockers:\n" + "\n".join(f"- {item}" for item in blockers))
    elif not packet.get("allow_recommendations", True):
        summary_parts.append(
            "Recommendation policy: rankings/top-lists were withheld due to verification quality constraints."
        )

    return _clean_section_content("\n\n".join(part for part in summary_parts if part)).strip()


def _build_orchestration_metadata(state: AgentState, context: dict) -> dict:
    branch_states = list(state.get("branch_states", []) or [])
    tasks = list(state.get("research_tasks", []) or [])
    contradiction_log = list(state.get("contradiction_log", []) or [])
    citation = state.get("citation_verification")
    critique = state.get("research_critique_result")
    reflect = state.get("reflect_result")

    return {
        "orchestration": {
            "task_count": len(tasks),
            "branch_count": len(branch_states),
            "actionable_branches": [
                branch.question
                for branch in branch_states
                if branch.next_action in {"deepen", "widen", "verify"}
            ],
            "branch_states": [
                {
                    "task_id": branch.task_id,
                    "question": branch.question,
                    "status": branch.status,
                    "next_action": branch.next_action,
                    "action_reason": branch.action_reason,
                    "evidence_count": branch.evidence_count,
                    "source_count": branch.source_count,
                    "confidence": branch.confidence,
                    "source_strategy": branch.source_strategy,
                    "gaps": branch.gaps,
                    "contradiction_notes": branch.contradiction_notes,
                }
                for branch in branch_states
            ],
            "contradictions": contradiction_log,
            "reflection": {
                "quality_score": reflect.quality_score if reflect else None,
                "needs_more_research": reflect.needs_more_research if reflect else None,
                "gaps": reflect.gaps if reflect else [],
            },
            "critique": {
                "verdict": critique.verdict if critique else None,
                "overall_score": critique.overall_score if critique else None,
                "blocking_issues": critique.blocking_issues if critique else [],
                "follow_up_queries": critique.follow_up_queries if critique else [],
            },
            "citations": {
                "passed": citation.passed if citation else None,
                "verified_count": citation.verified_count if citation else 0,
                "total": citation.total if citation else 0,
            },
            "synthesis_gate": context.get("synthesis_payload", {}),
            "claim_table_stats": {
                "total": len(context.get("claim_table", [])),
                "usable": sum(1 for claim in context.get("claim_table", []) if claim.get("usable")),
            },
            "available_charts": context.get("available_charts", []),
        }
    }


def _normalize_rendered_sections(parsed_sections: list[dict], section_packets: list[dict]) -> list[ReportSection]:
    normalized_by_title: dict[str, dict] = {
        str(section.get("title", "")).strip().lower(): section
        for section in parsed_sections
        if str(section.get("title", "")).strip()
    }

    output_sections: list[ReportSection] = []
    if section_packets:
        for packet in section_packets:
            packet_title = packet.get("title", "")
            llm_section = normalized_by_title.get(packet_title.strip().lower(), {})
            allowed_sources = packet.get("allowed_sources", [])
            raw_sources = list(llm_section.get("sources", []) or [])
            filtered_sources = [source for source in raw_sources if source in allowed_sources]
            if not filtered_sources:
                filtered_sources = allowed_sources[:6]

            content = _clean_section_content(str(llm_section.get("content", "")))
            if not content:
                branch_notes: list[str] = []
                for claim in packet.get("matched_claims", []):
                    sanitized = _sanitize_signal(claim.get("claim", ""), max_len=320)
                    if sanitized:
                        branch_notes.append(sanitized)
                for branch in packet.get("matched_branches", []):
                    for finding in branch.get("findings", [])[:3]:
                        sanitized = _sanitize_signal(finding, max_len=320)
                        if sanitized:
                            branch_notes.append(sanitized)
                content = "\n\n".join(branch_notes) or "Evidence for this section is limited; uncertainty is noted."

            output_sections.append(
                ReportSection(
                    title=packet_title,
                    content=content,
                    order=int(packet.get("order", len(output_sections) + 1)),
                    sources=filtered_sources,
                )
            )
        return output_sections

    for index, section in enumerate(parsed_sections, start=1):
        output_sections.append(
            ReportSection(
                title=section.get("title", f"Section {index}"),
                content=section.get("content", ""),
                order=section.get("order", index),
                sources=list(section.get("sources", []) or []),
            )
        )
    return output_sections


def _build_render_context(state: AgentState) -> dict:
    results = state.get("research_results", [])
    intake = state.get("intake_result")
    master_prompt = state.get("master_prompt")
    chart_paths = state.get("chart_paths", [])
    citation = state.get("citation_verification")
    critique = state.get("research_critique_result")
    reflect = state.get("reflect_result")
    evidence_items = list(state.get("evidence_items", []) or [])
    research_tasks = list(state.get("research_tasks", []) or [])
    claim_table = list(state.get("claim_table", []) or [])
    synthesis_payload = dict(state.get("synthesis_payload") or {})
    allow_recommendations = bool(state.get("allow_recommendations", True))

    section_targets: list[dict] = []
    if master_prompt and master_prompt.report_schema and master_prompt.report_schema.sections:
        for section in master_prompt.report_schema.sections:
            section_targets.append(
                {
                    "title": section.title,
                    "description": section.description,
                    "required": section.required,
                    "min_words": section.min_words,
                    "subsections": section.subsections,
                }
            )

    branch_packets: list[dict] = []
    for result in results:
        branch_evidence = [
            item.model_dump(mode="json")
            for item in evidence_items
            if result.query in item.claim or item.source_url in {source.url for source in result.sources}
        ][:8]
        branch_packets.append(
            {
                "query": result.query,
                "findings": result.findings[:6],
                "sources": [source.model_dump(mode="json") for source in result.sources[:8]],
                "gaps": result.gaps,
                "confidence": result.confidence,
                "evidence": branch_evidence,
            }
        )

    section_packets = _build_section_packets(section_targets, branch_packets, claim_table)
    citation_payload = citation.model_dump(mode="json") if citation else None
    critique_payload = critique.model_dump(mode="json") if critique else None
    reflect_payload = reflect.model_dump(mode="json") if reflect else None
    evidence_payload = [item.model_dump(mode="json") for item in evidence_items[:24]]
    executive_summary_packet = _build_executive_summary_packet(
        branch_packets,
        claim_table,
        evidence_payload,
        critique_payload,
        citation_payload,
        reflect_payload,
        synthesis_payload,
    )

    return {
        "intake": intake.model_dump() if intake else {},
        "report_schema": {
            "title_template": master_prompt.report_schema.title_template if master_prompt and master_prompt.report_schema else "",
            "constraints": master_prompt.report_schema.constraints if master_prompt and master_prompt.report_schema else [],
            "sections": section_targets,
        },
        "research_tasks": [task.model_dump(mode="json") for task in research_tasks],
        "research_branches": branch_packets,
        "section_packets": section_packets,
        "claim_table": claim_table[:80],
        "synthesis_payload": synthesis_payload,
        "allow_recommendations": allow_recommendations,
        "evidence_items": evidence_payload,
        "research_brief": state.get("research_brief", ""),
        "reflection": reflect_payload,
        "research_critique": critique_payload,
        "citation_verification": citation_payload,
        "executive_summary_packet": executive_summary_packet,
        "available_charts": [Path(p).name for p in chart_paths],
        "instructions": {
            "primary_rule": "Write from evidence and cited branches, not from generic prose.",
            "if_evidence_missing": "Explicitly note uncertainty instead of inventing detail.",
            "section_source_rule": "Every section should inherit only sources that support its claims.",
            "claim_table_rule": "Only use claim_table items where usable=true for decision-grade claims.",
            "recommendation_rule": (
                "Recommendations/top rankings are allowed."
                if allow_recommendations
                else "Do not output Top-N/rankings/best-choice recommendations; provide neutral analysis only."
            ),
        },
    }


@llm_retry()
async def _call_llm(system: str, user: str, model: str) -> str:
    base_messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]
    payload: dict = {
        "model": model,
        "messages": base_messages,
        "temperature": 0.3,
        "max_tokens": 12000,
    }
    if supports_json_mode(model):
        payload["response_format"] = {"type": "json_object"}

    fallback_payloads: list[dict] = [payload]
    if "response_format" in payload:
        fallback_payloads.append(
            {
                **payload,
                "response_format": None,
            }
        )
    fallback_payloads.append(
        {
            "model": "openai/gpt-4.1-mini",
            "messages": base_messages,
            "temperature": 0.2,
            "max_tokens": 8000,
            "response_format": {"type": "json_object"},
        }
    )
    fallback_payloads.append(
        {
            "model": "openai/gpt-4.1-mini",
            "messages": base_messages,
            "temperature": 0.2,
            "max_tokens": 8000,
        }
    )

    async with httpx.AsyncClient(timeout=300) as client:
        last_exc: Exception | None = None
        for index, candidate in enumerate(fallback_payloads, start=1):
            candidate = {k: v for k, v in candidate.items() if v is not None}
            try:
                resp = await client.post(
                    f"{settings.openrouter_base_url}/chat/completions",
                    headers={"Authorization": f"Bearer {settings.openrouter_api_key}"},
                    json=candidate,
                )
                resp.raise_for_status()
                return resp.json()["choices"][0]["message"]["content"]
            except httpx.HTTPStatusError as exc:
                last_exc = exc
                if exc.response.status_code == 400 and index < len(fallback_payloads):
                    logger.warning(
                        f"Renderer payload attempt {index} failed with 400; retrying with compatibility fallback"
                    )
                    continue
                raise
        if last_exc:
            raise last_exc
        raise RendererError("Renderer LLM call failed without a captured exception")


def _ensure_table_spacing(text: str) -> str:
    """Ensure blank lines before/after markdown tables so python-markdown renders them."""
    import re
    # Add blank line before table rows if missing
    text = re.sub(r'([^\n])\n(\|)', r'\1\n\n\2', text)
    # Add blank line after table block if missing
    text = re.sub(r'(\|[^\n]*\n)([^\|\n])', r'\1\n\2', text)
    return text


def _build_html(report: ReportOutput, chart_paths: list[str], lang: str = "en") -> str:
    """Build standalone HTML from report."""
    import markdown as md

    sections_html = ""
    for section in report.sections:
        content = _ensure_table_spacing(section.content)
        body = md.markdown(content, extensions=["tables", "fenced_code", "nl2br"])
        sources_links = "".join(
            f'<li><a href="{s}" target="_blank">{s}</a></li>' for s in section.sources
        )
        sources_block = f'<div class="sources"><strong>Sources</strong><ul>{sources_links}</ul></div>' if sources_links else ""
        sections_html += f"<section><h2>{section.title}</h2>{body}{sources_block}</section>\n"

    chart_images = ""
    for cp in chart_paths:
        fname = Path(cp).name
        chart_images += f'<figure><img src="charts/{fname}" alt="Chart" style="max-width:100%;border-radius:4px;"></figure>\n'

    exec_html = md.markdown(_ensure_table_spacing(report.executive_summary), extensions=["tables", "nl2br"])

    return f"""<!DOCTYPE html>
<html lang="{lang}">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{report.title}</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Georgia', serif; max-width: 960px; margin: 0 auto; padding: 48px 40px; color: #1a1a1a; line-height: 1.7; background: #fff; }}
  h1 {{ font-family: Arial, sans-serif; color: #003A70; font-size: 2em; border-bottom: 3px solid #0071CE; padding-bottom: 16px; margin-bottom: 24px; line-height: 1.3; }}
  h2 {{ font-family: Arial, sans-serif; color: #003A70; font-size: 1.35em; margin-top: 2.5em; margin-bottom: 0.75em; padding-bottom: 6px; border-bottom: 1px solid #cce0f5; }}
  h3 {{ font-family: Arial, sans-serif; color: #005A9C; font-size: 1.1em; margin-top: 1.5em; margin-bottom: 0.5em; }}
  p {{ margin-bottom: 1em; }}
  ul, ol {{ margin: 0.75em 0 1em 1.5em; }}
  li {{ margin-bottom: 0.4em; }}
  .executive-summary {{ background: #f0f6fd; padding: 28px 32px; border-left: 5px solid #0071CE; margin: 28px 0 40px; border-radius: 0 4px 4px 0; }}
  .executive-summary h3 {{ color: #003A70; margin-top: 0; margin-bottom: 12px; font-size: 1em; text-transform: uppercase; letter-spacing: 0.08em; }}
  table {{ border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 0.9em; font-family: Arial, sans-serif; }}
  thead th {{ background: #003A70; color: #fff; padding: 10px 14px; text-align: left; font-weight: 600; }}
  tbody tr:nth-child(even) {{ background: #f4f8fc; }}
  tbody tr:hover {{ background: #e8f0fa; }}
  td {{ border: 1px solid #d5e3f0; padding: 9px 14px; vertical-align: top; }}
  th {{ border: 1px solid #003A70; }}
  section {{ margin-bottom: 48px; }}
  .sources {{ margin-top: 20px; padding: 14px 18px; background: #f9f9f9; border-radius: 4px; font-size: 0.82em; }}
  .sources ul {{ margin: 6px 0 0 1.2em; }}
  .sources li {{ word-break: break-all; margin-bottom: 4px; }}
  a {{ color: #0071CE; text-decoration: none; }}
  a:hover {{ text-decoration: underline; }}
  figure {{ margin: 28px 0; text-align: center; }}
  strong {{ color: #003A70; }}
  code {{ background: #f0f4f8; padding: 2px 5px; border-radius: 3px; font-size: 0.88em; }}
  blockquote {{ border-left: 3px solid #0071CE; margin: 1em 0; padding: 0.5em 1em; color: #444; background: #f8fbff; }}
  @media print {{ body {{ padding: 20px; }} section {{ page-break-inside: avoid; }} }}
</style>
</head>
<body>
<h1>{report.title}</h1>
<div class="executive-summary"><h3>Executive Summary</h3>{exec_html}</div>
{chart_images}
{sections_html}
</body>
</html>"""


def _build_pdf(html_content: str, output_path: str) -> str:
    """Generate PDF from HTML using WeasyPrint."""
    from weasyprint import HTML

    HTML(string=html_content).write_pdf(output_path)
    logger.info(f"PDF saved: {output_path}")
    return output_path


def _build_docx(report: ReportOutput, chart_paths: list[str], output_path: str) -> str:
    """Generate DOCX from report using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Title"]
    style.font.color.rgb = RGBColor(0, 0x3A, 0x70)

    doc.add_heading(report.title, level=0)

    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph(report.executive_summary)
    p.style.font.size = Pt(11)

    for cp in chart_paths:
        if os.path.exists(cp):
            doc.add_picture(cp, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in report.sections:
        doc.add_heading(section.title, level=1)

        for block in section.content.split("\n\n"):
            text = block.strip()
            if not text:
                continue

            lines = text.split("\n")
            # Detect markdown table (starts with |)
            if lines[0].startswith("|") and len(lines) >= 2:
                # Parse header and rows (skip separator line)
                rows = [l for l in lines if l.strip() and not set(l.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set()]
                if rows:
                    cols = [c.strip() for c in rows[0].strip("|").split("|")]
                    table = doc.add_table(rows=len(rows), cols=len(cols))
                    table.style = "Table Grid"
                    for ri, row_text in enumerate(rows):
                        cells = [c.strip() for c in row_text.strip("|").split("|")]
                        for ci, cell_text in enumerate(cells[:len(cols)]):
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_text
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                continue

            if text.startswith("# "):
                doc.add_heading(text.lstrip("# ").strip(), level=2)
            elif text.startswith("## "):
                doc.add_heading(text.lstrip("## ").strip(), level=3)
            elif text.startswith("### "):
                doc.add_heading(text.lstrip("### ").strip(), level=4)
            elif lines[0].startswith("- ") or lines[0].startswith("* "):
                for line in lines:
                    line = line.lstrip("-* ").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(text)

        if section.sources:
            doc.add_heading("Sources", level=2)
            for src in section.sources:
                doc.add_paragraph(src, style="List Bullet")

    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")
    return output_path


def _write_synthesis_artifacts(out_dir: str, context: dict) -> dict[str, str]:
    claim_path = os.path.join(out_dir, "claim_table.json")
    synthesis_path = os.path.join(out_dir, "synthesis.md")

    claim_table = list(context.get("claim_table", []) or [])
    with open(claim_path, "w", encoding="utf-8") as f:
        json.dump(claim_table, f, ensure_ascii=False, indent=2)

    payload = dict(context.get("synthesis_payload") or {})
    metrics = dict(payload.get("metrics") or {})
    lines = [
        "# Synthesis Gate",
        "",
        f"- ready: `{bool(payload.get('ready'))}`",
        f"- allow_recommendations: `{bool(payload.get('allow_recommendations'))}`",
        f"- total_claims: `{int(metrics.get('total_claims', 0))}`",
        f"- verified_claims: `{int(metrics.get('verified_claims', 0))}`",
        f"- usable_claims: `{int(metrics.get('usable_claims', 0))}`",
        f"- verified_ratio: `{float(metrics.get('verified_ratio', 0.0)):.2f}`",
        f"- low_authority_ratio: `{float(metrics.get('low_authority_ratio', 0.0)):.2f}`",
        "",
        "## Blocking Reasons",
    ]
    blockers = list(payload.get("blocking_reasons") or [])
    if blockers:
        lines.extend([f"- {item}" for item in blockers])
    else:
        lines.append("- none")

    lines.extend(["", "## Top Claims", ""])
    for claim in list(payload.get("top_claims") or [])[:12]:
        text = _sanitize_signal(claim.get("claim", ""), max_len=240) or ""
        if not text:
            continue
        lines.append(
            f"- ({claim.get('citation_status', 'UNVERIFIED')}, score={float(claim.get('claim_score', 0.0)):.2f}) {text}"
        )

    with open(synthesis_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines).strip() + "\n")

    return {"claim_table_path": claim_path, "synthesis_path": synthesis_path}


@traceable(name="renderer")
async def run_renderer(state: AgentState) -> dict:
    logger.info("Renderer agent started")
    model = get_model(AgentTask.RENDERING)
    session_id = state.get("session_id", state.get("report_id", "default"))
    out_dir = os.path.join(settings.outputs_dir, session_id)
    os.makedirs(out_dir, exist_ok=True)

    intake = state.get("intake_result")
    chart_paths = state.get("chart_paths", [])
    lang = intake.language if intake else "en"
    context = _build_render_context(state)
    synthesis_artifacts = _write_synthesis_artifacts(out_dir, context)

    context_json = json.dumps(context, default=str)
    raw = await _call_llm(_get_renderer_prompt(lang), context_json, model)
    if not raw or not raw.strip():
        raise RendererError("LLM returned empty content")
    try:
        parsed = parse_llm_json(raw, context="renderer")
    except ValueError as exc:
        raise RendererError(str(exc)) from exc

    sections = _normalize_rendered_sections(parsed.get("sections", []), context.get("section_packets", []))
    sections = _apply_recommendation_policy(
        sections,
        allow_recommendations=bool(context.get("allow_recommendations", True)),
    )
    sections = _dedupe_and_prune_sections(sections)
    metadata = _build_orchestration_metadata(state, context)
    metadata["process_artifacts"] = synthesis_artifacts
    report = ReportOutput(
        id=state.get("report_id", ""),
        title=parsed.get("title", "Untitled Report"),
        executive_summary=_compose_executive_summary(
            context.get("executive_summary_packet", {}),
            parsed.get("executive_summary", ""),
        ),
        sections=sections,
        total_cost_usd=state.get("cost_usd", 0),
        metadata=metadata,
    )

    html_content = _build_html(report, chart_paths, lang=lang)
    html_path = os.path.join(out_dir, "report.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    logger.info(f"HTML saved: {html_path}")

    output_paths = [html_path]

    try:
        pdf_path = _build_pdf(html_content, os.path.join(out_dir, "report.pdf"))
        output_paths.append(pdf_path)
    except Exception as e:
        logger.warning(f"PDF generation failed, skipping optional PDF output: {e}")

    try:
        docx_path = _build_docx(report, chart_paths, os.path.join(out_dir, "report.docx"))
        output_paths.append(docx_path)
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")

    cost = estimate_cost(AgentTask.RENDERING, len(context_json) // 4, len(raw) // 4)
    logger.info(f"Renderer produced report with {len(sections)} sections, {len(output_paths)} files")

    return {
        "report": report,
        "final_report_paths": output_paths,
        "cost_usd": state.get("cost_usd", 0) + cost,
        "current_agent": "renderer",
        "status": "rendering",
    }
