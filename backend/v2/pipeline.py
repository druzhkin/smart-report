from __future__ import annotations

import asyncio
import io
import json
import re
import tempfile
from collections import defaultdict
from pathlib import Path
from typing import Awaitable, Callable

import httpx

from backend.agents.research_agent import _research_single
from backend.agents.presentation_agent import (
    _build_markdown as _build_presentation_markdown,
    _gamma_create,
    _generate_slide_json,
    _parse_slides_json,
)
from backend.agents.renderer import _build_docx as _build_rich_docx
from backend.agents.renderer import _build_html as _build_rich_html
from backend.agents.renderer import _call_llm as _call_renderer_llm
from backend.config import settings
from backend.pipeline.model_router import AgentTask, estimate_cost, estimate_cost_for_model, get_model
from backend.schemas.report_schema import ReportOutput, ReportSection, ReportStatus
from backend.utils.json_parse import parse_llm_json
from backend.v2.audit import audit_report_package
from backend.v2.grounding import (
    detect_contradictions as _detect_grounded_contradictions,
    extract_numeric_facts,
    find_unsupported_precise_numbers,
    sanitize_unsupported_precise_numbers,
)
from backend.v2.intake import build_depth_profile, build_request_spec
from backend.v2.materials import load_material_text
from backend.v2.models import (
    AnalysisBrief,
    AdjacentQuestionCandidate,
    ArtifactFormat,
    AuditSummary,
    CritiqueFinding,
    CritiqueKind,
    CoverageQuestionStatus,
    CoverageReport,
    DecisionTrigger,
    DepthProfile,
    MaterialRecord,
    PerplexityHandoffPrompt,
    ResearchPlan,
    ResearchQuestion,
    QuestionKind,
    RunEvent,
    RunStatus,
    RunSummary,
    SourceSnapshot,
    SourceLedgerEntry,
    SourceType,
    SpendCategory,
    SpendEntry,
    TaskSpec,
    ClaimRecord,
    EvidenceRecord,
    QualityAssessment,
    QualityIteration,
)
from backend.v2.quality import assess_report_quality, build_quality_iteration, build_revision_focus
from backend.v2.reference_data import match_reference_pack
from backend.v2.repository import FileRunRepository
from backend.v2.search import (
    DuckDuckGoSearchProvider,
    SeededSearchProvider,
    SearchProvider,
    classify_source_type,
    score_source,
    select_sources,
)


EmitFn = Callable[[RunEvent], Awaitable[None]]

LIVE_REPORT_MODEL = "anthropic/claude-sonnet-4"
REVIEW_MODEL_DIRECT = "sonar-pro"
REVIEW_MODEL_ROUTED = "perplexity/sonar-pro"
PERPLEXITY_CHAT_URL = "https://api.perplexity.ai/chat/completions"
QUALITY_REVISION_TARGET = 3
QUALITY_MAX_REVISION_ROUNDS = 4
QUALITY_MIN_IMPROVEMENT_DELTA = 0.25
QUALITY_HARMFUL_DELTA = -1.0
QUALITY_MAX_HARMFUL_ROUNDS = 2


def _write_audit_snapshot(run_id: str, audit: AuditSummary) -> None:
    audits_dir = Path(settings.reports_audits_dir)
    audits_dir.mkdir(parents=True, exist_ok=True)
    audits_dir.joinpath(f"{run_id}.json").write_text(
        json.dumps(audit.model_dump(mode="json"), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

LONGFORM_REPORT_PROMPT = """You are a BCG-grade research writer producing a client-ready decision report.
Return valid JSON with:
{{
  "title": "...",
  "subtitle": "...",
  "facts_line": "...",
  "executive_summary": "...",
  "sections": [
    {{"title": "...", "content": "...", "order": 1, "sources": ["https://..."]}}
  ]
}}

Mandatory requirements:
- Write the entire report in {language_name}.
- Produce a long-form report, not a memo: target 4,500-6,500 words total.
- Executive summary: 450-650 words with the decision, tradeoffs, and what to do next.
- Each section should usually land in the 350-650 word range unless the evidence base is genuinely thin.
- Produce exactly 8 substantial sections after the executive summary.
- Every section must be analytical, information-dense, and materially different from the others.
- Include 4-6 exhibits labeled "Exhibit 1", "Exhibit 2", etc. Qualitative comparison matrices count; do not force numeric tables where the evidence does not support them.
- Format exhibit labels as a plain line like `Exhibit 1` followed by a short subtitle line and then the table.
- Use only claims supported by the supplied evidence; if evidence is weak, say so directly.
- Do not invent benchmark deltas, uptime numbers, TCO figures, or precise percentages. If the supplied evidence does not support a number, use qualitative wording instead.
- Use precise numbers only from the supplied `numeric_claims` whitelist. Do not derive break-even points, staffing budgets, volume thresholds, or unit-economics tables unless those exact metrics are already supported there.
- Treat model version numbers such as `Claude 4.6` or `DeepSeek-V3.2` as labels, not evidence.
- Make the alternative space explicit. Do not behave as if the focal subject is the only serious option.
- Include at least one strong anti-thesis or counterargument against the current leading recommendation.
- Explain what could change the recommendation and which unknowns still need validation.
- Recommendation bullets must include evidence linkage either as [Evidence: C-01, C-02] or markdown source links.
- Use markdown bullets starting with `- `, not typographic bullets like `•`.
- Never cite internal payload fields or placeholders such as `[research_rows q2]`, `[critique]`, or `[request_spec]`.
- Never use bare numeric footnotes like `[1]` or synthetic markers like `[1 from q2]`; use markdown links or `[Evidence: C-01]`.
- Recommendation section must end with 5-7 concrete bullets and each bullet must carry evidence linkage.
- The roadmap section must be phased and operational, not generic.
- Keep the tone professional, sharp, and synthesis-heavy. Do not narrate the research process.
- Do not output placeholder text or meta commentary.
"""

QUALITY_REVISION_PROMPT = """You are revising specific parts of a client-ready analytical report to materially improve its value, depth, and decision usefulness.
Return valid JSON with:
{{
  "subtitle": "...",
  "facts_line": "...",
  "executive_summary": "...",
  "section_updates": [
    {{
      "target_title": "...",
      "title": "...",
      "content": "...",
      "sources": ["https://..."]
    }}
  ],
  "new_sections": [
    {{
      "title": "...",
      "content": "...",
      "sources": ["https://..."]
    }}
  ]
}}

Rules:
- Write all strings in {language_name}.
- Treat the current draft as a baseline to surpass, not defend.
- Revise only the executive summary when requested and only the listed target sections. Do not rewrite the full report.
- Improve the dimensions listed in `revision_focus` and address the stated weaknesses directly.
- Expand revised sections materially when `target_min_words` indicates the current section is too thin.
- Increase analytical density, comparative honesty, tradeoff clarity, boundary conditions, and decision usefulness.
- Do not introduce new precise numbers unless they are already supported by the provided claims or sources.
- Use the supplied `numeric_claims` whitelist as the only allowed source of exact metrics. If a number is not listed there, replace it with qualitative wording.
- Preserve valid evidence linkage in recommendation bullets and strengthen weak bullets rather than deleting them.
- Use markdown bullets starting with `- `, not typographic bullets like `•`.
- Format exhibits as plain lines like `Exhibit 1` followed by a short subtitle line and then the table.
- Never cite internal payload fields or placeholders such as `[research_rows q2]`, `[critique]`, or `[request_spec]`.
- Never use bare numeric footnotes like `[1]` or pseudo-citations like `[1 from q2]`.
- If a required section is missing, add it under `new_sections`.
- Do not narrate the revision process or mention scoring inside the report.
"""

LATERAL_REVIEW_PROMPT = """You are a skeptical principal analyst improving a decision-grade report before release.
Return valid JSON only with this shape:
{{
  "adjacent_questions": [
    {{
      "question": "...",
      "kind": "adjacent_alternative|adjacent_counterargument|adjacent_hidden_variable|adjacent_boundary|adjacent_stakeholder|adjacent_time_shift",
      "decision_impact": 0.0,
      "coverage_gap": 0.0,
      "novelty": 0.0,
      "comparative_value": 0.0,
      "research_cost": 0.0,
      "selection_reason": "..."
    }}
  ],
  "critique_findings": [
    {{
      "kind": "weak_evidence|omitted_question|missing_comparator|boundary_condition|decision_risk",
      "severity": "high|medium|low",
      "summary": "...",
      "rationale": "..."
    }}
  ],
  "decision_triggers": [
    {{
      "label": "...",
      "condition": "...",
      "implication": "...",
      "confidence": 0.0
    }}
  ]
}}

Rules:
- Write all strings in {language_name}.
- Propose 4-6 bounded adjacent questions, not an open-ended brainstorm.
- You must include at least one alternatives question and one counterargument question.
- Focus on what could materially change the recommendation, not generic curiosity.
- Critique findings should be sharp, concrete, and decision-relevant.
- Decision triggers should describe when the recommendation would switch.
- Stay close to the user's real objective; do not drift into unrelated research.
"""

LANGUAGE_NAMES = {
    "ru": "Russian",
    "en": "English",
}


def _language_name(language_code: str) -> str:
    return LANGUAGE_NAMES.get(language_code, language_code)


def _depth_profile(task_spec: TaskSpec) -> DepthProfile:
    return build_depth_profile(task_spec.request_spec.budget_tier)


def _estimate_tokens(text: str) -> int:
    normalized = " ".join(str(text or "").split()).strip()
    if not normalized:
        return 0
    return max(1, len(normalized) // 4)


def _make_spend_entry(
    *,
    category: SpendCategory,
    stage: str,
    provider: str,
    model: str,
    input_tokens: int = 0,
    output_tokens: int = 0,
    cost_usd: float | None = None,
    pricing_basis: str = "estimated",
    notes: str = "",
) -> SpendEntry:
    return SpendEntry(
        category=category,
        stage=stage,
        provider=provider,
        model=model,
        input_tokens=input_tokens,
        output_tokens=output_tokens,
        cost_usd=round(cost_usd if cost_usd is not None else estimate_cost_for_model(model, input_tokens, output_tokens), 6),
        pricing_basis=pricing_basis,
        notes=notes,
    )


def _coerce_spend_result(
    result: str | tuple[str, SpendEntry],
    *,
    category: SpendCategory,
    stage: str,
    provider: str,
    model: str,
) -> tuple[str, SpendEntry]:
    if isinstance(result, tuple) and len(result) == 2 and isinstance(result[0], str) and isinstance(result[1], SpendEntry):
        return result[0], result[1]
    content = result[0] if isinstance(result, tuple) else result
    return str(content), _make_spend_entry(
        category=category,
        stage=stage,
        provider=provider,
        model=model,
        cost_usd=0.0,
        pricing_basis="stubbed",
        notes="Compatibility path for test stub or legacy monkeypatch",
    )


def _record_spend(entries: list[SpendEntry], entry: SpendEntry) -> None:
    entries.append(entry)


def _spend_totals(entries: list[SpendEntry]) -> tuple[float, int]:
    total_cost = round(sum(item.cost_usd for item in entries), 6)
    total_tokens = sum(item.input_tokens + item.output_tokens for item in entries)
    return total_cost, total_tokens


def _aggregate_research_spend(
    branch_meta: list[dict],
    *,
    category: SpendCategory,
    stage: str,
    fallback_provider: str,
    fallback_model: str,
    branch_count: int,
    notes: str,
) -> SpendEntry | None:
    if not branch_meta:
        return None
    providers = sorted({str(item.get("provider") or fallback_provider) for item in branch_meta})
    models = sorted({str(item.get("model") or fallback_model) for item in branch_meta})
    pricing_bases = sorted({str(item.get("pricing_basis") or "estimated") for item in branch_meta})
    input_tokens = sum(int(item.get("input_tokens") or 0) for item in branch_meta)
    output_tokens = sum(int(item.get("output_tokens") or 0) for item in branch_meta)
    total_cost = round(sum(float(item.get("cost_usd") or 0.0) for item in branch_meta), 6)
    provider = providers[0] if len(providers) == 1 else "mixed"
    model = models[0] if len(models) == 1 else ", ".join(models)
    pricing_basis = pricing_bases[0] if len(pricing_bases) == 1 else "mixed"
    provider_note = f"providers={', '.join(providers)}; branches={branch_count}"
    merged_notes = f"{notes}; {provider_note}" if notes else provider_note
    return _make_spend_entry(
        category=category,
        stage=stage,
        provider=provider,
        model=model or fallback_model,
        input_tokens=input_tokens,
        output_tokens=output_tokens,
        cost_usd=total_cost,
        pricing_basis=pricing_basis,
        notes=merged_notes,
    )


def _sanitize_report_grounding(parsed_report: dict, claim_texts: list[str]) -> dict:
    sanitized = json.loads(json.dumps(parsed_report, ensure_ascii=False))
    for field in ("subtitle", "facts_line", "executive_summary"):
        value = str(sanitized.get(field, "")).strip()
        if value:
            sanitized[field] = sanitize_unsupported_precise_numbers(value, claim_texts)
    sanitized_sections = []
    for section in sanitized.get("sections", []) or []:
        next_section = dict(section)
        content = str(next_section.get("content", "")).strip()
        if content:
            next_section["content"] = sanitize_unsupported_precise_numbers(content, claim_texts)
        sanitized_sections.append(next_section)
    sanitized["sections"] = sanitized_sections
    return sanitized


def _strip_unlinked_recommendation_bullets(report_markdown: str) -> str:
    section_titles = (
        "Recommendation and Decision Posture",
        "Рекомендация и управленческая позиция",
    )
    cleaned = report_markdown
    for title in section_titles:
        pattern = rf"(## {re.escape(title)}\n\n)(.*?)(?=\n## |\Z)"
        match = re.search(pattern, cleaned, flags=re.DOTALL)
        if not match:
            continue
        body = match.group(2)
        rewritten_lines: list[str] = []
        for line in body.splitlines():
            stripped = line.strip()
            if stripped.startswith("- ") and "[Evidence:" not in stripped and "](" not in stripped:
                rewritten_lines.append(line.replace("- ", "", 1))
            else:
                rewritten_lines.append(line)
        cleaned = cleaned[: match.start(2)] + "\n".join(rewritten_lines) + cleaned[match.end(2) :]
    return cleaned


def _final_markdown_compliance_cleanup(report_markdown: str, claim_texts: list[str]) -> str:
    cleaned = sanitize_unsupported_precise_numbers(report_markdown, claim_texts)
    cleaned = re.sub(
        r"(?i)\bNet Promoter Score above\s+\d+(?:[.,]\d+)?\b",
        "Net Promoter Score above a strong satisfaction threshold",
        cleaned,
    )
    cleaned = _strip_unlinked_recommendation_bullets(cleaned)
    return cleaned


def _budget_to_research_depth(task_spec: TaskSpec) -> str:
    return _depth_profile(task_spec).research_depth


def _live_section_titles(language: str) -> list[str]:
    if language == "ru":
        return [
            "Контекст решения и почему вопрос важен сейчас",
            "Текущее состояние рынка и ключевые сигналы",
            "Сравнительный бенчмарк и пространство альтернатив",
            "Экономика решения и операционная модель",
            "Паттерны внедрения и кейсы",
            "Риски, ограничения и контраргументы",
            "Рекомендация и управленческая позиция",
            "Дорожная карта внедрения",
        ]
    return [
        "Decision Context and Why This Matters Now",
        "Baseline State and Market Signals",
        "Comparative Benchmark and Option Space",
        "Economics and Operating Model Implications",
        "Implementation Patterns and Case Studies",
        "Risks, Constraints, and Counterarguments",
        "Recommendation and Decision Posture",
        "Phased Roadmap",
    ]


def _budget_to_adjacent_question_limit(task_spec: TaskSpec) -> int:
    return _depth_profile(task_spec).adjacent_question_limit


def _has_any_token(values: list[str], tokens: tuple[str, ...]) -> bool:
    haystack = " ".join(values).lower()
    return any(token in haystack for token in tokens)


_TOPIC_STOPWORDS = {
    "about",
    "across",
    "after",
    "against",
    "agent",
    "agents",
    "analysis",
    "analyze",
    "architecture",
    "boundaries",
    "build",
    "buy",
    "candidate",
    "candidates",
    "choose",
    "combining",
    "compare",
    "comparison",
    "concrete",
    "consumer",
    "decision",
    "default",
    "design",
    "detailed",
    "evaluate",
    "evidence",
    "explicit",
    "first",
    "grade",
    "iteration",
    "managed",
    "must",
    "next",
    "operator",
    "outperform",
    "platform",
    "product",
    "products",
    "quality",
    "question",
    "questions",
    "real",
    "really",
    "recommend",
    "recommendation",
    "report",
    "researcher",
    "smart",
    "should",
    "stack",
    "stacks",
    "support",
    "system",
    "systems",
    "target",
    "that",
    "their",
    "this",
    "those",
    "traceability",
    "using",
    "what",
    "which",
    "with",
    "workflow",
    "workflows",
    "your",
}


def _topic_tokens(*texts: str) -> set[str]:
    tokens: set[str] = set()
    for text in texts:
        normalized = re.sub(r"[_./:-]+", " ", str(text or "").lower())
        for token in re.findall(r"[a-zа-я0-9+]{3,}", normalized):
            if token.isdigit() or token in _TOPIC_STOPWORDS:
                continue
            tokens.add(token)
    return tokens


def _task_topic_tokens(task_spec: TaskSpec, plan: ResearchPlan | None = None) -> set[str]:
    texts = [
        task_spec.request_spec.original_query,
        task_spec.request_spec.subject,
        task_spec.request_spec.decision_context,
        *task_spec.must_cover_questions[:4],
    ]
    if plan is not None:
        texts.extend(question.question for question in plan.primary_questions[:4])
        texts.extend(question.question for question in plan.selected_adjacent_questions[:4])
    return _topic_tokens(*texts)


_BUSINESS_TOPIC_TOKENS = (
    "market",
    "pricing",
    "price",
    "revenue",
    "monetization",
    "paid product",
    "subscription",
    "gtm",
    "go-to-market",
    "buyer",
    "customer",
    "consulting",
    "investment",
    "strategy team",
    "willingness to pay",
    "roi",
    "procurement",
    "freemium",
    "premium",
    "package",
    "packaging",
    "sales-led",
    "product-led",
    "рын",
    "монетиз",
    "цен",
    "выруч",
    "подпис",
    "платн",
    "клиент",
    "покуп",
    "спрос",
)

_BUSINESS_SOURCE_SIGNAL_TOKENS = (
    "pricing",
    "price",
    "monetization",
    "business model",
    "revenue",
    "subscription",
    "freemium",
    "premium",
    "go-to-market",
    "gtm",
    "packaging",
    "procurement",
    "buyer",
    "customer",
    "consulting",
    "investment",
    "roi",
    "tco",
    "sales",
    "willingness",
    "market",
    "adoption",
    "pricing strategy",
    "стоим",
    "цена",
    "монетиз",
    "бюджет",
    "выруч",
    "клиент",
    "покуп",
    "закуп",
)

_BUSINESS_Q2_SOURCE_SIGNAL_TOKENS = (
    "alternative",
    "alternatives",
    "competitor",
    "competitive",
    "market intelligence",
    "research platform",
    "research tool",
    "investment research",
    "consulting workflow",
    "perplexity",
    "alphasense",
    "cb insights",
    "pitchbook",
    "tegus",
    "hebbia",
    "gwi",
    "crunchbase",
    "free tool",
    "free substitute",
    "incumbent workflow",
    "buyer",
    "adoption",
    "switching",
    "quality",
    "cost",
    "операцион",
    "альтернатив",
    "конкурент",
    "workflow",
    "замен",
    "бесплат",
)

_BUSINESS_Q2_LOW_SIGNAL_PATTERNS = (
    "cost estimating",
    "assessment guide",
    "acquisition guide",
    "project governance",
    "program management",
)

_AMBIGUOUS_SMART_REPORT_PATTERNS = (
    "smart reporting",
    "smart reports",
    "purpose of smart reports",
    "bmc helix",
    "siebel",
    "oracle",
    "itsm",
)


def _is_business_topic_task(task_spec: TaskSpec) -> bool:
    if _is_stack_research_topic_query(task_spec.request_spec.original_query):
        return False
    texts = [
        task_spec.request_spec.original_query,
        task_spec.request_spec.subject,
        task_spec.request_spec.decision_context,
        *task_spec.must_cover_questions[:4],
    ]
    haystack = " ".join(texts).lower()
    return any(token in haystack for token in _BUSINESS_TOPIC_TOKENS)


def _business_source_signal_bonus(entry: SourceLedgerEntry, question_id: str, *, is_business_topic: bool) -> float:
    if not is_business_topic or question_id not in {"q1", "q2", "q3", "q4"}:
        return 0.0
    haystack = re.sub(r"[_./:-]+", " ", f"{entry.title} {entry.domain} {entry.url}".lower())
    signal_tokens = _BUSINESS_Q2_SOURCE_SIGNAL_TOKENS if question_id == "q2" else _BUSINESS_SOURCE_SIGNAL_TOKENS
    matches = sum(1 for token in signal_tokens if token in haystack)
    strong_types = {
        SourceType.RESEARCH_PAPER,
        SourceType.BENCHMARK,
        SourceType.USER_MATERIAL,
        SourceType.OFFICIAL_DOCUMENTATION,
    }
    if any(pattern in haystack for pattern in _AMBIGUOUS_SMART_REPORT_PATTERNS) and matches < 2:
        return -0.36 if question_id == "q2" else -0.28
    if question_id == "q2" and any(pattern in haystack for pattern in _BUSINESS_Q2_LOW_SIGNAL_PATTERNS) and matches < 2:
        return -0.28
    if question_id == "q2" and matches == 0 and entry.source_type in {SourceType.RESEARCH_PAPER, SourceType.GOVERNMENT, SourceType.OFFICIAL_DOCUMENTATION}:
        return -0.22
    if question_id == "q4" and matches == 0 and entry.source_type in {SourceType.RESEARCH_PAPER, SourceType.GOVERNMENT, SourceType.OFFICIAL_DOCUMENTATION}:
        return -0.16
    if matches >= 4:
        return 0.24 if question_id == "q2" else 0.18
    if matches >= 2:
        return 0.12 if question_id == "q2" else 0.1
    if entry.source_type in strong_types and matches == 1:
        return 0.05 if question_id == "q2" else 0.04
    if entry.source_type == SourceType.OFFICIAL_DOCUMENTATION:
        return -0.12 if question_id == "q2" else -0.08
    return 0.0


def _source_topic_alignment_score(entry: SourceLedgerEntry, topic_tokens: set[str]) -> float:
    if not topic_tokens:
        return 0.0
    haystack = re.sub(r"[_./:-]+", " ", f"{entry.title} {entry.domain} {entry.url}".lower())
    matched = sum(1 for token in topic_tokens if token in haystack)
    return min(1.0, matched / 3.0)


_Q4_SOURCE_SIGNAL_TOKENS = (
    "tradeoff",
    "trade-off",
    "risk",
    "pricing",
    "price",
    "cost",
    "roi",
    "budget",
    "willingness",
    "pay",
    "buyer",
    "procurement",
    "integration",
    "burden",
    "switch",
    "trigger",
    "condition",
    "boundary",
    "objection",
    "failure",
    "lock-in",
    "lock in",
    "latency",
    "reliability",
    "governance",
    "compliance",
    "free",
    "paid",
    "tco",
    "overhead",
    "rate limit",
    "citation",
    "hallucination",
    "self-host",
    "self host",
    "maintenance",
    "migration",
    "adoption",
    "стоим",
    "цена",
    "бюджет",
    "окуп",
    "риск",
    "возраж",
    "закуп",
    "интеграц",
    "барьер",
    "нагруз",
    "переключ",
    "услови",
    "огранич",
    "плат",
    "бесплат",
)

_Q4_LOW_SIGNAL_PATTERNS = (
    "decision-grade data",
    "operational excellence",
    "trusted data",
    "reinvention",
    "data foundation",
    "quality data",
)


def _question_source_signal_bonus(entry: SourceLedgerEntry, question_id: str) -> float:
    if question_id != "q4":
        return 0.0
    haystack = re.sub(r"[_./:-]+", " ", f"{entry.title} {entry.domain} {entry.url}".lower())
    matches = sum(1 for token in _Q4_SOURCE_SIGNAL_TOKENS if token in haystack)
    strong_types = {
        SourceType.OFFICIAL_DOCUMENTATION,
        SourceType.BENCHMARK,
        SourceType.RESEARCH_PAPER,
        SourceType.USER_MATERIAL,
    }
    if any(pattern in haystack for pattern in _Q4_LOW_SIGNAL_PATTERNS) and matches < 2:
        return -0.18
    if matches >= 4:
        return 0.18
    if matches >= 2:
        return 0.12
    if matches == 1 and entry.source_type in strong_types:
        return 0.04
    return -0.08


def _rank_live_sources(
    source_entries: list[SourceLedgerEntry],
    task_spec: TaskSpec,
    plan: ResearchPlan,
    *,
    limit: int = 16,
) -> list[SourceLedgerEntry]:
    topic_tokens = _task_topic_tokens(task_spec, plan)
    is_business_topic = _is_business_topic_task(task_spec)
    blocked_types = set(task_spec.blocked_source_types)
    type_caps = {
        SourceType.OFFICIAL_DOCUMENTATION: 8,
        SourceType.BENCHMARK: 4,
        SourceType.RESEARCH_PAPER: 3,
        SourceType.USER_MATERIAL: 4,
        SourceType.HIGH_QUALITY_SECONDARY: 3,
        SourceType.VENDOR_PAGE: 4,
        SourceType.GOVERNMENT: 2,
        SourceType.WEAK_SECONDARY: 0,
    }
    ranked: list[tuple[float, SourceLedgerEntry]] = []
    topical_candidates: list[tuple[float, SourceLedgerEntry]] = []
    best_by_primary_question: dict[str, tuple[float, SourceLedgerEntry]] = {}
    for entry in source_entries:
        topicality = _source_topic_alignment_score(entry, topic_tokens)
        question_bonus = max(
            (
                _question_source_signal_bonus(entry, question_id)
                + _business_source_signal_bonus(entry, question_id, is_business_topic=is_business_topic)
                for question_id in entry.question_links
            ),
            default=0.0,
        )
        score = entry.reliability_score + topicality * 0.28 + min(0.08, 0.03 * len(entry.question_links)) + question_bonus
        if entry.source_type == SourceType.WEAK_SECONDARY:
            score -= 0.35
        ranked.append((score, entry))
        if topicality >= 0.15 or entry.source_type in {SourceType.OFFICIAL_DOCUMENTATION, SourceType.BENCHMARK, SourceType.USER_MATERIAL}:
            topical_candidates.append((score, entry))
        if entry.source_type in blocked_types or entry.source_type == SourceType.WEAK_SECONDARY:
            continue
        for question in plan.primary_questions:
            if question.question_id not in entry.question_links:
                continue
            existing = best_by_primary_question.get(question.question_id)
            if existing is None or score > existing[0]:
                best_by_primary_question[question.question_id] = (score, entry)
    pool = topical_candidates if len(topical_candidates) >= min(8, max(4, limit // 2)) else ranked
    pool.sort(key=lambda item: item[0], reverse=True)
    selected: list[SourceLedgerEntry] = []
    selected_urls: set[str] = set()
    domain_counts: dict[str, int] = {}
    type_counts: dict[SourceType, int] = {}
    for question in plan.primary_questions:
        anchored = best_by_primary_question.get(question.question_id)
        if anchored is None:
            continue
        entry = anchored[1]
        if entry.url in selected_urls:
            continue
        if domain_counts.get(entry.domain, 0) >= 3:
            continue
        if type_counts.get(entry.source_type, 0) >= type_caps.get(entry.source_type, limit):
            continue
        selected.append(entry)
        selected_urls.add(entry.url)
        domain_counts[entry.domain] = domain_counts.get(entry.domain, 0) + 1
        type_counts[entry.source_type] = type_counts.get(entry.source_type, 0) + 1
        if len(selected) >= limit:
            return selected
    for _, entry in pool:
        if entry.source_type in blocked_types:
            continue
        if entry.url in selected_urls:
            continue
        if domain_counts.get(entry.domain, 0) >= 3:
            continue
        if type_counts.get(entry.source_type, 0) >= type_caps.get(entry.source_type, limit):
            continue
        selected.append(entry)
        selected_urls.add(entry.url)
        domain_counts[entry.domain] = domain_counts.get(entry.domain, 0) + 1
        type_counts[entry.source_type] = type_counts.get(entry.source_type, 0) + 1
        if len(selected) >= limit:
            break
    if len(selected) < min(8, limit):
        for _, entry in ranked:
            if entry.url in selected_urls or entry.source_type in blocked_types:
                continue
            if domain_counts.get(entry.domain, 0) >= 3:
                continue
            selected.append(entry)
            selected_urls.add(entry.url)
            domain_counts[entry.domain] = domain_counts.get(entry.domain, 0) + 1
            if len(selected) >= limit:
                break
    return selected


def _is_stack_research_topic_query(query: str) -> bool:
    lowered = (query or "").lower()
    return (
        any(token in lowered for token in ("llm", "model", "models", "gpt", "claude", "qwen", "deepseek", "gemini"))
        and any(token in lowered for token in ("github", "open-source", "opensource", "framework", "frameworks", "langchain", "haystack", "llamaindex", "orchestration"))
        and any(token in lowered for token in ("search", "research", "retrieval", "rag", "traceability", "revision"))
    )


def _stack_research_guardrail(question_id: str, angle: str) -> str:
    lowered = angle.lower()
    if question_id == "q2" or any(token in lowered for token in ("github", "repo", "repos", "project", "projects")):
        return (
            "Focus on repositories and first-party documentation directly relevant to web research, retrieval, citation, "
            "and long-form report orchestration. Exclude generic workflow engines such as Airflow, Temporal, or Conductor "
            "unless the source proves they are used as part of a deep-research agent stack."
        )
    if question_id == "q3" or any(token in lowered for token in ("architecture", "perplexity", "traceability", "controllability")):
        return (
            "Do not answer with standalone observability vendors, playground tools, or router products unless the source "
            "shows how they fit into an end-to-end report architecture. Prefer component-level evidence for search or extraction, "
            "stateful orchestration, tracing or evaluation, and report-synthesis or audit controls."
        )
    if question_id == "q4" or any(token in lowered for token in ("tradeoff", "risk", "failure", "condition", "switch")):
        return (
            "Tie tradeoffs and failure modes to concrete stack components and build-vs-buy boundaries. Reject generic AI-governance "
            "or data-architecture literature unless it materially changes the recommendation."
        )
    return ""


def _build_stack_backfill_queries(task_spec: TaskSpec) -> list[tuple[str, str]]:
    if not _is_stack_research_topic_query(task_spec.request_spec.original_query):
        return []
    return [
        ("q1", "Claude Opus 4.6 official docs pricing benchmark"),
        ("q1", "DeepSeek V3 official docs pricing"),
        ("q2", "gpt-researcher github official docs"),
        ("q2", "open_deep_research github official docs"),
        ("q3", "Tavily docs extract crawl pricing"),
        ("q3", "LangGraph docs durable execution human in the loop"),
        ("q3", "Langfuse docs traces evaluations self hosting"),
        ("q3", "Firecrawl docs markdown extraction search api"),
        ("q4", "Perplexity API pricing docs"),
        ("q4", "AWS agentic AI frameworks LangGraph tradeoffs"),
        ("q4", "Langfuse self-hosting docs observability"),
        ("q4", "Deep Research Bench leaderboard"),
    ]


def _build_business_backfill_queries(
    task_spec: TaskSpec,
    target_question_ids: set[str] | None = None,
) -> list[tuple[str, str]]:
    if not _is_business_topic_task(task_spec):
        return []
    query_specs = [
        ("q3", "AI pricing and monetization playbook SaaS packaging go to market"),
        ("q3", "freemium vs subscription pricing analytics software professional services"),
        ("q4", "pricing objections consulting software procurement switching costs"),
        ("q4", "willingness to pay enterprise analytics software ROI TCO"),
        ("q1", "professional services software monetization consulting investment market demand"),
        ("q1", "analytics reporting software buyer behavior consulting investment teams"),
        ("q2", "AlphaSense Tegus PitchBook CB Insights Perplexity Hebbia alternatives consulting investment teams market intelligence"),
    ]
    if not target_question_ids:
        return query_specs
    prioritized = [item for item in query_specs if item[0] in target_question_ids]
    return prioritized


def _coverage_gap_question_ids(coverage: CoverageReport) -> set[str]:
    return {
        item.question_id
        for item in coverage.questions
        if item.status != "covered"
    }


def _adjacent_candidate_score(candidate: AdjacentQuestionCandidate) -> float:
    score = (
        0.35 * candidate.decision_impact
        + 0.25 * candidate.coverage_gap
        + 0.20 * candidate.comparative_value
        + 0.15 * candidate.novelty
        - 0.15 * candidate.research_cost
    )
    return round(max(0.0, min(0.99, score)), 3)


def _clamp_score(value: object, default: float = 0.5) -> float:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        numeric = default
    return max(0.0, min(1.0, numeric))


def _normalize_question_kind(raw: object) -> QuestionKind:
    text = str(raw or "").strip().lower()
    mapping = {
        "adjacent_alternative": QuestionKind.ADJACENT_ALTERNATIVE,
        "alternative": QuestionKind.ADJACENT_ALTERNATIVE,
        "adjacent_counterargument": QuestionKind.ADJACENT_COUNTERARGUMENT,
        "counterargument": QuestionKind.ADJACENT_COUNTERARGUMENT,
        "adjacent_hidden_variable": QuestionKind.ADJACENT_HIDDEN_VARIABLE,
        "hidden_variable": QuestionKind.ADJACENT_HIDDEN_VARIABLE,
        "adjacent_boundary": QuestionKind.ADJACENT_BOUNDARY,
        "boundary": QuestionKind.ADJACENT_BOUNDARY,
        "adjacent_stakeholder": QuestionKind.ADJACENT_STAKEHOLDER,
        "stakeholder": QuestionKind.ADJACENT_STAKEHOLDER,
        "adjacent_time_shift": QuestionKind.ADJACENT_TIME_SHIFT,
        "time_shift": QuestionKind.ADJACENT_TIME_SHIFT,
    }
    return mapping.get(text, QuestionKind.ADJACENT_HIDDEN_VARIABLE)


def _normalize_critique_kind(raw: object) -> CritiqueKind:
    text = str(raw or "").strip().lower()
    mapping = {
        "weak_evidence": CritiqueKind.WEAK_EVIDENCE,
        "omitted_question": CritiqueKind.OMITTED_QUESTION,
        "missing_comparator": CritiqueKind.MISSING_COMPARATOR,
        "boundary_condition": CritiqueKind.BOUNDARY_CONDITION,
        "decision_risk": CritiqueKind.DECISION_RISK,
    }
    return mapping.get(text, CritiqueKind.DECISION_RISK)


def _normalize_severity(raw: object) -> str:
    text = str(raw or "").strip().lower()
    if text in {"high", "medium", "low"}:
        return text
    return "medium"


def build_adjacent_question_candidates(
    task_spec: TaskSpec,
    coverage: CoverageReport | None = None,
    claims: list[ClaimRecord] | None = None,
) -> list[AdjacentQuestionCandidate]:
    language = task_spec.request_spec.language
    subject = task_spec.request_spec.subject
    dimensions = task_spec.evaluation_dimensions[:3] or ["cost", "quality", "risk"]
    dimensions_text = ", ".join(dimensions)
    has_cost_axis = _has_any_token(dimensions, ("cost", "price", "roi", "эконом", "стоим", "марж", "окуп"))
    has_speed_axis = _has_any_token(dimensions, ("speed", "latency", "time", "rollout", "срок", "скорост"))
    has_reliability_axis = _has_any_token(dimensions, ("risk", "quality", "reliability", "governance", "надеж", "риск", "качеств", "управ"))
    gap_pressure = max(0.35, 1.0 - coverage.coverage_ratio) if coverage else 0.45
    contradiction_pressure = 0.15 if coverage and coverage.contradiction_count else 0.0
    claim_pressure = 0.08 if claims and len([item for item in claims if item.recommendation_safe]) < 3 else 0.0

    if language == "ru":
        templates = [
            (
                QuestionKind.ADJACENT_ALTERNATIVE,
                f"Какие реальные альтернативы {subject} нужно обязательно сравнить по осям {dimensions_text}, чтобы вывод был честным?",
                "Без явного option space итог легко превращается в адвокатскую записку про один объект.",
                0.96,
                max(0.55, gap_pressure),
                0.78,
                1.0,
                0.45,
            ),
            (
                QuestionKind.ADJACENT_COUNTERARGUMENT,
                f"В каких сценариях {subject} является слабым выбором или проигрывает альтернативам?",
                "Нужно заранее собрать strongest case against текущую гипотезу, иначе recommendation будет хрупкой.",
                0.93,
                max(0.52, gap_pressure + contradiction_pressure),
                0.82,
                0.88,
                0.4,
            ),
            (
                QuestionKind.ADJACENT_HIDDEN_VARIABLE,
                f"Какие скрытые факторы стоимости, latency, reliability, lock-in и операционной нагрузки могут изменить вывод по {subject}?",
                "Сильные решения часто ломаются не по headline-метрикам, а по скрытым переменным внедрения.",
                0.89,
                max(0.48, gap_pressure + (0.08 if has_cost_axis or has_reliability_axis else 0.0)),
                0.76,
                0.81,
                0.42,
            ),
            (
                QuestionKind.ADJACENT_BOUNDARY,
                f"При каких условиях и для каких команд текущая рекомендация по {subject} перестаёт быть верной?",
                "Нужно не только сказать, что выбирать, но и описать границы применимости вывода.",
                0.87,
                max(0.46, gap_pressure + contradiction_pressure + claim_pressure),
                0.74,
                0.77,
                0.36,
            ),
            (
                QuestionKind.ADJACENT_STAKEHOLDER,
                f"Какие сильные возражения по {subject} возникнут у CFO, CTO, procurement и операционной команды?",
                "В аналитическом документе важно увидеть конфликт интересов между ролями до внедрения.",
                0.8,
                max(0.4, gap_pressure),
                0.71,
                0.68,
                0.33,
            ),
            (
                QuestionKind.ADJACENT_TIME_SHIFT,
                f"Что в рынке или технологическом ландшафте вокруг {subject} может измениться в горизонте 6-12 месяцев и сдвинуть решение?",
                "В быстро меняющихся темах текущий winner может быстро устареть.",
                0.73,
                max(0.34, gap_pressure + (0.04 if has_speed_axis else 0.0)),
                0.67,
                0.6,
                0.28,
            ),
        ]
    else:
        templates = [
            (
                QuestionKind.ADJACENT_ALTERNATIVE,
                f"Which credible alternatives to {subject} must be compared across {dimensions_text} for the conclusion to be intellectually honest?",
                "Without an explicit option space, the report drifts into a brief for the focal subject rather than a decision document.",
                0.96,
                max(0.55, gap_pressure),
                0.78,
                1.0,
                0.45,
            ),
            (
                QuestionKind.ADJACENT_COUNTERARGUMENT,
                f"In which scenarios is {subject} the wrong choice or materially weaker than alternatives?",
                "The workflow should collect the strongest case against the current hypothesis before writing recommendations.",
                0.93,
                max(0.52, gap_pressure + contradiction_pressure),
                0.82,
                0.88,
                0.4,
            ),
            (
                QuestionKind.ADJACENT_HIDDEN_VARIABLE,
                f"What hidden variables around cost, latency, reliability, lock-in, and operating burden could change the decision on {subject}?",
                "Strong-looking options often fail on hidden implementation variables rather than headline benchmarks.",
                0.89,
                max(0.48, gap_pressure + (0.08 if has_cost_axis or has_reliability_axis else 0.0)),
                0.76,
                0.81,
                0.42,
            ),
            (
                QuestionKind.ADJACENT_BOUNDARY,
                f"Under what conditions does the current recommendation for {subject} stop being valid?",
                "A useful decision report needs applicability boundaries, not only a central recommendation.",
                0.87,
                max(0.46, gap_pressure + contradiction_pressure + claim_pressure),
                0.74,
                0.77,
                0.36,
            ),
            (
                QuestionKind.ADJACENT_STAKEHOLDER,
                f"What strong objections would CFO, CTO, procurement, and operations raise about {subject}?",
                "Decision-grade analysis should surface stakeholder tension before rollout rather than after.",
                0.8,
                max(0.4, gap_pressure),
                0.71,
                0.68,
                0.33,
            ),
            (
                QuestionKind.ADJACENT_TIME_SHIFT,
                f"What could change around {subject} over the next 6-12 months and shift the recommendation?",
                "In fast-moving categories, a current winner can become stale quickly.",
                0.73,
                max(0.34, gap_pressure + (0.04 if has_speed_axis else 0.0)),
                0.67,
                0.6,
                0.28,
            ),
        ]

    candidates: list[AdjacentQuestionCandidate] = []
    for kind, question, reason, decision_impact, coverage_gap, novelty, comparative_value, research_cost in templates:
        candidate = AdjacentQuestionCandidate(
            question=question,
            kind=kind,
            decision_impact=decision_impact,
            coverage_gap=coverage_gap,
            novelty=novelty,
            comparative_value=comparative_value,
            research_cost=research_cost,
            selection_reason=reason,
        )
        candidate.composite_score = _adjacent_candidate_score(candidate)
        candidates.append(candidate)
    return candidates


def _merge_adjacent_candidates(
    model_candidates: list[AdjacentQuestionCandidate],
    heuristic_candidates: list[AdjacentQuestionCandidate],
) -> list[AdjacentQuestionCandidate]:
    merged: list[AdjacentQuestionCandidate] = []
    seen_questions: set[str] = set()
    for candidate in sorted(model_candidates + heuristic_candidates, key=lambda item: item.composite_score, reverse=True):
        normalized = " ".join(candidate.question.lower().split())
        if not normalized or normalized in seen_questions:
            continue
        seen_questions.add(normalized)
        merged.append(candidate)
    return merged


def select_adjacent_questions(task_spec: TaskSpec, candidates: list[AdjacentQuestionCandidate]) -> list[ResearchQuestion]:
    limit = _budget_to_adjacent_question_limit(task_spec)
    if limit <= 0 or not candidates:
        return []

    by_kind: dict[QuestionKind, AdjacentQuestionCandidate] = {}
    for candidate in sorted(candidates, key=lambda item: item.composite_score, reverse=True):
        by_kind.setdefault(candidate.kind, candidate)

    required_kinds = [
        QuestionKind.ADJACENT_ALTERNATIVE,
        QuestionKind.ADJACENT_COUNTERARGUMENT,
    ]
    if limit >= 4:
        required_kinds.extend(
            [
                QuestionKind.ADJACENT_HIDDEN_VARIABLE,
                QuestionKind.ADJACENT_BOUNDARY,
            ]
        )
    if limit >= 5:
        required_kinds.append(QuestionKind.ADJACENT_STAKEHOLDER)
    if limit >= 6:
        required_kinds.append(QuestionKind.ADJACENT_TIME_SHIFT)

    selected_candidates: list[AdjacentQuestionCandidate] = []
    used_ids: set[str] = set()
    for kind in required_kinds:
        candidate = by_kind.get(kind)
        if candidate and candidate.candidate_id not in used_ids and len(selected_candidates) < limit:
            selected_candidates.append(candidate)
            used_ids.add(candidate.candidate_id)

    for candidate in sorted(candidates, key=lambda item: item.composite_score, reverse=True):
        if candidate.candidate_id in used_ids or len(selected_candidates) >= limit:
            continue
        selected_candidates.append(candidate)
        used_ids.add(candidate.candidate_id)

    return [
        ResearchQuestion(
            question_id=f"aq{index}",
            question=candidate.question,
            kind=candidate.kind,
            priority=index,
            required_evidence_count=1,
        )
        for index, candidate in enumerate(selected_candidates, start=1)
    ]


def build_critique_findings(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    claims: list[ClaimRecord],
    coverage: CoverageReport,
    adjacent_questions: list[ResearchQuestion],
) -> list[CritiqueFinding]:
    findings: list[CritiqueFinding] = []
    weak_claims = [claim for claim in claims[:10] if claim.confidence < 0.78 or len(claim.source_ids) < 2]
    adjacent_ids = [question.question_id for question in adjacent_questions]

    if weak_claims:
        findings.append(
            CritiqueFinding(
                kind=CritiqueKind.WEAK_EVIDENCE,
                severity="high" if len(weak_claims) >= 3 else "medium",
                summary=(
                    "The current draft still relies on thinly supported claims and should not finalize an unqualified recommendation."
                    if task_spec.request_spec.language != "ru"
                    else "Текущий draft всё ещё опирается на тонко подтверждённые claims и не должен заканчивать безусловной рекомендацией."
                ),
                rationale=(
                    "Several top claims have either low confidence or single-source support."
                    if task_spec.request_spec.language != "ru"
                    else "Несколько верхних claims имеют либо низкую уверенность, либо опираются на один источник."
                ),
                affected_claim_ids=[claim.claim_id for claim in weak_claims[:5]],
                follow_up_question_ids=adjacent_ids[:2],
            )
        )

    if coverage.gaps:
        findings.append(
            CritiqueFinding(
                kind=CritiqueKind.OMITTED_QUESTION,
                severity="high",
                summary=(
                    f"At least {len(coverage.gaps)} core questions remain under-covered and need explicit closure."
                    if task_spec.request_spec.language != "ru"
                    else f"Как минимум {len(coverage.gaps)} ключевых вопросов остаются недопокрытыми и требуют явного закрытия."
                ),
                rationale="; ".join(coverage.gaps[:3]),
                follow_up_question_ids=adjacent_ids[:3],
            )
        )

    alternative_question = next(
        (question for question in adjacent_questions if question.kind == QuestionKind.ADJACENT_ALTERNATIVE),
        None,
    )
    findings.append(
        CritiqueFinding(
            kind=CritiqueKind.MISSING_COMPARATOR,
            severity="high",
            summary=(
                "The report must make the option space explicit instead of treating the focal subject as the only serious path."
                if task_spec.request_spec.language != "ru"
                else "Отчёт должен явно показать пространство альтернатив, а не вести себя так, будто рассматриваемый объект является единственным серьёзным вариантом."
            ),
            rationale=(
                "Decision quality drops sharply when alternatives are implied rather than compared."
                if task_spec.request_spec.language != "ru"
                else "Качество решения резко падает, когда альтернативы подразумеваются, а не сравниваются явно."
            ),
            follow_up_question_ids=[alternative_question.question_id] if alternative_question else adjacent_ids[:1],
        )
    )

    if coverage.contradiction_count:
        findings.append(
            CritiqueFinding(
                kind=CritiqueKind.DECISION_RISK,
                severity="high",
                summary=(
                    "At least one contradiction cluster remains unresolved and can flip the final recommendation."
                    if task_spec.request_spec.language != "ru"
                    else "Как минимум один кластер противоречий остаётся неразрешённым и может перевернуть итоговую рекомендацию."
                ),
                rationale=(
                    f"Contradiction count: {coverage.contradiction_count}"
                    if task_spec.request_spec.language != "ru"
                    else f"Число противоречий: {coverage.contradiction_count}"
                ),
                affected_claim_ids=[claim.claim_id for claim in claims if claim.contradiction_notes][:4],
                follow_up_question_ids=adjacent_ids[:2],
            )
        )

    boundary_question = next(
        (question for question in adjacent_questions if question.kind == QuestionKind.ADJACENT_BOUNDARY),
        None,
    )
    findings.append(
        CritiqueFinding(
            kind=CritiqueKind.BOUNDARY_CONDITION,
            severity="medium",
            summary=(
                "The recommendation still needs explicit failure conditions and a statement of where it stops applying."
                if task_spec.request_spec.language != "ru"
                else "Рекомендации всё ещё нужны явные условия отказа и описание границ, где вывод перестаёт работать."
            ),
            rationale=(
                "A client-ready document should tell the reader not only what to do, but also when to switch strategy."
                if task_spec.request_spec.language != "ru"
                else "Клиентский документ должен говорить не только что делать, но и когда переключать стратегию."
            ),
            follow_up_question_ids=[boundary_question.question_id] if boundary_question else adjacent_ids[:1],
        )
    )

    return findings


def build_decision_triggers(task_spec: TaskSpec) -> list[DecisionTrigger]:
    language = task_spec.request_spec.language
    dimensions = task_spec.evaluation_dimensions[:4]
    lower_dims = [item.lower() for item in dimensions]
    triggers: list[DecisionTrigger] = []

    def add_trigger(label: str, condition: str, implication: str, confidence: float) -> None:
        triggers.append(
            DecisionTrigger(
                label=label,
                condition=condition,
                implication=implication,
                confidence=confidence,
            )
        )

    if _has_any_token(lower_dims, ("cost", "price", "roi", "стоим", "эконом", "окуп")):
        add_trigger(
            "Budget Dominance" if language != "ru" else "Доминирование бюджета",
            (
                "If total cost of ownership becomes the dominant filter, favor the lowest-burden option over the richest feature set."
                if language != "ru"
                else "Если совокупная стоимость владения становится главным фильтром, нужно предпочесть вариант с минимальной операционной нагрузкой, а не самый функционально насыщенный путь."
            ),
            (
                "The recommendation may shift toward a cheaper, simpler stack even if it is weaker on absolute quality."
                if language != "ru"
                else "Рекомендация может сместиться в сторону более дешёвого и простого стека, даже если он слабее по абсолютному качеству."
            ),
            0.74,
        )

    if _has_any_token(lower_dims, ("speed", "latency", "time", "rollout", "срок", "скорост")):
        add_trigger(
            "Speed-To-Value" if language != "ru" else "Скорость до ценности",
            (
                "If time-to-value and rollout speed dominate, retest whether a simpler packaged option beats the current recommendation."
                if language != "ru"
                else "Если доминируют time-to-value и скорость rollout, нужно перепроверить, не обгоняет ли текущую рекомендацию более простой и более упакованный вариант."
            ),
            (
                "The winning option may shift toward the fastest deployable path rather than the deepest capability set."
                if language != "ru"
                else "Победитель может сместиться в сторону самого быстрого в развёртывании пути, а не самого глубокого по возможностям."
            ),
            0.71,
        )

    if _has_any_token(lower_dims, ("risk", "quality", "reliability", "governance", "надеж", "риск", "качеств", "управ")):
        add_trigger(
            "Reliability / Governance" if language != "ru" else "Надёжность / governance",
            (
                "If reliability, governance, or auditability become non-negotiable, retest the recommendation against stricter control requirements."
                if language != "ru"
                else "Если надёжность, governance или auditability становятся non-negotiable, нужно перепроверить рекомендацию при более жёстких требованиях к контролю."
            ),
            (
                "The recommendation may shift toward the option with the cleanest operational controls rather than the strongest headline performance."
                if language != "ru"
                else "Рекомендация может сместиться в сторону варианта с самым чистым операционным контролем, а не с самым сильным headline-performance."
            ),
            0.78,
        )

    if not triggers:
        add_trigger(
            "Decision Boundary" if language != "ru" else "Граница решения",
            (
                "If one constraint becomes dominant, the current recommendation should be re-tested against credible alternatives."
                if language != "ru"
                else "Если одно из ограничений становится доминирующим, текущую рекомендацию нужно переоценить против сильных альтернатив."
            ),
            (
                "The best path is likely conditional rather than universal."
                if language != "ru"
                else "Лучший путь, скорее всего, условный, а не универсальный."
            ),
            0.62,
        )
    return triggers[:4]


def _review_model_depth(task_spec: TaskSpec) -> str:
    if task_spec.request_spec.budget_tier.value == "light":
        return "standard"
    return "deep"


def _build_lateral_review_payload(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    research_rows: list[dict],
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    coverage: CoverageReport,
    heuristic_candidates: list[AdjacentQuestionCandidate],
) -> dict:
    return {
        "request": task_spec.request_spec.original_query,
        "subject": task_spec.request_spec.subject,
        "language": task_spec.request_spec.language,
        "decision_context": task_spec.request_spec.decision_context,
        "evaluation_dimensions": task_spec.evaluation_dimensions,
        "primary_questions": [item.model_dump(mode="json") for item in plan.primary_questions],
        "coverage_report": coverage.model_dump(mode="json"),
        "top_claims": [
            {
                "claim_id": claim.claim_id,
                "statement": claim.statement,
                "question_id": claim.question_id,
                "confidence": claim.confidence,
                "source_count": len(claim.source_ids),
                "contradictions": claim.contradiction_notes,
            }
            for claim in claims[:12]
        ],
        "research_rows": research_rows[:8],
        "sources": [
            {
                "title": item.title,
                "url": item.url,
                "source_type": item.source_type.value,
                "reliability_score": item.reliability_score,
            }
            for item in source_ledger[:12]
        ],
        "seed_adjacent_questions": [item.model_dump(mode="json") for item in heuristic_candidates[:6]],
        "objective": (
            "Improve the current report by identifying the most decision-relevant side questions, weaknesses, and recommendation-switch conditions."
            if task_spec.request_spec.language != "ru"
            else "Усилить текущий отчёт, выделив самые значимые боковые вопросы, слабые места и условия, при которых рекомендация должна переключиться."
        ),
    }


async def _call_lateral_review_model(system_prompt: str, user_payload: dict, task_spec: TaskSpec) -> tuple[str, SpendEntry]:
    review_depth = _review_model_depth(task_spec)
    direct_model = "sonar" if review_depth == "standard" else REVIEW_MODEL_DIRECT
    routed_model = "perplexity/sonar" if review_depth == "standard" else REVIEW_MODEL_ROUTED
    serialized_payload = json.dumps(user_payload, ensure_ascii=False, indent=2)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": serialized_payload},
    ]
    if settings.perplexity_api_key:
        try:
            async with httpx.AsyncClient(timeout=240) as client:
                response = await client.post(
                    PERPLEXITY_CHAT_URL,
                    headers={
                        "Authorization": f"Bearer {settings.perplexity_api_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": direct_model,
                        "messages": messages,
                        "temperature": 0.1,
                        "return_citations": True,
                        "return_related_questions": False,
                    },
                )
                response.raise_for_status()
                body = response.json()
                content = (((body.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
                usage = body.get("usage") or {}
                input_tokens = int(usage.get("prompt_tokens") or _estimate_tokens(system_prompt + serialized_payload))
                output_tokens = int(usage.get("completion_tokens") or _estimate_tokens(content))
                return content, _make_spend_entry(
                    category=SpendCategory.REVIEW,
                    stage="critique",
                    provider="perplexity",
                    model=direct_model,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    pricing_basis="usage" if usage else "estimated_chars",
                    notes="Model-driven lateral review",
                )
        except Exception:
            pass
    content = await _call_renderer_llm(system_prompt, serialized_payload, routed_model)
    return content, _make_spend_entry(
        category=SpendCategory.REVIEW,
        stage="critique",
        provider="openrouter",
        model=routed_model,
        input_tokens=_estimate_tokens(system_prompt + serialized_payload),
        output_tokens=_estimate_tokens(content),
        pricing_basis="estimated_chars",
        notes="Model-driven lateral review fallback",
    )


async def _call_report_writer_model(
    system_prompt: str,
    user_payload: dict,
    model: str,
    *,
    timeout_seconds: int = 180,
    allow_fallback: bool = True,
    prefer_perplexity: bool = False,
    direct_model: str = REVIEW_MODEL_DIRECT,
) -> tuple[str, SpendEntry]:
    serialized_payload = json.dumps(user_payload, ensure_ascii=False, indent=2)
    last_error: Exception | None = None
    renderer_timeout = max(120, timeout_seconds)

    async def call_renderer() -> tuple[str, SpendEntry]:
        content = await asyncio.wait_for(_call_renderer_llm(system_prompt, serialized_payload, model), timeout=renderer_timeout)
        return content, _make_spend_entry(
            category=SpendCategory.WRITER,
            stage="report_writer",
            provider="openrouter",
            model=model,
            input_tokens=_estimate_tokens(system_prompt + serialized_payload),
            output_tokens=_estimate_tokens(content),
            pricing_basis="estimated_chars",
            notes="Long-form synthesis or revision via OpenRouter",
        )

    if not prefer_perplexity:
        try:
            return await call_renderer()
        except Exception as exc:
            last_error = exc

    if settings.perplexity_api_key:
        try:
            async with httpx.AsyncClient(timeout=timeout_seconds) as client:
                response = await client.post(
                    PERPLEXITY_CHAT_URL,
                    headers={
                        "Authorization": f"Bearer {settings.perplexity_api_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": direct_model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": serialized_payload},
                        ],
                        "temperature": 0.2,
                        "return_citations": False,
                        "return_related_questions": False,
                    },
                )
                response.raise_for_status()
                body = response.json()
                content = (((body.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
                if content:
                    usage = body.get("usage") or {}
                    input_tokens = int(usage.get("prompt_tokens") or _estimate_tokens(system_prompt + serialized_payload))
                    output_tokens = int(usage.get("completion_tokens") or _estimate_tokens(content))
                    return content, _make_spend_entry(
                        category=SpendCategory.WRITER,
                        stage="report_writer",
                        provider="perplexity",
                        model=direct_model,
                        input_tokens=input_tokens,
                        output_tokens=output_tokens,
                        pricing_basis="usage" if usage else "estimated_chars",
                        notes="Long-form synthesis or revision via Perplexity",
                    )
                last_error = RuntimeError("Perplexity writer returned empty content.")
        except Exception as exc:
            last_error = exc
    if allow_fallback and prefer_perplexity:
        try:
            return await call_renderer()
        except Exception as exc:
            last_error = exc
    detail = repr(last_error) if last_error else "unknown error"
    raise RuntimeError(f"Report writer call failed: {detail}")


def _parse_model_adjacent_candidates(raw_items: object) -> list[AdjacentQuestionCandidate]:
    if not isinstance(raw_items, list):
        return []
    parsed: list[AdjacentQuestionCandidate] = []
    for item in raw_items[:8]:
        if not isinstance(item, dict):
            continue
        question = " ".join(str(item.get("question", "")).split()).strip()
        if len(question) < 20:
            continue
        candidate = AdjacentQuestionCandidate(
            question=question,
            kind=_normalize_question_kind(item.get("kind")),
            decision_impact=_clamp_score(item.get("decision_impact"), 0.75),
            coverage_gap=_clamp_score(item.get("coverage_gap"), 0.55),
            novelty=_clamp_score(item.get("novelty"), 0.7),
            comparative_value=_clamp_score(item.get("comparative_value"), 0.75),
            research_cost=_clamp_score(item.get("research_cost"), 0.35),
            selection_reason=" ".join(str(item.get("selection_reason", "")).split()).strip(),
        )
        candidate.composite_score = _adjacent_candidate_score(candidate)
        parsed.append(candidate)
    return parsed


def _parse_model_critique_findings(raw_items: object) -> list[CritiqueFinding]:
    if not isinstance(raw_items, list):
        return []
    findings: list[CritiqueFinding] = []
    for item in raw_items[:8]:
        if not isinstance(item, dict):
            continue
        summary = " ".join(str(item.get("summary", "")).split()).strip()
        if len(summary) < 20:
            continue
        findings.append(
            CritiqueFinding(
                kind=_normalize_critique_kind(item.get("kind")),
                severity=_normalize_severity(item.get("severity")),
                summary=summary,
                rationale=" ".join(str(item.get("rationale", "")).split()).strip(),
            )
        )
    return findings


def _parse_model_decision_triggers(raw_items: object) -> list[DecisionTrigger]:
    if not isinstance(raw_items, list):
        return []
    triggers: list[DecisionTrigger] = []
    for item in raw_items[:6]:
        if not isinstance(item, dict):
            continue
        label = " ".join(str(item.get("label", "")).split()).strip()
        condition = " ".join(str(item.get("condition", "")).split()).strip()
        implication = " ".join(str(item.get("implication", "")).split()).strip()
        if len(label) < 4 or len(condition) < 12 or len(implication) < 12:
            continue
        triggers.append(
            DecisionTrigger(
                label=label,
                condition=condition,
                implication=implication,
                confidence=_clamp_score(item.get("confidence"), 0.72),
            )
        )
    return triggers


async def _generate_model_driven_review(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    research_rows: list[dict],
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    coverage: CoverageReport,
    *,
    record_spend: Callable[[SpendEntry], None] | None = None,
) -> tuple[list[AdjacentQuestionCandidate], list[CritiqueFinding], list[DecisionTrigger], dict]:
    heuristic_candidates = build_adjacent_question_candidates(task_spec, coverage, claims)
    heuristic_findings = build_critique_findings(task_spec, plan, claims, coverage, [])
    heuristic_triggers = build_decision_triggers(task_spec)
    prompt = LATERAL_REVIEW_PROMPT.format(language_name=_language_name(task_spec.request_spec.language))
    payload = _build_lateral_review_payload(
        task_spec,
        plan,
        research_rows,
        source_ledger,
        claims,
        coverage,
        heuristic_candidates,
    )
    raw, spend_entry = _coerce_spend_result(
        await _call_lateral_review_model(prompt, payload, task_spec),
        category=SpendCategory.REVIEW,
        stage="critique",
        provider="perplexity",
        model=REVIEW_MODEL_DIRECT,
    )
    if record_spend is not None:
        spend_entry.stage = "critique"
        spend_entry.category = SpendCategory.REVIEW
        record_spend(spend_entry)
    parsed = parse_llm_json(raw, context="v2_lateral_review")

    model_candidates = _parse_model_adjacent_candidates(parsed.get("adjacent_questions"))
    merged_candidates = _merge_adjacent_candidates(model_candidates, heuristic_candidates)

    model_findings = _parse_model_critique_findings(parsed.get("critique_findings"))
    merged_findings = model_findings or heuristic_findings

    model_triggers = _parse_model_decision_triggers(parsed.get("decision_triggers"))
    merged_triggers = model_triggers or heuristic_triggers

    review_artifact = {
        "source": "model",
        "raw_response": raw,
        "parsed": parsed,
    }
    return merged_candidates, merged_findings, merged_triggers, review_artifact


def _build_live_research_queries(task_spec: TaskSpec, plan: ResearchPlan) -> list[tuple[str, str]]:
    original_query = task_spec.request_spec.original_query.strip()
    original_query_lower = original_query.lower()
    geography = task_spec.request_spec.geography.replace("_", " ")
    dimensions = ", ".join(task_spec.evaluation_dimensions[:4])
    language = task_spec.request_spec.language
    prompts: list[tuple[str, str]] = []
    is_stack_research_topic = _is_stack_research_topic_query(original_query_lower)
    is_business_topic = _is_business_topic_task(task_spec)
    stack_hint = (
        "Explicitly compare managed search APIs (for example Tavily, Exa, Google Custom Search, Perplexity APIs), "
        "open-source orchestration frameworks (LangChain/LangGraph, Haystack, LlamaIndex), and mature GitHub projects "
        "(such as GPT Researcher or open_deep_research). Ignore generic product-architecture literature unless it directly "
        "supports build-vs-buy, provenance, or evidence-traceability decisions. Reject listicles, trend roundups, "
        "community posts, and generic vendor blogs unless no stronger source exists."
        if language != "ru"
        else "Явно сравни managed search APIs (например Tavily, Exa, Google Custom Search, Perplexity APIs), "
        "open-source orchestration frameworks (LangChain/LangGraph, Haystack, LlamaIndex) и зрелые GitHub-проекты "
        "(например GPT Researcher или open_deep_research). Игнорируй общий дискурс про product architecture, "
        "если он прямо не помогает по build-vs-buy, provenance или evidence-traceability. Отбрасывай listicles, "
        "trend roundups, community posts и generic vendor blogs, если доступны более сильные источники."
    )
    business_hint_en = (
        "Treat Smart Report as the name of the user's product, not as unrelated Smart Reporting modules or generic reporting documentation. "
        "Exclude Oracle, BMC, or similar product docs unless they directly support monetization, pricing, packaging, GTM, or buyer behavior decisions. "
        "Prefer sources on SaaS monetization, pricing strategy, consulting or investment software buying behavior, freemium vs paid conversion, enterprise procurement, and GTM execution."
    )
    business_hint_ru = (
        "Считай Smart Report названием продукта пользователя, а не ссылкой на чужие модули Smart Reporting или общую документацию по отчётности. "
        "Исключай Oracle, BMC и похожие product docs, если они прямо не помогают по монетизации, pricing, packaging, GTM или buyer behavior. "
        "Предпочитай источники по SaaS-монетизации, pricing strategy, buying behavior у consulting/investment команд, freemium-vs-paid conversion, enterprise procurement и GTM execution."
    )

    def compose(question_id: str, angle: str) -> tuple[str, str]:
        guardrail = _stack_research_guardrail(question_id, angle) if is_stack_research_topic else ""
        q2_focus_ru = (
            "Для этого угла нужны реальные competing products, incumbent workflows и бесплатные substitutes, а не абстрактные документы про reporting, cost estimating или project governance."
        )
        q2_focus_en = (
            "For this angle, pull real competing products, incumbent workflows, and free substitutes. Avoid generic reporting documentation, cost-estimating guides, or project-governance literature unless they directly compare the option set."
        )
        q4_focus_ru = (
            "Для этого угла нужны не общие рассуждения о decision quality, а конкретные trade-offs, buyer objections, "
            "willingness-to-pay constraints, procurement friction, integration burden и условия, при которых рекомендация должна переключиться."
        )
        q4_focus_en = (
            "For this angle, avoid generic decision-quality commentary. Pull concrete tradeoffs, buyer objections, "
            "willingness-to-pay constraints, procurement friction, integration burden, and explicit recommendation-switch conditions."
        )
        if language == "ru":
            if is_stack_research_topic:
                prompt = (
                    f"Primary research question: {angle}\n"
                    f"Original decision prompt: {original_query}\n"
                    f"География: {geography}\n"
                    f"Критерии решения: {dimensions}\n"
                    f"{stack_hint}\n"
                    "Нужны свежие официальные docs, GitHub repos, benchmark pages, pricing pages, maintenance/activity signals, "
                    "production case studies и явные trade-offs."
                )
            else:
                prompt = (
                    f"{original_query}\n\n"
                    f"Фокус исследования: {angle}\n"
                    f"География: {geography}\n"
                    f"Критерии решения: {dimensions}\n"
                    "Нужны свежие источники, конкретные цифры, кейсы внедрения, стоимость, риски и выводы для управленческого решения. "
                    "Предпочитай официальные документы, серьёзную аналитику и кейсы компаний."
                )
                if is_business_topic:
                    prompt += f"\n{business_hint_ru}"
                if is_business_topic and question_id == "q2":
                    prompt += f"\n{q2_focus_ru}"
                    prompt += "\nСравни Explicitly Perplexity, AlphaSense, PitchBook, CB Insights, Hebbia и текущий ручной workflow команды, если они релевантны теме."
                if question_id == "q4":
                    prompt += f"\n{q4_focus_ru}"
                    prompt += "\nFocus on pricing, ROI, procurement objections, integration burden, and explicit switch conditions."
        else:
            if is_stack_research_topic:
                prompt = (
                    f"Primary research question: {angle}\n"
                    f"Original decision prompt: {original_query}\n"
                    f"Geography: {geography}\n"
                    f"Decision criteria: {dimensions}\n"
                    f"{stack_hint}\n"
                    "Need fresh official docs, GitHub repo evidence, benchmark pages, pricing pages, maintenance and activity signals, "
                    "production case studies, and explicit tradeoffs."
                )
                if guardrail:
                    prompt += f"\n{guardrail}"
                if question_id == "q4":
                    prompt += f"\n{q4_focus_en}"
                    prompt += "\nLook for pricing, ROI, procurement friction, integration burden, and explicit switch conditions."
            else:
                prompt = (
                    f"{original_query}\n\n"
                    f"Research angle: {angle}\n"
                    f"Geography: {geography}\n"
                    f"Decision criteria: {dimensions}\n"
                    "Need fresh sources, concrete metrics, implementation case studies, cost structure, risks, and management implications. "
                    "Prefer official documentation, serious research, and company case evidence."
                )
                if is_business_topic:
                    prompt += f"\n{business_hint_en}"
                if is_business_topic and question_id == "q2":
                    prompt += f"\n{q2_focus_en}"
                    prompt += "\nCompare Perplexity, AlphaSense, PitchBook, CB Insights, Hebbia, and the current manual workflow if they are relevant to the decision."
                if question_id == "q4":
                    prompt += f"\n{q4_focus_en}"
                    prompt += "\nLook for pricing, ROI, procurement friction, integration burden, and explicit switch conditions."
        return question_id, prompt

    for question in plan.primary_questions[:4]:
        prompts.append(compose(question.question_id, question.question))

    extra_angles = (
        [
            "Кейсы внедрения, паттерны rollout и типовые ошибки",
            "Экономика решения, unit economics, ROI и скрытые операционные затраты",
            "Риски, ограничения, регуляторные и организационные барьеры",
        ]
        if language == "ru"
        else [
            "Implementation case studies, rollout patterns, and common failure modes",
            "Economics, unit economics, ROI, and hidden operating costs",
            "Risks, constraints, regulatory concerns, and organizational blockers",
        ]
    )
    for index, angle in enumerate(extra_angles, start=1):
        prompts.append(compose(f"l{index}", angle))

    deduped: list[tuple[str, str]] = []
    seen: set[str] = set()
    for question_id, prompt in prompts:
        if prompt in seen:
            continue
        seen.add(prompt)
        deduped.append((question_id, prompt))
    return deduped[:6]


def _build_fallback_search_queries(task_spec: TaskSpec, plan: ResearchPlan) -> list[str]:
    language = task_spec.request_spec.language
    geography = task_spec.request_spec.geography.replace("_", " ")
    subject_tokens = re.findall(r"[\w-]+", task_spec.request_spec.subject)[:8]
    subject_short = " ".join(subject_tokens) or task_spec.request_spec.subject
    dimensions = task_spec.evaluation_dimensions[:3]
    topic = task_spec.request_spec.original_query.lower()
    if language == "ru":
        templates = [
            f"{subject_short} {geography} исследование рынок кейс",
            f"{subject_short} {geography} стоимость качество риски",
            f"{subject_short} {geography} международный опыт бенчмарк",
            f"{subject_short} {geography} практика внедрения девелопер",
        ]
    else:
        templates = [
            f"{subject_short} {geography} market study case",
            f"{subject_short} {geography} cost quality risk",
            f"{subject_short} {geography} international benchmark",
            f"{subject_short} {geography} implementation case developer",
        ]
    if all(token in topic for token in ("llm", "github")) and ("search" in topic or "research" in topic):
        if language == "ru":
            templates.extend(
                [
                    f"лучшие llm модели web search deep research 2026 benchmark",
                    f"gpt researcher open deep research deep-searcher github comparison",
                    f"Perplexity alternative stack llm search critique synthesis github",
                    f"llm leaderboard web search deep research github production stack",
                ]
            )
        else:
            templates.extend(
                [
                    "best llm models web search deep research 2026 benchmark",
                    "gpt researcher open deep research deep-searcher github comparison",
                    "perplexity alternative stack llm search critique synthesis github",
                    "llm leaderboard web search deep research github production stack",
                ]
            )
    for question in plan.primary_questions[:2]:
        templates.append(f"{subject_short} {question.question}")
    for dimension in dimensions:
        templates.append(f"{subject_short} {dimension} {geography}")

    deduped: list[str] = []
    seen: set[str] = set()
    for query in templates:
        normalized = " ".join(query.split()).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)
    return deduped[:8]


def _build_question_fallback_queries(task_spec: TaskSpec, question: ResearchQuestion) -> list[str]:
    subject_tokens = re.findall(r"[\w-]+", task_spec.request_spec.subject)[:8]
    subject_short = " ".join(subject_tokens) or task_spec.request_spec.subject
    geography = task_spec.request_spec.geography.replace("_", " ")
    question_tokens = re.findall(r"[\w-]+", question.question)[:16]
    condensed_question = " ".join(question_tokens) or question.question
    topic = task_spec.request_spec.original_query.lower()
    question_lower = question.question.lower()
    is_business_topic = _is_business_topic_task(task_spec)
    queries = [
        f"{subject_short} {condensed_question}",
        f"{subject_short} {condensed_question} {geography}",
    ]

    if all(token in topic for token in ("llm", "github")) and ("search" in topic or "research" in topic):
        if any(token in question_lower for token in ("github", "project", "orchestration", "framework")):
            queries.extend(
                [
                    "assafelovic gpt researcher github repo",
                    "langchain ai open deep research github repo",
                    "gpt researcher docs gptr",
                    "deep-searcher github zilliz",
                    "langgraph deep research agent github",
                ]
            )
        if any(token in question_lower for token in ("tradeoff", "risk", "condition", "switch")):
            queries.extend(
                [
                    "llm citation accuracy hallucination web search benchmark",
                    "open source vs closed source llm cost comparison",
                    "perplexity limitations citation accuracy",
                    "agentic search latency cost tradeoffs",
                    "tavily serper search api pricing",
                ]
            )
        if any(token in question_lower for token in ("stack", "architecture", "perplexity")):
            queries.extend(
                [
                    "perplexity api pricing search citations",
                    "langgraph docs durable execution human in the loop",
                    "langfuse docs traces evaluations self hosting",
                    "tavily docs extract crawl pricing",
                    "firecrawl docs markdown extraction search api",
                ]
            )
        if any(token in question_lower for token in ("llm", "model", "models")):
            queries.extend(
                [
                    "best llm models web search deep research benchmark",
                    "open-source vs closed-source llm research comparison",
                ]
            )
    if is_business_topic and question.question_id == "q2":
        queries.extend(
            [
                "Perplexity AlphaSense PitchBook CB Insights Hebbia alternatives consulting investment teams",
                "market intelligence tools consulting investment teams comparison pricing workflow",
                "free substitutes manual workflow vs AI research tools consulting teams",
            ]
        )
    if is_business_topic and any(
        token in question_lower for token in ("monetization", "pricing", "price", "market", "go-to-market", "gtm", "risk", "switch")
    ):
        queries.extend(
            [
                "ai saas monetization consulting investment teams pricing",
                "freemium vs paid analytics software professional services",
                "enterprise software pricing procurement objections consulting firms",
                "go to market packaging ai workflow software professional services",
            ]
        )
    if any(token in question_lower for token in ("tradeoff", "risk", "condition", "switch", "риск", "услов", "переключ", "возраж", "закуп")):
        queries.extend(
            [
                f"{subject_short} willingness to pay pricing budget objections",
                f"{subject_short} procurement integration burden risks",
                f"{subject_short} free alternative vs paid tool objections",
                f"{subject_short} ROI TCO switching conditions",
            ]
        )

    deduped: list[str] = []
    seen: set[str] = set()
    for query in queries:
        normalized = " ".join(query.split()).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)
    return deduped[:6]


def _build_live_adjacent_query(task_spec: TaskSpec, question: ResearchQuestion) -> str:
    geography = task_spec.request_spec.geography.replace("_", " ")
    dimensions = ", ".join(task_spec.evaluation_dimensions[:4])
    if task_spec.request_spec.language == "ru":
        return (
            f"{task_spec.request_spec.original_query}\n\n"
            f"Критический боковой вопрос: {question.question}\n"
            f"География: {geography}\n"
            f"Оси решения: {dimensions}\n"
            "Нужно честно сравнить альтернативы, явно назвать trade-offs, привести количественные сигналы, "
            "описать сценарии, где основной вариант слаб, и не потерять связь с исходной управленческой целью."
        )
    return (
        f"{task_spec.request_spec.original_query}\n\n"
        f"Critical adjacent question: {question.question}\n"
        f"Geography: {geography}\n"
        f"Decision axes: {dimensions}\n"
        "Need an explicit alternative space, quantified tradeoffs, scenarios where the focal option underperforms, "
        "and implications that remain tied to the original decision goal."
    )


def _build_fallback_adjacent_query(task_spec: TaskSpec, question: ResearchQuestion) -> str:
    subject = task_spec.request_spec.subject
    geography = task_spec.request_spec.geography.replace("_", " ")
    suffix = (
        "official benchmark case study cost risk"
        if task_spec.request_spec.language != "ru"
        else "официальные документы бенчмарк кейс стоимость риски"
    )
    return f"{subject} {geography} {question.question} {suffix}".strip()


def _upsert_live_source(
    source_map: dict[str, SourceLedgerEntry],
    *,
    url: str,
    title: str,
    domain: str,
    question_id: str,
    preferred_domains: list[str],
    selection_reason: str,
) -> None:
    existing = source_map.get(url)
    if existing is not None:
        if question_id not in existing.question_links:
            existing.question_links.append(question_id)
        return

    source_type = classify_source_type(url)
    source_map[url] = SourceLedgerEntry(
        url=url,
        title=title or domain or url,
        domain=domain,
        source_type=source_type,
        publisher=domain,
        reliability_score=score_source(url, source_type, preferred_domains),
        selection_reason=selection_reason,
        question_links=[question_id],
    )


def _extract_markdown_links(text: str) -> list[str]:
    return re.findall(r"\[[^\]]+\]\((https?://[^)]+)\)", text or "")


_MOJIBAKE_REPLACEMENTS = {
    "вЂ”": "-",
    "вЂ": "-",
    "в€’": "-",
    "в†’": "->",
    "в‰Ґ": ">=",
    "в‰¤": "<=",
    "вЂ™": "'",
    "вЂњ": '"',
    "вЂќ": '"',
    "â€”": "-",
    "â€“": "-",
    "â€™": "'",
    "â€œ": '"',
    "â€": '"',
    "â†’": "->",
    "â‰¥": ">=",
    "â‰¤": "<=",
}


def _sanitize_llm_markdown(text: str) -> str:
    cleaned = str(text or "").replace("\r\n", "\n").strip()
    for broken, fixed in _MOJIBAKE_REPLACEMENTS.items():
        cleaned = cleaned.replace(broken, fixed)
    cleaned = re.sub(r"(?m)^[ \t]*[•*–—]\s+", "- ", cleaned)
    cleaned = re.sub(r"\*\*(Exhibit\s+\d+):\s*([^*\n]+)\*\*", r"\1\n\n_\2_", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\*\*(Таблица\s+\d+):\s*([^*\n]+)\*\*", r"\1\n\n_\2_", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\[\s*\.\.\.\s*\]", "", cleaned)
    cleaned = re.sub(
        r"\[(?:\s*(?:request_spec|coverage_report|critique(?:_findings)?|decision_triggers|research_rows)\b[^\]]*|\s*[alq]q?\d+[^\]]*)\](?!\()",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(
        r"\[Evidence:\s*(?!\s*C-\d+(?:\s*,\s*C-\d+)*\s*\])[^\]]+\]",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r"\[(?:\d+(?:\s+from\s+[^\]]+)?)\](?!\()", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"(?im)^\(word count:[^)]+\)\s*$", "", cleaned)
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = _normalize_markdown_tables(cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _normalize_markdown_tables(text: str) -> str:
    if "|" not in text:
        return text

    def is_table_line(value: str) -> bool:
        stripped = value.strip()
        return bool(stripped) and stripped.startswith("|") and stripped.endswith("|")

    lines = text.split("\n")
    normalized: list[str] = []
    total = len(lines)
    for index, line in enumerate(lines):
        if line.strip():
            normalized.append(line)
            continue

        previous_index = index - 1
        while previous_index >= 0 and not lines[previous_index].strip():
            previous_index -= 1

        next_index = index + 1
        while next_index < total and not lines[next_index].strip():
            next_index += 1

        previous_line = lines[previous_index] if previous_index >= 0 else ""
        next_line = lines[next_index] if next_index < total else ""
        if is_table_line(previous_line) and is_table_line(next_line):
            continue
        normalized.append(line)

    return "\n".join(normalized)


_LIVE_FINDING_NOISE_PATTERNS = (
    "key strengths by use case",
    "recommended projects by capability",
    "architecture recommendations",
    "output format:",
    "exhibit ",
    "| project |",
    "| stack |",
    "body application/json",
    "skip to main content",
    "home page v2",
    "query string required",
    "maximum string length",
    "maximum number of results",
    "sources to search",
    "timeout integer default",
    "api reference",
    "on this page",
    "quick start",
    "basic scrape",
)


def _normalize_live_finding_text(text: str) -> str:
    cleaned = _sanitize_llm_markdown(text)
    cleaned = re.sub(r"(?im)^#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"(?im)^\s*(?:\d+[.)]|[-*])\s+", "", cleaned)
    cleaned = cleaned.replace("**", "").replace("`", "")
    cleaned = re.sub(r"\s*\|\s*", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" -")
    return cleaned


def _is_live_finding_usable(text: str) -> bool:
    lowered = text.lower()
    if len(text) < 50 or len(text) > 560:
        return False
    if _is_low_signal_claim_text(text):
        return False
    if any(marker in lowered for marker in _LIVE_FINDING_NOISE_PATTERNS):
        return False
    if lowered.count(" - ") >= 3:
        return False
    if lowered.count(":") >= 4:
        return False
    if re.search(r"\b(?:table|exhibit)\s+\d+\b", lowered):
        return False
    return True


def _stack_question_live_tokens(question_id: str) -> tuple[str, ...]:
    if question_id == "q2":
        return (
            "github",
            "repo",
            "research",
            "retrieval",
            "citation",
            "rag",
            "agent",
            "workflow",
            "llm",
            "search",
        )
    if question_id == "q3":
        return (
            "langgraph",
            "llamaindex",
            "ragflow",
            "tavily",
            "firecrawl",
            "langfuse",
            "perplexity",
            "retrieval",
            "citation",
            "trace",
            "audit",
            "orchestration",
            "workflow",
            "agent",
            "search",
        )
    if question_id == "q4":
        return (
            "pricing",
            "cost",
            "roi",
            "latency",
            "citation",
            "hallucination",
            "lock-in",
            "lock in",
            "self-host",
            "self host",
            "maintenance",
            "integration",
            "burden",
            "rate limit",
            "benchmark",
            "switch",
            "risk",
            "tradeoff",
            "trade-off",
            "compliance",
            "governance",
        )
    return ()


def _stack_live_finding_relevant(question_id: str, source: SourceLedgerEntry, text: str) -> bool:
    tokens = _stack_question_live_tokens(question_id)
    if not tokens:
        return True
    haystack = f"{source.title} {source.url} {text}".lower()
    matches = sum(1 for token in tokens if token in haystack)
    return matches >= 2


def _q4_live_finding_relevant(task_spec: TaskSpec | None, source: SourceLedgerEntry, text: str, query: str) -> bool:
    haystack = f"{source.title} {source.url} {text}".lower()
    strong_types = {
        SourceType.OFFICIAL_DOCUMENTATION,
        SourceType.BENCHMARK,
        SourceType.RESEARCH_PAPER,
        SourceType.USER_MATERIAL,
    }
    tokens = list(_Q4_SOURCE_SIGNAL_TOKENS)
    if task_spec and _is_stack_research_topic_query(task_spec.request_spec.original_query):
        tokens.extend(_stack_question_live_tokens("q4"))
    matches = sum(1 for token in tokens if token in haystack)
    if any(pattern in haystack for pattern in _Q4_LOW_SIGNAL_PATTERNS) and matches < 2:
        return False
    if matches >= 2:
        return True
    if source.source_type in strong_types and matches >= 1:
        return True
    return False


def _report_word_count(payload: dict) -> int:
    parts = [payload.get("executive_summary", "")]
    parts.extend(section.get("content", "") for section in payload.get("sections", []))
    return sum(len(str(part).split()) for part in parts)


def _build_live_evidence(
    research_rows: list[dict],
    source_ledger: list[SourceLedgerEntry],
    task_spec: TaskSpec | None = None,
) -> list[EvidenceRecord]:
    source_by_url = {entry.url: entry for entry in source_ledger}
    evidence: list[EvidenceRecord] = []
    is_stack_topic = bool(task_spec and _is_stack_research_topic_query(task_spec.request_spec.original_query))
    for row in research_rows:
        supporting_urls = [url for url in row.get("source_urls", []) if url in source_by_url]
        if not supporting_urls:
            continue
        effective_question_id = row.get("primary_question_id") or row["question_id"]
        query_tokens = _topic_tokens(str(row.get("query", "")))
        for index, finding in enumerate(row.get("findings", []), start=1):
            normalized = _normalize_live_finding_text(str(finding))
            if not _is_live_finding_usable(normalized):
                continue
            source_url = supporting_urls[(index - 1) % len(supporting_urls)]
            source = source_by_url[source_url]
            topicality = _source_topic_alignment_score(source, query_tokens)
            if query_tokens and topicality < 0.15 and not any(token in normalized.lower() for token in sorted(query_tokens)[:8]):
                continue
            if effective_question_id == "q4" and not _q4_live_finding_relevant(task_spec, source, normalized, str(row.get("query", ""))):
                continue
            if is_stack_topic and not _stack_live_finding_relevant(effective_question_id, source, normalized):
                continue
            confidence = min(0.97, source.reliability_score + (0.08 if re.search(r"\d", normalized) else 0.0))
            evidence.append(
                EvidenceRecord(
                    question_id=effective_question_id,
                    source_id=source.source_id,
                    claim=normalized,
                    snippet=normalized[:320],
                    confidence=confidence,
                    extraction_method="live_research",
                )
            )
    return evidence


def _build_live_report_payload(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    research_rows: list[dict],
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    coverage: CoverageReport,
    adjacent_questions: list[ResearchQuestion],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
) -> dict:
    source_by_id = {entry.source_id: entry for entry in source_ledger}
    report_claims = _select_report_worthy_claims(claims)
    numeric_claims = _numeric_claim_catalog(report_claims, source_by_id)
    return {
        "request": task_spec.request_spec.original_query,
        "request_spec": task_spec.request_spec.model_dump(mode="json"),
        "evaluation_dimensions": task_spec.evaluation_dimensions,
        "must_cover_questions": [question.question for question in plan.primary_questions],
        "required_section_titles": _live_section_titles(task_spec.request_spec.language),
        "coverage_report": coverage.model_dump(mode="json"),
        "adjacent_questions": [item.model_dump(mode="json") for item in adjacent_questions],
        "critique_findings": [item.model_dump(mode="json") for item in critique_findings],
        "decision_triggers": [item.model_dump(mode="json") for item in decision_triggers],
        "sources": [
            {
                "source_id": entry.source_id,
                "title": entry.title,
                "url": entry.url,
                "domain": entry.domain,
                "source_type": entry.source_type.value,
                "reliability_score": entry.reliability_score,
                "selection_reason": entry.selection_reason,
            }
            for entry in source_ledger[:16]
        ],
        "research_rows": research_rows[:8],
        "claims": [
            {
                "claim_id": claim.claim_id,
                "statement": claim.statement,
                "question_id": claim.question_id,
                "confidence": claim.confidence,
                "sources": [
                    {
                        "title": source_by_id[source_id].title,
                        "url": source_by_id[source_id].url,
                    }
                    for source_id in claim.source_ids
                    if source_id in source_by_id
                ],
            }
            for claim in report_claims[:20]
        ],
        "numeric_claims": numeric_claims,
        "grounding_contract": {
            "allowed_numeric_claim_ids": [item["claim_id"] for item in numeric_claims],
            "rules": [
                "Use exact numbers only when they appear in numeric_claims.",
                "Do not derive break-even thresholds, staffing budgets, TCO models, or volume curves unless those exact numbers already exist in numeric_claims.",
                "If the evidence is directional but not numeric, write the point qualitatively.",
            ],
        },
        "writing_brief": {
            "tone": "BCG-grade, synthesis-heavy, evidence-first, client-ready",
            "forbidden": ["generic filler", "meta commentary", "process narration", "empty platitudes"],
            "section_rule": "Each section must materially advance the decision, not restate the same facts.",
            "mandatory_moves": [
                "Make the alternative space explicit, not implied",
                "State the strongest counterargument against the current recommendation",
                "Explain what could change the recommendation",
                "Name the unknowns that still require validation",
            ],
        },
    }


def _truncate_revision_text(text: str, max_words: int = 240) -> str:
    words = str(text or "").split()
    if len(words) <= max_words:
        return str(text or "").strip()
    head_words = max(80, max_words // 2)
    tail_words = max(40, max_words // 4)
    return " ".join(words[:head_words] + ["[...]"] + words[-tail_words:])


def _title_key(title: str) -> str:
    return " ".join(re.findall(r"[\wа-яё]+", title.lower()))


def _titles_match(left: str, right: str) -> bool:
    left_key = _title_key(left)
    right_key = _title_key(right)
    if not left_key or not right_key:
        return False
    if left_key == right_key or left_key in right_key or right_key in left_key:
        return True
    left_tokens = set(left_key.split())
    right_tokens = set(right_key.split())
    overlap = len(left_tokens & right_tokens)
    return overlap >= max(2, min(len(left_tokens), len(right_tokens)) - 1)


def _numeric_claim_catalog(claims: list[ClaimRecord], source_by_id: dict[str, SourceLedgerEntry]) -> list[dict]:
    catalog: list[dict] = []
    for claim in claims:
        numeric_facts = extract_numeric_facts(claim.statement)
        if not numeric_facts:
            continue
        catalog.append(
            {
                "claim_id": claim.claim_id,
                "statement": claim.statement,
                "question_id": claim.question_id,
                "allowed_numbers": [fact.raw for fact in numeric_facts[:4]],
                "sources": [
                    {
                        "title": source_by_id[source_id].title,
                        "url": source_by_id[source_id].url,
                    }
                    for source_id in claim.source_ids[:2]
                    if source_id in source_by_id
                ],
            }
        )
        if len(catalog) >= 14:
            break
    return catalog


def _pick_revision_targets(language: str, current_draft: dict, revision_focus: list[str]) -> list[dict[str, str]]:
    required_titles = _live_section_titles(language)
    section_preferences = {
        "claim depth": [required_titles[2], required_titles[3]],
        "evidence density": [required_titles[2], required_titles[6]],
        "grounding discipline": [required_titles[3], required_titles[6]],
        "lateral breadth": [required_titles[2], required_titles[5]],
        "decision usefulness": [required_titles[6], required_titles[7]],
        "presentation depth": [required_titles[1], required_titles[4]],
    }
    reason_map = {
        "claim depth": "Increase non-trivial comparative reasoning, quantified tradeoffs, and boundary conditions.",
        "evidence density": "Tighten evidence linkage, source-backed claims, and explicit proof points inside the section.",
        "grounding discipline": "Remove unsupported exact metrics, keep only claim-backed numbers, and replace invented economics with qualitative tradeoff language.",
        "lateral breadth": "Make alternatives, counterarguments, hidden variables, and switch conditions explicit.",
        "decision usefulness": "Strengthen recommendation logic, operating implications, and next-step usability.",
        "presentation depth": "Increase structure, section density, and exhibit quality without adding filler.",
    }
    current_sections = current_draft.get("sections", []) or []
    targets: list[dict[str, str]] = []
    seen_titles: set[str] = set()
    if "presentation depth" in revision_focus and current_sections:
        thin_sections = sorted(
            current_sections,
            key=lambda section: len(str(section.get("content", "")).split()),
        )
        for section in thin_sections[:3]:
            resolved_title = str(section.get("title", "")).strip()
            title_key = _title_key(resolved_title)
            if not title_key or title_key in seen_titles:
                continue
            seen_titles.add(title_key)
            targets.append(
                {
                    "target_title": resolved_title,
                    "reason": "Expand this thin section materially, add at least one concrete comparison or exhibit, and make the section decision-useful rather than summary-like.",
                }
            )
    for focus in revision_focus:
        for preferred_title in section_preferences.get(focus, []):
            matching_section = next(
                (section for section in current_sections if _titles_match(preferred_title, str(section.get("title", "")))),
                None,
            )
            resolved_title = str((matching_section or {}).get("title", "")).strip() or preferred_title
            title_key = _title_key(resolved_title)
            if not title_key or title_key in seen_titles:
                continue
            seen_titles.add(title_key)
            targets.append(
                {
                    "target_title": resolved_title,
                    "reason": reason_map.get(focus, "Improve analytical quality and decision value."),
                }
            )
            break
        if len(targets) >= 4:
            break
    if not targets and current_sections:
        first_title = str(current_sections[0].get("title", "")).strip()
        if first_title:
            targets.append(
                {
                    "target_title": first_title,
                    "reason": "Strengthen the current weakest section without rewriting the entire report.",
                }
            )
    return targets


def _build_quality_revision_payload(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    current_draft: dict,
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    assessment: QualityAssessment,
    revision_focus: list[str],
) -> dict:
    source_by_id = {entry.source_id: entry for entry in source_ledger}
    selected_claims = _select_report_worthy_claims(claims)[:10]
    numeric_claims = _numeric_claim_catalog(_select_report_worthy_claims(claims), source_by_id)
    current_sections = current_draft.get("sections", []) or []
    required_titles = _live_section_titles(task_spec.request_spec.language)
    target_sections = _pick_revision_targets(task_spec.request_spec.language, current_draft, revision_focus)
    missing_required_sections = [
        title
        for title in required_titles
        if not any(_titles_match(title, str(section.get("title", ""))) for section in current_sections)
    ]
    revise_executive_summary = any(
        focus in {"decision usefulness", "presentation depth", "claim depth"} for focus in revision_focus
    )
    target_section_payload: list[dict] = []
    report_word_count = _report_word_count(current_draft)
    for target in target_sections:
        current_section = next(
            (section for section in current_sections if _titles_match(target["target_title"], str(section.get("title", "")))),
            None,
        )
        if current_section is None:
            continue
        current_content = str(current_section.get("content", "")).strip()
        target_section_payload.append(
            {
                "target_title": str(current_section.get("title", "")).strip() or target["target_title"],
                "reason": target["reason"],
                "current_word_count": len(current_content.split()),
                "target_min_words": max(420 if report_word_count < 2600 else 340, len(current_content.split()) + 260),
                "current_content_excerpt": _truncate_revision_text(current_content, max_words=260),
                "current_sources": list(dict.fromkeys(current_section.get("sources", [])))[:4],
            }
        )
    return {
        "request": task_spec.request_spec.original_query,
        "subject": task_spec.request_spec.subject,
        "decision_context": task_spec.request_spec.decision_context,
        "evaluation_dimensions": task_spec.evaluation_dimensions[:6],
        "must_cover_questions": [question.question for question in plan.primary_questions],
        "required_section_titles": required_titles,
        "revision_focus": revision_focus,
        "report_shape": {
            "title": current_draft.get("title", ""),
            "subtitle": current_draft.get("subtitle", ""),
            "facts_line": current_draft.get("facts_line", ""),
            "section_count": len(current_sections),
            "word_count": _report_word_count(current_draft),
        },
        "quality_assessment": {
            "overall_score": assessment.overall_score,
            "verdict": assessment.verdict,
            "weaknesses": assessment.weaknesses[:6],
            "rewrite_priorities": assessment.rewrite_priorities[:6],
            "lowest_dimensions": [
                {
                    "dimension": item.dimension,
                    "score": item.score,
                    "rationale": item.rationale,
                }
                for item in sorted(assessment.dimensions, key=lambda dimension: dimension.score)[:4]
            ],
        },
        "revise_executive_summary": revise_executive_summary,
        "current_executive_summary": _truncate_revision_text(current_draft.get("executive_summary", ""), max_words=260),
        "target_sections": target_section_payload,
        "missing_required_sections": missing_required_sections[:2],
        "critique_findings": [item.summary for item in critique_findings[:8]],
        "decision_triggers": [item.model_dump(mode="json") for item in decision_triggers[:6]],
        "numeric_claims": numeric_claims,
        "grounding_contract": {
            "allowed_numeric_claim_ids": [item["claim_id"] for item in numeric_claims],
            "rules": [
                "Delete or rewrite unsupported exact numbers instead of paraphrasing them.",
                "Keep only exact metrics that appear in numeric_claims.",
                "Prefer qualitative tradeoffs over synthetic economics when the evidence is thin.",
            ],
        },
        "sources": [
            {
                "title": entry.title,
                "url": entry.url,
                "source_type": entry.source_type.value,
                "reliability_score": entry.reliability_score,
            }
            for entry in source_ledger[:8]
        ],
        "claims": [
            {
                "claim_id": claim.claim_id,
                "statement": claim.statement,
                "question_id": claim.question_id,
                "confidence": claim.confidence,
                "sources": [
                    {
                        "title": source_by_id[source_id].title,
                        "url": source_by_id[source_id].url,
                    }
                    for source_id in claim.source_ids[:2]
                    if source_id in source_by_id
                ],
            }
            for claim in selected_claims
        ],
    }


def _merge_live_report_revision(current_draft: dict, revision_patch: dict) -> dict:
    merged = json.loads(json.dumps(current_draft, ensure_ascii=False))
    valid_updates = revision_patch if isinstance(revision_patch, dict) else {}
    if str(valid_updates.get("subtitle", "")).strip():
        merged["subtitle"] = str(valid_updates["subtitle"]).strip()
    if str(valid_updates.get("facts_line", "")).strip():
        merged["facts_line"] = str(valid_updates["facts_line"]).strip()
    if str(valid_updates.get("executive_summary", "")).strip():
        merged["executive_summary"] = str(valid_updates["executive_summary"]).strip()

    sections = [dict(section) for section in merged.get("sections", [])]
    valid_urls = {
        url
        for section in sections
        for url in ([item for item in section.get("sources", []) if isinstance(item, str)] + _extract_markdown_links(str(section.get("content", ""))))
        if url.startswith("http")
    }

    def normalize_sources(raw_sources: object, content: str) -> list[str]:
        sources = [url for url in raw_sources or [] if isinstance(url, str) and url.startswith("http")]
        linked = _extract_markdown_links(content)
        combined = list(dict.fromkeys([*sources, *linked]))
        if combined:
            return combined[:6]
        return list(valid_urls)[:4]

    for update in valid_updates.get("section_updates", []) or []:
        if not isinstance(update, dict):
            continue
        content = str(update.get("content", "")).strip()
        title = str(update.get("title", "")).strip()
        target_title = str(update.get("target_title", "")).strip() or title
        if not content or not target_title:
            continue
        existing_index = next(
            (index for index, section in enumerate(sections) if _titles_match(target_title, str(section.get("title", "")))),
            None,
        )
        new_section = {
            "title": title or target_title,
            "content": content,
            "order": (sections[existing_index].get("order", existing_index + 1) if existing_index is not None else len(sections) + 1),
            "sources": normalize_sources(update.get("sources", []), content),
        }
        if existing_index is None:
            sections.append(new_section)
        else:
            sections[existing_index] = new_section

    for new_section_raw in valid_updates.get("new_sections", []) or []:
        if not isinstance(new_section_raw, dict):
            continue
        title = str(new_section_raw.get("title", "")).strip()
        content = str(new_section_raw.get("content", "")).strip()
        if not title or not content:
            continue
        if any(_titles_match(title, str(section.get("title", ""))) for section in sections):
            continue
        sections.append(
            {
                "title": title,
                "content": content,
                "order": len(sections) + 1,
                "sources": normalize_sources(new_section_raw.get("sources", []), content),
            }
        )

    for index, section in enumerate(sections, start=1):
        section["order"] = index
    merged["sections"] = sections
    return merged


async def _synthesize_longform_report(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    research_rows: list[dict],
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    coverage: CoverageReport,
    adjacent_questions: list[ResearchQuestion],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    *,
    record_spend: Callable[[SpendEntry], None] | None = None,
) -> dict:
    payload = _build_live_report_payload(
        task_spec,
        plan,
        research_rows,
        source_ledger,
        claims,
        coverage,
        adjacent_questions,
        critique_findings,
        decision_triggers,
    )
    prompt = LONGFORM_REPORT_PROMPT.format(language_name=_language_name(task_spec.request_spec.language))
    raw, spend_entry = _coerce_spend_result(
        await _call_report_writer_model(
            prompt,
            payload,
            LIVE_REPORT_MODEL,
            prefer_perplexity=_depth_profile(task_spec).prefer_perplexity_writer,
            direct_model=REVIEW_MODEL_DIRECT,
        ),
        category=SpendCategory.WRITER,
        stage="report_writer",
        provider="openrouter",
        model=LIVE_REPORT_MODEL,
    )
    if record_spend is not None:
        spend_entry.stage = "report_writer"
        spend_entry.category = SpendCategory.WRITER
        record_spend(spend_entry)
    parsed = parse_llm_json(raw, context="v2_longform_report")
    if _report_word_count(parsed) < 4200 or len(parsed.get("sections", [])) < 7:
        payload["revision_request"] = (
            "The previous draft was too short. Expand the report to at least 4,500 words with denser evidence, richer exhibits, and more operational detail in every section, but do not add unsupported exact numbers."
        )
        raw, retry_spend = _coerce_spend_result(
            await _call_report_writer_model(
                prompt,
                payload,
                LIVE_REPORT_MODEL,
                prefer_perplexity=_depth_profile(task_spec).prefer_perplexity_writer,
                direct_model=REVIEW_MODEL_DIRECT,
            ),
            category=SpendCategory.WRITER,
            stage="report_writer_retry",
            provider="openrouter",
            model=LIVE_REPORT_MODEL,
        )
        if record_spend is not None:
            retry_spend.stage = "report_writer_retry"
            retry_spend.category = SpendCategory.WRITER
            record_spend(retry_spend)
        parsed = parse_llm_json(raw, context="v2_longform_report_retry")
    return parsed


def _normalize_live_sections(parsed_sections: list[dict], source_ledger: list[SourceLedgerEntry]) -> list[ReportSection]:
    valid_urls = {entry.url for entry in source_ledger}
    fallback_urls = [entry.url for entry in source_ledger[:4]]
    sections: list[ReportSection] = []
    for index, raw_section in enumerate(parsed_sections, start=1):
        title = str(raw_section.get("title", "")).strip() or f"Section {index}"
        content = _sanitize_llm_markdown(raw_section.get("content", ""))
        if not content:
            continue
        sources = [url for url in raw_section.get("sources", []) if url in valid_urls]
        if not sources:
            sources = [url for url in _extract_markdown_links(content) if url in valid_urls]
        if not sources:
            sources = fallback_urls[:]
        sections.append(
            ReportSection(
                title=title,
                content=content,
                order=index,
                sources=list(dict.fromkeys(sources)),
            )
        )
    return sections


def _materialize_live_report_candidate(
    run_id: str,
    parsed_report: dict,
    source_ledger: list[SourceLedgerEntry],
    coverage: CoverageReport,
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    total_cost: float,
    brief: AnalysisBrief,
    language: str,
) -> tuple[ReportOutput, str, str, str]:
    report = _build_live_report_output(run_id, parsed_report, source_ledger, total_cost)
    report = _append_decision_addendum_sections(
        report,
        brief,
        source_ledger,
        coverage,
        critique_findings,
        decision_triggers,
        language,
    )
    subtitle = str(report.metadata.get("subtitle", "")).strip()
    facts_line = str(report.metadata.get("facts_line", "")).strip()
    markdown_text = _build_markdown_from_report(report, source_ledger, subtitle=subtitle, facts_line=facts_line)
    return report, subtitle, facts_line, markdown_text


def _assess_live_report_candidate(
    task_spec: TaskSpec,
    report: ReportOutput,
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    evidence: list[EvidenceRecord],
    coverage: CoverageReport,
    adjacent_questions: list[ResearchQuestion],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
) -> QualityAssessment:
    return assess_report_quality(
        task_spec,
        report,
        source_ledger,
        claims,
        evidence,
        coverage,
        adjacent_questions,
        critique_findings,
        decision_triggers,
    )


_AUDIT_RECOMMENDATION_PATTERNS = (
    r"\brecommend",
    r"\bshould\b",
    r"\bpriority\b",
    r"\baction\b",
    r"\bвывод",
    r"\bрекоменд",
)


def _report_body_for_grounding_scan(report_md: str) -> str:
    start_match = re.search(r"^##\s+Executive Summary\s*$", report_md, flags=re.MULTILINE)
    if not start_match:
        return report_md
    end_match = re.search(r"^##\s+Evidence Coverage and Source Quality\s*$", report_md, flags=re.MULTILINE)
    if not end_match:
        end_match = re.search(r"^##\s+Sources\s*$", report_md, flags=re.MULTILINE)
    start = start_match.start()
    end = end_match.start() if end_match else len(report_md)
    body = report_md[start:end]
    body = re.sub(r"(?im)^Exhibit\s+\d+.*$", " ", body)
    body = re.sub(r"(?im)^\|.*\|\s*$", " ", body)
    body = re.sub(r"\[Evidence:\s*(?!\s*C-\d+(?:\s*,\s*C-\d+)*\s*\])[^\]]+\]", " ", body, flags=re.IGNORECASE)
    body = re.sub(r"\[[alq]q?\d+[^\]]*\]", " ", body, flags=re.IGNORECASE)
    body = re.sub(r"\(\d+\)", " ", body)
    body = re.sub(r"\(Word count:[^)]+\)", " ", body, flags=re.IGNORECASE)
    return body


def _section_body_md(report_md: str, heading: str) -> str:
    match = re.search(
        rf"^##\s+{re.escape(heading)}\s*$([\s\S]*?)(?=^##\s+.+$|\Z)",
        report_md,
        flags=re.MULTILINE,
    )
    return match.group(1) if match else ""


def _recommendation_bullets_missing_evidence(report_md: str) -> list[str]:
    recommendation_block = _section_body_md(report_md, "Recommendation and Decision Posture")
    recommendation_lines = [
        line.strip()
        for line in recommendation_block.splitlines()
        if line.strip().startswith("- ")
        and any(re.search(pattern, line.lower()) for pattern in _AUDIT_RECOMMENDATION_PATTERNS)
    ]
    return [
        line
        for line in recommendation_lines
        if "[Evidence:" not in line
        and not re.search(r"\[[^\]]+\]\(https?://[^)]+\)", line)
        and "bounded" not in line.lower()
    ]


async def _run_live_compliance_revision(
    *,
    run_id: str,
    task_spec: TaskSpec,
    plan: ResearchPlan,
    parsed_report: dict,
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    evidence: list[EvidenceRecord],
    coverage: CoverageReport,
    adjacent_questions: list[ResearchQuestion],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    brief: AnalysisBrief,
    total_cost: float,
    assessment: QualityAssessment,
    iterations: list[QualityIteration],
    emit: EmitFn,
    record_spend: Callable[[SpendEntry], None] | None = None,
) -> tuple[dict, ReportOutput, str, str, str, QualityAssessment, list[QualityIteration]]:
    report, subtitle, facts_line, markdown_text = _materialize_live_report_candidate(
        run_id,
        parsed_report,
        source_ledger,
        coverage,
        critique_findings,
        decision_triggers,
        total_cost,
        brief,
        task_spec.request_spec.language,
    )
    unsupported = find_unsupported_precise_numbers(
        _report_body_for_grounding_scan(markdown_text),
        [claim.statement for claim in claims],
    )
    missing_evidence = _recommendation_bullets_missing_evidence(markdown_text)
    if not unsupported and not missing_evidence:
        return parsed_report, report, subtitle, facts_line, markdown_text, assessment, iterations

    revision_focus = ["grounding discipline"]
    if missing_evidence:
        revision_focus.append("decision usefulness")
    await emit(
        RunEvent(
            step="quality",
            status="started",
            message=(
                "Compliance revision: tightening grounding discipline"
                + (" and evidence-linked recommendations" if missing_evidence else "")
            ),
        )
    )
    try:
        revised_parsed = await _revise_longform_report(
            task_spec,
            plan,
            parsed_report,
            source_ledger,
            claims,
            critique_findings,
            decision_triggers,
            assessment,
            revision_focus,
            record_spend=record_spend,
            spend_category=SpendCategory.COMPLIANCE_REVISION,
            spend_stage="compliance_revision",
        )
    except Exception as exc:
        await emit(
            RunEvent(
                step="quality",
                status="warning",
                message=f"Compliance revision failed and was skipped: {exc}",
            )
        )
        return parsed_report, report, subtitle, facts_line, markdown_text, assessment, iterations

    candidate_report, candidate_subtitle, candidate_facts_line, candidate_markdown = _materialize_live_report_candidate(
        run_id,
        revised_parsed,
        source_ledger,
        coverage,
        critique_findings,
        decision_triggers,
        total_cost,
        brief,
        task_spec.request_spec.language,
    )
    candidate_assessment = _assess_live_report_candidate(
        task_spec,
        candidate_report,
        source_ledger,
        claims,
        evidence,
        coverage,
        adjacent_questions,
        critique_findings,
        decision_triggers,
    )
    candidate_unsupported = find_unsupported_precise_numbers(
        _report_body_for_grounding_scan(candidate_markdown),
        [claim.statement for claim in claims],
    )
    candidate_missing = _recommendation_bullets_missing_evidence(candidate_markdown)
    current_issue_count = len(unsupported) + len(missing_evidence)
    candidate_issue_count = len(candidate_unsupported) + len(candidate_missing)
    if candidate_issue_count < current_issue_count and candidate_assessment.overall_score >= assessment.overall_score - 1.0:
        iterations.append(
            build_quality_iteration(
                len(iterations),
                candidate_assessment,
                previous_score=assessment.overall_score,
                improved=candidate_assessment.overall_score > assessment.overall_score,
                revision_focus=revision_focus,
                consecutive_improvements=0,
                notes=[
                    f"Compliance revision reduced unsupported-or-unlinked issues from {current_issue_count} to {candidate_issue_count}.",
                ],
            )
        )
        await emit(
            RunEvent(
                step="quality",
                status="done",
                message=(
                    f"Compliance revision accepted: unsupported numbers {len(unsupported)} -> {len(candidate_unsupported)}, "
                    f"missing evidence bullets {len(missing_evidence)} -> {len(candidate_missing)}"
                ),
            )
        )
        return (
            revised_parsed,
            candidate_report,
            candidate_subtitle,
            candidate_facts_line,
            candidate_markdown,
            candidate_assessment,
            iterations,
        )

    iterations.append(
        build_quality_iteration(
            len(iterations),
            assessment,
            previous_score=assessment.overall_score,
            improved=False,
            revision_focus=revision_focus,
            consecutive_improvements=0,
            notes=[
                "Compliance revision was discarded because it did not materially reduce unsupported numbers or evidence-linkage gaps.",
            ],
        )
    )
    await emit(
        RunEvent(
            step="quality",
            status="warning",
            message="Compliance revision did not materially improve grounding discipline; kept the stronger draft",
        )
    )
    return parsed_report, report, subtitle, facts_line, markdown_text, assessment, iterations


async def _revise_longform_report(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    current_draft: dict,
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    assessment: QualityAssessment,
    revision_focus: list[str],
    *,
    record_spend: Callable[[SpendEntry], None] | None = None,
    spend_category: SpendCategory = SpendCategory.QUALITY_REVISION,
    spend_stage: str = "quality_revision",
) -> dict:
    revision_payload = _build_quality_revision_payload(
        task_spec,
        plan,
        current_draft,
        source_ledger,
        claims,
        critique_findings,
        decision_triggers,
        assessment,
        revision_focus,
    )
    prompt = QUALITY_REVISION_PROMPT.format(language_name=_language_name(task_spec.request_spec.language))
    raw, spend_entry = _coerce_spend_result(
        await _call_report_writer_model(
            prompt,
            revision_payload,
            LIVE_REPORT_MODEL,
            timeout_seconds=120,
            allow_fallback=True,
            prefer_perplexity=_depth_profile(task_spec).prefer_perplexity_writer,
            direct_model=REVIEW_MODEL_DIRECT,
        ),
        category=spend_category,
        stage=spend_stage,
        provider="openrouter",
        model=LIVE_REPORT_MODEL,
    )
    if record_spend is not None:
        spend_entry.category = spend_category
        spend_entry.stage = spend_stage
        record_spend(spend_entry)
    parsed = parse_llm_json(raw, context="v2_longform_revision")
    if isinstance(parsed, dict) and isinstance(parsed.get("sections"), list):
        return parsed
    return _merge_live_report_revision(current_draft, parsed)


async def _run_live_quality_revision_loop(
    *,
    run_id: str,
    task_spec: TaskSpec,
    plan: ResearchPlan,
    research_rows: list[dict],
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    evidence: list[EvidenceRecord],
    coverage: CoverageReport,
    adjacent_questions: list[ResearchQuestion],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    brief: AnalysisBrief,
    total_cost: float,
    initial_parsed_report: dict,
    emit: EmitFn,
    record_spend: Callable[[SpendEntry], None] | None = None,
) -> tuple[dict, ReportOutput, str, str, str, QualityAssessment, list[QualityIteration]]:
    profile = _depth_profile(task_spec)
    best_parsed = initial_parsed_report
    harmful_rounds = 0
    best_report, best_subtitle, best_facts_line, best_markdown = _materialize_live_report_candidate(
        run_id,
        best_parsed,
        source_ledger,
        coverage,
        critique_findings,
        decision_triggers,
        total_cost,
        brief,
        task_spec.request_spec.language,
    )
    best_assessment = _assess_live_report_candidate(
        task_spec,
        best_report,
        source_ledger,
        claims,
        evidence,
        coverage,
        adjacent_questions,
        critique_findings,
        decision_triggers,
    )

    iterations: list[QualityIteration] = [
        build_quality_iteration(
            0,
            best_assessment,
            revision_focus=[],
            consecutive_improvements=0,
            notes=["Initial long-form draft assessment."],
        )
    ]
    consecutive_improvements = 0

    for revision_round in range(1, profile.quality_max_rounds + 1):
        if consecutive_improvements >= profile.quality_revision_target:
            break
        revision_focus = build_revision_focus(best_assessment)
        await emit(
            RunEvent(
                step="quality",
                status="started",
                message=f"Quality revision round {revision_round}: focusing on {', '.join(revision_focus[:3])}",
            )
        )
        try:
            candidate_parsed = await _revise_longform_report(
                task_spec,
                plan,
                best_parsed,
                source_ledger,
                claims,
                critique_findings,
                decision_triggers,
                best_assessment,
                revision_focus,
                record_spend=record_spend,
                spend_category=SpendCategory.QUALITY_REVISION,
                spend_stage=f"quality_revision_{revision_round}",
            )
        except Exception as exc:
            iterations.append(
                build_quality_iteration(
                    revision_round,
                    best_assessment,
                    previous_score=best_assessment.overall_score,
                    revision_focus=revision_focus,
                    consecutive_improvements=consecutive_improvements,
                    notes=[f"Revision round failed: {exc}"],
                )
            )
            await emit(
                RunEvent(
                    step="quality",
                    status="warning",
                    message=f"Quality revision round {revision_round} failed: {exc}",
                )
            )
            break

        candidate_report, candidate_subtitle, candidate_facts_line, candidate_markdown = _materialize_live_report_candidate(
            run_id,
            candidate_parsed,
            source_ledger,
            coverage,
            critique_findings,
            decision_triggers,
            total_cost,
            brief,
            task_spec.request_spec.language,
        )
        candidate_assessment = _assess_live_report_candidate(
            task_spec,
            candidate_report,
            source_ledger,
            claims,
            evidence,
            coverage,
            adjacent_questions,
            critique_findings,
            decision_triggers,
        )
        delta = round(candidate_assessment.overall_score - best_assessment.overall_score, 2)
        improved = delta >= QUALITY_MIN_IMPROVEMENT_DELTA
        harmful_rounds = harmful_rounds + 1 if delta <= QUALITY_HARMFUL_DELTA else 0
        consecutive_improvements = consecutive_improvements + 1 if improved else 0
        iterations.append(
            build_quality_iteration(
                revision_round,
                candidate_assessment,
                previous_score=best_assessment.overall_score,
                improved=improved,
                revision_focus=revision_focus,
                consecutive_improvements=consecutive_improvements,
                notes=[f"Round delta: {delta:+.2f}"],
            )
        )
        if improved:
            best_parsed = candidate_parsed
            best_report = candidate_report
            best_subtitle = candidate_subtitle
            best_facts_line = candidate_facts_line
            best_markdown = candidate_markdown
            best_assessment = candidate_assessment
            await emit(
                RunEvent(
                    step="quality",
                    status="done",
                    message=f"Quality revision round {revision_round} improved overall score by {delta:+.2f} to {candidate_assessment.overall_score:.2f}",
                )
            )
        else:
            await emit(
                RunEvent(
                    step="quality",
                    status="warning",
                    message=f"Quality revision round {revision_round} did not improve enough ({delta:+.2f}); keeping the stronger draft",
                )
            )
            if harmful_rounds >= QUALITY_MAX_HARMFUL_ROUNDS:
                iterations[-1].notes.append("Quality loop stopped early after repeated harmful revisions.")
                await emit(
                    RunEvent(
                        step="quality",
                        status="warning",
                        message="Stopping quality loop after repeated harmful revisions to preserve the stronger draft",
                    )
                )
                break

    return best_parsed, best_report, best_subtitle, best_facts_line, best_markdown, best_assessment, iterations


def _build_markdown_from_report(
    report: ReportOutput,
    source_ledger: list[SourceLedgerEntry],
    *,
    subtitle: str = "",
    facts_line: str = "",
) -> str:
    lines = [f"# {report.title}", ""]
    if subtitle:
        lines.extend([subtitle, ""])
    if facts_line:
        lines.extend([facts_line, ""])
    lines.extend(["## Executive Summary", "", _normalize_markdown_tables(report.executive_summary), ""])
    for section in sorted(report.sections, key=lambda item: item.order):
        lines.extend([f"## {section.title}", "", _normalize_markdown_tables(section.content.strip()), ""])
    lines.extend(["## Sources", ""])
    seen_urls: set[str] = set()
    for source in source_ledger:
        if source.url in seen_urls:
            continue
        seen_urls.add(source.url)
        lines.append(f"- [{source.title}]({source.url})")
    lines.append("")
    return "\n".join(lines)


def _claim_link_suffix(claim: ClaimRecord, source_by_id: dict[str, SourceLedgerEntry]) -> str:
    links = [
        f"[{source_by_id[source_id].title}]({source_by_id[source_id].url})"
        for source_id in claim.source_ids
        if source_id in source_by_id
    ]
    if not links:
        return f"[Evidence: {claim.claim_id}]"
    return f"[Evidence: {claim.claim_id}] " + ", ".join(links[:2])


def _claims_table(exhibit_no: int, claims: list[ClaimRecord], source_by_id: dict[str, SourceLedgerEntry]) -> str:
    rows = [
        f"| {claim.claim_id} | {claim.statement[:140]} | {_claim_link_suffix(claim, source_by_id)} |"
        for claim in claims[:4]
    ]
    if not rows:
        rows = ["| C-00 | Evidence was limited for this slice | [Evidence: pending] |"]
    return "\n".join(
        [
            f"Exhibit {exhibit_no}",
            "",
            "| Claim | Evidence signal | Linkage |",
            "|---|---|---|",
            *rows,
        ]
    )


def _markdown_table_block(headers: list[str], rows: list[str], right_aligned_columns: set[int] | None = None) -> str:
    header_row = "| " + " | ".join(headers) + " |"
    right_aligned_columns = right_aligned_columns or set()
    separator_row = "|" + "|".join("---:" if index in right_aligned_columns else "---" for index, _ in enumerate(headers)) + "|"
    normalized_rows = [row for row in rows if row.strip()]
    return "\n".join([header_row, separator_row, *normalized_rows])


def _heuristic_longform_report(
    task_spec: TaskSpec,
    source_ledger: list[SourceLedgerEntry],
    claims: list[ClaimRecord],
    coverage: CoverageReport,
) -> dict:
    language = task_spec.request_spec.language
    section_titles = _live_section_titles(language)
    source_by_id = {entry.source_id: entry for entry in source_ledger}
    top_claims = claims[: max(8, min(len(claims), 24))]
    if not top_claims:
        top_claims = [
            ClaimRecord(
                claim_id="C-00",
                statement=(
                    "Evidence collection was partial; the report should be read as a structured market note rather than a fully verified recommendation."
                    if language != "ru"
                    else "Сбор доказательств был частичным; отчёт следует читать как структурированную аналитическую записку, а не как полностью верифицированную рекомендацию."
                ),
                question_id="fallback",
                supporting_evidence_ids=[],
                source_ids=[],
                confidence=0.2,
            )
        ]

    intro_templates = (
        [
            "Этот раздел фиксирует, как текущая проблема выглядит с точки зрения решения, а не только с точки зрения описания рынка. Для девелопера здесь важны не абстрактные тренды, а то, как меняются продукт, себестоимость, контроль качества, скорость передачи квартир и модель работы с подрядчиками.",
            "Управленческий смысл собранных сигналов состоит в том, что зрелые игроки почти никогда не рассматривают отделку как изолированную опцию. Она работает как связка из стандарта продукта, механики кастомизации, гарантийного сервиса, контроля дефектов и цифрового контура сопровождения клиента.",
            "Поэтому даже при неполной полноте источников полезно смотреть на тему через операционный контур: какие решения масштабируются, какие создают конфликт с покупателем, где возникают скрытые расходы и какие практики дают повторяемый эффект на маржу и скорость оборота.",
        ]
        if language == "ru"
        else [
            "This section frames the issue as a decision problem rather than a descriptive market note. The relevant lens is not abstract trend commentary but what changes product design, operating cost, quality control, delivery speed, and contractor governance in practice.",
            "The management implication of the evidence is that mature operators rarely treat the topic as a standalone option. It behaves as a system that combines product standardization, managed customization, warranty service, defect control, and a digital service layer around the buyer.",
            "Even when the evidence is incomplete, it is still useful to interpret the market through an operating lens: which practices scale, which create buyer conflict, where hidden costs appear, and which interventions have repeatable effects on margin and velocity.",
        ]
    )

    section_claim_groups: list[list[ClaimRecord]] = []
    for index in range(len(section_titles)):
        group = top_claims[index::len(section_titles)]
        if not group:
            group = top_claims[index % len(top_claims): index % len(top_claims) + 3]
        section_claim_groups.append(group[:4])

    sections: list[dict] = []
    for index, title in enumerate(section_titles, start=1):
        section_claims = section_claim_groups[index - 1]
        bullets = "\n".join(
            f"- {claim.statement} {_claim_link_suffix(claim, source_by_id)}"
            for claim in section_claims
        )
        section_body = "\n\n".join(
            [
                intro_templates[0],
                intro_templates[1],
                intro_templates[2],
                _claims_table(index, section_claims, source_by_id),
                bullets,
            ]
        )
        sections.append(
            {
                "title": title,
                "content": section_body,
                "order": index,
                "sources": [
                    source_by_id[source_id].url
                    for claim in section_claims
                    for source_id in claim.source_ids
                    if source_id in source_by_id
                ][:4],
            }
        )

    executive_summary = "\n\n".join(
        [
            (
                f"Текущий отчёт по теме «{task_spec.request_spec.subject}» собран в режиме fallback, потому что LLM-провайдеры для глубокой веб-аналитики недоступны. Это не снимает ценности материала: даже без синтетического long-form writer база из {len(source_ledger)} источников и {len(claims)} claims уже позволяет зафиксировать рабочие управленческие выводы."
                if language == "ru"
                else f"This report on '{task_spec.request_spec.subject}' was assembled in fallback mode because the external LLM providers for deep web research were unavailable. That does not make the output useless: even without a synthetic long-form writer, a base of {len(source_ledger)} sources and {len(claims)} claims is still enough to produce working management conclusions."
            ),
            (
                f"Покрытие ключевых вопросов составляет {coverage.covered_questions}/{coverage.total_questions}. Это означает, что по основным веткам можно сформулировать направление действий, но чувствительные зоны нужно дочищать точечной проверкой, особенно там, где остаются пробелы по стоимости, правовой модели или операционным ограничениям."
                if language == "ru"
                else f"Coverage of the core questions is {coverage.covered_questions}/{coverage.total_questions}. That is sufficient for a directional position, but the most sensitive areas still require targeted follow-up, especially where cost, legal model, or operating constraints remain under-specified."
            ),
            "\n".join(
                f"- {claim.statement} {_claim_link_suffix(claim, source_by_id)}"
                for claim in top_claims[:6]
            ),
        ]
    )

    return {
        "title": (
            f"{task_spec.request_spec.subject}: стратегический отчёт"
            if language == "ru"
            else f"{task_spec.request_spec.subject}: strategic report"
        ),
        "subtitle": (
            "Evidence-first long-form report assembled under degraded external provider conditions"
            if language != "ru"
            else "Evidence-first отчёт, собранный в режиме деградации внешних провайдеров"
        ),
        "facts_line": f"{len(source_ledger)} sources | {len(claims)} claims | {coverage.covered_questions}/{coverage.total_questions} core questions covered",
        "executive_summary": executive_summary,
        "sections": sections,
    }


async def _fallback_live_research(
    task_spec: TaskSpec,
    plan: ResearchPlan,
) -> tuple[list[dict], list[SourceLedgerEntry], list[EvidenceRecord], list[SourceSnapshot]]:
    provider = DuckDuckGoSearchProvider()
    source_map: dict[str, SourceLedgerEntry] = {}
    snapshot_by_url: dict[str, SourceSnapshot] = {}
    research_rows: list[dict] = []
    for question in plan.primary_questions:
        single_plan = ResearchPlan(
            primary_questions=[question],
            preferred_domains=plan.preferred_domains,
            required_source_mix=plan.required_source_mix,
        )
        all_candidates: list[SearchCandidate] = []
        for query in _build_question_fallback_queries(task_spec, question):
            all_candidates.extend(await provider.search(query, single_plan))
        if not all_candidates:
            broad_query = " ".join(re.findall(r"[\w-]+", question.question)[:12])
            if broad_query:
                all_candidates.extend(await provider.search(broad_query, single_plan))

        selected_sources = select_sources(all_candidates, single_plan)
        row_sources: list[SourceLedgerEntry] = []
        row_snapshots: list[SourceSnapshot] = []
        for source in selected_sources[:4]:
            existing = source_map.get(source.url)
            if existing is None:
                source_map[source.url] = source
                existing = source
            elif question.question_id not in existing.question_links:
                existing.question_links.append(question.question_id)
            row_sources.append(existing)

            snapshot = snapshot_by_url.get(existing.url)
            if snapshot is None:
                snapshot = await provider.fetch(existing)
                snapshot_by_url[existing.url] = snapshot
            if snapshot.fetch_status == "ok" and snapshot.content:
                row_snapshots.append(snapshot)

        row_lookup = {source.source_id: source.reliability_score for source in row_sources}
        row_question_links = {source.source_id: set(source.question_links) for source in row_sources}
        question_evidence = build_evidence_ledger(single_plan, row_snapshots, row_lookup, row_question_links)[:6]
        source_urls = [
            source.url
            for item in question_evidence
            for source in row_sources
            if source.source_id == item.source_id
        ]
        research_rows.append(
            {
                "question_id": question.question_id,
                "query": question.question,
                "confidence": 0.55 if question_evidence else 0.25,
                "gaps": [] if question_evidence else ["No usable evidence found for this branch"],
                "findings": [item.claim for item in question_evidence],
                "source_urls": list(dict.fromkeys(source_urls)),
                "sources": [
                    source.model_dump(mode="json")
                    for source in row_sources
                    if source.url in source_urls
                ],
            }
        )
    source_ledger = _rank_live_sources(list(source_map.values()), task_spec, plan, limit=16)
    allowlist = {entry.url for entry in source_ledger}
    snapshots = [snapshot for url, snapshot in snapshot_by_url.items() if url in allowlist]
    usable_snapshots = [item for item in snapshots if item.fetch_status == "ok" and item.content]
    source_lookup = {source.source_id: source.reliability_score for source in source_ledger}
    source_question_links = {source.source_id: set(source.question_links) for source in source_ledger}
    evidence = build_evidence_ledger(plan, usable_snapshots, source_lookup, source_question_links)
    for row in research_rows:
        row["source_urls"] = [url for url in row["source_urls"] if url in allowlist]
        row["sources"] = [source for source in row["sources"] if source.get("url") in allowlist]
    return research_rows, source_ledger, evidence, snapshots


def build_research_plan(task_spec: TaskSpec) -> ResearchPlan:
    pack = match_reference_pack(task_spec.request_spec.original_query)
    primary_questions = [
        ResearchQuestion(
            question_id=f"q{index + 1}",
            question=question,
            kind=QuestionKind.PRIMARY,
            priority=index + 1,
            required_evidence_count=2,
        )
        for index, question in enumerate(task_spec.must_cover_questions[:4])
    ]
    secondary_questions = [
        ResearchQuestion(
            question_id=f"s{index + 1}",
            question=f"Assess {dimension} implications for {task_spec.request_spec.subject}",
            kind=QuestionKind.SECONDARY,
            priority=index + 1,
            required_evidence_count=1,
        )
        for index, dimension in enumerate(task_spec.evaluation_dimensions[:4])
    ]
    suggested_search_queries = [
        f"{task_spec.request_spec.subject} {question.question} official docs benchmark"
        for question in primary_questions
    ]
    preferred_domains = sorted({source.domain for source in pack.sources}) if pack else []
    adjacent_candidates = build_adjacent_question_candidates(task_spec)
    return ResearchPlan(
        primary_questions=primary_questions,
        secondary_questions=secondary_questions,
        adjacent_question_candidates=adjacent_candidates,
        claims_to_validate=[f"{task_spec.request_spec.subject}: {dimension}" for dimension in task_spec.evaluation_dimensions[:5]],
        claims_to_disprove=[f"Unqualified winner claim for {task_spec.request_spec.subject}"],
        required_evidence_per_question={question.question_id: question.required_evidence_count for question in primary_questions},
        suggested_search_queries=suggested_search_queries,
        preferred_domains=preferred_domains,
        required_source_mix=task_spec.allowed_source_types,
        chart_candidates=["evidence_coverage", "source_mix"],
        stop_conditions=["At least 2 sources per primary question", "No unresolved critical contradiction cluster"],
    )


def build_perplexity_handoff_prompts(task_spec: TaskSpec, plan: ResearchPlan) -> list[PerplexityHandoffPrompt]:
    language = task_spec.request_spec.language
    dimensions = ", ".join(task_spec.evaluation_dimensions[:4])
    primary_questions = "\n".join(f"- {question.question}" for question in plan.primary_questions[:4])
    if language == "ru":
        prompts = [
            PerplexityHandoffPrompt(
                stage="initial_landscape",
                title="Core deep research prompt",
                rationale="Собирает основное пространство вариантов, факторы выбора и сильнейшие первичные источники.",
                prompt=(
                    f"Подготовь decision-grade deep research по теме: {task_spec.request_spec.original_query}\n\n"
                    f"Контекст решения: {task_spec.request_spec.decision_context}\n"
                    f"География: {task_spec.request_spec.geography}\n"
                    f"Критерии выбора: {dimensions}\n"
                    "Обязательно покрой вопросы:\n"
                    f"{primary_questions}\n\n"
                    "Требования к ответу:\n"
                    "- опирайся на official docs, pricing pages, GitHub repos, benchmark pages и сильные кейсы\n"
                    "- явно покажи option space, а не только один фокусный вариант\n"
                    "- отмечай неизвестности и спорные места\n"
                    "- не выдумывай точные цифры без опоры на источник\n"
                    "- в конце дай structured findings, alternatives, risks, decision triggers и references"
                ),
            ),
            PerplexityHandoffPrompt(
                stage="alternatives_and_tradeoffs",
                title="Alternative space and trade-offs",
                rationale="Выдёргивает альтернативы и контраргументы, которые модели часто пропускают в первом проходе.",
                prompt=(
                    f"Критически проверь пространство альтернатив по теме: {task_spec.request_spec.subject}\n\n"
                    f"Исходный запрос: {task_spec.request_spec.original_query}\n"
                    f"Оси сравнения: {dimensions}\n"
                    "Сделай акцент на:\n"
                    "- какие реальные альтернативы чаще всего недооценивают\n"
                    "- где текущая рекомендация может быть ошибочной\n"
                    "- какие trade-offs, hidden variables и switch conditions реально меняют выбор\n"
                    "- какие вопросы нужно явно дозакрыть до финальной рекомендации"
                ),
            ),
            PerplexityHandoffPrompt(
                stage="validation_backlog",
                title="Validation backlog and objections",
                rationale="Формирует explicit validation backlog, чтобы не выпускать красивый, но хрупкий вывод.",
                prompt=(
                    f"Ты читаешь почти готовый аналитический отчёт по теме: {task_spec.request_spec.subject}\n"
                    f"Решение: {task_spec.request_spec.decision_context}\n\n"
                    "Дай именно senior-analyst critique:\n"
                    "- что недодоказано\n"
                    "- какие objections самые опасные\n"
                    "- какие 5 проверок сильнее всего изменят качество решения\n"
                    "- при каких условиях рекомендация должна переключиться на альтернативу\n"
                    "- какие данные или внутренние материалы клиенту нужно добавить в следующий проход"
                ),
            ),
        ]
    else:
        prompts = [
            PerplexityHandoffPrompt(
                stage="initial_landscape",
                title="Core deep research prompt",
                rationale="Maps the option space, evidence base, and strongest primary sources.",
                prompt=(
                    f"Prepare a decision-grade deep research package for: {task_spec.request_spec.original_query}\n\n"
                    f"Decision context: {task_spec.request_spec.decision_context}\n"
                    f"Geography: {task_spec.request_spec.geography}\n"
                    f"Decision criteria: {dimensions}\n"
                    "Must cover:\n"
                    f"{primary_questions}\n\n"
                    "Requirements:\n"
                    "- prioritize official docs, pricing pages, GitHub repos, benchmark pages, and serious case evidence\n"
                    "- show the true option space, not just the focal candidate\n"
                    "- flag unresolved unknowns and contradictions\n"
                    "- avoid unsupported precise numbers\n"
                    "- finish with structured findings, alternatives, risks, decision triggers, and references"
                ),
            ),
            PerplexityHandoffPrompt(
                stage="alternatives_and_tradeoffs",
                title="Alternative space and trade-offs",
                rationale="Pulls out alternative paths and counterarguments that often get missed in the first pass.",
                prompt=(
                    f"Critically challenge the current option space for: {task_spec.request_spec.subject}\n\n"
                    f"Original ask: {task_spec.request_spec.original_query}\n"
                    f"Comparison axes: {dimensions}\n"
                    "Focus on:\n"
                    "- credible alternatives that are often overlooked\n"
                    "- where the current recommendation is likely wrong\n"
                    "- tradeoffs, hidden variables, and switch conditions that materially change the decision\n"
                    "- which questions must be explicitly closed before a final recommendation"
                ),
            ),
            PerplexityHandoffPrompt(
                stage="validation_backlog",
                title="Validation backlog and objections",
                rationale="Builds an explicit validation backlog instead of a fragile polished conclusion.",
                prompt=(
                    f"You are reviewing a near-final analytical report on: {task_spec.request_spec.subject}\n"
                    f"Decision to support: {task_spec.request_spec.decision_context}\n\n"
                    "Provide senior-analyst critique on:\n"
                    "- what remains under-evidenced\n"
                    "- the most dangerous objections\n"
                    "- the 5 checks that would most improve decision quality\n"
                    "- the conditions that should flip the recommendation\n"
                    "- which internal materials or data should be added next"
                ),
            ),
        ]
    return prompts


def _build_validation_questions(
    task_spec: TaskSpec,
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    coverage: CoverageReport | None = None,
) -> list[ResearchQuestion]:
    profile = _depth_profile(task_spec)
    limit = profile.validation_research_branches
    if limit <= 0:
        return []
    questions: list[ResearchQuestion] = []
    if coverage is not None:
        gap_questions = [item for item in coverage.questions if item.status != "covered"]
        for gap in gap_questions:
            if len(questions) >= limit:
                break
            if gap.question_id == "q4":
                question_text = (
                    "Close the uncovered core question on tradeoffs, risks, buyer objections, pricing pressure, procurement friction, and recommendation-switch conditions for the final recommendation."
                    if task_spec.request_spec.language != "ru"
                    else "Закрой незакрытый core-вопрос про trade-offs, риски, buyer objections, ценовое давление, procurement friction и условия переключения рекомендации."
                )
                kind = QuestionKind.ADJACENT_BOUNDARY
            elif gap.question_id == "q2":
                question_text = (
                    "Close the uncovered core question on competitors, incumbent workflows, and free substitutes that should be compared against the focal product."
                    if task_spec.request_spec.language != "ru"
                    else "Закрой незакрытый core-вопрос про конкурентов, incumbent workflows и бесплатные substitutes, с которыми нужно сравнить продукт."
                )
                kind = QuestionKind.ADJACENT_ALTERNATIVE
            else:
                question_text = (
                    f"Close this uncovered core question with concrete evidence for the final recommendation: {gap.question}"
                    if task_spec.request_spec.language != "ru"
                    else f"Закрой этот незакрытый core-вопрос конкретными фактами для финальной рекомендации: {gap.question}"
                )
                kind = QuestionKind.ADJACENT_HIDDEN_VARIABLE
            questions.append(
                ResearchQuestion(
                    question_id=f"vg{len(questions) + 1}",
                    question=question_text,
                    kind=kind,
                    priority=len(questions) + 1,
                    required_evidence_count=2,
                )
            )
    for index, finding in enumerate(critique_findings[:limit], start=1):
        if len(questions) >= limit:
            break
        summary = " ".join(finding.summary.split()).strip()
        if not summary:
            continue
        question_text = (
            f"Validate or disprove this unresolved issue for the final recommendation: {summary}"
            if task_spec.request_spec.language != "ru"
            else f"Проверь и по возможности опровергни этот критичный незакрытый вопрос для финальной рекомендации: {summary}"
        )
        questions.append(
            ResearchQuestion(
                question_id=f"v{index}",
                question=question_text,
                kind=QuestionKind.ADJACENT_BOUNDARY,
                priority=index,
                required_evidence_count=2,
            )
        )
    trigger_slots = max(0, limit - len(questions))
    for offset, trigger in enumerate(decision_triggers[:trigger_slots], start=1):
        question_text = (
            f"Test the switch condition '{trigger.label}': {trigger.condition}. What evidence would confirm or weaken it?"
            if task_spec.request_spec.language != "ru"
            else f"Проверь decision trigger '{trigger.label}': {trigger.condition}. Какие факты подтверждают его или ослабляют?"
        )
        questions.append(
            ResearchQuestion(
                question_id=f"vt{offset}",
                question=question_text,
                kind=QuestionKind.ADJACENT_BOUNDARY,
                priority=len(questions) + 1,
                required_evidence_count=2,
            )
        )
    return questions[:limit]


def _truncate_phrase(text: str, max_chars: int = 140) -> str:
    normalized = " ".join(str(text or "").split()).strip()
    if len(normalized) <= max_chars:
        return normalized
    cutoff = normalized.rfind(" ", 0, max_chars - 1)
    if cutoff < max_chars // 2:
        cutoff = max_chars - 1
    return normalized[:cutoff].rstrip(" ,;:-") + "…"


def _extract_parenthetical_examples(text: str) -> str:
    match = re.search(r"\((?:e\.g\.,?|for example|например)\s*([^)]+)\)", text, flags=re.IGNORECASE)
    if not match:
        return ""
    examples = re.sub(r"\b(and|or)\b", ",", match.group(1), flags=re.IGNORECASE)
    parts = [part.strip(" .") for part in examples.split(",") if part.strip(" .")]
    if not parts:
        return ""
    return ", ".join(parts[:4])


def _adjacent_question_brief_line(question: ResearchQuestion, language: str = "en") -> str:
    text = " ".join(question.question.split()).strip()
    lowered = text.lower()
    examples = _extract_parenthetical_examples(text)

    if language == "ru":
        if question.kind == QuestionKind.ADJACENT_ALTERNATIVE:
            if examples:
                return f"Явный comparison set должен включать {examples}, а не только фокальный продукт."
            return "Отчёт должен явно сравнить фокальный вариант с сильнейшими реальными альтернативами."
        if question.kind == QuestionKind.ADJACENT_COUNTERARGUMENT:
            if any(marker in lowered for marker in ("chatgpt", "perplexity", "free", "бесплат")):
                return "Нужно честно проверить strongest counter-case: останутся ли покупатели на ChatGPT, Perplexity или собранном вручную бесплатном стеке вместо платного продукта."
            return "Отчёт должен явно разобрать strongest counter-case, а не только защищать основной вывод."
        if question.kind == QuestionKind.ADJACENT_HIDDEN_VARIABLE:
            return "Скрытые переменные вроде стоимости внедрения, procurement friction и операционной нагрузки всё ещё могут перевернуть вывод."
        if question.kind == QuestionKind.ADJACENT_BOUNDARY:
            return "Нужно явно описать boundary conditions: при каких порогах рынка, ROI или сложности рекомендация перестаёт работать."
        if question.kind == QuestionKind.ADJACENT_STAKEHOLDER:
            return "Возражения CFO, procurement и владельцев процесса по ROI и integration risk должны быть закрыты явно."
        if question.kind == QuestionKind.ADJACENT_TIME_SHIFT:
            return "Вывод нужно проверить на сдвиг горизонта 6-12 месяцев: рынок и бесплатные альтернативы могут быстро поменять картину."
        return _truncate_phrase(text, 160)

    if question.kind == QuestionKind.ADJACENT_ALTERNATIVE:
        if examples:
            return f"The comparison set should explicitly include {examples}, not just the focal product."
        return "The report should explicitly compare the focal option against the strongest credible alternatives."
    if question.kind == QuestionKind.ADJACENT_COUNTERARGUMENT:
        if any(marker in lowered for marker in ("chatgpt", "perplexity", "free", "free options")):
            return "The strongest counter-case is that buyers may stay with ChatGPT, Perplexity, or stitched free workflows instead of paying for an integrated product."
        return "The report must test the strongest counter-case instead of only defending the current thesis."
    if question.kind == QuestionKind.ADJACENT_HIDDEN_VARIABLE:
        return "Hidden variables such as implementation burden, procurement friction, and operating cost may still overturn the recommendation."
    if question.kind == QuestionKind.ADJACENT_BOUNDARY:
        return "The recommendation still needs explicit boundary conditions showing when market size, ROI, or complexity make it uneconomic."
    if question.kind == QuestionKind.ADJACENT_STAKEHOLDER:
        return "CFO, procurement, and workflow-owner objections around ROI and integration risk still need explicit closure."
    if question.kind == QuestionKind.ADJACENT_TIME_SHIFT:
        return "The conclusion should be stress-tested against a 6-12 month market shift, especially if free tools keep improving."
    return _truncate_phrase(text, 160)


def _coverage_gap_brief_line(gap: str, language: str = "en") -> str:
    normalized = " ".join(str(gap or "").split()).strip().rstrip("?")
    lowered = normalized.lower()
    if language == "ru":
        if lowered.startswith("what are the strongest evidence-backed tradeoffs") or "tradeoffs" in lowered:
            return "Tradeoffs, risks и decision triggers всё ещё недодоказаны и требуют отдельного закрытия."
        if lowered.startswith("what credible alternatives should be compared"):
            return "Пространство альтернатив всё ещё недосравнено по качеству, стоимости и операционному риску."
        if lowered.startswith("which option or stack best supports"):
            return "Текущая рекомендация всё ещё слабо привязана к контексту решения и требует более жёсткого option-to-context fit."
        if lowered.startswith("what concrete decision should this report support"):
            return "Формулировка самого решения всё ещё нуждается в более чёткой рамке."
        return _truncate_phrase(normalized + ".", 160)

    if lowered.startswith("what are the strongest evidence-backed tradeoffs") or "tradeoffs" in lowered:
        return "Tradeoffs, risks, and recommendation-switch conditions remain under-evidenced."
    if lowered.startswith("what credible alternatives should be compared"):
        return "The alternative space still needs explicit comparison across quality, cost, and operating risk."
    if lowered.startswith("which option or stack best supports"):
        return "The current recommendation still needs a cleaner option-to-context fit."
    if lowered.startswith("what concrete decision should this report support"):
        return "The decision frame still needs sharper definition."
    return _truncate_phrase(normalized + ".", 160)


def _first_number(text: str) -> float | None:
    matches = re.findall(r"(\d+(?:[.,]\d+)?)", text)
    if not matches:
        return None
    parsed: list[float] = []
    for value in matches:
        try:
            parsed.append(float(value.replace(",", ".")))
        except ValueError:
            continue
    if not parsed:
        return None
    non_year_values = [value for value in parsed if not (1900 <= value <= 2100 and float(value).is_integer())]
    candidates = non_year_values or []
    if not candidates:
        return None
    return candidates[0]


_CONTRADICTION_STOPWORDS = {
    "this",
    "that",
    "with",
    "from",
    "have",
    "will",
    "into",
    "their",
    "about",
    "which",
    "where",
    "what",
    "using",
    "than",
    "between",
    "across",
    "through",
    "enterprise",
    "market",
    "intelligence",
    "architecture",
    "decision",
}


def _claim_numeric_signature(text: str) -> str:
    lowered = text.lower()
    if "%" in lowered or "percent" in lowered or "cagr" in lowered or "roi" in lowered:
        return "ratio"
    if any(token in lowered for token in ("$", "usd", "eur", "€", "million", "billion", "млн", "млрд")):
        return "money"
    if any(token in lowered for token in ("month", "months", "year", "years", "quarter", "quarters", "месяц", "год", "квартал")):
        return "time"
    return "plain"


def _claim_contradiction_tokens(text: str) -> set[str]:
    lowered = text.lower()
    if any(marker in lowered for marker in ("|", "###", "<br", "**", "`")):
        return set()
    return {
        token
        for token in re.findall(r"[a-zA-Zа-яА-Я]{4,}", lowered)
        if token not in _CONTRADICTION_STOPWORDS
    }


_LOW_SIGNAL_CLAIM_PATTERNS = (
    "welcome to",
    "if you're interested",
    "follow these steps",
    "fork this repository",
    "curated collection",
    "original source attribution",
    "last updated:",
    "how we picked our top",
    "copy chevron-down",
    "clone the repository",
    "usage option comparison table",
    "out-of-the-box ui",
    "the multi agent team",
    "deployments and usage",
    "follow the quickstart",
    "your question here",
    "response = client.responses.create",
    "if you've got a moment",
    "the following table compares",
    "for more information, see",
)


def _is_low_signal_claim_text(text: str) -> bool:
    lowered = text.lower()
    if any(marker in lowered for marker in _LOW_SIGNAL_CLAIM_PATTERNS):
        return True
    if any(marker in text for marker in ("вЂ", "рџ", "\ufffd")):
        return True
    if re.search(r"\bstep\s+\d+\b", lowered):
        return True
    if re.search(r"\b(git clone|pip install|npm install|cd [a-z0-9_./-]+)\b", lowered):
        return True
    if re.search(r"\b(skip to content|sign in|privacy|terms|cookies|copy|chevron|hashtag)\b", lowered):
        return True
    return False


def _select_report_worthy_claims(claims: list[ClaimRecord]) -> list[ClaimRecord]:
    preferred = [
        claim
        for claim in claims
        if not claim.contradiction_notes and len(claim.source_ids) >= 1 and not _is_low_signal_claim_text(claim.statement)
    ]
    if preferred:
        return preferred
    return [claim for claim in claims if not _is_low_signal_claim_text(claim.statement)] or claims


def _sentence_candidates(snapshot: SourceSnapshot, question: ResearchQuestion) -> list[str]:
    parts = [snapshot.title, snapshot.excerpt]
    parts.extend(re.split(r"(?<=[.!?])\s+", snapshot.content))
    question_tokens = {token for token in re.findall(r"[a-zA-Zа-яА-Я0-9]+", question.question.lower()) if len(token) > 3}
    ranked: list[tuple[int, str]] = []
    for sentence in parts:
        normalized = " ".join(sentence.split()).strip()
        if len(normalized) < 40 or len(normalized) > 320:
            continue
        lowered = normalized.lower()
        if _is_low_signal_claim_text(normalized):
            continue
        if any(
            marker in lowered
            for marker in (
                "watchers",
                "forks",
                "stars",
                "footer ©",
                "releases no releases",
                "packages 0",
                "sign in",
                "skip to content",
                "github, inc",
                "activity stars",
                "topics",
                "cookies",
                "privacy",
                "terms",
                "welcome to",
                "if you're interested",
                "follow these steps",
                "fork this repository",
                "curated collection",
                "original source attribution",
                "last updated:",
                "how we picked our top",
                "clone the repository",
                "usage option comparison table",
                "out-of-the-box ui",
                "the multi agent team",
                "deployments and usage",
                "follow the quickstart",
                "your question here",
                "response = client.responses.create",
                "for more information, see",
                "if you've got a moment",
                "the following table compares",
            )
        ):
            continue
        score = 0
        if any(token in lowered for token in question_tokens):
            score += 2
        if re.search(r"\d", normalized):
            score += 2
        if any(token in lowered for token in ("tradeoff", "risk", "strong", "governance", "cost", "deploy", "quality", "benchmark", "leaderboard", "compare", "stack", "citation")):
            score += 1
        if score > 0:
            ranked.append((score, normalized))
    ranked.sort(key=lambda item: item[0], reverse=True)
    return [sentence for _, sentence in ranked[:3]]


def build_evidence_ledger(
    plan: ResearchPlan,
    snapshots: list[SourceSnapshot],
    source_lookup: dict[str, float],
    source_question_links: dict[str, set[str]] | None = None,
) -> list[EvidenceRecord]:
    evidence: list[EvidenceRecord] = []
    for question in plan.primary_questions:
        for snapshot in snapshots:
            if source_question_links:
                linked_questions = source_question_links.get(snapshot.source_id)
                if linked_questions and question.question_id not in linked_questions:
                    continue
            for sentence in _sentence_candidates(snapshot, question):
                confidence = min(0.95, source_lookup.get(snapshot.source_id, 0.6) + (0.1 if re.search(r"\d", sentence) else 0.0))
                evidence.append(
                    EvidenceRecord(
                        question_id=question.question_id,
                        source_id=snapshot.source_id,
                        claim=sentence,
                        snippet=sentence,
                        confidence=confidence,
                    )
                )
    return evidence


def detect_contradictions(claims: list[ClaimRecord]) -> list[str]:
    return _detect_grounded_contradictions(claims)


def build_claim_table(evidence: list[EvidenceRecord]) -> list[ClaimRecord]:
    claims: list[ClaimRecord] = []
    by_fingerprint: dict[str, ClaimRecord] = {}
    for index, item in enumerate(evidence, start=1):
        normalized_claim = re.sub(r"\W+", " ", item.claim.lower()).strip()
        fingerprint = f"{item.question_id}::{normalized_claim}"
        claim = by_fingerprint.get(fingerprint)
        if claim is None:
            claim = ClaimRecord(
                claim_id=f"C-{index:02d}",
                statement=item.claim,
                question_id=item.question_id,
                supporting_evidence_ids=[item.evidence_id],
                source_ids=[item.source_id],
                confidence=item.confidence,
            )
            by_fingerprint[fingerprint] = claim
            claims.append(claim)
        else:
            claim.supporting_evidence_ids.append(item.evidence_id)
            if item.source_id not in claim.source_ids:
                claim.source_ids.append(item.source_id)
            claim.confidence = max(claim.confidence, item.confidence)
    contradictions = detect_contradictions(claims)
    for claim in claims:
        claim.recommendation_safe = claim.confidence >= 0.75 and not claim.contradiction_notes
    claims.sort(key=lambda item: item.confidence, reverse=True)
    if contradictions:
        claims.sort(key=lambda item: (len(item.contradiction_notes), -item.confidence))
    return claims


def build_coverage_report(plan: ResearchPlan, claims: list[ClaimRecord], source_count: int, contradiction_notes: list[str]) -> CoverageReport:
    question_map: list[CoverageQuestionStatus] = []
    covered = 0
    for question in plan.primary_questions:
        question_claims = [claim for claim in claims if claim.question_id == question.question_id]
        status = "covered" if len(question_claims) >= question.required_evidence_count else "gap"
        if status == "covered":
            covered += 1
        question_map.append(
            CoverageQuestionStatus(
                question_id=question.question_id,
                question=question.question,
                evidence_count=sum(len(claim.supporting_evidence_ids) for claim in question_claims),
                source_count=sum(len(claim.source_ids) for claim in question_claims),
                status=status,
            )
        )
    total = max(1, len(plan.primary_questions))
    strong_source_ratio = min(1.0, source_count / max(1, len(plan.primary_questions) * 2))
    return CoverageReport(
        total_questions=len(plan.primary_questions),
        covered_questions=covered,
        coverage_ratio=covered / total,
        strong_source_ratio=strong_source_ratio,
        contradiction_count=len(contradiction_notes),
        questions=question_map,
        gaps=[item.question for item in question_map if item.status != "covered"],
    )


def build_analysis_brief(
    task_spec: TaskSpec,
    claims: list[ClaimRecord],
    coverage: CoverageReport,
    adjacent_questions: list[ResearchQuestion] | None = None,
    critique_findings: list[CritiqueFinding] | None = None,
    decision_triggers: list[DecisionTrigger] | None = None,
) -> AnalysisBrief:
    adjacent_questions = adjacent_questions or []
    critique_findings = critique_findings or []
    decision_triggers = decision_triggers or []
    safe_claims = [claim for claim in _select_report_worthy_claims(claims) if claim.recommendation_safe][:5]
    findings = [f"{claim.statement} [Evidence: {', '.join(claim.supporting_evidence_ids[:2])}]" for claim in safe_claims]
    language = task_spec.request_spec.language
    limitations = [_coverage_gap_brief_line(gap, language) for gap in coverage.gaps]
    if coverage.contradiction_count:
        limitations.append("At least one contradiction cluster remains unresolved.")
    critique_priorities = [item.summary for item in critique_findings[:4]]
    option_space: list[str] = []
    for question in adjacent_questions:
        if question.kind not in {QuestionKind.ADJACENT_ALTERNATIVE, QuestionKind.ADJACENT_COUNTERARGUMENT}:
            continue
        line = _adjacent_question_brief_line(question, language)
        if line:
            option_space.append(line)
    critical_unknowns = limitations + critique_priorities[:2]
    trigger_lines = [f"{item.label}: {item.condition} -> {item.implication}" for item in decision_triggers[:4]]
    recommendation_posture = "bounded_analysis_only"
    if coverage.coverage_ratio >= 0.66 and coverage.contradiction_count == 0 and len(safe_claims) >= 3:
        recommendation_posture = "evidence_backed_recommendations_allowed"
    summary = (
        f"{task_spec.request_spec.subject}: coverage {coverage.covered_questions}/{coverage.total_questions} primary questions, "
        f"{len(safe_claims)} recommendation-safe claims, contradiction count {coverage.contradiction_count}."
    )
    risks = list(task_spec.constraints)
    if not risks:
        risks.append("Final choice still depends on team-specific workload validation.")
    return AnalysisBrief(
        title=f"{task_spec.request_spec.subject}: Decision Brief",
        executive_summary=summary,
        decision_context=task_spec.request_spec.decision_context,
        recommendation_posture=recommendation_posture,
        key_findings=findings or ["Evidence is still too thin for a strong recommendation."],
        key_risks=risks,
        option_space=option_space or ["Alternative space still needs explicit comparison."],
        critical_unknowns=critical_unknowns or ["The recommendation is still sensitive to untested conditions."],
        decision_triggers=trigger_lines,
        improvement_priorities=critique_priorities or ["Explicitly compare alternatives and define failure conditions."],
        limitations=limitations or ["No major limitations recorded."],
        uncertainty_statement=(
            "Recommendations are bounded because the evidence base is incomplete."
            if recommendation_posture != "evidence_backed_recommendations_allowed"
            else "Evidence quality is sufficient for a directional recommendation, but local validation is still advised."
        ),
        chart_candidates=["evidence_coverage"],
    )


def _decision_addendum_sections(brief: AnalysisBrief, language: str) -> list[tuple[str, list[str]]]:
    if language == "ru":
        return [
            ("Пространство альтернатив", brief.option_space or ["Явное сравнение альтернатив пока недосформировано."]),
            ("Что может изменить рекомендацию", brief.decision_triggers or ["Рекомендация чувствительна к смене доминирующего ограничения."]),
            ("Неизвестное и следующие вопросы", brief.critical_unknowns or ["Есть открытые вопросы, которые ещё могут изменить вывод."]),
        ]
    return [
        ("Option Space", brief.option_space or ["The explicit alternative space is still underdeveloped."]),
        ("What Could Change The Recommendation", brief.decision_triggers or ["The recommendation is still sensitive to a change in the dominant constraint."]),
        ("Unknowns and Next Questions", brief.critical_unknowns or ["Open questions still remain and can alter the conclusion."]),
    ]


def _table_cell(value: object, max_chars: int | None = None) -> str:
    normalized = " ".join(str(value or "").split()).strip()
    normalized = normalized.replace("|", "/")
    if max_chars:
        normalized = _truncate_phrase(normalized, max_chars)
    return normalized or "-"


def _source_type_label(source_type: SourceType, language: str) -> str:
    if language == "ru":
        labels = {
            SourceType.OFFICIAL_DOCUMENTATION: "official doc",
            SourceType.VENDOR_PAGE: "vendor page",
            SourceType.GOVERNMENT: "government",
            SourceType.RESEARCH_PAPER: "research paper",
            SourceType.BENCHMARK: "benchmark",
            SourceType.USER_MATERIAL: "user material",
            SourceType.HIGH_QUALITY_SECONDARY: "high-quality secondary",
            SourceType.WEAK_SECONDARY: "weak secondary",
        }
        return labels.get(source_type, source_type.value.replace("_", " "))
    return source_type.value.replace("_", " ")


def _coverage_status_label(status: str, language: str) -> str:
    if language == "ru":
        return "покрыто" if status == "covered" else "gap"
    return status


def _traceability_appendix_sections(
    *,
    coverage: CoverageReport,
    source_ledger: list[SourceLedgerEntry],
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    language: str,
) -> list[tuple[str, str]]:
    top_sources = source_ledger[:6]
    coverage_rows = [
        f"| {question.question_id.upper()} | {_table_cell(question.question, 120)} | {_coverage_status_label(question.status, language)} | {question.evidence_count} | {question.source_count} |"
        for question in coverage.questions
    ] or ["| Q-? | - | gap | 0 | 0 |"]
    source_rows = [
        f"| [{_table_cell(source.title, 72)}]({source.url}) | {_source_type_label(source.source_type, language)} | {source.reliability_score:.2f} | {_table_cell(', '.join(source.question_links[:4]), 48)} |"
        for source in top_sources
    ] or ["| - | - | 0.00 | - |"]
    critique_rows = [
        f"| {_table_cell(item.kind.value)} | {_table_cell(item.severity)} | {_table_cell(item.summary, 140)} |"
        for item in critique_findings[:4]
    ] or ["| decision_risk | medium | Additional validation priorities were not captured in time. |"]
    trigger_rows = [
        f"| {_table_cell(item.label)} | {_table_cell(item.condition, 140)} | {_table_cell(item.implication, 140)} | {item.confidence:.2f} |"
        for item in decision_triggers[:4]
    ] or ["| No trigger captured | Further validation needed | Keep recommendation bounded | 0.00 |"]

    if language == "ru":
        coverage_title = "Покрытие доказательств и качество источников"
        validation_title = "Матрица валидации и decision triggers"
        coverage_intro = (
            "Этот appendix делает рекомендацию проверяемой: показывает, какие core questions реально покрыты, "
            "какие источники несут наибольший вес, и сколько противоречий осталось открытыми."
        )
        coverage_summary = (
            f"- Core questions covered: {coverage.covered_questions}/{coverage.total_questions}\n"
            f"- Open contradiction clusters: {coverage.contradiction_count}\n"
            f"- Source pack kept intentionally narrow to avoid noisy SEO and market-report filler."
        )
        validation_intro = (
            "Сильный аналитический отчёт не заканчивается красивым выводом. Он оставляет explicit validation backlog: "
            "что ещё перепроверить, какие objections самые опасные и при каких условиях recommendation должна переключиться."
        )
    else:
        coverage_title = "Evidence Coverage and Source Quality"
        validation_title = "Validation Priorities and Decision Triggers"
        coverage_intro = (
            "This appendix makes the recommendation auditable rather than merely persuasive. "
            "It shows which core questions are actually covered, which sources carry the most weight, and how much contradiction remains unresolved."
        )
        coverage_summary = (
            f"- Core questions covered: {coverage.covered_questions}/{coverage.total_questions}\n"
            f"- Open contradiction clusters: {coverage.contradiction_count}\n"
            f"- Source pack kept intentionally narrow to avoid noisy SEO and market-report filler."
        )
        validation_intro = (
            "A serious analytical report does not stop at a polished conclusion. "
            "It leaves an explicit validation backlog: what still needs to be checked, which objections matter most, and which conditions should flip the recommendation."
        )

    coverage_content = "\n\n".join(
        [
            coverage_intro,
            coverage_summary,
            "Exhibit 7",
            _markdown_table_block(
                ["Core question", "Focus", "Status", "Evidence records", "Source links"],
                coverage_rows,
                right_aligned_columns={3, 4},
            ),
            "Exhibit 8",
            _markdown_table_block(
                ["Source", "Type", "Reliability", "Linked questions"],
                source_rows,
                right_aligned_columns={2},
            ),
        ]
    )
    validation_content = "\n\n".join(
        [
            validation_intro,
            "Exhibit 9",
            _markdown_table_block(
                ["Validation issue", "Severity", "Why it still matters"],
                critique_rows,
            ),
            "Exhibit 10",
            _markdown_table_block(
                ["Trigger", "Condition", "Recommendation shift", "Confidence"],
                trigger_rows,
                right_aligned_columns={3},
            ),
        ]
    )
    return [
        (coverage_title, coverage_content),
        (validation_title, validation_content),
    ]


def _append_decision_addendum_sections(
    report: ReportOutput,
    brief: AnalysisBrief,
    source_ledger: list[SourceLedgerEntry],
    coverage: CoverageReport,
    critique_findings: list[CritiqueFinding],
    decision_triggers: list[DecisionTrigger],
    language: str,
) -> ReportOutput:
    order = max((section.order for section in report.sections), default=0)
    source_urls = [item.url for item in source_ledger[:4]]
    for title, lines in _decision_addendum_sections(brief, language):
        normalized_title = title.strip().lower()
        deterministic_content = "\n".join(f"- {line}" for line in lines if line)
        existing_section = next((section for section in report.sections if section.title.strip().lower() == normalized_title), None)
        if existing_section is not None:
            existing_section.title = title
            existing_section.content = deterministic_content
            existing_section.sources = source_urls
            continue
        order += 1
        report.sections.append(
            ReportSection(
                title=title,
                content=deterministic_content,
                order=order,
                sources=source_urls,
            )
        )
    for title, content in _traceability_appendix_sections(
        coverage=coverage,
        source_ledger=source_ledger,
        critique_findings=critique_findings,
        decision_triggers=decision_triggers,
        language=language,
    ):
        normalized_title = title.strip().lower()
        existing_section = next((section for section in report.sections if section.title.strip().lower() == normalized_title), None)
        if existing_section is not None:
            existing_section.title = title
            existing_section.content = content
            existing_section.sources = source_urls
            continue
        order += 1
        report.sections.append(
            ReportSection(
                title=title,
                content=content,
                order=order,
                sources=source_urls,
            )
        )
    return report


def build_report_markdown(
    task_spec: TaskSpec,
    brief: AnalysisBrief,
    plan: ResearchPlan,
    claims: list[ClaimRecord],
    coverage: CoverageReport,
    source_ledger: list[dict],
) -> str:
    top_claims = _select_report_worthy_claims(claims)[:6]
    decision_sections = _decision_addendum_sections(brief, task_spec.request_spec.language)
    lines = [
        f"# {brief.title}",
        "",
        "## Executive Summary",
        "",
        brief.executive_summary,
        "",
        "## Decision Context",
        "",
        f"- Goal: {task_spec.request_spec.goal}",
        f"- Subject: {task_spec.request_spec.subject}",
        f"- Decision context: {brief.decision_context}",
        f"- Geography: {task_spec.request_spec.geography}",
        f"- Time horizon: {task_spec.request_spec.time_horizon}",
        "",
        "## Evaluation Frame",
        "",
        *[f"- {item}" for item in task_spec.evaluation_dimensions],
        "",
        "## Key Findings",
        "",
        *[f"- {claim.statement} [Evidence: {', '.join(claim.supporting_evidence_ids[:2])}]" for claim in top_claims],
        "",
        "## Comparative Analysis",
        "",
        *[f"- {question.question}" for question in plan.primary_questions],
        *[f"- {question.question}" for question in plan.selected_adjacent_questions],
        "",
        "## Recommendation and Decision Posture",
        "",
    ]
    if brief.recommendation_posture == "evidence_backed_recommendations_allowed":
        lines.extend([
            f"- Prioritize options that score well on {', '.join(task_spec.evaluation_dimensions[:3])}. [Evidence: {', '.join(claims[0].supporting_evidence_ids[:2]) if claims else 'n/a'}]",
            f"- Treat vendor claims as directional until validated against the target workload. [Evidence: {', '.join(claims[1].supporting_evidence_ids[:2]) if len(claims) > 1 else 'n/a'}]",
        ])
    else:
        lines.append("- Bounded recommendation: evidence is informative but not strong enough for an unqualified winner call.")
    lines.extend([
        "",
        "## Gaps & Risks",
        "",
        *[f"- {item}" for item in brief.limitations],
        *[f"- {item}" for item in brief.key_risks],
        "",
    ])
    for title, items in decision_sections:
        lines.extend(
            [
                f"## {title}",
                "",
                *[f"- {item}" for item in items],
                "",
            ]
        )
    lines.extend([
        "## Evidence Coverage",
        "",
        f"- Covered questions: {coverage.covered_questions}/{coverage.total_questions}",
        f"- Contradiction count: {coverage.contradiction_count}",
        "",
        "## Sources",
        "",
        *[f"- [{source['title']}]({source['url']})" for source in source_ledger],
        "",
    ])
    return "\n".join(lines).strip() + "\n"


def _build_adjacent_research_row(
    question: ResearchQuestion,
    primary_question_id: str | None,
    findings: list[str],
    source_urls: list[str],
    sources: list[SourceLedgerEntry],
) -> dict:
    deduped_urls = list(dict.fromkeys(source_urls))
    return {
        "question_id": question.question_id,
        "primary_question_id": primary_question_id,
        "query": question.question,
        "confidence": 0.62 if findings else 0.28,
        "gaps": [] if findings else ["No usable evidence found for this branch"],
        "findings": findings[:8],
        "source_urls": deduped_urls,
        "sources": [source.model_dump(mode="json") for source in sources[:6]],
    }


def _has_usable_research_rows(rows: list[dict]) -> bool:
    return any(row.get("findings") for row in rows)


def _adjacent_to_primary_question_id(plan: ResearchPlan, question: ResearchQuestion) -> str | None:
    if not plan.primary_questions:
        return None
    question_tokens = _topic_tokens(question.question)
    lowered_question = question.question.lower()
    best_score = -1.0
    best_question_id = plan.primary_questions[0].question_id
    for index, primary_question in enumerate(plan.primary_questions):
        score = float(len(question_tokens & _topic_tokens(primary_question.question)))
        if question.kind == QuestionKind.ADJACENT_ALTERNATIVE and index == 1:
            score += 2.5
        if question.kind == QuestionKind.ADJACENT_COUNTERARGUMENT and index in {1, 3}:
            score += 1.75
        if question.kind == QuestionKind.ADJACENT_HIDDEN_VARIABLE and index == 3:
            score += 2.0
        if question.kind == QuestionKind.ADJACENT_BOUNDARY:
            if any(
                marker in lowered_question
                for marker in ("under what", "when ", "threshold", "switch", "stop", "cease", "boundary")
            ):
                if index == 3:
                    score += 3.0
            else:
                if index == 2:
                    score += 2.0
                if index == 3:
                    score += 1.0
        if question.kind == QuestionKind.ADJACENT_STAKEHOLDER and index == 3:
            score += 2.0
        if question.kind == QuestionKind.ADJACENT_TIME_SHIFT and index == 3:
            score += 2.0
        if score > best_score:
            best_score = score
            best_question_id = primary_question.question_id
    return best_question_id


async def _run_live_adjacent_research(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    adjacent_questions: list[ResearchQuestion],
    source_ledger: list[SourceLedgerEntry],
) -> tuple[list[dict], list[SourceLedgerEntry], float, list[dict]]:
    if not adjacent_questions:
        return [], source_ledger, 0.0, []

    depth = _budget_to_research_depth(task_spec)
    source_map = {entry.url: entry.model_copy(deep=True) for entry in source_ledger}
    research_rows: list[dict] = []
    total_cost = 0.0
    branch_meta: list[dict] = []
    profile = _depth_profile(task_spec)
    research_runs = await asyncio.gather(
        *[
            _research_single(_build_live_adjacent_query(task_spec, question), iteration=index, depth=depth)
            for index, question in enumerate(adjacent_questions[: profile.adjacent_research_branches], start=1)
        ]
    )
    for question, run_result in zip(adjacent_questions[: profile.adjacent_research_branches], research_runs):
        result, branch_cost, _evidence_items = run_result[:3]
        meta = run_result[3] if len(run_result) > 3 else {}
        total_cost += branch_cost
        branch_meta.append({**meta, "cost_usd": branch_cost})
        primary_question_id = _adjacent_to_primary_question_id(plan, question)
        source_urls: list[str] = []
        row_sources: list[SourceLedgerEntry] = []
        for source in result.sources:
            _upsert_live_source(
                source_map,
                url=source.url,
                title=source.title,
                domain=source.domain,
                question_id=question.question_id,
                preferred_domains=plan.preferred_domains,
                selection_reason=f"Returned by adjacent research branch for {question.question_id}",
            )
            if source.url in source_map:
                row_sources.append(source_map[source.url])
            source_urls.append(source.url)
        research_rows.append(_build_adjacent_research_row(question, primary_question_id, result.findings[:8], source_urls, row_sources))
    merged_sources = _rank_live_sources(list(source_map.values()), task_spec, plan, limit=profile.source_limit)
    allowlist = {entry.url for entry in merged_sources}
    for row in research_rows:
        row["source_urls"] = [url for url in row["source_urls"] if url in allowlist]
        row["sources"] = [source for source in row["sources"] if source.get("url") in allowlist]
    return research_rows, merged_sources, total_cost, branch_meta


async def _run_fallback_adjacent_research(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    adjacent_questions: list[ResearchQuestion],
    source_ledger: list[SourceLedgerEntry],
    snapshots: list[SourceSnapshot],
) -> tuple[list[dict], list[SourceLedgerEntry], list[SourceSnapshot]]:
    if not adjacent_questions:
        return [], source_ledger, snapshots

    provider = DuckDuckGoSearchProvider()
    source_map = {entry.url: entry.model_copy(deep=True) for entry in source_ledger}
    snapshot_by_url = {snapshot.url: snapshot for snapshot in snapshots}
    research_rows: list[dict] = []

    profile = _depth_profile(task_spec)
    for question in adjacent_questions[: profile.adjacent_research_branches]:
        single_plan = ResearchPlan(primary_questions=[question], preferred_domains=plan.preferred_domains)
        primary_question_id = _adjacent_to_primary_question_id(plan, question)
        candidates = await provider.search(_build_fallback_adjacent_query(task_spec, question), single_plan)
        selected_sources = select_sources(candidates, single_plan)
        row_sources: list[SourceLedgerEntry] = []
        row_snapshots: list[SourceSnapshot] = []
        for source in selected_sources[:4]:
            existing = source_map.get(source.url)
            if existing is None:
                source_map[source.url] = source
                existing = source
            if question.question_id not in existing.question_links:
                existing.question_links.append(question.question_id)
            row_sources.append(existing)
            if source.url not in snapshot_by_url:
                fetched = await provider.fetch(existing)
                snapshot_by_url[source.url] = fetched
            row_snapshots.append(snapshot_by_url[source.url])

        findings: list[str] = []
        source_urls: list[str] = []
        for snapshot in row_snapshots:
            if snapshot.fetch_status != "ok" or not snapshot.content:
                continue
            source_urls.append(snapshot.url)
            findings.extend(_sentence_candidates(snapshot, question))
        research_rows.append(_build_adjacent_research_row(question, primary_question_id, findings[:8], source_urls, row_sources))

    merged_sources = _rank_live_sources(list(source_map.values()), task_spec, plan, limit=profile.source_limit)
    merged_snapshots = list(snapshot_by_url.values())
    allowlist = {entry.url for entry in merged_sources}
    for row in research_rows:
        row["source_urls"] = [url for url in row["source_urls"] if url in allowlist]
        row["sources"] = [source for source in row["sources"] if source.get("url") in allowlist]
    merged_snapshots = [snapshot for snapshot in merged_snapshots if snapshot.url in allowlist]
    return research_rows, merged_sources, merged_snapshots


async def _run_stack_source_backfill(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    source_ledger: list[SourceLedgerEntry],
    snapshots: list[SourceSnapshot],
) -> tuple[list[dict], list[SourceLedgerEntry], list[SourceSnapshot]]:
    query_specs = _build_stack_backfill_queries(task_spec)
    if not query_specs:
        return [], source_ledger, snapshots
    profile = _depth_profile(task_spec)
    query_specs = query_specs[: profile.stack_backfill_limit]

    provider = DuckDuckGoSearchProvider()
    source_map = {entry.url: entry.model_copy(deep=True) for entry in source_ledger}
    snapshot_by_url = {snapshot.url: snapshot for snapshot in snapshots}
    question_by_id = {question.question_id: question for question in plan.primary_questions}
    research_rows: list[dict] = []

    for question_id, query in query_specs:
        question = question_by_id.get(question_id)
        if question is None:
            continue
        single_plan = ResearchPlan(
            primary_questions=[question],
            preferred_domains=plan.preferred_domains,
            required_source_mix=plan.required_source_mix,
        )
        candidates = await provider.search(query, single_plan)
        selected_sources = select_sources(candidates, single_plan)[:2]
        if not selected_sources:
            continue
        findings: list[str] = []
        source_urls: list[str] = []
        row_sources: list[SourceLedgerEntry] = []
        for source in selected_sources:
            existing = source_map.get(source.url)
            if existing is None:
                source_map[source.url] = source
                existing = source
            if question_id not in existing.question_links:
                existing.question_links.append(question_id)
            row_sources.append(existing)
            if source.url not in snapshot_by_url:
                snapshot_by_url[source.url] = await provider.fetch(existing)
            snapshot = snapshot_by_url[source.url]
            if snapshot.fetch_status != "ok" or not snapshot.content:
                continue
            source_urls.append(snapshot.url)
            findings.extend(_sentence_candidates(snapshot, question))
        if not findings:
            continue
        research_rows.append(
            {
                "question_id": question_id,
                "query": query,
                "confidence": 0.66,
                "gaps": [],
                "findings": findings[:6],
                "source_urls": list(dict.fromkeys(source_urls)),
                "sources": [source.model_dump(mode="json") for source in row_sources[:4]],
            }
        )

    merged_sources = _rank_live_sources(list(source_map.values()), task_spec, plan, limit=profile.source_limit)
    allowlist = {entry.url for entry in merged_sources}
    for row in research_rows:
        row["source_urls"] = [url for url in row["source_urls"] if url in allowlist]
        row["sources"] = [source for source in row["sources"] if source.get("url") in allowlist]
    merged_snapshots = [snapshot for snapshot in snapshot_by_url.values() if snapshot.url in allowlist]
    filtered_rows = [row for row in research_rows if row["source_urls"] and row["findings"]]
    return filtered_rows, merged_sources, merged_snapshots


async def _run_business_source_backfill(
    task_spec: TaskSpec,
    plan: ResearchPlan,
    source_ledger: list[SourceLedgerEntry],
    snapshots: list[SourceSnapshot],
    target_question_ids: set[str] | None = None,
) -> tuple[list[dict], list[SourceLedgerEntry], list[SourceSnapshot]]:
    query_specs = _build_business_backfill_queries(task_spec, target_question_ids)
    if not query_specs:
        return [], source_ledger, snapshots
    profile = _depth_profile(task_spec)
    query_specs = query_specs[: max(6, profile.stack_backfill_limit)]

    provider = DuckDuckGoSearchProvider()
    source_map = {entry.url: entry.model_copy(deep=True) for entry in source_ledger}
    snapshot_by_url = {snapshot.url: snapshot for snapshot in snapshots}
    question_by_id = {question.question_id: question for question in plan.primary_questions}
    research_rows: list[dict] = []

    for question_id, query in query_specs:
        question = question_by_id.get(question_id)
        if question is None:
            continue
        single_plan = ResearchPlan(
            primary_questions=[question],
            preferred_domains=plan.preferred_domains,
            required_source_mix=plan.required_source_mix,
        )
        candidates = await provider.search(query, single_plan)
        selected_sources = select_sources(candidates, single_plan)[:2]
        if not selected_sources:
            continue
        findings: list[str] = []
        source_urls: list[str] = []
        row_sources: list[SourceLedgerEntry] = []
        for source in selected_sources:
            existing = source_map.get(source.url)
            if existing is None:
                source_map[source.url] = source
                existing = source
            if question_id not in existing.question_links:
                existing.question_links.append(question_id)
            row_sources.append(existing)
            if source.url not in snapshot_by_url:
                snapshot_by_url[source.url] = await provider.fetch(existing)
            snapshot = snapshot_by_url[source.url]
            if snapshot.fetch_status != "ok" or not snapshot.content:
                continue
            source_urls.append(snapshot.url)
            findings.extend(_sentence_candidates(snapshot, question))
        if not findings:
            continue
        research_rows.append(
            {
                "question_id": question_id,
                "query": query,
                "confidence": 0.66,
                "gaps": [],
                "findings": findings[:6],
                "source_urls": list(dict.fromkeys(source_urls)),
                "sources": [source.model_dump(mode="json") for source in row_sources[:4]],
            }
        )

    merged_sources = _rank_live_sources(list(source_map.values()), task_spec, plan, limit=profile.source_limit)
    allowlist = {entry.url for entry in merged_sources}
    for row in research_rows:
        row["source_urls"] = [url for url in row["source_urls"] if url in allowlist]
        row["sources"] = [source for source in row["sources"] if source.get("url") in allowlist]
    merged_snapshots = [snapshot for snapshot in snapshot_by_url.values() if snapshot.url in allowlist]
    filtered_rows = [row for row in research_rows if row["source_urls"] and row["findings"]]
    return filtered_rows, merged_sources, merged_snapshots


def render_html(report_markdown: str, title: str) -> str:
    try:
        import markdown

        body = markdown.markdown(report_markdown, extensions=["tables", "fenced_code"])
    except Exception:
        escaped = report_markdown.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        body = f"<pre>{escaped}</pre>"
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <style>
    body {{ font-family: Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 40px 24px; color: #18212f; }}
    h1, h2 {{ color: #0f3d68; }}
    pre {{ white-space: pre-wrap; }}
    a {{ color: #0f5ea8; }}
  </style>
</head>
<body>{body}</body>
</html>"""


def _wrap_pdf_lines(text: str, width: int = 92) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines():
        if not raw_line.strip():
            lines.append("")
            continue
        current = ""
        for word in raw_line.split():
            candidate = word if not current else f"{current} {word}"
            if len(candidate) > width and current:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
    return lines


def _escape_pdf_text(text: str) -> str:
    safe = text.encode("latin-1", "replace").decode("latin-1")
    return safe.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _render_simple_pdf(text: str, title: str) -> bytes:
    body_lines = _wrap_pdf_lines(f"{title}\n\n{text}")
    lines_per_page = 46
    pages = [body_lines[index:index + lines_per_page] for index in range(0, len(body_lines), lines_per_page)] or [[]]

    objects: list[str] = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    page_refs: list[str] = []

    for page_lines in pages:
        commands = ["BT", "/F1 10 Tf", "72 770 Td", "14 TL"]
        for line in page_lines:
            if line:
                commands.append(f"({_escape_pdf_text(line)}) Tj")
            commands.append("T*")
        commands.append("ET")
        stream = "\n".join(commands)
        stream_bytes = stream.encode("latin-1", "replace")
        content_object_number = len(objects) + 1
        objects.append(f"<< /Length {len(stream_bytes)} >>\nstream\n{stream}\nendstream")
        page_object_number = len(objects) + 1
        objects.append(
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Resources << /Font << /F1 3 0 R >> >> "
            f"/Contents {content_object_number} 0 R >>"
        )
        page_refs.append(f"{page_object_number} 0 R")

    objects[1] = f"<< /Type /Pages /Count {len(page_refs)} /Kids [{' '.join(page_refs)}] >>"

    pdf = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n{obj}\nendobj\n".encode("latin-1", "replace"))

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
    pdf.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode("latin-1")
    )
    return bytes(pdf)


def _markdown_inline_to_pdf_text(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", text)
    return text


def _render_reportlab_pdf(report_markdown: str, title: str, subtitle: str = "", facts_line: str = "") -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except Exception:
        return _render_simple_pdf(report_markdown, title)

    font_name = "Helvetica"
    font_bold = "Helvetica-Bold"
    for candidate in (
        Path("C:/Windows/Fonts/ARIALUNI.TTF"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ):
        if candidate.exists():
            try:
                pdfmetrics.registerFont(TTFont("SmartReportBase", str(candidate)))
                font_name = "SmartReportBase"
                font_bold = "SmartReportBase"
                break
            except Exception:
                continue

    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "SmartTitle",
        parent=styles["Title"],
        fontName=font_bold,
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0F3D68"),
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "SmartSubtitle",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#3F5468"),
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "SmartHeading",
        parent=styles["Heading2"],
        fontName=font_bold,
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#0F3D68"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "SmartBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#1D2733"),
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "SmartBullet",
        parent=body_style,
        leftIndent=10,
        firstLineIndent=0,
    )
    facts_style = ParagraphStyle(
        "SmartFacts",
        parent=subtitle_style,
        fontName=font_bold,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#5A6E82"),
        spaceAfter=14,
    )

    story = [Paragraph(_markdown_inline_to_pdf_text(title), title_style)]
    if subtitle:
        story.append(Paragraph(_markdown_inline_to_pdf_text(subtitle), subtitle_style))
    if facts_line:
        story.append(Paragraph(_markdown_inline_to_pdf_text(facts_line), facts_style))
    story.append(Spacer(1, 6))

    for block in [part.strip() for part in report_markdown.split("\n\n") if part.strip()]:
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        first_line = lines[0].strip()
        if first_line.startswith("# "):
            continue
        if first_line.startswith("## "):
            story.append(Paragraph(_markdown_inline_to_pdf_text(first_line[3:].strip()), heading_style))
            if len(lines) > 1:
                story.append(Paragraph(_markdown_inline_to_pdf_text("<br/>".join(lines[1:])), body_style))
            continue
        if first_line.startswith("### "):
            story.append(Paragraph(_markdown_inline_to_pdf_text(first_line[4:].strip()), heading_style))
            if len(lines) > 1:
                story.append(Paragraph(_markdown_inline_to_pdf_text("<br/>".join(lines[1:])), body_style))
            continue
        if all(line.startswith("|") and line.endswith("|") for line in lines) and len(lines) >= 2:
            table_rows: list[list[str]] = []
            for line in lines:
                cells = [cell.strip() for cell in line.strip("|").split("|")]
                if set("".join(cells).replace("-", "").replace(":", "").strip()) == set():
                    continue
                table_rows.append([_markdown_inline_to_pdf_text(cell) for cell in cells])
            if table_rows:
                table = Table(table_rows, repeatRows=1, hAlign="LEFT")
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F3D68")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("FONTNAME", (0, 0), (-1, -1), font_name),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9EB4C8")),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FB")]),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEADING", (0, 0), (-1, -1), 11),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 8))
            continue
        if all(line.startswith("- ") or line.startswith("* ") for line in lines):
            for line in lines:
                story.append(
                    Paragraph(
                        _markdown_inline_to_pdf_text(line[2:].strip()),
                        bullet_style,
                        bulletText="-",
                    )
                )
            story.append(Spacer(1, 4))
            continue
        story.append(Paragraph(_markdown_inline_to_pdf_text("<br/>".join(lines)), body_style))

    document.build(story)
    return buffer.getvalue()


def render_pdf(html_content: str, report_markdown: str, title: str, subtitle: str = "", facts_line: str = "") -> bytes | None:
    try:
        from weasyprint import HTML

        buffer = io.BytesIO()
        HTML(string=html_content).write_pdf(buffer)
        return buffer.getvalue()
    except Exception:
        return _render_reportlab_pdf(report_markdown, title, subtitle=subtitle, facts_line=facts_line)


def render_docx(report_markdown: str, title: str) -> bytes | None:
    try:
        from docx import Document

        document = Document()
        document.add_heading(title, level=0)
        for block in report_markdown.split("\n\n"):
            text = block.strip()
            if not text:
                continue
            if text.startswith("#"):
                document.add_heading(text.lstrip("# ").strip(), level=1)
            else:
                document.add_paragraph(text)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception:
        return None


def _build_live_report_output(
    run_id: str,
    parsed_report: dict,
    source_ledger: list[SourceLedgerEntry],
    total_cost: float,
) -> ReportOutput:
    sections = _normalize_live_sections(parsed_report.get("sections", []), source_ledger)
    title = _sanitize_llm_markdown(parsed_report.get("title", "")).strip() or "Untitled report"
    executive_summary = _sanitize_llm_markdown(parsed_report.get("executive_summary", ""))
    if not executive_summary:
        executive_summary = "Evidence was collected, but the report writer did not return a usable executive summary."
    return ReportOutput(
        id=run_id,
        title=title,
        executive_summary=executive_summary,
        sections=sections,
        status=ReportStatus.COMPLETED,
        total_cost_usd=total_cost,
        metadata={
            "subtitle": _sanitize_llm_markdown(parsed_report.get("subtitle", "")).strip(),
            "facts_line": _sanitize_llm_markdown(parsed_report.get("facts_line", "")).strip(),
        },
    )


def _render_rich_docx_bytes(report: ReportOutput) -> bytes | None:
    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "report.docx"
        try:
            _build_rich_docx(report, [], str(output_path))
        except Exception:
            return None
        return output_path.read_bytes() if output_path.exists() else None


def _material_url(run_id: str, material_id: str) -> str:
    return f"material://{run_id}/{material_id}"


def _materialize_material_rows(
    repo: FileRunRepository,
    run_id: str,
    task_spec: TaskSpec,
    plan: ResearchPlan,
    materials: list[MaterialRecord],
) -> tuple[list[dict], list[SourceLedgerEntry], list[SourceSnapshot]]:
    if not materials:
        return [], [], []

    source_entries: list[SourceLedgerEntry] = []
    snapshots: list[SourceSnapshot] = []
    question_rows: list[dict] = []

    for material in materials:
        text = load_material_text(repo, run_id, material)
        if not text.strip():
            continue
        material_url = _material_url(run_id, material.material_id)
        source_entry = SourceLedgerEntry(
            url=material_url,
            title=material.title,
            domain="user-material",
            source_type=SourceType.USER_MATERIAL,
            publisher="user upload",
            reliability_score=0.92,
            selection_reason="Run-scoped material supplied by the user",
            question_links=[],
        )
        snapshot = SourceSnapshot(
            source_id=source_entry.source_id,
            url=material_url,
            title=material.title,
            content=text,
            excerpt=text[:400],
            provider="user-material",
            fetch_status="ok",
        )
        source_entries.append(source_entry)
        snapshots.append(snapshot)

    for question in plan.primary_questions:
        findings: list[str] = []
        row_sources: list[SourceLedgerEntry] = []
        source_urls: list[str] = []
        for source_entry, snapshot in zip(source_entries, snapshots):
            candidates = _sentence_candidates(snapshot, question)
            if not candidates:
                continue
            if question.question_id not in source_entry.question_links:
                source_entry.question_links.append(question.question_id)
            row_sources.append(source_entry)
            source_urls.append(source_entry.url)
            findings.extend(candidates[:3])
        if not findings:
            continue
        question_rows.append(
            {
                "question_id": question.question_id,
                "query": f"User materials for {question.question}",
                "confidence": 0.74,
                "gaps": [],
                "findings": findings[:8],
                "source_urls": list(dict.fromkeys(source_urls)),
                "sources": [item.model_dump(mode="json") for item in row_sources],
            }
        )
    return question_rows, source_entries, snapshots


async def _generate_presentation_package(
    repo: FileRunRepository,
    run_id: str,
    report: ReportOutput,
    language: str,
) -> tuple[dict, list[SpendEntry], str | None]:
    input_payload = {
        "executive_summary": report.executive_summary,
        "key_data": [
            {"section": section.title, "source_count": len(section.sources)}
            for section in report.sections
        ],
        "report_schema": {
            "title": report.title,
            "sections": [section.title for section in report.sections],
        },
        "chart_refs": [],
    }
    serialized = json.dumps(input_payload, ensure_ascii=False)
    model = get_model(AgentTask.PRESENTATION)
    raw = await _generate_slide_json(serialized, model, lang=language)
    slides_json = _parse_slides_json(raw, report.title)
    slides_path = repo.write_report_file(run_id, "slides.json", json.dumps(slides_json, ensure_ascii=False, indent=2))
    markdown = _build_presentation_markdown(slides_json)
    repo.write_report_file(run_id, "presentation.md", markdown)

    spend_entries = [
        _make_spend_entry(
            category=SpendCategory.PRESENTATION,
            stage="presentation_outline",
            provider="openrouter",
            model=model,
            input_tokens=_estimate_tokens(serialized),
            output_tokens=_estimate_tokens(raw),
            cost_usd=estimate_cost(AgentTask.PRESENTATION, _estimate_tokens(serialized), _estimate_tokens(raw)),
            pricing_basis="estimated_chars",
            notes="Presentation structure and slide plan",
        )
    ]

    pptx_path: str | None = None
    if settings.gamma_api_key:
        target = repo.report_dir(run_id) / "report.pptx"
        await _gamma_create(markdown, str(target))
        pptx_path = str(target)
        spend_entries.append(
            _make_spend_entry(
                category=SpendCategory.PRESENTATION,
                stage="presentation_export",
                provider="gamma",
                model="gamma-export",
                cost_usd=0.0,
                pricing_basis="external_service_unpriced",
                notes="Gamma export pricing is not modeled locally",
            )
        )

    return slides_json, spend_entries, pptx_path or str(slides_path)


async def _execute_live_report_run(
    repo: FileRunRepository,
    summary: RunSummary,
    task_spec: TaskSpec,
    emit: EmitFn,
) -> RunSummary:
    profile = _depth_profile(task_spec)
    plan = build_research_plan(task_spec)
    summary.depth_profile = profile
    if task_spec.allow_perplexity_handoff and not summary.handoff_prompts:
        summary.handoff_prompts = build_perplexity_handoff_prompts(task_spec, plan)
    spend_entries = list(summary.spend_breakdown)

    def record_spend(entry: SpendEntry) -> None:
        _record_spend(spend_entries, entry)

    repo.save_artifact(summary.run_id, "depth_profile.json", profile.model_dump(mode="json"))
    if summary.handoff_prompts:
        repo.save_artifact(
            summary.run_id,
            "handoff_prompts.json",
            [item.model_dump(mode="json") for item in summary.handoff_prompts],
        )
    if summary.materials:
        repo.save_artifact(
            summary.run_id,
            "materials.json",
            [item.model_dump(mode="json") for item in summary.materials],
        )
    await emit(RunEvent(step="planning", status="started", message="Building long-form research plan"))
    repo.save_artifact(summary.run_id, "request_spec.json", task_spec.request_spec.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "task_spec.json", task_spec.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "research_plan.json", plan.model_dump(mode="json"))
    await emit(RunEvent(step="planning", status="done", message="Long-form plan ready"))

    await emit(RunEvent(step="research", status="started", message="Running deep research branches"))
    research_queries = _build_live_research_queries(task_spec, plan)[: profile.initial_research_branches]
    depth = _budget_to_research_depth(task_spec)
    research_rows: list[dict] = []
    source_ledger: list[SourceLedgerEntry] = []
    snapshots: list[SourceSnapshot] = []
    evidence: list[EvidenceRecord] | None = None
    used_fallback = False

    try:
        research_runs = await asyncio.gather(
            *[
                _research_single(query, iteration=index, depth=depth)
                for index, (_question_id, query) in enumerate(research_queries, start=1)
            ]
        )

        source_map: dict[str, SourceLedgerEntry] = {}
        initial_branch_meta: list[dict] = []
        for (question_id, query), run_result in zip(research_queries, research_runs):
            result, branch_cost, _evidence_items = run_result[:3]
            meta = run_result[3] if len(run_result) > 3 else {}
            initial_branch_meta.append({**meta, "cost_usd": branch_cost})
            source_urls: list[str] = []
            for source in result.sources:
                _upsert_live_source(
                    source_map,
                    url=source.url,
                    title=source.title,
                    domain=source.domain,
                    question_id=question_id,
                    preferred_domains=plan.preferred_domains,
                    selection_reason=f"Returned by live research branch for {question_id}",
                )
                source_urls.append(source.url)
            research_rows.append(
                {
                    "question_id": question_id,
                    "query": query,
                    "confidence": result.confidence,
                    "gaps": result.gaps,
                    "findings": result.findings[:8],
                    "source_urls": list(dict.fromkeys(source_urls)),
                    "sources": [source.model_dump(mode="json") for source in result.sources[:8]],
                }
            )

        initial_spend_entry = _aggregate_research_spend(
            initial_branch_meta,
            category=SpendCategory.RESEARCH,
            stage="initial_research",
            fallback_provider="perplexity",
            fallback_model=depth,
            branch_count=len(research_queries),
            notes="Primary research pass",
        )
        if initial_spend_entry is not None:
            record_spend(initial_spend_entry)

        source_ledger = _rank_live_sources(list(source_map.values()), task_spec, plan, limit=profile.source_limit)
        if not source_ledger:
            raise RuntimeError("Live research returned no usable sources")
        source_url_allowlist = {entry.url for entry in source_ledger}
        for row in research_rows:
            row["source_urls"] = [url for url in row["source_urls"] if url in source_url_allowlist]
            row["sources"] = [source for source in row["sources"] if source.get("url") in source_url_allowlist]
    except Exception as exc:
        used_fallback = True
        await emit(
            RunEvent(
                step="research",
                status="warning",
                message=f"Paid live research unavailable, falling back to web search: {exc}",
            )
        )
        research_rows, source_ledger, evidence, snapshots = await _fallback_live_research(task_spec, plan)

    try:
        backfill_rows, source_ledger, snapshots = await _run_stack_source_backfill(
            task_spec,
            plan,
            source_ledger,
            snapshots,
        )
        if backfill_rows:
            research_rows.extend(backfill_rows)
            await emit(
                RunEvent(
                    step="research",
                    status="done",
                    message=f"Official-doc backfill added {len(backfill_rows)} targeted branches",
                )
            )
    except Exception as exc:
        await emit(
            RunEvent(
                step="research",
                status="warning",
                message=f"Official-doc backfill was skipped: {exc}",
            )
        )

    if summary.materials:
        material_rows, material_sources, material_snapshots = _materialize_material_rows(
            repo,
            summary.run_id,
            task_spec,
            plan,
            summary.materials,
        )
        if material_sources:
            source_map = {entry.url: entry for entry in source_ledger}
            for source in material_sources:
                existing = source_map.get(source.url)
                if existing is None:
                    source_map[source.url] = source
                    existing = source
                else:
                    for question_id in source.question_links:
                        if question_id not in existing.question_links:
                            existing.question_links.append(question_id)
            source_ledger = _rank_live_sources(list(source_map.values()), task_spec, plan, limit=profile.source_limit)
            allowlist = {entry.url for entry in source_ledger}
            snapshots.extend([snapshot for snapshot in material_snapshots if snapshot.url in allowlist])
            research_rows.extend(
                [
                    {
                        **row,
                        "source_urls": [url for url in row["source_urls"] if url in allowlist],
                        "sources": [source for source in row["sources"] if source.get("url") in allowlist],
                    }
                    for row in material_rows
                    if any(url in allowlist for url in row["source_urls"])
                ]
            )
            await emit(
                RunEvent(
                    step="research",
                    status="done",
                    message=f"Added {len(material_sources)} user materials into the evidence pack",
                )
            )

    repo.save_artifact(summary.run_id, "live_research_results.json", research_rows)
    repo.save_artifact(summary.run_id, "source_ledger.json", [item.model_dump(mode="json") for item in source_ledger])
    if snapshots:
        repo.save_artifact(summary.run_id, "source_snapshots.json", [item.model_dump(mode="json") for item in snapshots])
    await emit(
        RunEvent(
            step="research",
            status="done",
            message=f"Completed {len(research_rows)} research branches with {len(source_ledger)} selected sources",
        )
    )

    await emit(RunEvent(step="evidence", status="started", message="Assembling claim and evidence ledger"))
    if evidence is None:
        evidence = _build_live_evidence(research_rows, source_ledger, task_spec)
    repo.save_artifact(summary.run_id, "evidence_ledger.json", [item.model_dump(mode="json") for item in evidence])
    claims = build_claim_table(evidence)
    repo.save_artifact(summary.run_id, "claim_table.json", [item.model_dump(mode="json") for item in claims])
    contradiction_notes = sorted({note for claim in claims for note in claim.contradiction_notes})
    coverage = build_coverage_report(plan, claims, len(source_ledger), contradiction_notes)
    repo.save_artifact(summary.run_id, "coverage_report.json", coverage.model_dump(mode="json"))
    await emit(
        RunEvent(
            step="evidence",
            status="done",
            message=f"Built {len(claims)} claims from {len(evidence)} evidence records",
        )
    )

    business_gap_question_ids = _coverage_gap_question_ids(coverage) & {"q1", "q2", "q3", "q4"}
    if _is_business_topic_task(task_spec) and business_gap_question_ids:
        await emit(
            RunEvent(
                step="research",
                status="started",
                message=(
                    "Running targeted business backfill only for uncovered questions: "
                    + ", ".join(sorted(business_gap_question_ids))
                ),
            )
        )
        try:
            business_backfill_rows, source_ledger, snapshots = await _run_business_source_backfill(
                task_spec,
                plan,
                source_ledger,
                snapshots,
                target_question_ids=business_gap_question_ids,
            )
            if business_backfill_rows:
                research_rows.extend(business_backfill_rows)
                repo.save_artifact(summary.run_id, "live_research_results.json", research_rows)
                repo.save_artifact(summary.run_id, "source_ledger.json", [item.model_dump(mode="json") for item in source_ledger])
                if snapshots:
                    repo.save_artifact(summary.run_id, "source_snapshots.json", [item.model_dump(mode="json") for item in snapshots])
                evidence = _build_live_evidence(research_rows, source_ledger, task_spec)
                repo.save_artifact(summary.run_id, "evidence_ledger.json", [item.model_dump(mode="json") for item in evidence])
                claims = build_claim_table(evidence)
                repo.save_artifact(summary.run_id, "claim_table.json", [item.model_dump(mode="json") for item in claims])
                contradiction_notes = sorted({note for claim in claims for note in claim.contradiction_notes})
                coverage = build_coverage_report(plan, claims, len(source_ledger), contradiction_notes)
                repo.save_artifact(summary.run_id, "coverage_report.json", coverage.model_dump(mode="json"))
                await emit(
                    RunEvent(
                        step="research",
                        status="done",
                        message=f"Business backfill added {len(business_backfill_rows)} targeted branches for {', '.join(sorted(business_gap_question_ids))}",
                    )
                )
            else:
                await emit(
                    RunEvent(
                        step="research",
                        status="done",
                        message="Business backfill found no stronger targeted branches for the uncovered questions",
                    )
                )
        except Exception as exc:
            await emit(
                RunEvent(
                    step="research",
                    status="warning",
                    message=f"Business backfill was skipped: {exc}",
                )
            )

    await emit(RunEvent(step="critique", status="started", message="Generating side questions and critique findings"))
    review_artifact: dict = {"source": "heuristic"}
    try:
        (
            plan.adjacent_question_candidates,
            critique_findings,
            decision_triggers,
            review_artifact,
        ) = await _generate_model_driven_review(
            task_spec,
            plan,
            research_rows,
            source_ledger,
            claims,
            coverage,
            record_spend=record_spend,
        )
    except Exception as exc:
        review_artifact = {"source": "heuristic", "error": str(exc)}
        plan.adjacent_question_candidates = build_adjacent_question_candidates(task_spec, coverage, claims)
        critique_findings = build_critique_findings(task_spec, plan, claims, coverage, [])
        decision_triggers = build_decision_triggers(task_spec)
    plan.selected_adjacent_questions = select_adjacent_questions(task_spec, plan.adjacent_question_candidates)
    if review_artifact.get("source") != "model":
        critique_findings = build_critique_findings(task_spec, plan, claims, coverage, plan.selected_adjacent_questions)
    repo.save_artifact(summary.run_id, "research_plan.json", plan.model_dump(mode="json"))
    repo.save_artifact(
        summary.run_id,
        "adjacent_questions.json",
        [item.model_dump(mode="json") for item in plan.selected_adjacent_questions],
    )
    repo.save_artifact(
        summary.run_id,
        "critique_findings.json",
        [item.model_dump(mode="json") for item in critique_findings],
    )
    repo.save_artifact(
        summary.run_id,
        "decision_triggers.json",
        [item.model_dump(mode="json") for item in decision_triggers],
    )
    repo.save_artifact(summary.run_id, "lateral_review.json", review_artifact)
    await emit(
        RunEvent(
            step="critique",
            status="done",
            message=f"Prepared {len(plan.selected_adjacent_questions)} bounded side-questions and {len(critique_findings)} critique findings via {review_artifact.get('source', 'heuristic')}",
        )
    )

    selected_adjacent_questions = plan.selected_adjacent_questions[: profile.adjacent_research_branches]
    if selected_adjacent_questions:
        await emit(
            RunEvent(
                step="research",
                status="started",
                message=f"Running targeted side-question pass across {len(selected_adjacent_questions)} branches",
            )
        )
        try:
            if used_fallback:
                adjacent_rows, source_ledger, snapshots = await _run_fallback_adjacent_research(
                    task_spec,
                    plan,
                    selected_adjacent_questions,
                    source_ledger,
                    snapshots,
                )
            else:
                adjacent_rows, source_ledger, adjacent_cost, adjacent_meta = await _run_live_adjacent_research(
                    task_spec,
                    plan,
                    selected_adjacent_questions,
                    source_ledger,
                )
                adjacent_spend_entry = _aggregate_research_spend(
                    adjacent_meta,
                    category=SpendCategory.RESEARCH,
                    stage="adjacent_research",
                    fallback_provider="perplexity",
                    fallback_model=depth,
                    branch_count=len(selected_adjacent_questions),
                    notes="Adjacent-question pass",
                )
                if adjacent_spend_entry is not None:
                    record_spend(adjacent_spend_entry)
                if not _has_usable_research_rows(adjacent_rows):
                    await emit(
                        RunEvent(
                            step="research",
                            status="warning",
                            message="Live side-question pass returned no usable findings; retrying with fallback retrieval",
                        )
                    )
                    adjacent_rows, source_ledger, snapshots = await _run_fallback_adjacent_research(
                        task_spec,
                        plan,
                        selected_adjacent_questions,
                        source_ledger,
                        snapshots,
                    )
            if adjacent_rows:
                research_rows.extend(adjacent_rows)
                evidence = _build_live_evidence(research_rows, source_ledger, task_spec)
                claims = build_claim_table(evidence)
                contradiction_notes = sorted({note for claim in claims for note in claim.contradiction_notes})
                coverage = build_coverage_report(plan, claims, len(source_ledger), contradiction_notes)
                critique_findings = build_critique_findings(task_spec, plan, claims, coverage, selected_adjacent_questions)
                decision_triggers = build_decision_triggers(task_spec)
                repo.save_artifact(summary.run_id, "live_research_results.json", research_rows)
                repo.save_artifact(summary.run_id, "source_ledger.json", [item.model_dump(mode="json") for item in source_ledger])
                if snapshots:
                    repo.save_artifact(summary.run_id, "source_snapshots.json", [item.model_dump(mode="json") for item in snapshots])
                repo.save_artifact(summary.run_id, "evidence_ledger.json", [item.model_dump(mode="json") for item in evidence])
                repo.save_artifact(summary.run_id, "claim_table.json", [item.model_dump(mode="json") for item in claims])
                repo.save_artifact(summary.run_id, "coverage_report.json", coverage.model_dump(mode="json"))
                repo.save_artifact(
                    summary.run_id,
                    "critique_findings.json",
                    [item.model_dump(mode="json") for item in critique_findings],
                )
                repo.save_artifact(
                    summary.run_id,
                    "decision_triggers.json",
                    [item.model_dump(mode="json") for item in decision_triggers],
                )
                repo.save_artifact(summary.run_id, "lateral_review.json", review_artifact)
            await emit(
                RunEvent(
                    step="research",
                    status="done",
                    message=f"Targeted side-question pass completed with {len(adjacent_rows)} additional branches",
                )
            )
        except Exception as exc:
            await emit(
                RunEvent(
                    step="research",
                    status="warning",
                    message=f"Live side-question pass failed ({exc}); retrying with fallback retrieval",
                )
            )
            try:
                adjacent_rows, source_ledger, snapshots = await _run_fallback_adjacent_research(
                    task_spec,
                    plan,
                    selected_adjacent_questions,
                    source_ledger,
                    snapshots,
                )
                if adjacent_rows:
                    research_rows.extend(adjacent_rows)
                    evidence = _build_live_evidence(research_rows, source_ledger, task_spec)
                    claims = build_claim_table(evidence)
                    contradiction_notes = sorted({note for claim in claims for note in claim.contradiction_notes})
                    coverage = build_coverage_report(plan, claims, len(source_ledger), contradiction_notes)
                    critique_findings = build_critique_findings(task_spec, plan, claims, coverage, selected_adjacent_questions)
                    decision_triggers = build_decision_triggers(task_spec)
                    repo.save_artifact(summary.run_id, "live_research_results.json", research_rows)
                    repo.save_artifact(summary.run_id, "source_ledger.json", [item.model_dump(mode="json") for item in source_ledger])
                    if snapshots:
                        repo.save_artifact(summary.run_id, "source_snapshots.json", [item.model_dump(mode="json") for item in snapshots])
                    repo.save_artifact(summary.run_id, "evidence_ledger.json", [item.model_dump(mode="json") for item in evidence])
                    repo.save_artifact(summary.run_id, "claim_table.json", [item.model_dump(mode="json") for item in claims])
                    repo.save_artifact(summary.run_id, "coverage_report.json", coverage.model_dump(mode="json"))
                    repo.save_artifact(
                        summary.run_id,
                        "critique_findings.json",
                        [item.model_dump(mode="json") for item in critique_findings],
                    )
                    repo.save_artifact(
                        summary.run_id,
                        "decision_triggers.json",
                        [item.model_dump(mode="json") for item in decision_triggers],
                    )
                await emit(
                    RunEvent(
                        step="research",
                        status="done",
                        message=f"Fallback side-question pass completed with {len(adjacent_rows)} additional branches",
                    )
                )
            except Exception as fallback_exc:
                await emit(
                    RunEvent(
                        step="research",
                        status="warning",
                        message=f"Targeted side-question pass failed and fallback retrieval also failed: {fallback_exc}",
                    )
                )

    validation_questions = _build_validation_questions(task_spec, critique_findings, decision_triggers, coverage)
    if validation_questions:
        await emit(
            RunEvent(
                step="research",
                status="started",
                message=f"Running validation pass across {len(validation_questions)} contradiction and trigger checks",
            )
        )
        try:
            if used_fallback:
                validation_rows, source_ledger, snapshots = await _run_fallback_adjacent_research(
                    task_spec,
                    plan,
                    validation_questions,
                    source_ledger,
                    snapshots,
                )
            else:
                validation_rows, source_ledger, validation_cost, validation_meta = await _run_live_adjacent_research(
                    task_spec,
                    plan,
                    validation_questions,
                    source_ledger,
                )
                validation_spend_entry = _aggregate_research_spend(
                    validation_meta,
                    category=SpendCategory.RESEARCH,
                    stage="validation_research",
                    fallback_provider="perplexity",
                    fallback_model=depth,
                    branch_count=len(validation_questions),
                    notes="Validation pass",
                )
                if validation_spend_entry is not None:
                    record_spend(validation_spend_entry)
                if not _has_usable_research_rows(validation_rows):
                    await emit(
                        RunEvent(
                            step="research",
                            status="warning",
                            message="Live validation pass returned no usable findings; retrying with fallback retrieval",
                        )
                    )
                    validation_rows, source_ledger, snapshots = await _run_fallback_adjacent_research(
                        task_spec,
                        plan,
                        validation_questions,
                        source_ledger,
                        snapshots,
                    )
            if validation_rows:
                research_rows.extend(validation_rows)
                evidence = _build_live_evidence(research_rows, source_ledger, task_spec)
                claims = build_claim_table(evidence)
                contradiction_notes = sorted({note for claim in claims for note in claim.contradiction_notes})
                coverage = build_coverage_report(plan, claims, len(source_ledger), contradiction_notes)
                critique_findings = build_critique_findings(task_spec, plan, claims, coverage, selected_adjacent_questions)
                decision_triggers = build_decision_triggers(task_spec)
                repo.save_artifact(summary.run_id, "live_research_results.json", research_rows)
                repo.save_artifact(summary.run_id, "source_ledger.json", [item.model_dump(mode="json") for item in source_ledger])
                if snapshots:
                    repo.save_artifact(summary.run_id, "source_snapshots.json", [item.model_dump(mode="json") for item in snapshots])
                repo.save_artifact(summary.run_id, "evidence_ledger.json", [item.model_dump(mode="json") for item in evidence])
                repo.save_artifact(summary.run_id, "claim_table.json", [item.model_dump(mode="json") for item in claims])
                repo.save_artifact(summary.run_id, "coverage_report.json", coverage.model_dump(mode="json"))
            await emit(
                RunEvent(
                    step="research",
                    status="done",
                    message=f"Validation pass completed with {len(validation_rows)} additional branches",
                )
            )
        except Exception as exc:
            await emit(
                RunEvent(
                    step="research",
                    status="warning",
                    message=f"Live validation pass failed ({exc}); retrying with fallback retrieval",
                )
            )
            try:
                validation_rows, source_ledger, snapshots = await _run_fallback_adjacent_research(
                    task_spec,
                    plan,
                    validation_questions,
                    source_ledger,
                    snapshots,
                )
                if validation_rows:
                    research_rows.extend(validation_rows)
                    evidence = _build_live_evidence(research_rows, source_ledger, task_spec)
                    claims = build_claim_table(evidence)
                    contradiction_notes = sorted({note for claim in claims for note in claim.contradiction_notes})
                    coverage = build_coverage_report(plan, claims, len(source_ledger), contradiction_notes)
                    critique_findings = build_critique_findings(task_spec, plan, claims, coverage, selected_adjacent_questions)
                    decision_triggers = build_decision_triggers(task_spec)
                    repo.save_artifact(summary.run_id, "live_research_results.json", research_rows)
                    repo.save_artifact(summary.run_id, "source_ledger.json", [item.model_dump(mode="json") for item in source_ledger])
                    if snapshots:
                        repo.save_artifact(summary.run_id, "source_snapshots.json", [item.model_dump(mode="json") for item in snapshots])
                    repo.save_artifact(summary.run_id, "evidence_ledger.json", [item.model_dump(mode="json") for item in evidence])
                    repo.save_artifact(summary.run_id, "claim_table.json", [item.model_dump(mode="json") for item in claims])
                    repo.save_artifact(summary.run_id, "coverage_report.json", coverage.model_dump(mode="json"))
                    repo.save_artifact(
                        summary.run_id,
                        "critique_findings.json",
                        [item.model_dump(mode="json") for item in critique_findings],
                    )
                    repo.save_artifact(
                        summary.run_id,
                        "decision_triggers.json",
                        [item.model_dump(mode="json") for item in decision_triggers],
                    )
                await emit(
                    RunEvent(
                        step="research",
                        status="done",
                        message=f"Fallback validation pass completed with {len(validation_rows)} additional branches",
                    )
                )
            except Exception as fallback_exc:
                await emit(
                    RunEvent(
                        step="research",
                        status="warning",
                        message=f"Validation pass failed and fallback retrieval also failed: {fallback_exc}",
                    )
                )

    brief = build_analysis_brief(
        task_spec,
        claims,
        coverage,
        plan.selected_adjacent_questions,
        critique_findings,
        decision_triggers,
    )
    await emit(RunEvent(step="report", status="started", message="Synthesizing long-form report"))
    try:
        parsed_report = await _synthesize_longform_report(
            task_spec,
            plan,
            research_rows,
            source_ledger,
            claims,
            coverage,
            plan.selected_adjacent_questions,
            critique_findings,
            decision_triggers,
            record_spend=record_spend,
        )
    except Exception as exc:
        await emit(
            RunEvent(
                step="report",
                status="warning",
                message=f"LLM synthesis unavailable, using heuristic long-form writer: {exc}",
            )
        )
        parsed_report = _heuristic_longform_report(task_spec, source_ledger, claims, coverage)
    (
        parsed_report,
        report,
        subtitle,
        facts_line,
        markdown_text,
        quality_assessment,
        quality_iterations,
    ) = await _run_live_quality_revision_loop(
        run_id=summary.run_id,
        task_spec=task_spec,
        plan=plan,
        research_rows=research_rows,
        source_ledger=source_ledger,
        claims=claims,
        evidence=evidence,
        coverage=coverage,
        adjacent_questions=plan.selected_adjacent_questions,
        critique_findings=critique_findings,
        decision_triggers=decision_triggers,
        brief=brief,
        total_cost=_spend_totals(spend_entries)[0],
        initial_parsed_report=parsed_report,
        emit=emit,
        record_spend=record_spend,
    )
    (
        parsed_report,
        report,
        subtitle,
        facts_line,
        markdown_text,
        quality_assessment,
        quality_iterations,
    ) = await _run_live_compliance_revision(
        run_id=summary.run_id,
        task_spec=task_spec,
        plan=plan,
        parsed_report=parsed_report,
        source_ledger=source_ledger,
        claims=claims,
        evidence=evidence,
        coverage=coverage,
        adjacent_questions=plan.selected_adjacent_questions,
        critique_findings=critique_findings,
        decision_triggers=decision_triggers,
        brief=brief,
        total_cost=_spend_totals(spend_entries)[0],
        assessment=quality_assessment,
        iterations=quality_iterations,
        emit=emit,
        record_spend=record_spend,
    )
    remaining_unsupported = find_unsupported_precise_numbers(
        _report_body_for_grounding_scan(markdown_text),
        [claim.statement for claim in claims],
    )
    if remaining_unsupported:
        sanitized_parsed_report = _sanitize_report_grounding(parsed_report, [claim.statement for claim in claims])
        if sanitized_parsed_report != parsed_report:
            previous_quality_score = quality_assessment.overall_score
            parsed_report = sanitized_parsed_report
            report, subtitle, facts_line, markdown_text = _materialize_live_report_candidate(
                summary.run_id,
                parsed_report,
                source_ledger,
                coverage,
                critique_findings,
                decision_triggers,
                _spend_totals(spend_entries)[0],
                brief,
                task_spec.request_spec.language,
            )
            quality_assessment = _assess_live_report_candidate(
                task_spec,
                report,
                source_ledger,
                claims,
                evidence,
                coverage,
                plan.selected_adjacent_questions,
                critique_findings,
                decision_triggers,
            )
            quality_iterations.append(
                build_quality_iteration(
                    len(quality_iterations),
                    quality_assessment,
                    previous_score=previous_quality_score,
                    improved=False,
                    revision_focus=["grounding discipline"],
                    consecutive_improvements=0,
                    notes=[
                        f"Deterministic sanitization removed unsupported precise numbers: {', '.join(remaining_unsupported[:4])}",
                    ],
                )
            )
            await emit(
                RunEvent(
                    step="quality",
                    status="done",
                    message="Applied deterministic grounding cleanup to strip unsupported precise numbers from the final draft",
                )
            )
    report.metadata["quality_overall_score"] = quality_assessment.overall_score
    report.metadata["quality_verdict"] = quality_assessment.verdict
    report.metadata["quality_improvement_rounds"] = len([item for item in quality_iterations if item.improved])
    if not facts_line:
        facts_line = (
            f"{len(source_ledger)} sources | {len(claims)} claims | "
            f"{coverage.covered_questions}/{coverage.total_questions} core questions covered | "
            f"quality {quality_assessment.overall_score:.1f}/100"
        )
        report.metadata["facts_line"] = facts_line
    markdown_text = _build_markdown_from_report(report, source_ledger, subtitle=subtitle, facts_line=facts_line)

    brief.title = report.title
    brief.executive_summary = report.executive_summary
    brief.improvement_priorities = quality_assessment.rewrite_priorities[:]
    if quality_assessment.weaknesses:
        brief.limitations = list(dict.fromkeys(brief.limitations + quality_assessment.weaknesses[:3]))
    summary.quality_assessment = quality_assessment
    repo.save_artifact(summary.run_id, "analysis_brief.json", brief.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "report_synthesis.json", parsed_report)
    repo.save_artifact(summary.run_id, "quality_assessment.json", quality_assessment.model_dump(mode="json"))
    repo.save_artifact(
        summary.run_id,
        "quality_iterations.json",
        [item.model_dump(mode="json") for item in quality_iterations],
    )
    for item in quality_iterations:
        repo.save_artifact(
            summary.run_id,
            f"quality_round_{item.iteration}.json",
            item.model_dump(mode="json"),
        )
    total_cost, total_tokens = _spend_totals(spend_entries)
    report.total_cost_usd = total_cost
    report.metadata["spend_breakdown"] = [item.model_dump(mode="json") for item in spend_entries]
    report.metadata["depth_profile"] = profile.model_dump(mode="json")
    report.metadata["material_count"] = len(summary.materials)
    markdown_text = _final_markdown_compliance_cleanup(markdown_text, [claim.statement for claim in claims])
    html_text = _build_rich_html(report, [], lang=task_spec.request_spec.language)
    repo.write_report_file(summary.run_id, "report.md", markdown_text)
    if ArtifactFormat.HTML in task_spec.output_package:
        repo.write_report_file(summary.run_id, "report.html", html_text)
    if ArtifactFormat.PDF in task_spec.output_package:
        pdf_bytes = render_pdf(html_text, markdown_text, report.title, subtitle=subtitle, facts_line=facts_line)
        if pdf_bytes:
            repo.write_report_file(summary.run_id, "report.pdf", pdf_bytes)
    if ArtifactFormat.DOCX in task_spec.output_package:
        docx_bytes = _render_rich_docx_bytes(report)
        if docx_bytes:
            repo.write_report_file(summary.run_id, "report.docx", docx_bytes)
    if ArtifactFormat.PPTX in task_spec.output_package:
        try:
            _slides_json, presentation_spend, _presentation_path = await _generate_presentation_package(
                repo,
                summary.run_id,
                report,
                task_spec.request_spec.language,
            )
            for entry in presentation_spend:
                record_spend(entry)
            total_cost, total_tokens = _spend_totals(spend_entries)
            report.total_cost_usd = total_cost
            report.metadata["spend_breakdown"] = [item.model_dump(mode="json") for item in spend_entries]
        except Exception as exc:
            await emit(
                RunEvent(
                    step="report",
                    status="warning",
                    message=f"Presentation export was skipped: {exc}",
                )
            )

    package_payloads = {
        "request_spec.json": task_spec.request_spec.model_dump(mode="json"),
        "task_spec.json": task_spec.model_dump(mode="json"),
        "depth_profile.json": profile.model_dump(mode="json"),
        "research_plan.json": plan.model_dump(mode="json"),
        "sources.json": [item.model_dump(mode="json") for item in source_ledger],
        "live_research_results.json": research_rows,
        "evidence_ledger.json": [item.model_dump(mode="json") for item in evidence],
        "claim_table.json": [item.model_dump(mode="json") for item in claims],
        "adjacent_questions.json": [item.model_dump(mode="json") for item in plan.selected_adjacent_questions],
        "critique_findings.json": [item.model_dump(mode="json") for item in critique_findings],
        "decision_triggers.json": [item.model_dump(mode="json") for item in decision_triggers],
        "lateral_review.json": review_artifact,
        "analysis_brief.json": brief.model_dump(mode="json"),
        "coverage_report.json": coverage.model_dump(mode="json"),
        "quality_assessment.json": quality_assessment.model_dump(mode="json"),
        "quality_iterations.json": [item.model_dump(mode="json") for item in quality_iterations],
        "spend_breakdown.json": [item.model_dump(mode="json") for item in spend_entries],
        "materials.json": [item.model_dump(mode="json") for item in summary.materials],
        "handoff_prompts.json": [item.model_dump(mode="json") for item in summary.handoff_prompts],
        "report_output.json": report.model_dump(mode="json"),
    }
    for filename, payload in package_payloads.items():
        repo.write_report_file(
            summary.run_id,
            filename,
            json.dumps(payload, ensure_ascii=False, indent=2),
        )

    audit = audit_report_package(repo.report_dir(summary.run_id))
    repo.write_report_file(
        summary.run_id,
        "audit_summary.json",
        json.dumps(audit.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    _write_audit_snapshot(summary.run_id, audit)
    repo.save_artifact(summary.run_id, "audit_summary.json", audit.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "report_output.json", report.model_dump(mode="json"))

    await emit(RunEvent(step="report", status="done", message="Long-form package compiled"))
    await emit(RunEvent(step="audit", status="started", message="Running release gate"))
    await emit(RunEvent(step="audit", status="done", message=audit.release_status))

    report.status = ReportStatus.COMPLETED
    report.metadata["release_status"] = audit.release_status
    repo.save_artifact(summary.run_id, "report_output.json", report.model_dump(mode="json"))
    summary.title = report.title
    summary.cost_usd = total_cost
    summary.tokens_used = total_tokens
    summary.spend_breakdown = spend_entries
    summary.analysis_brief = brief
    summary.coverage_report = coverage
    summary.audit_summary = audit
    summary.status = RunStatus.COMPLETED
    return repo.save_run(summary)


async def choose_provider(task_spec: TaskSpec) -> SearchProvider:
    if match_reference_pack(task_spec.request_spec.original_query):
        return SeededSearchProvider()
    return DuckDuckGoSearchProvider()


async def execute_report_run(repo: FileRunRepository, summary: RunSummary, task_spec: TaskSpec, emit: EmitFn) -> RunSummary:
    if not match_reference_pack(task_spec.request_spec.original_query):
        return await _execute_live_report_run(repo, summary, task_spec, emit)

    plan = build_research_plan(task_spec)
    await emit(RunEvent(step="planning", status="started", message="Building research plan"))
    repo.save_artifact(summary.run_id, "request_spec.json", task_spec.request_spec.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "task_spec.json", task_spec.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "research_plan.json", plan.model_dump(mode="json"))
    await emit(RunEvent(step="planning", status="done", message="Research plan ready"))

    provider = await choose_provider(task_spec)
    await emit(RunEvent(step="search", status="started", message=f"Searching with {provider.name}"))
    all_candidates = []
    for query in plan.suggested_search_queries[:4]:
        all_candidates.extend(await provider.search(query, plan))
    source_ledger = select_sources(all_candidates, plan)
    repo.save_artifact(summary.run_id, "source_ledger.json", [item.model_dump(mode="json") for item in source_ledger])
    await emit(RunEvent(step="search", status="done", message=f"Selected {len(source_ledger)} sources"))

    await emit(RunEvent(step="evidence", status="started", message="Fetching and extracting evidence"))
    snapshots = [await provider.fetch(source) for source in source_ledger]
    repo.save_artifact(summary.run_id, "source_snapshots.json", [item.model_dump(mode="json") for item in snapshots])
    usable_snapshots = [item for item in snapshots if item.fetch_status == "ok" and item.content]
    failed_snapshot_count = len(snapshots) - len(usable_snapshots)
    source_lookup = {source.source_id: source.reliability_score for source in source_ledger}
    source_question_links = {source.source_id: set(source.question_links) for source in source_ledger}
    evidence = build_evidence_ledger(plan, usable_snapshots, source_lookup, source_question_links)
    repo.save_artifact(summary.run_id, "evidence_ledger.json", [item.model_dump(mode="json") for item in evidence])
    claims = build_claim_table(evidence)
    repo.save_artifact(summary.run_id, "claim_table.json", [item.model_dump(mode="json") for item in claims])
    contradiction_notes = sorted({note for claim in claims for note in claim.contradiction_notes})
    coverage = build_coverage_report(plan, claims, len(usable_snapshots), contradiction_notes)
    repo.save_artifact(summary.run_id, "coverage_report.json", coverage.model_dump(mode="json"))
    await emit(
        RunEvent(
            step="evidence",
            status="done",
            message=f"Built {len(claims)} claims from {len(usable_snapshots)} fetched sources ({failed_snapshot_count} failed)",
        )
    )

    await emit(RunEvent(step="critique", status="started", message="Generating bounded side questions and critique findings"))
    review_artifact: dict = {"source": "heuristic"}
    try:
        (
            plan.adjacent_question_candidates,
            critique_findings,
            decision_triggers,
            review_artifact,
        ) = await _generate_model_driven_review(task_spec, plan, [], source_ledger, claims, coverage)
    except Exception as exc:
        review_artifact = {"source": "heuristic", "error": str(exc)}
        plan.adjacent_question_candidates = build_adjacent_question_candidates(task_spec, coverage, claims)
        critique_findings = build_critique_findings(task_spec, plan, claims, coverage, [])
        decision_triggers = build_decision_triggers(task_spec)
    plan.selected_adjacent_questions = select_adjacent_questions(task_spec, plan.adjacent_question_candidates)
    if review_artifact.get("source") != "model":
        critique_findings = build_critique_findings(task_spec, plan, claims, coverage, plan.selected_adjacent_questions)
    repo.save_artifact(summary.run_id, "research_plan.json", plan.model_dump(mode="json"))
    repo.save_artifact(
        summary.run_id,
        "adjacent_questions.json",
        [item.model_dump(mode="json") for item in plan.selected_adjacent_questions],
    )
    repo.save_artifact(
        summary.run_id,
        "critique_findings.json",
        [item.model_dump(mode="json") for item in critique_findings],
    )
    repo.save_artifact(
        summary.run_id,
        "decision_triggers.json",
        [item.model_dump(mode="json") for item in decision_triggers],
    )
    repo.save_artifact(summary.run_id, "lateral_review.json", review_artifact)
    await emit(
        RunEvent(
            step="critique",
            status="done",
            message=f"Prepared {len(plan.selected_adjacent_questions)} side questions and {len(critique_findings)} critique findings via {review_artifact.get('source', 'heuristic')}",
        )
    )

    await emit(RunEvent(step="report", status="started", message="Compiling report package"))
    brief = build_analysis_brief(
        task_spec,
        claims,
        coverage,
        plan.selected_adjacent_questions,
        critique_findings,
        decision_triggers,
    )
    repo.save_artifact(summary.run_id, "analysis_brief.json", brief.model_dump(mode="json"))
    markdown_text = build_report_markdown(
        task_spec,
        brief,
        plan,
        claims,
        coverage,
        [item.model_dump(mode="json") for item in source_ledger],
    )
    html_text = render_html(markdown_text, brief.title)
    report_dir = repo.report_dir(summary.run_id)
    repo.write_report_file(summary.run_id, "report.md", markdown_text)
    repo.write_report_file(summary.run_id, "report.html", html_text)
    pdf_bytes = render_pdf(html_text, markdown_text, brief.title)
    if pdf_bytes:
        repo.write_report_file(summary.run_id, "report.pdf", pdf_bytes)
    docx_bytes = render_docx(markdown_text, brief.title)
    if docx_bytes:
        repo.write_report_file(summary.run_id, "report.docx", docx_bytes)
    package_payloads = {
        "request_spec.json": task_spec.request_spec.model_dump(mode="json"),
        "task_spec.json": task_spec.model_dump(mode="json"),
        "research_plan.json": plan.model_dump(mode="json"),
        "sources.json": [item.model_dump(mode="json") for item in source_ledger],
        "source_snapshots.json": [item.model_dump(mode="json") for item in snapshots],
        "evidence_ledger.json": [item.model_dump(mode="json") for item in evidence],
        "claim_table.json": [item.model_dump(mode="json") for item in claims],
        "adjacent_questions.json": [item.model_dump(mode="json") for item in plan.selected_adjacent_questions],
        "critique_findings.json": [item.model_dump(mode="json") for item in critique_findings],
        "decision_triggers.json": [item.model_dump(mode="json") for item in decision_triggers],
        "lateral_review.json": review_artifact,
        "analysis_brief.json": brief.model_dump(mode="json"),
        "coverage_report.json": coverage.model_dump(mode="json"),
    }
    for filename, payload in package_payloads.items():
        repo.write_report_file(
            summary.run_id,
            filename,
            json.dumps(payload, ensure_ascii=False, indent=2),
        )
    audit = audit_report_package(report_dir)
    repo.write_report_file(
        summary.run_id,
        "audit_summary.json",
        json.dumps(audit.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    _write_audit_snapshot(summary.run_id, audit)
    repo.save_artifact(summary.run_id, "audit_summary.json", audit.model_dump(mode="json"))
    await emit(RunEvent(step="report", status="done", message="Report package compiled"))
    await emit(RunEvent(step="audit", status="started", message="Running release gate"))
    await emit(RunEvent(step="audit", status="done", message=audit.release_status))

    report = ReportOutput(
        id=summary.run_id,
        title=brief.title,
        executive_summary=brief.executive_summary,
        status=ReportStatus.COMPLETED,
        sections=[
            ReportSection(
                title="Decision Context",
                content=brief.decision_context,
                order=1,
                sources=[item.url for item in source_ledger[:2]],
            ),
            ReportSection(
                title="Key Findings",
                content="\n".join(f"- {item}" for item in brief.key_findings),
                order=2,
                sources=[item.url for item in source_ledger[:4]],
            ),
            ReportSection(
                title="Gaps & Risks",
                content="\n".join(f"- {item}" for item in brief.limitations + brief.key_risks),
                order=3,
                sources=[item.url for item in source_ledger[:2]],
            ),
        ],
        total_cost_usd=0.0,
        metadata={
            "analysis_brief": brief.model_dump(mode="json"),
            "coverage_report": coverage.model_dump(mode="json"),
            "audit_summary": audit.model_dump(mode="json"),
            "release_status": audit.release_status,
            "claim_count": len(claims),
            "source_count": len(source_ledger),
        },
    )
    report = _append_decision_addendum_sections(
        report,
        brief,
        source_ledger,
        coverage,
        critique_findings,
        decision_triggers,
        task_spec.request_spec.language,
    )
    quality_assessment = _assess_live_report_candidate(
        task_spec,
        report,
        source_ledger,
        claims,
        evidence,
        coverage,
        plan.selected_adjacent_questions,
        critique_findings,
        decision_triggers,
    )
    quality_iterations = [
        build_quality_iteration(
            0,
            quality_assessment,
            revision_focus=[],
            consecutive_improvements=0,
            notes=["Single-pass deterministic package assessment."],
        )
    ]
    report.metadata["quality_overall_score"] = quality_assessment.overall_score
    report.metadata["quality_verdict"] = quality_assessment.verdict
    facts_line = (
        f"{len(source_ledger)} sources | {len(claims)} claims | "
        f"{coverage.covered_questions}/{coverage.total_questions} core questions covered | "
        f"quality {quality_assessment.overall_score:.1f}/100"
    )
    report.metadata["facts_line"] = facts_line
    brief.title = report.title
    brief.executive_summary = report.executive_summary
    markdown_text = _build_markdown_from_report(report, source_ledger, facts_line=facts_line)
    markdown_text = _final_markdown_compliance_cleanup(markdown_text, [claim.statement for claim in claims])
    html_text = _build_rich_html(report, [], lang=task_spec.request_spec.language)
    repo.write_report_file(summary.run_id, "report.md", markdown_text)
    repo.write_report_file(summary.run_id, "report.html", html_text)
    pdf_bytes = render_pdf(html_text, markdown_text, report.title, facts_line=facts_line)
    if pdf_bytes:
        repo.write_report_file(summary.run_id, "report.pdf", pdf_bytes)
    docx_bytes = _render_rich_docx_bytes(report)
    if docx_bytes:
        repo.write_report_file(summary.run_id, "report.docx", docx_bytes)
    brief.improvement_priorities = quality_assessment.rewrite_priorities[:]
    if quality_assessment.weaknesses:
        brief.limitations = list(dict.fromkeys(brief.limitations + quality_assessment.weaknesses[:3]))
    summary.quality_assessment = quality_assessment
    repo.write_report_file(
        summary.run_id,
        "analysis_brief.json",
        json.dumps(brief.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    repo.write_report_file(
        summary.run_id,
        "quality_assessment.json",
        json.dumps(quality_assessment.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    repo.write_report_file(
        summary.run_id,
        "quality_iterations.json",
        json.dumps([item.model_dump(mode="json") for item in quality_iterations], ensure_ascii=False, indent=2),
    )
    audit = audit_report_package(report_dir)
    report.metadata["audit_summary"] = audit.model_dump(mode="json")
    repo.write_report_file(
        summary.run_id,
        "audit_summary.json",
        json.dumps(audit.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    _write_audit_snapshot(summary.run_id, audit)
    repo.save_artifact(summary.run_id, "audit_summary.json", audit.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "report_output.json", report.model_dump(mode="json"))
    repo.save_artifact(summary.run_id, "quality_assessment.json", quality_assessment.model_dump(mode="json"))
    repo.save_artifact(
        summary.run_id,
        "quality_iterations.json",
        [item.model_dump(mode="json") for item in quality_iterations],
    )

    summary.title = brief.title
    summary.analysis_brief = brief
    summary.coverage_report = coverage
    summary.audit_summary = audit
    report.metadata["release_status"] = audit.release_status
    summary.status = RunStatus.COMPLETED
    return repo.save_run(summary)


def build_draft_run(
    run_id: str,
    query: str,
    *,
    depth: str | None = None,
    output_formats: list[ArtifactFormat] | None = None,
    allow_perplexity_handoff: bool = False,
) -> RunSummary:
    request_spec = build_request_spec(query, depth=depth)
    return RunSummary(
        run_id=run_id,
        request=query,
        budget_tier=request_spec.budget_tier,
        request_spec=request_spec,
        requested_output_formats=output_formats or [
            ArtifactFormat.MARKDOWN,
            ArtifactFormat.HTML,
            ArtifactFormat.PDF,
            ArtifactFormat.DOCX,
            ArtifactFormat.JSON,
        ],
        allow_perplexity_handoff=allow_perplexity_handoff,
        depth_profile=build_depth_profile(request_spec.budget_tier),
    )
