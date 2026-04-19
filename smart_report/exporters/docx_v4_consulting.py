"""Professional DOCX renderer — consulting-grade output.

Consumes a FinalReport (v4 schema with Track A structured output fields) and
produces a multi-page docx resembling a consulting deliverable:

  1. Cover page      — large serif title, subtitle, meta, accent bar, page break
  2. Executive Summary (pp 2-3)
       • Q&A block from qa_section (direct answers to sub-questions)
       • Key Numbers grid (key_numbers_highlight), 24pt bold values
       • Ranking visual (ranking list with weight bars)
  3. TOC             — via Word field code (auto-numbering by Word/LibreOffice)
  4. Main chapters   — sections parsed from main_synthesis, tables, callouts, chart PNGs
  5. Sources section — grouped by tool
  6. Footer          — page number + "Smart Report · YYYY-MM-DD"

CLI usage:
    python -m smart_report.exporters.docx_v4_consulting <fixture.json> output.docx

Stack: python-docx 1.2.0 (no node.js).
"""

from __future__ import annotations

import io
import json
import re
import sys
from datetime import date
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

from ..models import (
    CalloutBlock,
    ChartSpec,
    FinalReport,
    KeyNumberHighlight,
    QAItem,
    RankingItem,
    Table,
)

# ---------------------------------------------------------------------------
# Design constants
# ---------------------------------------------------------------------------

ACCENT_COLOR = RGBColor(0xB8, 0x86, 0x2E)  # #B8862E — amber gold
ACCENT_HEX = "B8862E"
DARK_COLOR = RGBColor(0x0A, 0x0A, 0x0A)  # #0A0A0A — near-black
LIGHT_BG_HEX = "FEF9EE"  # callout background tint
HEADER_BG_HEX = "F5F0E8"  # table header shading
ALT_ROW_HEX = "FAFAF5"  # zebra alt row
WHITE_HEX = "FFFFFF"

FONT_DISPLAY = "Georgia"   # serif display — Spectral/Source Serif preferred but rarely embedded
FONT_BODY = "Calibri"      # sans body — Inter preferred, Calibri widely available
FONT_MONO = "Courier New"

PT_TITLE = 44         # cover title
PT_SUBTITLE = 18      # cover subtitle
PT_H1 = 24            # chapter heading
PT_H2 = 16            # section heading
PT_H3 = 13            # sub-heading
PT_BODY = 11          # body text
PT_SMALL = 9          # captions, footnotes
PT_FOOTER = 9
PT_KEY_VALUE = 26     # key number value
PT_KEY_LABEL = 9      # key number label

MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(2.0)
MARGIN_RIGHT = Cm(2.0)

KIND_ICONS = {
    "insight": "⚡",
    "warning": "⚠",
    "key_number": "📊",
    "note": "ℹ",
}


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------


def _make_element(tag: str, **attrs) -> OxmlElement:
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k) if ":" in k else k, v)
    return el


def _set_run_font(run: "Run", name: str, size_pt: float, bold: bool = False,
                  italic: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Also set east-asian and complex-script font so Cyrillic renders properly
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _set_para_spacing(para: "Paragraph", before_pt: float = 0, after_pt: float = 6,
                      line_spacing: float = 1.5) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    spacing.set(qn("w:line"), str(int(line_spacing * 240)))
    spacing.set(qn("w:lineRule"), "auto")


def _set_shading(element, fill_hex: str, color_hex: str = "auto") -> None:
    """Apply background shading to a cell or paragraph element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), color_hex)
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_padding(cell, top_pt: float = 3, bottom_pt: float = 3,
                      left_pt: float = 5, right_pt: float = 5) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top_pt), ("bottom", bottom_pt),
                      ("left", left_pt), ("right", right_pt)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(int(val * 20)))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _add_border_left(para: "Paragraph", color_hex: str = ACCENT_HEX,
                     size_eighths: int = 32) -> None:
    """Add a 4pt left border to a paragraph (callout style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighths))   # size in 1/8 pt → 32 = 4pt
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)


def _add_para_background(para: "Paragraph", fill_hex: str) -> None:
    """Add a shading/background to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is not None:
        pPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _insert_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type("page"))


def docx_break_type(kind: str):
    """Return the WD_BREAK enum for the given kind string."""
    from docx.enum.text import WD_BREAK
    mapping = {"page": WD_BREAK.PAGE, "column": WD_BREAK.COLUMN}
    return mapping.get(kind, WD_BREAK.PAGE)


def _set_outline_level(para: "Paragraph", level: int) -> None:
    """Set outline level so Word/LibreOffice builds TOC from it (0 = Heading 1)."""
    pPr = para._p.get_or_add_pPr()
    outlineLvl = pPr.find(qn("w:outlineLvl"))
    if outlineLvl is None:
        outlineLvl = OxmlElement("w:outlineLvl")
        pPr.append(outlineLvl)
    outlineLvl.set(qn("w:val"), str(level))


def _add_toc_field(doc: Document) -> None:
    """Insert a TOC field code that Word/LibreOffice renders on open/update."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_sep)
    r.append(fldChar_end)


def _add_page_numbers_footer(doc: Document, date_str: str) -> None:
    """Add footer with page number and date to every section."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear default paragraph
    for para in footer.paragraphs:
        para.clear()
        break

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run_label = para.add_run(f"Smart Report · {date_str}    ")
    _set_run_font(run_label, FONT_BODY, PT_FOOTER, color=RGBColor(0x88, 0x88, 0x88))

    # Page number field
    run_num = para.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    r = run_num._r
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_end)

    _set_run_font(run_num, FONT_BODY, PT_FOOTER, color=RGBColor(0x88, 0x88, 0x88))

    # Separator between page num and total
    run_sep = para.add_run(" / ")
    _set_run_font(run_sep, FONT_BODY, PT_FOOTER, color=RGBColor(0x88, 0x88, 0x88))

    run_total = para.add_run()
    fldChar_begin2 = OxmlElement("w:fldChar")
    fldChar_begin2.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = " NUMPAGES "
    fldChar_end2 = OxmlElement("w:fldChar")
    fldChar_end2.set(qn("w:fldCharType"), "end")
    r2 = run_total._r
    r2.append(fldChar_begin2)
    r2.append(instrText2)
    r2.append(fldChar_end2)
    _set_run_font(run_total, FONT_BODY, PT_FOOTER, color=RGBColor(0x88, 0x88, 0x88))


def _set_document_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------


def _render_cover(doc: Document, report: FinalReport, date_str: str) -> None:
    """Render the cover page with title, subtitle, meta block, accent bar."""
    # Accent bar at top — a coloured single-row table spanning the full width
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, ACCENT_HEX)
    cell.text = ""
    # Set row height to simulate a thick accent stripe (1 cm)
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(Cm(0.8).pt * 20)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    # Remove table borders so only shading shows
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    existing_borders = tblPr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)

    # Large spacer
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Title
    title_text = _derive_title(report)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run(title_text)
    _set_run_font(run, FONT_DISPLAY, PT_TITLE, bold=True, color=DARK_COLOR)
    _set_para_spacing(p_title, before_pt=12, after_pt=16)

    # Subtitle
    subtitle = _derive_subtitle(report)
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_sub.add_run(subtitle)
        _set_run_font(run, FONT_BODY, PT_SUBTITLE, color=RGBColor(0x44, 0x44, 0x44))
        _set_para_spacing(p_sub, before_pt=0, after_pt=24)

    # Spacer
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Meta block
    meta_lines = [
        f"Дата: {date_str}",
        f"Сессия: {report.session_id}",
        "Smart Report — аналитическая система",
    ]
    for line in meta_lines:
        p_meta = doc.add_paragraph()
        run = p_meta.add_run(line)
        _set_run_font(run, FONT_BODY, PT_SMALL, color=RGBColor(0x77, 0x77, 0x77))
        _set_para_spacing(p_meta, before_pt=0, after_pt=3)

    # Page break
    _insert_page_break(doc)


def _derive_title(report: FinalReport) -> str:
    q = report.question.strip()
    # Shorten for display: take first 80 chars, break at word
    if len(q) > 80:
        q = q[:77] + "..."
    return q


def _derive_subtitle(report: FinalReport) -> str:
    if report.executive_summary and report.executive_summary.main_answer:
        first_sentence = report.executive_summary.main_answer.split(".")[0].strip()
        if first_sentence and len(first_sentence) < 160:
            return first_sentence + "."
    return ""


# ---------------------------------------------------------------------------
# Executive Summary
# ---------------------------------------------------------------------------


def _render_exec_summary(doc: Document, report: FinalReport) -> None:
    """Render the executive summary section (pages 2-3)."""
    # Section heading
    p_h = doc.add_heading("Executive Summary", level=1)
    _style_heading(p_h, level=1)

    # Q&A block
    if report.qa_section:
        p_qa_h = doc.add_heading("Прямые ответы на ключевые вопросы", level=2)
        _style_heading(p_qa_h, level=2)
        for item in report.qa_section:
            _render_qa_item(doc, item)

    # Key Numbers highlight grid
    if report.key_numbers_highlight:
        p_kn_h = doc.add_heading("Ключевые цифры", level=2)
        _style_heading(p_kn_h, level=2)
        _render_key_numbers_grid(doc, report.key_numbers_highlight)

    # Ranking visual
    if report.ranking:
        p_r_h = doc.add_heading("Ранжирование", level=2)
        _style_heading(p_r_h, level=2)
        _render_ranking(doc, report.ranking)

    # Legacy key numbers from executive_summary (if no structured ones)
    if not report.key_numbers_highlight and report.executive_summary.key_numbers:
        p_kn_h = doc.add_heading("Ключевые цифры", level=2)
        _style_heading(p_kn_h, level=2)
        for kn in report.executive_summary.key_numbers:
            p = doc.add_paragraph()
            run_val = p.add_run(f"{kn.value}")
            _set_run_font(run_val, FONT_BODY, 16, bold=True, color=ACCENT_COLOR)
            p.add_run(f" · {kn.metric}")
            if kn.subject:
                p.add_run(f" — {kn.subject}")
            _set_para_spacing(p, before_pt=2, after_pt=2)

    # Main answer if no qa_section
    if not report.qa_section and report.executive_summary.main_answer:
        p_ma = doc.add_paragraph()
        p_ma.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_ma.add_run(report.executive_summary.main_answer)
        _set_run_font(run, FONT_BODY, PT_BODY)
        _set_para_spacing(p_ma, before_pt=6, after_pt=6, line_spacing=1.5)
        _add_border_left(p_ma)
        _add_para_background(p_ma, LIGHT_BG_HEX)

    # Top findings
    if report.executive_summary.top_findings:
        p_tf_h = doc.add_heading("Ключевые находки", level=2)
        _style_heading(p_tf_h, level=2)
        for tf in report.executive_summary.top_findings:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(tf))
            _set_run_font(run, FONT_BODY, PT_BODY)
            _set_para_spacing(p, before_pt=0, after_pt=4)


def _render_qa_item(doc: Document, item: QAItem) -> None:
    """Render a single Q&A pair as bold question + regular answer."""
    p = doc.add_paragraph()
    run_q = p.add_run(f"В: {item.question}")
    _set_run_font(run_q, FONT_BODY, PT_BODY, bold=True, color=DARK_COLOR)
    _set_para_spacing(p, before_pt=8, after_pt=2)

    p2 = doc.add_paragraph()
    run_a = p2.add_run(f"О: {item.answer}")
    _set_run_font(run_a, FONT_BODY, PT_BODY)
    _set_para_spacing(p2, before_pt=0, after_pt=4)

    if item.details_ref:
        p3 = doc.add_paragraph()
        run_ref = p3.add_run(f"Подробнее: {item.details_ref}")
        _set_run_font(run_ref, FONT_BODY, PT_SMALL, italic=True,
                      color=RGBColor(0x77, 0x77, 0x77))
        _set_para_spacing(p3, before_pt=0, after_pt=8)


def _render_key_numbers_grid(doc: Document,
                              key_numbers: list[KeyNumberHighlight]) -> None:
    """Render key numbers as a 3-column grid table with large value text."""
    if not key_numbers:
        return

    # Group in rows of 3
    cols = 3
    rows_data = []
    row_buf = []
    for kn in key_numbers:
        row_buf.append(kn)
        if len(row_buf) == cols:
            rows_data.append(row_buf)
            row_buf = []
    if row_buf:
        # Pad with None to complete the last row
        while len(row_buf) < cols:
            row_buf.append(None)
        rows_data.append(row_buf)

    tbl = doc.add_table(rows=len(rows_data), cols=cols)
    tbl.style = "Table Grid"

    # Remove all outer borders
    _remove_table_outer_borders(tbl)

    for ri, row_items in enumerate(rows_data):
        for ci, kn in enumerate(row_items):
            cell = tbl.cell(ri, ci)
            _set_cell_shading(cell, LIGHT_BG_HEX)
            _set_cell_padding(cell, top_pt=8, bottom_pt=8, left_pt=10, right_pt=10)
            if kn is None:
                cell.text = ""
                continue
            # Value (large bold)
            p_val = cell.paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_val = p_val.add_run(kn.value)
            _set_run_font(run_val, FONT_DISPLAY, PT_KEY_VALUE, bold=True,
                          color=ACCENT_COLOR)
            _set_para_spacing(p_val, before_pt=0, after_pt=4)

            # Label
            p_lbl = cell.add_paragraph()
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_lbl = p_lbl.add_run(kn.label)
            _set_run_font(run_lbl, FONT_BODY, PT_KEY_LABEL, color=DARK_COLOR)
            _set_para_spacing(p_lbl, before_pt=0, after_pt=2)

            # Source ref
            if kn.source_ref:
                p_src = cell.add_paragraph()
                p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_src = p_src.add_run(kn.source_ref)
                _set_run_font(run_src, FONT_BODY, PT_SMALL, italic=True,
                              color=RGBColor(0x99, 0x99, 0x99))
                _set_para_spacing(p_src, before_pt=0, after_pt=0)

    doc.add_paragraph()  # spacer after grid


def _render_ranking(doc: Document, ranking: list[RankingItem]) -> None:
    """Render ranking as a visual list with weight bars and evidence labels."""
    if not ranking:
        return

    max_weight = max((r.weight or 0 for r in ranking), default=0)

    for item in ranking:
        # Bar indicator
        weight = item.weight
        if max_weight > 0 and weight is not None:
            bar_filled = int((weight / max_weight) * 20)
            bar = "█" * bar_filled + "░" * (20 - bar_filled)
            weight_str = f"{weight}%"
        else:
            bar = "░" * 20
            weight_str = ""

        p = doc.add_paragraph()
        # Label bold
        run_label = p.add_run(f"{item.label}")
        _set_run_font(run_label, FONT_BODY, PT_BODY, bold=True, color=DARK_COLOR)

        if weight_str:
            run_w = p.add_run(f"  {weight_str}")
            _set_run_font(run_w, FONT_BODY, PT_SMALL, color=ACCENT_COLOR)

        run_bar = p.add_run(f"\n{bar}")
        _set_run_font(run_bar, FONT_MONO, PT_SMALL, color=ACCENT_COLOR)

        run_rat = p.add_run(f"\n{item.rationale}")
        _set_run_font(run_rat, FONT_BODY, PT_SMALL, color=RGBColor(0x55, 0x55, 0x55))

        evid_labels = {"high": "✓ высокая надёжность", "medium": "~ средняя надёжность",
                       "low": "? низкая надёжность"}
        evid_text = evid_labels.get(item.evidence_strength, "")
        if evid_text:
            run_evid = p.add_run(f"  [{evid_text}]")
            _set_run_font(run_evid, FONT_BODY, PT_SMALL,
                          color=RGBColor(0x77, 0x77, 0x77))

        _set_para_spacing(p, before_pt=4, after_pt=8)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# TOC
# ---------------------------------------------------------------------------


def _render_toc(doc: Document) -> None:
    """Insert TOC heading and field code."""
    p_h = doc.add_heading("Содержание", level=1)
    _style_heading(p_h, level=1)
    _set_outline_level(p_h, 9)  # exclude TOC heading from TOC itself

    p_toc = doc.add_paragraph()
    p_toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_toc.add_run()
    _set_run_font(run, FONT_BODY, PT_BODY)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")

    # Placeholder text that Word replaces on update
    run2 = p_toc.add_run("(обновить оглавление в Word: Ctrl+A → F9)")
    _set_run_font(run2, FONT_BODY, PT_SMALL, italic=True,
                  color=RGBColor(0xAA, 0xAA, 0xAA))

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_sep)
    r.append(fldChar_end)

    _insert_page_break(doc)


# ---------------------------------------------------------------------------
# Main chapters
# ---------------------------------------------------------------------------


def _render_main_chapters(doc: Document, report: FinalReport,
                           chart_dir: Path | None = None) -> None:
    """Parse main_synthesis markdown sections and render each as a chapter."""
    synthesis = report.main_synthesis.strip()
    if not synthesis:
        return

    # Split by ## headings (level 2 becomes H1 in docx, level 3 → H2)
    chapters = _split_markdown_sections(synthesis)

    # Map tables and callouts by index for inline placement heuristics
    table_idx = 0
    callout_idx = 0
    chart_idx = 0

    for chapter_title, chapter_body in chapters:
        # Chapter heading
        p_h = doc.add_heading(chapter_title, level=1)
        _style_heading(p_h, level=1)

        # Render body paragraphs
        _render_markdown_body(doc, chapter_body)

        # Embed tables inline after their chapter (matching by sequential index)
        if table_idx < len(report.tables):
            _render_table(doc, report.tables[table_idx])
            table_idx += 1

        # Embed callout inline
        if callout_idx < len(report.callouts):
            _render_callout(doc, report.callouts[callout_idx])
            callout_idx += 1

        # Embed chart PNG
        if chart_idx < len(report.charts):
            _render_chart_placeholder(doc, report.charts[chart_idx],
                                      chart_idx, chart_dir)
            chart_idx += 1

    # Render remaining tables, callouts, charts at end of synthesis section
    while table_idx < len(report.tables):
        _render_table(doc, report.tables[table_idx])
        table_idx += 1

    while callout_idx < len(report.callouts):
        _render_callout(doc, report.callouts[callout_idx])
        callout_idx += 1

    while chart_idx < len(report.charts):
        _render_chart_placeholder(doc, report.charts[chart_idx],
                                  chart_idx, chart_dir)
        chart_idx += 1


def _split_markdown_sections(text: str) -> list[tuple[str, str]]:
    """Split markdown into (title, body) pairs by ## headings."""
    # Match ## or ### headings
    pattern = re.compile(r"^#{1,3}\s+(.+)$", re.MULTILINE)
    positions = [(m.start(), m.group(1).strip()) for m in pattern.finditer(text)]

    if not positions:
        return [("Основной синтез", text)]

    sections = []
    for i, (start, title) in enumerate(positions):
        body_start = text.index("\n", start) + 1 if "\n" in text[start:] else len(text)
        body_end = positions[i + 1][0] if i + 1 < len(positions) else len(text)
        body = text[body_start:body_end].strip()
        sections.append((title, body))
    return sections


def _render_markdown_body(doc: Document, text: str) -> None:
    """Render markdown body text into docx paragraphs.

    Handles:
    - ### subheadings → H2
    - **bold** inline
    - Bullet list items starting with - or *
    - Plain paragraphs
    """
    paras = text.split("\n\n")
    for block in paras:
        block = block.strip()
        if not block:
            continue

        # Sub-heading (### level)
        if block.startswith("###"):
            heading_text = block.lstrip("#").strip()
            p_h = doc.add_heading(heading_text, level=2)
            _style_heading(p_h, level=2)
            continue

        # Sub-heading ## level inside body
        if block.startswith("##"):
            heading_text = block.lstrip("#").strip()
            p_h = doc.add_heading(heading_text, level=2)
            _style_heading(p_h, level=2)
            continue

        # Bullet list block
        lines = block.split("\n")
        if all(l.strip().startswith(("-", "*", "•")) for l in lines if l.strip()):
            for line in lines:
                line = line.strip().lstrip("-*• ").strip()
                if line:
                    p = doc.add_paragraph(style="List Bullet")
                    _render_inline_md(p, line)
                    _set_para_spacing(p, before_pt=0, after_pt=3)
            continue

        # Horizontal rule
        if block.strip() in ("---", "***", "___"):
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _render_inline_md(p, block)
        _set_para_spacing(p, before_pt=0, after_pt=6, line_spacing=1.5)


def _render_inline_md(para: "Paragraph", text: str) -> None:
    """Add runs to para, handling **bold** and *italic* inline markdown."""
    # Split on bold/italic markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            _set_run_font(run, FONT_BODY, PT_BODY, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            _set_run_font(run, FONT_BODY, PT_BODY, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            _set_run_font(run, FONT_MONO, PT_BODY)
        elif part:
            run = para.add_run(part)
            _set_run_font(run, FONT_BODY, PT_BODY)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


def _render_table(doc: Document, table: Table) -> None:
    """Render a structured Table as a styled docx table."""
    # Title
    p_title = doc.add_paragraph()
    run = p_title.add_run(table.title)
    _set_run_font(run, FONT_BODY, PT_SMALL + 1, bold=True, color=DARK_COLOR)
    _set_para_spacing(p_title, before_pt=10, after_pt=3)

    n_cols = len(table.columns)
    if n_cols == 0:
        return

    n_rows = 1 + len(table.rows)  # header + data
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"

    # Header row
    header_row = tbl.rows[0]
    for ci, col_name in enumerate(table.columns):
        cell = header_row.cells[ci]
        _set_cell_shading(cell, HEADER_BG_HEX)
        _set_cell_padding(cell, top_pt=4, bottom_pt=4, left_pt=6, right_pt=6)
        p = cell.paragraphs[0]
        run = p.add_run(str(col_name))
        _set_run_font(run, FONT_BODY, PT_SMALL + 1, bold=True, color=DARK_COLOR)

    # Data rows
    for ri, row_data in enumerate(table.rows):
        docx_row = tbl.rows[ri + 1]
        bg = ALT_ROW_HEX if ri % 2 == 1 else WHITE_HEX
        for ci, cell_val in enumerate(row_data):
            if ci >= n_cols:
                break
            cell = docx_row.cells[ci]
            _set_cell_shading(cell, bg)
            _set_cell_padding(cell, top_pt=3, bottom_pt=3, left_pt=6, right_pt=6)
            p = cell.paragraphs[0]
            text = str(cell_val) if cell_val is not None else ""
            # Right-align if looks numeric
            if _is_numeric_str(text):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text)
            _set_run_font(run, FONT_BODY, PT_SMALL + 1)

    # Add hairlines only (between rows, not outer borders)
    _set_table_hairlines_only(tbl)

    # Caption
    if table.caption:
        p_cap = doc.add_paragraph()
        run = p_cap.add_run(table.caption)
        _set_run_font(run, FONT_BODY, PT_SMALL, italic=True,
                      color=RGBColor(0x77, 0x77, 0x77))
        _set_para_spacing(p_cap, before_pt=2, after_pt=2)

    if table.source_ref:
        p_src = doc.add_paragraph()
        run = p_src.add_run(f"Источник: {table.source_ref}")
        _set_run_font(run, FONT_BODY, PT_SMALL, italic=True,
                      color=RGBColor(0x99, 0x99, 0x99))
        _set_para_spacing(p_src, before_pt=0, after_pt=8)

    doc.add_paragraph()  # spacer


def _is_numeric_str(s: str) -> bool:
    """Return True if string looks like a number (possibly with units)."""
    clean = re.sub(r"[%₽$€\s,.]", "", s.strip())
    return bool(re.match(r"^-?\d+$", clean)) and len(clean) > 0


def _set_table_hairlines_only(tbl) -> None:
    """Set only horizontal hairlines between rows; remove outer borders."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement("w:tblBorders")
    # Outer borders: none
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    # Vertical inside: none
    insideV = OxmlElement("w:insideV")
    insideV.set(qn("w:val"), "none")
    tblBorders.append(insideV)
    # Horizontal inside: hairline
    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "single")
    insideH.set(qn("w:sz"), "2")   # 0.25pt hairline
    insideH.set(qn("w:color"), "CCCCCC")
    tblBorders.append(insideH)
    tblPr.append(tblBorders)


def _remove_table_outer_borders(tbl) -> None:
    """Remove all borders from a table (used for key numbers grid)."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Callouts
# ---------------------------------------------------------------------------


def _render_callout(doc: Document, callout: CalloutBlock) -> None:
    """Render a callout block: border-left accent, light background, icon + title."""
    icon = KIND_ICONS.get(callout.kind, "•")

    # Title paragraph
    p_title = doc.add_paragraph()
    run_icon = p_title.add_run(f"{icon} ")
    _set_run_font(run_icon, FONT_BODY, PT_BODY, color=ACCENT_COLOR)
    run_title = p_title.add_run(callout.title)
    _set_run_font(run_title, FONT_BODY, PT_BODY, bold=True, color=DARK_COLOR)
    _add_border_left(p_title, color_hex=ACCENT_HEX, size_eighths=32)
    _add_para_background(p_title, LIGHT_BG_HEX)
    _set_para_spacing(p_title, before_pt=8, after_pt=0)

    # Body paragraph
    p_body = doc.add_paragraph()
    run_body = p_body.add_run(callout.body)
    _set_run_font(run_body, FONT_BODY, PT_BODY)
    _add_border_left(p_body, color_hex=ACCENT_HEX, size_eighths=32)
    _add_para_background(p_body, LIGHT_BG_HEX)
    _set_para_spacing(p_body, before_pt=0, after_pt=10, line_spacing=1.4)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Charts
# ---------------------------------------------------------------------------


def _render_chart_placeholder(doc: Document, chart: ChartSpec,
                               idx: int, chart_dir: Path | None) -> None:
    """Embed chart PNG if available, otherwise insert a placeholder paragraph."""
    png_path: Path | None = None

    if chart_dir is not None:
        # Try to find a PNG from Track C output
        candidates = [
            chart_dir / f"chart_{idx}.png",
            chart_dir / f"chart_{idx:02d}.png",
            chart_dir / f"{chart.title.replace(' ', '_')[:40]}.png",
        ]
        for c in candidates:
            if c.exists():
                png_path = c
                break

    if png_path is not None and png_path.exists():
        # Add chart title
        p_title = doc.add_paragraph()
        run = p_title.add_run(chart.title)
        _set_run_font(run, FONT_BODY, PT_SMALL + 1, bold=True)
        _set_para_spacing(p_title, before_pt=8, after_pt=3)
        # Embed image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run_img = p_img.add_run()
            run_img.add_picture(str(png_path), width=Cm(14))
        except Exception:
            run_img.add_run(f"[Ошибка вставки PNG: {png_path}]")
        if chart.caption:
            p_cap = doc.add_paragraph()
            run = p_cap.add_run(chart.caption)
            _set_run_font(run, FONT_BODY, PT_SMALL, italic=True,
                          color=RGBColor(0x77, 0x77, 0x77))
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Placeholder block
        p_ph = doc.add_paragraph()
        p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon = "📊"
        run = p_ph.add_run(f"{icon} [{chart.chart_type.upper()}] {chart.title}")
        _set_run_font(run, FONT_BODY, PT_SMALL, italic=True,
                      color=RGBColor(0xAA, 0xAA, 0xAA))
        _add_border_left(p_ph, color_hex=ACCENT_HEX, size_eighths=16)
        _add_para_background(p_ph, LIGHT_BG_HEX)
        if chart.caption:
            p_cap = doc.add_paragraph()
            run = p_cap.add_run(chart.caption)
            _set_run_font(run, FONT_BODY, PT_SMALL, italic=True,
                          color=RGBColor(0xAA, 0xAA, 0xAA))
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p_ph, before_pt=6, after_pt=10)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Consensus / Conflicts / Gaps sections
# ---------------------------------------------------------------------------


def _render_extra_sections(doc: Document, report: FinalReport) -> None:
    """Render consensus, conflicts, gaps_filled sections if non-empty."""
    sections = [
        ("Консенсус источников", report.consensus_section),
        ("Противоречия между источниками", report.conflicts_section),
        ("Закрытые и оставшиеся пробелы", report.gaps_filled_section),
    ]
    for title, content in sections:
        if content and content.strip():
            p_h = doc.add_heading(title, level=1)
            _style_heading(p_h, level=1)
            _render_markdown_body(doc, content.strip())


# ---------------------------------------------------------------------------
# Sources section
# ---------------------------------------------------------------------------


def _render_sources(doc: Document, report: FinalReport) -> None:
    """Render sources section grouped by tool."""
    if not report.all_sources:
        return

    p_h = doc.add_heading("Источники", level=1)
    _style_heading(p_h, level=1)

    # Group by tool
    tool_order = ["perplexity", "openai_dr", "claude", "other"]
    tool_labels = {
        "perplexity": "Perplexity",
        "openai_dr": "OpenAI Deep Research",
        "claude": "Claude",
        "other": "Прочие источники",
    }

    by_tool: dict[str, list] = {t: [] for t in tool_order}
    for src in report.all_sources:
        tool = src.tool if src.tool in by_tool else "other"
        by_tool[tool].append(src)

    for tool in tool_order:
        sources = by_tool.get(tool, [])
        if not sources:
            continue
        label = tool_labels.get(tool, tool)
        p_sub = doc.add_heading(label, level=2)
        _style_heading(p_sub, level=2)

        for src in sources:
            p = doc.add_paragraph(style="List Bullet")
            run_title = p.add_run(src.title or "(без названия)")
            _set_run_font(run_title, FONT_BODY, PT_BODY, bold=bool(src.url))
            if src.url:
                run_url = p.add_run(f"\n{src.url}")
                _set_run_font(run_url, FONT_BODY, PT_SMALL,
                              color=RGBColor(0x20, 0x60, 0xAA))
            if src.reliability:
                rel_label = {"high": "высокая", "medium": "средняя",
                             "low": "низкая"}.get(src.reliability, src.reliability)
                run_rel = p.add_run(f"  [{rel_label}]")
                _set_run_font(run_rel, FONT_BODY, PT_SMALL,
                              color=RGBColor(0x88, 0x88, 0x88))
            _set_para_spacing(p, before_pt=2, after_pt=4)


# ---------------------------------------------------------------------------
# Consistency appendix (non-invasive, only material/minor issues)
# ---------------------------------------------------------------------------


def _render_consistency_appendix(doc: Document, report: FinalReport) -> None:
    """Render an optional appendix with methodological notes from the Critic.

    Only renders if report.metadata["consistency_check"] exists and has
    material or minor issues. Critical issues are expected to be resolved
    in the retry pass and should NOT appear here.

    This appendix turns unresolved nuances into explicit honesty about edge
    cases, improving trust rather than hiding imperfections.
    """
    cc = report.metadata.get("consistency_check")
    if not isinstance(cc, dict):
        return

    issues_raw = cc.get("issues", [])
    if not isinstance(issues_raw, list):
        return

    # Filter to material/minor only (critical should have been resolved)
    non_critical = [
        i for i in issues_raw
        if isinstance(i, dict) and i.get("severity") in ("material", "minor")
    ]
    if not non_critical:
        return

    _insert_page_break(doc)

    p_h = doc.add_heading("Приложение: Методологические замечания", level=1)
    _style_heading(p_h, level=1)

    intro = doc.add_paragraph(
        "В ходе внутреннего аудита отчёта выявлены следующие nuance-моменты, "
        "которые не являются ошибками, но требуют внимания читателя "
        "при интерпретации данных:"
    )
    _set_run_font(intro.runs[0] if intro.runs else intro.add_run(""), FONT_BODY, PT_BODY)
    _set_para_spacing(intro, before_pt=6, after_pt=12)

    sev_labels = {"material": "Существенное", "minor": "Незначительное"}
    cat_labels = {
        "number_conflict": "Числовое расхождение",
        "ranking_qa_mismatch": "Расхождение ранжирования и Q&A",
        "verdict_evidence_gap": "Вердикт vs данные",
        "table_prose_disagreement": "Таблица vs текст",
        "source_attribution_inconsistency": "Атрибуция источников",
    }

    for i, issue in enumerate(non_critical, 1):
        sev = issue.get("severity", "minor")
        cat = issue.get("category", "")
        sev_label = sev_labels.get(sev, sev)
        cat_label = cat_labels.get(cat, cat)

        p_item_h = doc.add_heading(f"Замечание {i}: {cat_label} ({sev_label})", level=2)
        _style_heading(p_item_h, level=2)

        loc_a = issue.get("location_a", "")
        stmt_a = issue.get("statement_a", "")
        loc_b = issue.get("location_b", "")
        stmt_b = issue.get("statement_b", "")
        why = issue.get("why_inconsistent", "")
        fix = issue.get("suggested_fix", "")

        if loc_a and stmt_a:
            p = doc.add_paragraph()
            run_loc = p.add_run(f"{loc_a}: ")
            _set_run_font(run_loc, FONT_BODY, PT_BODY, bold=True)
            run_stmt = p.add_run(f'"{stmt_a}"')
            _set_run_font(run_stmt, FONT_BODY, PT_BODY, italic=True)
            _set_para_spacing(p, before_pt=4, after_pt=2)

        if loc_b and stmt_b:
            p = doc.add_paragraph()
            run_loc = p.add_run(f"{loc_b}: ")
            _set_run_font(run_loc, FONT_BODY, PT_BODY, bold=True)
            run_stmt = p.add_run(f'"{stmt_b}"')
            _set_run_font(run_stmt, FONT_BODY, PT_BODY, italic=True)
            _set_para_spacing(p, before_pt=2, after_pt=2)

        if why:
            p = doc.add_paragraph()
            run_lbl = p.add_run("Суть: ")
            _set_run_font(run_lbl, FONT_BODY, PT_BODY, bold=True)
            run_why = p.add_run(why)
            _set_run_font(run_why, FONT_BODY, PT_BODY)
            _set_para_spacing(p, before_pt=2, after_pt=2)

        if fix:
            p = doc.add_paragraph()
            run_lbl = p.add_run("Рекомендация: ")
            _set_run_font(run_lbl, FONT_BODY, PT_SMALL, bold=True,
                          color=ACCENT_COLOR)
            run_fix = p.add_run(fix)
            _set_run_font(run_fix, FONT_BODY, PT_SMALL)
            _set_para_spacing(p, before_pt=2, after_pt=8)


# ---------------------------------------------------------------------------
# Heading styles
# ---------------------------------------------------------------------------


def _style_heading(para: "Paragraph", level: int) -> None:
    """Apply custom styling on top of Word's built-in heading."""
    for run in para.runs:
        if level == 1:
            _set_run_font(run, FONT_DISPLAY, PT_H1, bold=True, color=DARK_COLOR)
        elif level == 2:
            _set_run_font(run, FONT_DISPLAY, PT_H2, bold=True, color=DARK_COLOR)
        else:
            _set_run_font(run, FONT_BODY, PT_H3, bold=True, color=DARK_COLOR)
    _set_para_spacing(para, before_pt=16, after_pt=6)
    # Ensure heading is in TOC outline level
    _set_outline_level(para, level - 1)


# ---------------------------------------------------------------------------
# Main render function
# ---------------------------------------------------------------------------


def render_consulting_docx(
    report: FinalReport,
    output_path: Path,
    chart_dir: Path | None = None,
) -> Path:
    """Render a FinalReport to a consulting-grade DOCX at output_path.

    Args:
        report:       The structured FinalReport (with qa_section, tables, etc.)
        output_path:  Where to write the .docx file.
        chart_dir:    Optional directory containing chart PNGs from Track C.
                      If None or charts not found, placeholders are used.

    Returns:
        The output_path (for chaining).
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    date_str = date.today().isoformat()

    doc = Document()
    _set_document_margins(doc)

    # Footer with page numbers
    _add_page_numbers_footer(doc, date_str)

    # 1. Cover page
    _render_cover(doc, report, date_str)

    # 2. Executive Summary (Q&A, Key Numbers, Ranking)
    _render_exec_summary(doc, report)
    _insert_page_break(doc)

    # 3. TOC (field code — Word/LibreOffice renders on open/update)
    _render_toc(doc)

    # 4. Main chapters from main_synthesis + tables/callouts/charts
    _render_main_chapters(doc, report, chart_dir=chart_dir)

    # 5. Additional sections
    _render_extra_sections(doc, report)

    # 6. Sources
    _render_sources(doc, report)

    # 7. Methodological notes appendix (non-invasive — only if consistency_check has
    #    material/minor issues; critical issues should have been resolved in retry)
    _render_consistency_appendix(doc, report)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def _load_fixture(path: Path) -> FinalReport:
    """Load a FinalReport from a JSON fixture file."""
    raw = json.loads(path.read_text(encoding="utf-8"))
    return FinalReport.model_validate(raw)


def main(argv: list[str] | None = None) -> None:
    import argparse

    parser = argparse.ArgumentParser(
        description="Render a FinalReport JSON to a consulting DOCX."
    )
    parser.add_argument("fixture", type=Path, help="Path to FinalReport JSON")
    parser.add_argument("output", type=Path, help="Output .docx path")
    parser.add_argument(
        "--chart-dir",
        type=Path,
        default=None,
        help="Directory with chart PNG files (Track C output)",
    )
    args = parser.parse_args(argv)

    report = _load_fixture(args.fixture)
    out = render_consulting_docx(report, args.output, chart_dir=args.chart_dir)
    print(f"Rendered: {out} ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
