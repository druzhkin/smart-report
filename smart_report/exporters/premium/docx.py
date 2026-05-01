"""Premium long-form DOCX renderer.

This renderer is intentionally separate from the legacy DOCX exporters. It
takes the renderer-neutral PremiumReportDocument and produces a consulting
report artifact: cover, evidence scorecard, long-form sections, visual tables,
and appendices. It is domain-neutral; all topic-specific content must already
be present in the prepared document model.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import PremiumPreparedBlock, PremiumPreparedSection, PremiumReportDocument

NAVY = "152238"
INK = "1F2933"
MUTED = "667085"
GOLD = "B08D57"
PAPER = "F7F4EE"
LINE = "D9DEE7"


def render_premium_docx(
    document: PremiumReportDocument,
    path: Path,
    *,
    include_internal_audit: bool = False,
) -> Path:
    """Render a premium long-form report DOCX."""

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = document.title
    doc.core_properties.subject = document.subtitle
    doc.core_properties.author = "Smart Report"
    _setup_document(doc)
    _render_cover(doc, document)
    _render_decision_dashboard(doc, document, include_internal_audit=include_internal_audit)
    _render_client_evidence_snapshot(doc, document)
    if include_internal_audit:
        _render_scorecard(doc, document)
        _render_readiness_gate(doc, document)
    _render_toc_placeholder(doc, document)
    for section in document.sections:
        _render_section(doc, section)
    if document.appendices:
        _render_appendix_divider(doc, document)
    for appendix in document.appendices:
        _render_section(doc, appendix, appendix=True)
    doc.save(path)
    return path


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)
    _setup_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 24, NAVY, 0, 10),
        ("Heading 1", 17, NAVY, 16, 7),
        ("Heading 2", 12, NAVY, 10, 5),
        ("Heading 3", 10.5, INK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _setup_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("SMART REPORT")
    _set_run(hr, size=7.5, color=GOLD, bold=True, all_caps=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("Премиальный аналитический отчёт | стр. ")
    _set_run(run, size=7.5, color=MUTED)
    _add_page_number(fp)


def _render_cover(doc: Document, report: PremiumReportDocument) -> None:
    _add_rule(doc, NAVY, 30)
    kicker = doc.add_paragraph("SMART REPORT | ПРЕМИАЛЬНЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ")
    kicker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run(kicker.runs[0], size=8, color=GOLD, bold=True, all_caps=True)

    title = doc.add_paragraph(report.title, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = doc.add_paragraph(report.subtitle)
    _set_run(subtitle.runs[0], size=11, color=MUTED)

    doc.add_paragraph()
    meta = doc.add_table(rows=1, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(meta, header=False)
    values = [
        ("Тип отчёта", _report_type_label(report.plan.report_type)),
        ("Аудитория", _audience_label(report.plan.audience)),
        ("Мин. объём", f"{report.plan.deliverables.report_min_pages} стр."),
        ("Доказательства", f"{report.source_count} источников / {report.numeric_fact_count} фактов"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.rows[0].cells[idx]
        _shade(cell, PAPER)
        p = cell.paragraphs[0]
        p.add_run(label + "\n").bold = True
        p.add_run(value)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    _style_table(callout, header=False)
    cell = callout.rows[0].cells[0]
    _shade(cell, NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    cover_note = (
        f"{_report_type_label(report.plan.report_type)} для аудитории: "
        f"{_audience_label(report.plan.audience)}. "
        f"Доказательная база: {report.source_count} источников / "
        f"{report.numeric_fact_count} фактов."
    )
    run = p.add_run(cover_note)
    _set_run(run, size=9.2, color="FFFFFF", bold=True)
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def _render_decision_dashboard(
    doc: Document,
    report: PremiumReportDocument,
    *,
    include_internal_audit: bool = False,
) -> None:
    doc.add_heading("Резюме для решения", level=1)
    intro = doc.add_paragraph(
        "Эта страница фиксирует управленческий ответ, уровень доказательной базы и практическое "
        "следствие для решения. Ограничения и спорные места вынесены явно, чтобы читатель видел "
        "не только вывод, но и границы его применимости."
    )
    _set_run(intro.runs[0], size=9.5, color=MUTED, italic=True)

    readiness = report.premium_readiness or {}
    ready = bool(readiness.get("ready"))
    score = readiness.get("score", "?")
    issues = readiness.get("issues") or []

    cards = [
        (
            "Короткий ответ",
            _first_section_block_body(report, "executive_summary", "Короткий ответ")
            or report.title,
            NAVY,
        ),
        (
            "Глубина доказательств",
            f"{report.source_count} источников; {report.numeric_fact_count} числовых фактов.",
            GOLD,
        ),
        (
            "Уровень доверия",
            (
                f"Внутренний score {score}/100; открытых вопросов: {len(issues)}."
                if include_internal_audit
                else _client_confidence_label(report)
            ),
            "2E7D32" if ready else GOLD,
        ),
        (
            "Практическое следствие",
            report.plan.decision_context,
            INK,
        ),
    ]

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(table, header=False)
    for idx, (label, value, accent) in enumerate(cards):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = table.rows[row_idx].cells[col_idx]
        _shade(cell, "FFFFFF")
        _cell_border(cell, accent)
        p = cell.paragraphs[0]
        title = p.add_run(label + "\n")
        _set_run(title, size=8.5, color=accent, bold=True, all_caps=True)
        body = p.add_run(_clip(value, 360))
        _set_run(body, size=9.2, color=INK)

    doc.add_paragraph()


def _render_client_evidence_snapshot(doc: Document, report: PremiumReportDocument) -> None:
    doc.add_heading("Карта доказательств", level=1)
    rows = [
        ["Покрытие источниками", f"{report.source_count} источников"],
        ["Числовая база", f"{report.numeric_fact_count} числовых фактов"],
        ["Связка факт-источник", "Ключевые утверждения сопровождаются ссылками и вынесены в приложения."],
        ["Ограничения", "Неполные или спорные данные явно отмечены в разделах рисков и ограничений."],
    ]
    _render_key_value_table(doc, rows)
    note = doc.add_paragraph(
        "Принцип чтения: выводы в основном тексте отделены от приложений. Детальные реестры "
        "источников, фактов и спорных утверждений сохранены ниже для проверки, но не заменяют "
        "управленческий синтез."
    )
    _set_run(note.runs[0], size=9.2, color=MUTED, italic=True)


def _render_scorecard(doc: Document, report: PremiumReportDocument) -> None:
    doc.add_heading("Карта доказательной базы", level=1)
    rows = [
        ["Порог качества", report.plan.quality_bar],
        ["Требуемые доказательства", f"{report.plan.evidence.min_sources}+ источников, {report.plan.evidence.min_numeric_facts}+ числовых фактов"],
        ["Собранные доказательства", f"{report.source_count} источников, {report.numeric_fact_count} числовых фактов"],
        ["Материалы", _deliverables(report)],
    ]
    _render_key_value_table(doc, rows)
    doc.add_paragraph(
        "Эта карта является контролем качества выдачи, а не декором. Если доказательная база тонкая, "
        "отчёт должен явно показать ограничение, а не прятать неопределённость."
    )


def _render_readiness_gate(doc: Document, report: PremiumReportDocument) -> None:
    readiness = report.premium_readiness
    if not readiness:
        return

    ready = bool(readiness.get("ready"))
    score = readiness.get("score", "?")
    issues = readiness.get("issues") or []
    strengths = readiness.get("strengths") or []

    doc.add_heading("Гейт готовности к платной выдаче", level=1)
    status = "ГОТОВ К ПЛАТНОЙ ВЫДАЧЕ КЛИЕНТУ" if ready else "НЕ ГОТОВ К ПЛАТНОЙ ВЫДАЧЕ КЛИЕНТУ"
    status_line = doc.add_paragraph(status)
    _set_run(status_line.runs[0], size=11, color=NAVY if ready else "B42318", bold=True)
    rows = [
        ["Статус", status],
        ["Оценка", f"{score}/100"],
        ["Открытые проблемы", str(len(issues))],
        ["Сильные стороны", str(len(strengths))],
    ]
    _render_key_value_table(doc, rows)

    if not ready:
        warning = doc.add_table(rows=1, cols=1)
        _style_table(warning, header=False)
        cell = warning.rows[0].cells[0]
        _shade(cell, "FDECEC")
        p = cell.paragraphs[0]
        p.add_run("Предупреждение по выдаче: ").bold = True
        p.add_run(
            "Этот документ пока является премиальным черновиком. Его нельзя продавать или показывать "
            "как финальный платный отчёт, пока блокеры ниже не закрыты."
        )

    visible_issues = [
        issue for issue in issues
        if isinstance(issue, dict)
    ][:12]
    if visible_issues:
        _render_table(
            doc,
            ["Критичность", "Код", "Проблема", "Что исправить"],
            [
                [
                    str(issue.get("severity", "")),
                    str(issue.get("code", "")),
                    str(issue.get("message", "")),
                    str(issue.get("recommendation", "")),
                ]
                for issue in visible_issues
            ],
        )
    if strengths:
        _render_notes(doc, [str(item) for item in strengths[:8]])


def _render_toc_placeholder(doc: Document, report: PremiumReportDocument) -> None:
    doc.add_heading("Как читать отчёт", level=1)
    lead = doc.add_paragraph(
        "Отчёт устроен как управленческий документ: сначала вывод и логика решения, "
        "затем доказательства, сценарии, риски и приложения. Сырые реестры не должны "
        "перебивать чтение - они вынесены в приложения и пакет данных."
    )
    _set_run(lead.runs[0], size=10.5, color=INK)
    rows = [
        ["1", "Вывод", "Короткий ответ, уровень доверия, практическое следствие."],
        ["2", "Доказательства", "Ключевые числа, согласия, противоречия и ограничения."],
        ["3", "Сценарии", "Базовый, позитивный и негативный исходы с триггерами."],
        ["4", "Действия", "Риски, условия входа/отказа и что мониторить дальше."],
        ["A-C", "Приложения", "Источники, факты и ограничения для проверки."],
    ]
    _render_table(doc, ["Блок", "Что внутри", "Как использовать"], rows, max_rows=8)


def _render_appendix_divider(doc: Document, report: PremiumReportDocument) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    _add_rule(doc, GOLD, 12)
    doc.add_heading("Приложения и пакет данных", level=1)
    lead = doc.add_paragraph(
        "Приложения нужны не для линейного чтения, а для проверки: откуда взялись "
        "ключевые числа, какие источники использованы и какие ограничения остаются. "
        "Основной вывод уже дан выше; ниже сохранён аудиторский след."
    )
    _set_run(lead.runs[0], size=10.5, color=INK)
    rows = [
        [f"A{idx}", appendix.title, appendix.purpose]
        for idx, appendix in enumerate(report.appendices, 1)
    ]
    _render_table(doc, ["Блок", "Приложение", "Что проверять"], rows, max_rows=8, max_text=130)
    note = doc.add_paragraph(
        "Полные источники и расширенные CSV-реестры остаются в пакете данных; "
        "в DOCX оставлены только строки, полезные для чтения и проверки позиции."
    )
    _set_run(note.runs[0], size=9, color=MUTED, italic=True)


def _render_section(doc: Document, section: PremiumPreparedSection, *, appendix: bool = False) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    _add_rule(doc, GOLD if appendix else NAVY, 10)
    doc.add_heading(section.title, level=1 if not appendix else 2)
    purpose = doc.add_paragraph(section.purpose)
    _set_run(purpose.runs[0], size=9.5, color=MUTED, italic=True)
    for block in section.blocks:
        _render_block(doc, block, appendix=appendix)


def _render_block(doc: Document, block: PremiumPreparedBlock, *, appendix: bool = False) -> None:
    doc.add_heading(block.title, level=2)
    if block.body:
        _render_markdown_like_body(doc, block.body)
    if block.rows:
        if block.kind == "kpi_grid":
            _render_metric_exhibit(doc, block)
        elif block.kind == "scenario_matrix":
            _render_scenario_exhibit(doc, block)
        elif block.kind == "risk_register":
            _render_risk_exhibit(doc, block, appendix=appendix)
        elif block.kind == "evidence_table" and len(block.columns) >= 6:
            _render_evidence_exhibit(doc, block, appendix=appendix)
        elif block.kind == "source_quality_table":
            _render_source_exhibit(doc, block, appendix=appendix)
        else:
            _render_table(
                doc,
                block.columns,
                block.rows,
                max_rows=12 if appendix else 7,
                max_text=160 if appendix else 120,
            )
    if block.notes:
        _render_notes(doc, block.notes)


def _render_notes(doc: Document, notes: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    _style_table(table, header=False)
    cell = table.rows[0].cells[0]
    _shade(cell, "FFF8E7")
    p = cell.paragraphs[0]
    p.add_run("Заметка аналитика: ").bold = True
    p.add_run(" ".join(notes))


def _render_metric_exhibit(doc: Document, block: PremiumPreparedBlock) -> None:
    rows = block.rows[:8]
    cards = rows[:4]
    if cards:
        table = doc.add_table(rows=1, cols=len(cards))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _style_table(table, header=False)
        for idx, row in enumerate(cards):
            cell = table.rows[0].cells[idx]
            _shade(cell, PAPER)
            metric = _short(row[0] if len(row) > 0 else "", 34)
            value = _short(row[1] if len(row) > 1 else "", 28)
            subject = _short(row[2] if len(row) > 2 else "", 44)
            p = cell.paragraphs[0]
            label = p.add_run(metric + "\n")
            _set_run(label, size=7.5, color=MUTED, bold=True, all_caps=True)
            val = p.add_run(value + "\n")
            _set_run(val, size=15, color=NAVY, bold=True)
            sub = p.add_run(subject)
            _set_run(sub, size=8, color=INK)
    remaining = rows[4:8]
    if remaining:
        compact = [
            [
                _short(row[0] if len(row) > 0 else "", 42),
                _short(row[1] if len(row) > 1 else "", 32),
                _short(row[2] if len(row) > 2 else "", 58),
            ]
            for row in remaining
        ]
        _render_table(doc, ["Метрика", "Значение", "Контекст"], compact, max_rows=4)


def _render_evidence_exhibit(
    doc: Document,
    block: PremiumPreparedBlock,
    *,
    appendix: bool = False,
) -> None:
    rows = block.rows[: 12 if appendix else 6]
    compact = []
    for row in rows:
        if len(row) >= 7:
            compact.append(
                [
                    _short(row[2], 42),
                    _short(row[1], 34),
                    _short(row[3], 62),
                    _source_label(row[6]),
                ]
            )
        else:
            compact.append([_short(value, 70) for value in row[:4]])
    _render_table(
        doc,
        ["Метрика", "Значение", "Объект", "Источник"],
        compact,
        max_rows=12 if appendix else 6,
        max_text=120,
    )
    _data_pack_note(doc, len(block.rows), len(rows))


def _render_source_exhibit(
    doc: Document,
    block: PremiumPreparedBlock,
    *,
    appendix: bool = False,
) -> None:
    rows = block.rows[: 12 if appendix else 6]
    compact = [
        [
            _short(row[0] if len(row) > 0 else "", 58),
            _source_label(row[1] if len(row) > 1 else ""),
            _short(row[3] if len(row) > 3 else "", 24),
        ]
        for row in rows
    ]
    _render_table(
        doc,
        ["Источник", "Домен", "Надёжность"],
        compact,
        max_rows=12 if appendix else 6,
        max_text=120,
    )
    _data_pack_note(doc, len(block.rows), len(rows))


def _render_scenario_exhibit(doc: Document, block: PremiumPreparedBlock) -> None:
    for row in block.rows[:3]:
        title = _short(row[0] if len(row) > 0 else "", 40)
        definition = _short(row[1] if len(row) > 1 else "", 140)
        implication = _short(row[2] if len(row) > 2 else "", 420)
        table = doc.add_table(rows=1, cols=1)
        _style_table(table, header=False)
        cell = table.rows[0].cells[0]
        _shade(cell, "FFFFFF")
        _cell_border(cell, GOLD if "Баз" in title else LINE)
        p = cell.paragraphs[0]
        t = p.add_run(title + "\n")
        _set_run(t, size=10.5, color=NAVY, bold=True)
        d = p.add_run(definition + "\n")
        _set_run(d, size=8.8, color=MUTED, italic=True)
        body = p.add_run(implication)
        _set_run(body, size=9.2, color=INK)


def _render_risk_exhibit(
    doc: Document,
    block: PremiumPreparedBlock,
    *,
    appendix: bool = False,
) -> None:
    rows = [
        [
            _short(row[0] if len(row) > 0 else "", 54),
            _short(row[1] if len(row) > 1 else "", 42),
            _short(row[3] if len(row) > 3 else row[2] if len(row) > 2 else "", 160),
        ]
        for row in block.rows[: 10 if appendix else 5]
    ]
    _render_table(doc, ["Риск / тема", "Тип", "Что делать"], rows, max_rows=len(rows), max_text=150)


def _data_pack_note(doc: Document, total: int, shown: int) -> None:
    if total <= shown:
        return
    note = doc.add_paragraph(
        f"Показаны {shown} наиболее полезных строк из {total}; полный реестр сохранён в пакете данных."
    )
    _set_run(note.runs[0], size=8.5, color=MUTED, italic=True)


def _render_key_value_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(table, header=False)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if col_idx == 0:
                _shade(cell, PAPER)
                cell.paragraphs[0].add_run(value).bold = True
            else:
                cell.paragraphs[0].add_run(value)


def _render_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    max_rows: int = 10,
    max_text: int = 140,
) -> None:
    if not headers:
        headers = [f"Колонка {idx + 1}" for idx in range(max((len(row) for row in rows), default=1))]
    headers = headers[:5]
    visible_rows = rows[:max_rows]
    table = doc.add_table(rows=len(visible_rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade(cell, NAVY)
        run = cell.paragraphs[0].add_run(header)
        _set_run(run, size=8, color="FFFFFF", bold=True)
    for row_idx, row in enumerate(visible_rows, start=1):
        for col_idx, value in enumerate(row[: len(headers)]):
            cell = table.rows[row_idx].cells[col_idx]
            if row_idx % 2 == 0:
                _shade(cell, "FAFBFC")
            cell.paragraphs[0].add_run(_table_cell_text(value, max_text))
    if len(rows) > len(visible_rows):
        note = doc.add_paragraph(
            f"Таблица сокращена для читаемости: ещё {len(rows) - len(visible_rows)} строк доступны в пакете данных."
        )
        _set_run(note.runs[0], size=8.5, color=MUTED, italic=True)


def _style_table(table, *, header: bool = True) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _cell_border(cell, LINE)
            _cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    _set_run(run, size=8.2)


def _add_rule(doc: Document, color: str, size: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    p_pr.append(border)


def _shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_border(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ["top", "left", "bottom", "right"]:
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    _set_run(run, size=7.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _set_run(run, *, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None, all_caps: bool = False) -> None:
    run.font.name = "Aptos"
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if all_caps:
        run.font.all_caps = True


def _paragraphs(text: str) -> list[str]:
    return [part.strip() for part in str(text).splitlines() if part.strip()]


def _render_markdown_like_body(doc: Document, text: str) -> None:
    """Render common markdown shapes as readable DOCX paragraphs."""

    for raw in _paragraphs(text):
        line = raw.strip()
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
            continue
        if line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)


def _deliverables(report: PremiumReportDocument) -> str:
    deliverables = report.plan.deliverables
    names = []
    if deliverables.require_docx:
        names.append("DOCX-отчёт")
    if deliverables.require_pdf:
        names.append("PDF")
    if deliverables.require_pptx:
        names.append("PPTX-презентация")
    if deliverables.require_data_pack:
        names.append("пакет данных")
    return ", ".join(names)


def _client_confidence_label(report: PremiumReportDocument) -> str:
    readiness = report.premium_readiness or {}
    score = readiness.get("score")
    if isinstance(score, (int, float)) and not isinstance(score, bool):
        if score >= 85:
            return "Высокий: доказательная база достаточна для управленческого решения."
        if score >= 70:
            return "Средний: вывод пригоден для решения, но требует чтения ограничений."
        return "Ограниченный: выводы нужно использовать как рабочую гипотезу."
    if report.source_count >= 15 and report.numeric_fact_count >= 60:
        return "Средний или высокий: база источников и числовых фактов достаточна."
    return "Ограниченный: доказательная база требует расширения."


def _audience_label(audience: str) -> str:
    return {
        "buyer": "покупатель",
        "investor": "инвестор",
        "executive": "руководитель",
        "operator": "оператор",
        "developer": "девелопер",
        "analyst": "аналитик",
        "technical_lead": "технический руководитель",
        "general_client": "клиент",
    }.get(audience, audience.replace("_", " "))


def _report_type_label(report_type: str) -> str:
    return {
        "market": "Рыночный анализ",
        "investment": "Инвестиционный анализ",
        "competitive": "Конкурентный анализ",
        "strategy": "Стратегический отчёт",
        "technical_audit": "Технический аудит",
        "legal_regulatory": "Правовой и регуляторный анализ",
        "due_diligence": "Due diligence",
        "general_research": "Исследование",
    }.get(report_type, report_type.replace("_", " ").title())


def _first_section_block_body(
    report: PremiumReportDocument,
    section_id: str,
    block_title: str,
) -> str:
    for section in report.sections:
        if section.id != section_id:
            continue
        for block in section.blocks:
            if block.title == block_title and block.body:
                return block.body
    return ""


def _clip(text: str, limit: int) -> str:
    cleaned = " ".join(str(text or "").split())
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: max(0, limit - 1)].rstrip() + "..."


def _short(value: object, limit: int) -> str:
    text = _clean_inline(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit].rsplit(" ", 1)[0].rstrip(" ,.;:")
    return clipped + "..."


def _table_cell_text(value: object, limit: int) -> str:
    text = _clean_inline(value)
    if text.startswith(("https://", "http://")):
        return _source_label(text)
    return _short(text, limit)


def _clean_inline(value: object) -> str:
    text = " ".join(str(value or "").split())
    for token in ("**", "__"):
        text = text.replace(token, "")
    return text


def _source_label(value: object) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    for prefix in ("https://", "http://"):
        if text.startswith(prefix):
            text = text[len(prefix) :]
            break
    domain = text.split("/", 1)[0]
    return _short(domain or text, 48)
