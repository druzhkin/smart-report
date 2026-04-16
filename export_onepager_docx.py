"""One-page DOCX export — strict single A4 page.

Margins 15 mm, tight line spacing, minimal font sizes. No images.
Content is capped to fit: 3 top findings, 4 priority blocks, 2 connections,
4 gaps. Degrades gracefully when exec_summary is absent.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from models import Report

# ---------- palette (matches export_docx.py) ----------

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
NAVY_HEX = "1B3A5C"
RED_HEX = "C0392B"
AMBER_HEX = "D97706"
GREEN_HEX = "27AE60"
GREY_BG_HEX = "F0F4F8"
WHITE_HEX = "FFFFFF"
GREY_TXT = RGBColor(0x55, 0x55, 0x55)
FONT = "Arial"

PRIORITY_COLOR = {"high": RED_HEX, "medium": AMBER_HEX, "low": GREEN_HEX}
PRIORITY_LABEL_RU = {"high": "Высок.", "medium": "Средн.", "low": "Низк."}


# ---------- helpers ----------

def _trunc(s: str, n: int) -> str:
    """Truncate string to n chars with ellipsis."""
    if not s:
        return ""
    s = s.strip()
    return s if len(s) <= n else s[: n - 1] + "…"


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, hex_color: str = "CCCCCC", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), hex_color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_no_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _tight_para(para, space_before: int = 0, space_after: int = 30) -> None:
    """Set tight spacing on a paragraph (twips: 20 twips = 1 pt)."""
    pf = para.paragraph_format
    pf.space_before = Pt(0) if space_before == 0 else Pt(space_before / 20)
    pf.space_after = Pt(space_after / 20)
    pf.line_spacing = 1.0


def _run(para, text: str, bold: bool = False, size: int = 9,
         color: RGBColor | None = None, italic: bool = False):
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _section_label(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _tight_para(para, space_before=0, space_after=20)
    run = para.add_run(text.upper())
    run.font.name = FONT
    run.font.size = Pt(7)
    run.font.bold = True
    run.font.color.rgb = NAVY
    pf = para.paragraph_format
    pf.space_before = Pt(3)
    # bottom rule via paragraph border
    p_pr = para._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def _extract_first_number(text: str) -> str:
    m = re.search(r"(\$\s?[\d.,]+[kKmMbBтмлрд]*|[\d][\d\s.,]*%)", text)
    return m.group(1).strip() if m else ""


# ---------- data extraction (mirrors export_onepager.py logic) ----------

def _get_top_findings(report: Report) -> list[str]:
    if report.exec_summary and report.exec_summary.top_findings:
        return [tf.headline for tf in report.exec_summary.top_findings[:3]]
    findings: list[str] = []
    for b in report.blocks:
        for f in b.findings or []:
            if f.has_numbers and f.claim:
                findings.append(f.claim)
            if len(findings) >= 3:
                return findings
    return findings


def _get_key_numbers(report: Report, seen: set[str]) -> list[tuple[str, str]]:
    kpis: list[tuple[str, str]] = []
    for b in report.blocks:
        for f in b.findings or []:
            if not f.has_numbers:
                continue
            key = (f.claim or "").strip()[:80]
            if not key or key in seen:
                continue
            seen.add(key)
            num = _extract_first_number(f.claim)
            if not num:
                continue
            ctx = _trunc(f.claim, 80)
            kpis.append((num, ctx))
            if len(kpis) >= 3:
                return kpis
    return kpis


def _get_priority_rows(report: Report) -> list[tuple[str, str, str, str]]:
    header_by_cell = {h.cell: h for h in (report.block_headers or [])}
    ordered = sorted(
        report.blocks,
        key=lambda b: {"high": 0, "medium": 1, "low": 2}.get(
            (header_by_cell[b.cell].priority if b.cell in header_by_cell else "low"), 3
        ),
    )
    rows: list[tuple[str, str, str, str]] = []
    for b in ordered[:4]:
        h = header_by_cell.get(b.cell)
        pri = h.priority if h else "low"
        one = _trunc(h.one_liner if h and h.one_liner else (b.summary or ""), 90)
        num = _trunc(h.strongest_number if h and h.strongest_number else "", 30)
        rows.append((_trunc(b.cell, 28), one, num, pri))
    return rows


def _get_gaps(report: Report) -> list[str]:
    if report.exec_summary and report.exec_summary.key_gaps:
        return [_trunc(g, 80) for g in report.exec_summary.key_gaps[:4]]
    gaps: list[str] = []
    for b in report.blocks:
        if b.gaps:
            gaps.append(_trunc(b.gaps[0], 80))
        if len(gaps) >= 4:
            break
    return gaps


# ---------- document assembly ----------

def export_onepager_docx(report: Report, out_path: Path) -> Path:
    """Build a strict one-A4-page DOCX one-pager and write to out_path."""
    doc = Document()

    # -- page setup: A4 with 15 mm margins --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(1.5))

    # -- base Normal style --
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    today = date.today().strftime("%d.%m.%Y")
    goal_text = ""
    if report.exec_summary and report.exec_summary.goal_restate:
        goal_text = report.exec_summary.goal_restate
    else:
        goal_text = report.goal or "Отчёт"

    n_blocks = len(report.blocks)
    n_conns = len(report.connections or [])

    # ---- HEADER STRIP (1-row table, full width, NAVY bg) ----
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.style = "Table Grid"
    hdr_tbl.autofit = False
    total_w = Cm(18)  # 21 - 2*1.5
    hdr_tbl.columns[0].width = int(total_w * 0.75)
    hdr_tbl.columns[1].width = int(total_w * 0.25)

    left_cell = hdr_tbl.cell(0, 0)
    right_cell = hdr_tbl.cell(0, 1)
    _shade_cell(left_cell, NAVY_HEX)
    _shade_cell(right_cell, NAVY_HEX)

    # remove all inner borders
    for c in (left_cell, right_cell):
        tc_pr = c._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "none")
            tc_borders.append(b)
        tc_pr.append(tc_borders)

    lp = left_cell.paragraphs[0]
    _tight_para(lp, space_before=0, space_after=20)
    kicker = lp.add_run("ONE-PAGER  ·  SMART REPORT")
    kicker.font.name = FONT
    kicker.font.size = Pt(7)
    kicker.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
    lp.add_run("\n")
    title_run = lp.add_run(_trunc(goal_text, 120))
    title_run.font.name = FONT
    title_run.font.size = Pt(13)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rp = right_cell.paragraphs[0]
    _tight_para(rp)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_run = rp.add_run(f"{today}\n{n_blocks} блоков · {n_conns} связей")
    meta_run.font.name = FONT
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

    # spacing after header
    spacer = doc.add_paragraph()
    _tight_para(spacer, space_before=0, space_after=20)

    # ---- TOP FINDINGS ----
    top_findings = _get_top_findings(report)
    seen_keys: set[str] = {t.strip()[:80] for t in top_findings}

    if top_findings:
        _section_label(doc, "Главные выводы")
        for i, tf in enumerate(top_findings, 1):
            p = doc.add_paragraph()
            _tight_para(p, space_after=16)
            _run(p, f"{i}. ", bold=True, size=9, color=NAVY)
            _run(p, _trunc(tf, 160), size=9)

    # ---- KEY NUMBERS (2-col mini table) ----
    kpis = _get_key_numbers(report, seen_keys)
    if kpis:
        _section_label(doc, "Ключевые цифры")
        cols = min(len(kpis), 3)
        kpi_tbl = doc.add_table(rows=1, cols=cols)
        kpi_tbl.style = "Table Grid"
        kpi_tbl.autofit = False
        col_w = total_w // cols
        for ci in range(cols):
            kpi_tbl.columns[ci].width = col_w
        for ci, (num, ctx) in enumerate(kpis[:cols]):
            cell = kpi_tbl.cell(0, ci)
            _shade_cell(cell, GREY_BG_HEX)
            _set_cell_borders(cell, "DDDDDD", 4)
            cp = cell.paragraphs[0]
            _tight_para(cp, space_after=12)
            _run(cp, num + "\n", bold=True, size=12, color=NAVY)
            _run(cp, ctx, size=7, color=GREY_TXT)
        sp2 = doc.add_paragraph()
        _tight_para(sp2, space_after=16)

    # ---- PRIORITY BLOCKS ----
    rows = _get_priority_rows(report)
    if rows:
        _section_label(doc, "Приоритеты по блокам")
        pri_tbl = doc.add_table(rows=1 + len(rows), cols=4)
        pri_tbl.style = "Table Grid"
        pri_tbl.autofit = False
        col_widths = [
            int(total_w * 0.22),
            int(total_w * 0.50),
            int(total_w * 0.17),
            int(total_w * 0.11),
        ]
        for ci, w in enumerate(col_widths):
            pri_tbl.columns[ci].width = w

        # header row
        hdrs = ["Блок", "Суть", "Цифра", "Приор."]
        for ci, hdr_text in enumerate(hdrs):
            cell = pri_tbl.cell(0, ci)
            _shade_cell(cell, NAVY_HEX)
            _set_cell_borders(cell, NAVY_HEX, 4)
            hp = cell.paragraphs[0]
            _tight_para(hp, space_after=12)
            _run(hp, hdr_text, bold=True, size=7,
                 color=RGBColor(0xFF, 0xFF, 0xFF))

        for ri, (cell_name, one, num, pri) in enumerate(rows, 1):
            bg = "FFFFFF" if ri % 2 == 0 else GREY_BG_HEX
            pri_color_hex = PRIORITY_COLOR.get(pri, "999999")
            pri_label = PRIORITY_LABEL_RU.get(pri, pri or "—")
            values = [cell_name, one, num or "—", pri_label]
            for ci, val in enumerate(values):
                cell = pri_tbl.cell(ri, ci)
                _shade_cell(cell, bg)
                _set_cell_borders(cell, "DDDDDD", 4)
                vp = cell.paragraphs[0]
                _tight_para(vp, space_after=10)
                color = RGBColor.from_string(pri_color_hex) if ci == 3 else None
                _run(vp, val, bold=(ci == 3), size=8, color=color)
        sp3 = doc.add_paragraph()
        _tight_para(sp3, space_after=16)

    # ---- CONNECTIONS ----
    conns = (report.connections or [])[:2]
    if conns:
        _section_label(doc, "Неожиданные связи")
        for c in conns:
            pair = " ↔ ".join(c.domains or [])
            desc = _trunc(c.description or "", 140)
            p = doc.add_paragraph()
            _tight_para(p, space_after=14)
            _run(p, _trunc(pair, 60) + ": ", bold=True, size=8, color=NAVY)
            _run(p, desc, size=8)

    # ---- ANALOGIES (top 2, only if block count <= 3 to preserve page budget) ----
    if len(report.blocks) <= 3:
        top_analogies: list[tuple[str, object]] = []
        for b in report.blocks:
            for a in (getattr(b, "analogies", None) or []):
                top_analogies.append((b.cell, a))
                if len(top_analogies) >= 2:
                    break
            if len(top_analogies) >= 2:
                break
        if top_analogies:
            _section_label(doc, "Аналогии")
            for cell, a in top_analogies:
                loc = getattr(a, "location", "") or ""
                lesson = getattr(a, "lesson", "") or ""
                label = f"{loc}: " if loc else f"{cell}: "
                p = doc.add_paragraph()
                _tight_para(p, space_after=10)
                _run(p, label, bold=True, size=8, color=NAVY)
                _run(p, _trunc(lesson, 120), size=8)

    # ---- INVERSIONS (top 2 critical only, page-budget aware) ----
    inv_blocks = getattr(report, "assumption_inversions", None) or []
    critical_inv = [
        (bi.block_cell, inv)
        for bi in inv_blocks
        for inv in bi.inversions
        if inv.dependency == "critical"
    ][:2]
    if critical_inv:
        _section_label(doc, "Проверка допущений")
        for cell, inv in critical_inv:
            p = doc.add_paragraph()
            _tight_para(p, space_after=10)
            _run(p, f"{cell}: ", bold=True, size=8, color=NAVY)
            _run(p, f"{inv.assumption} \u2192 ", size=8)
            _run(p, inv.inversion, bold=True, size=8)

    # ---- GAPS FOOTER ----
    gaps = _get_gaps(report)
    if gaps:
        _section_label(doc, "Пробелы / следующие шаги")
        for g in gaps:
            p = doc.add_paragraph()
            _tight_para(p, space_after=10)
            _run(p, "→ ", bold=True, size=8, color=NAVY)
            _run(p, g, size=8)

    # ---- SCENARIO CONE (conditional_verdict only — keep page tight) ----
    cone = getattr(report, "scenario_cone", None)
    if cone:
        cone_verdict = getattr(cone, "conditional_verdict", "") or ""
        cone_horizon = getattr(cone, "question_horizon", "12-24 месяцев") or "12-24 месяцев"
        cone_scenarios = getattr(cone, "scenarios", []) or []
        _section_label(doc, f"Конус сценариев · {cone_horizon}")
        if cone_scenarios:
            sc_cols = min(len(cone_scenarios), 3)
            sc_tbl = doc.add_table(rows=1, cols=sc_cols)
            sc_tbl.style = "Table Grid"
            sc_tbl.autofit = False
            sc_col_w = total_w // sc_cols
            for ci in range(sc_cols):
                sc_tbl.columns[ci].width = sc_col_w
            for ci, s in enumerate(cone_scenarios[:sc_cols]):
                cell = sc_tbl.cell(0, ci)
                _shade_cell(cell, GREY_BG_HEX)
                _set_cell_borders(cell, "DDDDDD", 4)
                cp = cell.paragraphs[0]
                _tight_para(cp, space_after=10)
                _run(cp, (getattr(s, "probability", "") or "") + "\n", bold=True, size=10, color=NAVY)
                _run(cp, _trunc(getattr(s, "name", "") or "", 30) + "\n", bold=True, size=8)
                _run(cp, _trunc(getattr(s, "description", "") or "", 80), size=7, color=GREY_TXT)
            sp_sc = doc.add_paragraph()
            _tight_para(sp_sc, space_after=10)
        if cone_verdict:
            p = doc.add_paragraph()
            _tight_para(p, space_after=12)
            _run(p, cone_verdict, size=8, italic=True, color=GREY_TXT)

    # ---- branding footer line ----
    fp = doc.add_paragraph()
    _tight_para(fp, space_before=0, space_after=0)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(fp, f"Smart Report · сгенерировано {today}", size=7, color=GREY_TXT,
         italic=True)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
