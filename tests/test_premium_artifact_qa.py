from __future__ import annotations

import pytest

from smart_report.exporters.premium import (
    assemble_premium_report_document,
    render_premium_docx,
    render_premium_pdf,
    render_premium_pptx,
)
from smart_report.exporters.premium.artifact_qa import run_qa
from smart_report.models import (
    AnalysisOutput,
    Conflict,
    ConsensusClaim,
    ExecutiveSummaryV4,
    FinalReport,
    Gap,
    NumericFact,
    Source,
    SourceRef,
)


def _report() -> FinalReport:
    return FinalReport(
        session_id="premium-artifact-qa",
        question="Forecast a market with scenario and risk recommendations",
        executive_summary=ExecutiveSummaryV4(
            main_answer="Base case is moderate growth with material downside triggers.",
            top_findings=["Demand is rate-sensitive.", "Supply remains constrained."],
            confidence_note="Confidence is medium because transaction data is partial.",
            what_meta_adds="Consensus, conflicts, and gaps are separated.",
        ),
        main_synthesis="The market baseline combines price, demand, supply, and policy drivers.",
        consensus_section="Sources agree on the direction but differ on magnitude.",
        conflicts_section="Sources disagree on timing.",
        gaps_filled_section="Project-level microdata remains unavailable.",
        all_sources=[
            Source(title="Official source", url="https://example.com/official", reliability="high"),
            Source(title="Industry source", url="https://example.com/industry", reliability="medium"),
        ],
    )


def _analysis() -> AnalysisOutput:
    facts = [
        NumericFact(
            fact_id=f"f{idx}",
            value=f"{idx}%",
            metric="growth",
            subject="market",
            relevance_to_question="high",
            sources=[SourceRef(url="https://example.com/official", title="Official")],
        )
        for idx in range(1, 8)
    ]
    return AnalysisOutput(
        consensus=[
            ConsensusClaim(claim="Demand is sensitive to financing conditions.", confidence="high"),
            ConsensusClaim(claim="Supply remains constrained.", confidence="medium"),
        ],
        conflicts=[
            Conflict(
                topic="Growth range",
                source_a="A",
                claim_a="High growth",
                source_b="B",
                claim_b="Moderate growth",
                importance="material",
            )
        ],
        gaps=[
            Gap(
                topic="Transaction microdata",
                why_critical="Needed for exact pricing",
                what_to_find="Closed transaction dataset",
            )
        ],
        all_numeric_facts=facts,
        high_relevance_facts=facts,
    )


@pytest.fixture()
def premium_artifacts(tmp_path):
    pytest.importorskip("docx")
    pytest.importorskip("pptx")
    document = assemble_premium_report_document(
        _report(),
        analysis=_analysis(),
        premium_readiness={
            "ready": False,
            "score": 67,
            "issues": [
                {
                    "code": "premium_visual_render_not_verified",
                    "severity": "major",
                    "message": "Visual rendering has not been completed.",
                    "recommendation": "Run DOCX/PPTX render QA before paid delivery.",
                }
            ],
            "strengths": ["Consensus layer is present."],
        },
    )
    docx_path = render_premium_docx(document, tmp_path / "premium_report.docx")
    pptx_path = render_premium_pptx(document, tmp_path / "premium_deck.pptx")
    return docx_path, pptx_path


def test_premium_artifact_qa_structural_checks_pass_without_render(premium_artifacts, tmp_path):
    docx_path, pptx_path = premium_artifacts

    report = run_qa(docx_path=docx_path, pptx_path=pptx_path, out_dir=tmp_path, render=False)

    assert report["status"] == "passed"
    assert report["summary"]["artifacts"] == 2
    assert report["summary"]["passed_structural"] == 2
    docx_result = next(item for item in report["results"] if item["kind"] == "docx")
    pptx_result = next(item for item in report["results"] if item["kind"] == "pptx")
    assert docx_result["metrics"]["tables"] >= 4
    assert docx_result["metrics"]["estimated_pages"] >= 1
    assert docx_result["metrics"]["narrative_chars"] >= 1800
    assert docx_result["metrics"]["source_reference_count"] >= 2
    assert docx_result["metrics"]["overlong_paragraphs"] == []
    assert docx_result["metrics"]["has_decision_dashboard"] is True
    assert docx_result["metrics"]["has_scorecard"] is True
    assert docx_result["metrics"]["has_readiness_gate"] is False
    assert pptx_result["metrics"]["slides"] >= 10
    assert pptx_result["metrics"]["has_client_position"] is True
    assert pptx_result["metrics"]["has_readiness"] is False


def test_premium_artifact_qa_rejects_untraceable_docx_text_dump(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    docx_path = tmp_path / "bad_report.docx"
    doc = Document()
    doc.add_paragraph("SMART REPORT | PREMIUM ANALYTICAL REPORT")
    doc.add_heading("Client Decision Dashboard", level=1)
    doc.add_heading("Executive Evidence Scorecard", level=1)
    doc.add_heading("Report Structure", level=1)
    for idx in range(26):
        doc.add_heading(f"Section {idx + 1}", level=2)
        body = (
            "This paragraph is intentionally long and undifferentiated. " * 18
            if idx == 0
            else "A short analytical paragraph with a business implication and no traceable source. "
            * 4
        )
        doc.add_paragraph(body)
    for table_idx in range(4):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = f"Table {table_idx + 1}"
        table.cell(1, 1).text = "Unreferenced"
    doc.save(docx_path)

    report = run_qa(docx_path=docx_path, out_dir=tmp_path, render=False)

    assert report["status"] == "failed"
    docx_result = next(item for item in report["results"] if item["kind"] == "docx")
    assert docx_result["metrics"]["source_reference_count"] == 0
    assert docx_result["metrics"]["overlong_paragraphs"]
    assert any("source references" in issue for issue in docx_result["issues"])
    assert any("overlong paragraph" in issue for issue in docx_result["issues"])


def test_premium_artifact_qa_checks_publication_pdf(tmp_path):
    pytest.importorskip("reportlab")
    pytest.importorskip("pypdf")
    document = assemble_premium_report_document(
        _report(),
        analysis=_analysis(),
        premium_readiness={
            "ready": True,
            "score": 91,
            "issues": [],
            "strengths": ["Publication-grade layout requirements are present."],
        },
    )
    pdf_path = render_premium_pdf(document, tmp_path / "premium_report.pdf")

    report = run_qa(pdf_path=pdf_path, out_dir=tmp_path, render=False)

    assert report["status"] == "passed"
    assert report["summary"]["artifacts"] == 1
    pdf_result = next(item for item in report["results"] if item["kind"] == "pdf")
    assert pdf_result["metrics"]["pages"] >= 20
    assert pdf_result["metrics"]["has_publication_marker"] is True
    assert pdf_result["metrics"]["has_exhibit_pages"] is True
    assert pdf_result["metrics"]["has_source_notes"] is True
    assert pdf_result["metrics"]["landscape_pages"] == []
    assert pdf_result["metrics"]["thin_pages_after_cover"] == []


def test_premium_artifact_qa_rejects_landscape_pdf(tmp_path):
    pytest.importorskip("reportlab")
    pytest.importorskip("pypdf")

    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "landscape_report.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=landscape(letter))
    for page in range(20):
        c.drawString(48, 540, "SMART REPORT | Publication-grade PDF")
        c.drawString(48, 512, f"EXHIBIT {page + 1}: Source: internal QA fixture.")
        c.drawString(
            48,
            484,
            "This page has enough text to avoid thin-page failure while keeping orientation invalid.",
        )
        c.showPage()
    c.save()

    report = run_qa(pdf_path=pdf_path, out_dir=tmp_path, render=False)

    assert report["status"] == "failed"
    pdf_result = next(item for item in report["results"] if item["kind"] == "pdf")
    assert pdf_result["metrics"]["landscape_pages"]
    assert any("landscape" in issue for issue in pdf_result["issues"])


def test_premium_artifact_qa_reports_missing_render_tools(monkeypatch, premium_artifacts, tmp_path):
    docx_path, pptx_path = premium_artifacts

    monkeypatch.setattr(
        "smart_report.exporters.premium.artifact_qa._find_tool",
        lambda _name, _candidates: None,
    )
    report = run_qa(docx_path=docx_path, pptx_path=pptx_path, out_dir=tmp_path, render=True)

    assert report["status"] == "blocked"
    assert report["summary"]["blocked_render"] == 2
    assert all("soffice" in item["missing_tools"] for item in report["results"])
    assert all("pdftoppm" in item["missing_tools"] for item in report["results"])


def test_premium_artifact_qa_records_rendered_docx_page_count(
    monkeypatch, premium_artifacts, tmp_path
):
    docx_path, _pptx_path = premium_artifacts

    def _fake_find_tool(name, _candidates):
        return name

    def _fake_run(cmd, **_kwargs):
        out_dir = tmp_path / f"{docx_path.stem}_docx"
        out_dir.mkdir(parents=True, exist_ok=True)
        if "--convert-to" in cmd:
            (out_dir / f"{docx_path.stem}.pdf").write_bytes(b"%PDF")
        else:
            for idx in range(1, 4):
                (out_dir / f"{docx_path.stem}-{idx:02d}.png").write_bytes(b"png")

    monkeypatch.setattr("smart_report.exporters.premium.artifact_qa._find_tool", _fake_find_tool)
    monkeypatch.setattr("smart_report.exporters.premium.artifact_qa.subprocess.run", _fake_run)

    report = run_qa(docx_path=docx_path, out_dir=tmp_path, render=True)

    docx_result = next(item for item in report["results"] if item["kind"] == "docx")
    assert docx_result["render_status"] == "passed"
    assert docx_result["metrics"]["rendered_pages"] == 3
    assert docx_result["metrics"]["rendered_total_bytes"] == 9
