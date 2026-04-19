"""Tests for the consulting DOCX renderer (Track B).

Tests cover:
1. Rendering from a rich FinalReport fixture (all new fields populated)
2. Rendering from a minimal FinalReport (backward-compat — only legacy fields)
3. Document structure: sections, tables, callouts, Q&A, key numbers, ranking, sources
4. DOCX file validity (opens cleanly via python-docx readback)
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

pytest.importorskip("docx")

from docx import Document

from smart_report.exporters.docx_v4_consulting import render_consulting_docx
from smart_report.models import (
    CalloutBlock,
    ChartSpec,
    ExecutiveSummaryV4,
    FinalReport,
    KeyNumber,
    KeyNumberHighlight,
    QAItem,
    RankingItem,
    Source,
    Table,
)

# ---------------------------------------------------------------------------
# Helpers / fixtures
# ---------------------------------------------------------------------------

FIXTURES_DIR = Path(__file__).parent / "fixtures"


def _rich_report() -> FinalReport:
    """Full FinalReport with all Track A structured output fields."""
    return FinalReport(
        session_id="test-consulting-001",
        question="Что входит в оптимальный пакет amenities для бизнес-класса Москвы?",
        research_prompt_used="Анализ amenities в бизнес и премиум новостройках Москвы.",
        executive_summary=ExecutiveSummaryV4(
            main_answer=(
                "Бассейн, фитнес и консьерж дают 8–14% чистой выручки поверх потери площади."
            ),
            ranking="Бассейн > фитнес > консьерж > SPA > коворкинг",
            top_findings=[
                "Проекты с бассейном продаются на 12–18% дороже аналогов.",
                "Консьерж окупается за 3–5 лет — лучший ROI из всех amenities.",
                "Сигарная теряет актуальность: спрос -40% с 2020 г.",
            ],
            key_numbers=[
                KeyNumber(value="12–18%", metric="премия за бассейн", subject="бизнес-класс"),
            ],
            confidence_note="Высокая уверенность по ценовым данным.",
            what_meta_adds="Разрешён конфликт: доля ипотеки 55% (ЕРЗ) принята вместо 68%.",
        ),
        main_synthesis=(
            "## Что реально пользуется спросом\n\n"
            "Покупатели бизнес-класса ставят на первое место **качество МОП** и wellness.\n\n"
            "- Indoor бассейн 25–35м\n"
            "- Фитнес-зал 400–600м²\n"
            "- Консьерж 24/7\n\n"
            "## Ценовая премия за amenities\n\n"
            "По данным анализа 47 проектов: бассейн даёт 12–18%, фитнес 5–8%.\n\n"
            "## Экономический баланс\n\n"
            "Потеря площади 7–9% при ценовой премии 15–22% = +8–14% чистой выручки."
        ),
        consensus_section="Три источника согласны: wellness растёт, МОП критичен.",
        conflicts_section="Ипотека: 55% (ЕРЗ) vs 68% (OpenAI DR) — принята цифра ЕРЗ.",
        gaps_filled_section="Размер бассейна: принят диапазон 25–35м.",
        all_sources=[
            Source(title="ЕРЗ 2025", url="https://erzrf.ru/", tool="perplexity",
                   reliability="high"),
            Source(title="Knight Frank 2024", url="https://knightfrank.ru/",
                   tool="openai_dr", reliability="high"),
            Source(title="CBRE 2023", url="https://cbre.com/", tool="claude",
                   reliability="medium"),
        ],
        metadata={"source_reports_count": 3},
        qa_section=[
            QAItem(
                question="Что пользуется наибольшим спросом?",
                answer="Бассейн, фитнес, консьерж — топ-3 по ценовому влиянию.",
                details_ref="Раздел «Что реально пользуется спросом», стр. 4",
            ),
            QAItem(
                question="Сколько платят за amenities?",
                answer="Совокупная ценовая премия пакета бассейн+фитнес+консьерж — 20–28%.",
                details_ref="Таблица 1, стр. 6",
            ),
            QAItem(
                question="Каков экономический баланс для застройщика?",
                answer="Потеря 7–9% площади компенсируется 15–22% ценовой премией (+8–14% выручки).",
                details_ref="Раздел «Экономический баланс», стр. 7",
            ),
        ],
        ranking=[
            RankingItem(label="Indoor бассейн", weight=45, rationale="12–18% премии",
                        evidence_strength="high"),
            RankingItem(label="Фитнес-зал", weight=35, rationale="5–8% премии",
                        evidence_strength="high"),
            RankingItem(label="Консьерж", weight=30, rationale="Лучший ROI",
                        evidence_strength="high"),
        ],
        tables=[
            Table(
                title="Ценовые премии по типам amenities",
                columns=["Amenity", "Премия", "Окупаемость"],
                rows=[
                    ["Бассейн 25м", "12–18%", "8–12 лет"],
                    ["Фитнес", "5–8%", "7–9 лет"],
                    ["Консьерж", "4–7%", "3–5 лет"],
                ],
                caption="Данные по 47 проектам Москвы.",
                source_ref="ЕРЗ 2025, Knight Frank 2024",
            ),
            Table(
                title="Топ проектов по цене",
                columns=["Проект", "Цена (тыс./м²)", "Бассейн"],
                rows=[
                    ["Prime Park", "883.8", "Да"],
                    ["Tsvetnoy & Co", "720", "Да"],
                ],
                caption="H1 2025.",
                source_ref="РБК Недвижимость",
            ),
        ],
        charts=[
            ChartSpec(
                chart_type="bar",
                title="Ценовая премия по amenities (%)",
                data={"labels": ["Бассейн", "Фитнес", "Консьерж"],
                      "values": [15, 6.5, 5.5]},
                x_label="Тип amenity",
                y_label="Премия (%)",
                caption="Средние значения по 47 проектам.",
            ),
        ],
        callouts=[
            CalloutBlock(
                kind="insight",
                title="Консьерж — лучший ROI",
                body="При минимальных потерях площади консьерж окупается за 3–5 лет.",
            ),
            CalloutBlock(
                kind="warning",
                title="Сигарная теряет актуальность",
                body="Спрос упал на 40% с 2020. Рекомендуется отказаться в пользу коворкинга.",
            ),
        ],
        key_numbers_highlight=[
            KeyNumberHighlight(
                value="883.8 тыс.",
                label="руб./м² — Prime Park H1 2025",
                source_ref="РБК Недвижимость",
                importance="headline",
            ),
            KeyNumberHighlight(
                value="12–18%",
                label="премия за indoor бассейн",
                source_ref="Knight Frank / ЕРЗ",
                importance="headline",
            ),
            KeyNumberHighlight(
                value="73%",
                label="покупателей ценят МОП",
                source_ref="Knight Frank 2024",
                importance="primary",
            ),
        ],
    )


def _minimal_report() -> FinalReport:
    """Minimal FinalReport with only legacy fields (backward-compat test)."""
    return FinalReport(
        session_id="test-minimal-001",
        question="Какой оптимальный набор amenities для бизнес-класса?",
        executive_summary=ExecutiveSummaryV4(
            main_answer="Бассейн + фитнес + консьерж — оптимальный набор.",
        ),
        main_synthesis="## Основной вывод\n\nBassein + fitnes + konsierj.",
        all_sources=[
            Source(title="ЕРЗ 2025", url="https://erzrf.ru/", tool="perplexity",
                   reliability="high"),
        ],
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestRichReport:
    """Tests using the rich fixture with all Track A fields populated."""

    def test_generates_file(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        assert out.exists()
        assert out.stat().st_size > 5_000  # at least 5KB — real content

    def test_docx_opens_without_error(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        assert len(doc.paragraphs) > 20

    def test_has_cover_heading(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Cover should contain a fragment of the question
        assert "amenities" in all_text.lower() or "бизнес" in all_text.lower()

    def test_has_exec_summary_heading(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Executive Summary" in headings

    def test_has_qa_questions(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Q&A items should appear as "В: ..." paragraphs
        assert "В:" in all_text

    def test_has_ranking(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Indoor бассейн" in all_text

    def test_has_toc_section(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Содержание" in headings

    def test_has_tables(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        assert len(doc.tables) >= 2  # at least key numbers grid + 1 data table

    def test_table_header_has_content(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        # Find a table with "Amenity" header
        found = False
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if "Amenity" in cell.text or "Премия" in cell.text:
                        found = True
        assert found, "Expected table with 'Amenity' or 'Премия' column not found"

    def test_has_callout_content(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Консьерж" in all_text  # from callout title
        assert "ROI" in all_text        # from callout body

    def test_has_sources_section(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Источники" in headings

    def test_sources_grouped_by_tool(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Perplexity" in headings
        assert "OpenAI Deep Research" in headings
        assert "Claude" in headings

    def test_has_main_chapter_headings(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Что реально пользуется спросом" in headings
        assert "Ценовая премия за amenities" in headings

    def test_chart_placeholder_present(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Chart placeholder contains [BAR] or chart title
        assert "BAR" in all_text or "Ценовая премия по amenities" in all_text

    def test_key_numbers_table_present(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        # Key numbers are rendered as a table — check table cells
        key_number_found = False
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if "883.8" in cell.text or "12–18%" in cell.text:
                        key_number_found = True
        assert key_number_found, "Key number values not found in any table cell"

    def test_consensus_section_present(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Консенсус источников" in headings

    def test_no_empty_file(self, tmp_path):
        report = _rich_report()
        out = render_consulting_docx(report, tmp_path / "consulting.docx")
        assert out.stat().st_size > 0


class TestMinimalReport:
    """Backward-compatibility: minimal FinalReport without new Track A fields."""

    def test_renders_without_error(self, tmp_path):
        report = _minimal_report()
        out = render_consulting_docx(report, tmp_path / "minimal.docx")
        assert out.exists()
        assert out.stat().st_size > 2_000

    def test_opens_cleanly(self, tmp_path):
        report = _minimal_report()
        out = render_consulting_docx(report, tmp_path / "minimal.docx")
        doc = Document(str(out))
        assert len(doc.paragraphs) > 5

    def test_has_exec_summary_without_qa(self, tmp_path):
        """Without qa_section, main_answer should appear as a callout block."""
        report = _minimal_report()
        out = render_consulting_docx(report, tmp_path / "minimal.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Бассейн" in all_text  # from main_answer

    def test_has_sources(self, tmp_path):
        report = _minimal_report()
        out = render_consulting_docx(report, tmp_path / "minimal.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ЕРЗ" in all_text


class TestFixtureFile:
    """Test against the JSON fixture file (full amenities sample)."""

    def test_fixture_renders(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_consulting_docx(report, tmp_path / "fixture_output.docx")
        assert out.exists()
        assert out.stat().st_size > 10_000

    def test_fixture_has_all_sections(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_consulting_docx(report, tmp_path / "fixture_output.docx")
        doc = Document(str(out))

        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Executive Summary" in headings
        assert "Содержание" in headings
        assert "Источники" in headings
        # Should have at least 5 main chapter headings + ES + TOC + Sources
        assert len(headings) >= 8

    def test_fixture_has_minimum_tables(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_consulting_docx(report, tmp_path / "fixture_output.docx")
        doc = Document(str(out))
        # 1 key_numbers grid + 3 data tables = 4 minimum
        assert len(doc.tables) >= 4

    def test_fixture_has_five_qa_items(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_consulting_docx(report, tmp_path / "fixture_output.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # 5 QA items → at least 5 "В:" prefixes
        qa_count = all_text.count("В:")
        assert qa_count >= 5, f"Expected >=5 Q&A items, got {qa_count}"
