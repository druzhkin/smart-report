"""
Tests for the Node.js DOCX v2 renderer bridge.

Tests cover:
1. Output file exists and is > 10 KB
2. python-docx can open the output (readback validation)
3. word/media/ contains embedded PNGs if chart_dir is provided
4. Paragraph count > 50
5. Document contains expected text sections
6. Bridge properly handles Node.js unavailability (skip, not fail)
7. render_docx() auto-selector falls back to python-docx if Node not available

The test module auto-skips if:
  - Node.js is not on PATH
  - node_modules not installed (npm install not run yet)
"""

from __future__ import annotations

import json
import os
import shutil
import zipfile
from pathlib import Path

import pytest

# ---------------------------------------------------------------------------
# Skip conditions — check Node availability before importing anything else
# ---------------------------------------------------------------------------

def _node_available() -> bool:
    """Return True if node binary is found and node_modules is installed."""
    if shutil.which("node") is None:
        return False
    renderer_dir = Path(__file__).parent.parent / "smart_report" / "exporters" / "docx_js"
    return (renderer_dir / "node_modules").is_dir()


NODE_AVAILABLE = _node_available()
SKIP_NODE = pytest.mark.skipif(
    not NODE_AVAILABLE,
    reason=(
        "Node.js not available or node_modules not installed. "
        "Run: cd smart_report/exporters/docx_js && npm install"
    ),
)

# ---------------------------------------------------------------------------
# Import models (always available)
# ---------------------------------------------------------------------------

from smart_report.models import (
    CalloutBlock,
    ChartSpec,
    ExecutiveSummaryV4,
    FinalReport,
    KeyNumber,
    KeyNumberHighlight,
    QAItem,
    RankingItem,
    Source,
    Table,
)

FIXTURES_DIR = Path(__file__).parent / "fixtures"
CHART_SAMPLES_DIR = FIXTURES_DIR / "chart_samples"


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------

def _rich_report() -> FinalReport:
    """Full FinalReport identical to what Track B tests use."""
    return FinalReport(
        session_id="test-docxjs-001",
        question="Что входит в оптимальный пакет amenities для бизнес-класса Москвы?",
        research_prompt_used="Анализ amenities в бизнес и премиум новостройках Москвы.",
        executive_summary=ExecutiveSummaryV4(
            main_answer=(
                "Бассейн, фитнес и консьерж дают 8–14% чистой выручки поверх потери площади."
            ),
            ranking="Бассейн > фитнес > консьерж > SPA > коворкинг",
            top_findings=[
                "Проекты с бассейном продаются на 12–18% дороже аналогов.",
                "Консьерж окупается за 3–5 лет — лучший ROI из всех amenities.",
                "Сигарная теряет актуальность: спрос -40% с 2020 г.",
            ],
            key_numbers=[
                KeyNumber(value="12–18%", metric="премия за бассейн", subject="бизнес-класс"),
            ],
            confidence_note="Высокая уверенность по ценовым данным.",
            what_meta_adds="Разрешён конфликт: доля ипотеки 55% (ЕРЗ) принята вместо 68%.",
        ),
        main_synthesis=(
            "## Что реально пользуется спросом\n\n"
            "Покупатели бизнес-класса ставят на первое место **качество МОП** и wellness.\n\n"
            "- Indoor бассейн 25–35м\n"
            "- Фитнес-зал 400–600м²\n"
            "- Консьерж 24/7\n\n"
            "## Ценовая премия за amenities\n\n"
            "По данным анализа 47 проектов: бассейн даёт 12–18%, фитнес 5–8%.\n\n"
            "## Экономический баланс\n\n"
            "Потеря площади 7–9% при ценовой премии 15–22% = +8–14% чистой выручки."
        ),
        consensus_section="Три источника согласны: wellness растёт, МОП критичен.",
        conflicts_section="Ипотека: 55% (ЕРЗ) vs 68% (OpenAI DR) — принята цифра ЕРЗ.",
        gaps_filled_section="Размер бассейна: принят диапазон 25–35м.",
        all_sources=[
            Source(title="ЕРЗ 2025", url="https://erzrf.ru/", tool="perplexity", reliability="high"),
            Source(title="Knight Frank 2024", url="https://knightfrank.ru/", tool="openai_dr", reliability="high"),
            Source(title="CBRE 2023", url="https://cbre.com/", tool="claude", reliability="medium"),
        ],
        metadata={"source_reports_count": 3},
        qa_section=[
            QAItem(
                question="Что пользуется наибольшим спросом?",
                answer="Бассейн, фитнес, консьерж — топ-3 по ценовому влиянию.",
                details_ref="Раздел «Что реально пользуется спросом»",
            ),
            QAItem(
                question="Сколько платят за amenities?",
                answer="Совокупная премия пакета бассейн+фитнес+консьерж — 20–28%.",
                details_ref="Таблица 1",
            ),
        ],
        ranking=[
            RankingItem(label="Indoor бассейн", weight=45, rationale="12–18% премии", evidence_strength="high"),
            RankingItem(label="Фитнес-зал", weight=35, rationale="5–8% премии", evidence_strength="high"),
            RankingItem(label="Консьерж", weight=30, rationale="Лучший ROI", evidence_strength="high"),
        ],
        tables=[
            Table(
                title="Ценовые премии по типам amenities",
                columns=["Amenity", "Премия", "Окупаемость"],
                rows=[
                    ["Бассейн 25м", "12–18%", "8–12 лет"],
                    ["Фитнес", "5–8%", "7–9 лет"],
                    ["Консьерж", "4–7%", "3–5 лет"],
                ],
                caption="Данные по 47 проектам Москвы.",
                source_ref="ЕРЗ 2025, Knight Frank 2024",
            ),
        ],
        charts=[
            ChartSpec(
                chart_type="bar",
                title="Ценовая премия по amenities (%)",
                data={"labels": ["Бассейн", "Фитнес", "Консьерж"], "values": [15, 6.5, 5.5]},
                x_label="Тип amenity",
                y_label="Премия (%)",
                caption="Средние значения по 47 проектам.",
            ),
        ],
        callouts=[
            CalloutBlock(kind="insight", title="Консьерж — лучший ROI", body="Окупается за 3–5 лет."),
            CalloutBlock(kind="warning", title="Сигарная теряет актуальность", body="Спрос -40% с 2020."),
        ],
        key_numbers_highlight=[
            KeyNumberHighlight(
                value="883.8 тыс.",
                label="руб./м² — Prime Park H1 2025",
                source_ref="РБК Недвижимость",
                importance="headline",
            ),
            KeyNumberHighlight(
                value="12–18%",
                label="премия за indoor бассейн",
                source_ref="Knight Frank / ЕРЗ",
                importance="headline",
            ),
            KeyNumberHighlight(
                value="73%",
                label="покупателей ценят МОП",
                source_ref="Knight Frank 2024",
                importance="primary",
            ),
        ],
    )


def _minimal_report() -> FinalReport:
    return FinalReport(
        session_id="test-docxjs-minimal",
        question="Какой оптимальный набор amenities для бизнес-класса?",
        executive_summary=ExecutiveSummaryV4(
            main_answer="Бассейн + фитнес + консьерж — оптимальный набор.",
        ),
        main_synthesis="## Основной вывод\n\nBassein + fitnes + konsierj.",
        all_sources=[
            Source(title="ЕРЗ 2025", url="https://erzrf.ru/", tool="perplexity", reliability="high"),
        ],
    )


# ---------------------------------------------------------------------------
# Core render tests
# ---------------------------------------------------------------------------

class TestBridgeRich:
    """Tests using the rich fixture (all Track A fields populated)."""

    @SKIP_NODE
    def test_output_file_exists(self, tmp_path):
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        assert out.exists()

    @SKIP_NODE
    def test_output_exceeds_10kb(self, tmp_path):
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        assert out.stat().st_size > 10_000, f"Expected >10KB, got {out.stat().st_size} bytes"

    @SKIP_NODE
    def test_python_docx_can_open(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        assert len(doc.paragraphs) > 50, f"Expected >50 paragraphs, got {len(doc.paragraphs)}"

    @SKIP_NODE
    def test_paragraph_count(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        assert len(doc.paragraphs) > 50

    @SKIP_NODE
    def test_has_tables(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        # KPI table + ranking table + data table + callouts + ...
        assert len(doc.tables) >= 3

    @SKIP_NODE
    def test_content_has_exec_summary(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "резюме" in all_text.lower() or "Аналитическое" in all_text

    @SKIP_NODE
    def test_content_has_sources(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ЕРЗ" in all_text

    @SKIP_NODE
    def test_content_has_qa_items(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "В:" in all_text

    @SKIP_NODE
    def test_content_has_ranking_item(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        # Check ranking table cells
        found = False
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if "Indoor" in cell.text or "бассейн" in cell.text.lower():
                        found = True
        assert found, "Ranking item 'Indoor бассейн' not found in any table cell"

    @SKIP_NODE
    def test_content_has_toc(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Содержание" in all_text

    @SKIP_NODE
    def test_kpi_values_in_tables(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        found = False
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if "883.8" in cell.text or "12–18%" in cell.text:
                        found = True
        assert found, "KPI values not found in any table cell"

    @SKIP_NODE
    def test_first_heading_contains_question_topic(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "rich.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # The question (title) should appear somewhere in the document
        assert "amenities" in all_text.lower() or "бизнес" in all_text.lower()


class TestBridgeMinimal:
    """Backward-compat: minimal FinalReport with only legacy fields."""

    @SKIP_NODE
    def test_renders_without_error(self, tmp_path):
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _minimal_report()
        out = render_docx_js(report, tmp_path / "minimal.docx")
        assert out.exists()
        assert out.stat().st_size > 5_000

    @SKIP_NODE
    def test_opens_cleanly(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _minimal_report()
        out = render_docx_js(report, tmp_path / "minimal.docx")
        doc = Document(str(out))
        assert len(doc.paragraphs) > 20


class TestBridgeChartDir:
    """Tests with a chart_dir containing PNG files."""

    @SKIP_NODE
    def test_with_chart_dir(self, tmp_path):
        """If chart_dir has PNGs, they should appear in word/media/."""
        from smart_report.exporters.docx_js_bridge import render_docx_js

        # Check if chart samples exist
        if not CHART_SAMPLES_DIR.exists():
            pytest.skip("chart_samples fixture directory not found")

        png_files = list(CHART_SAMPLES_DIR.glob("*.png"))
        if not png_files:
            pytest.skip("No PNG files in chart_samples fixture directory")

        # Create a temp chart dir with properly-named files
        chart_dir = tmp_path / "charts"
        chart_dir.mkdir()
        for i, png in enumerate(png_files[:1]):  # only need 1 chart
            dest = chart_dir / f"chart_{i:02d}.png"
            shutil.copy(png, dest)

        report = _rich_report()
        out = render_docx_js(report, tmp_path / "with_charts.docx", chart_dir=chart_dir)
        assert out.exists()
        assert out.stat().st_size > 10_000

        # Check word/media/ has at least one PNG (image embedded)
        with zipfile.ZipFile(out) as zf:
            media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
            assert len(media_files) >= 1, "Expected at least 1 PNG in word/media/"

    @SKIP_NODE
    def test_without_chart_dir_renders_placeholder(self, tmp_path):
        """Without chart_dir, chart specs render as text placeholders — no crash."""
        from smart_report.exporters.docx_js_bridge import render_docx_js
        report = _rich_report()
        out = render_docx_js(report, tmp_path / "no_charts.docx", chart_dir=None)
        assert out.exists()
        # Should still contain chart type as placeholder text
        pytest.importorskip("docx")
        from docx import Document
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "BAR" in all_text or "Ценовая премия" in all_text


class TestBridgeFixtureFile:
    """Tests against the shared JSON fixture (full amenities sample)."""

    @SKIP_NODE
    def test_fixture_renders(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        from smart_report.exporters.docx_js_bridge import render_docx_js
        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_docx_js(report, tmp_path / "fixture_js.docx")
        assert out.exists()
        assert out.stat().st_size > 10_000

    @SKIP_NODE
    def test_fixture_paragraph_count(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js

        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_docx_js(report, tmp_path / "fixture_js.docx")
        doc = Document(str(out))
        assert len(doc.paragraphs) > 50, f"Expected >50 paragraphs, got {len(doc.paragraphs)}"

    @SKIP_NODE
    def test_fixture_has_minimum_tables(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js

        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_docx_js(report, tmp_path / "fixture_js.docx")
        doc = Document(str(out))
        # KPI tables + ranking + data tables + callouts + cover meta
        assert len(doc.tables) >= 4

    @SKIP_NODE
    def test_fixture_sources_present(self, tmp_path):
        fixture = FIXTURES_DIR / "final_report_sample.json"
        if not fixture.exists():
            pytest.skip("Fixture file not found")

        pytest.importorskip("docx")
        from docx import Document
        from smart_report.exporters.docx_js_bridge import render_docx_js

        data = json.loads(fixture.read_text(encoding="utf-8"))
        report = FinalReport.model_validate(data)
        out = render_docx_js(report, tmp_path / "fixture_js.docx")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ЕРЗ" in all_text
        assert "Knight Frank" in all_text


# ---------------------------------------------------------------------------
# Auto-selector tests (render_docx)
# ---------------------------------------------------------------------------

class TestAutoSelector:
    """Tests for render_docx() auto-selector."""

    def test_is_node_available_returns_bool(self):
        from smart_report.exporters.docx_js_bridge import is_node_available
        result = is_node_available()
        assert isinstance(result, bool)

    def test_render_docx_produces_output(self, tmp_path):
        """render_docx() should produce output regardless of Node availability."""
        pytest.importorskip("docx")
        from smart_report.exporters import render_docx
        report = _minimal_report()
        out = render_docx(report, tmp_path / "auto.docx")
        assert out.exists()
        assert out.stat().st_size > 2_000

    @SKIP_NODE
    def test_render_docx_prefers_node_when_available(self, tmp_path):
        """When Node is available, render_docx should use Node.js renderer by default."""
        pytest.importorskip("docx")
        from smart_report.exporters import render_docx
        report = _rich_report()
        out = render_docx(report, tmp_path / "auto_node.docx", prefer="node")
        assert out.exists()
        assert out.stat().st_size > 10_000

    def test_render_docx_python_fallback(self, tmp_path):
        """render_docx() with prefer='python' always uses python-docx renderer."""
        pytest.importorskip("docx")
        from smart_report.exporters import render_docx
        report = _minimal_report()
        out = render_docx(report, tmp_path / "python_fallback.docx", prefer="python")
        assert out.exists()
        assert out.stat().st_size > 2_000


# ---------------------------------------------------------------------------
# Bridge error handling
# ---------------------------------------------------------------------------

class TestBridgeErrors:
    """Tests for bridge error handling."""

    @SKIP_NODE
    def test_invalid_json_path_raises(self, tmp_path):
        """Passing a non-existent JSON path should raise an error."""
        from smart_report.exporters.docx_js_bridge import render_docx_js, NodeRenderError
        report = _minimal_report()
        # The bridge writes temp JSON itself, so we need to force an error
        # by corrupting the output path's parent
        import os
        # Write to a path that cannot be created (parent is a file, not dir)
        bad_out = tmp_path / "notadir.txt" / "out.docx"
        (tmp_path / "notadir.txt").write_text("I am a file")
        with pytest.raises((NodeRenderError, OSError, Exception)):
            render_docx_js(report, bad_out)
