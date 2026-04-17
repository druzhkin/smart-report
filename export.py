"""Export Report → markdown / docx / json. Executive Summary first; blocks sorted by priority.

DOCX layout (business-grade): cover → TL;DR → Executive Summary → Matrix (table) →
Blocks (with priority callout + conclusion + key number + prose + evidence + risks) →
Cross-links → Scenarios → Assumption inversions → Sources appendix → Methodology.
"""
from __future__ import annotations

import json
import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from models import Block, BlockHeader, Finding, Report

PRIORITY_ORDER = {"high": 0, "medium": 1, "low": 2}
PRIORITY_EMOJI = {"high": "🔴", "medium": "🟡", "low": "🟢"}
PRIORITY_LABEL_RU = {
    "high": "Высокий приоритет",
    "medium": "Средний приоритет",
    "low": "Низкий приоритет",
}

SOURCE_TYPE_RU = {
    "primary_academic": "Академические источники",
    "primary_official": "Официальные источники",
    "primary_data": "Первичные данные",
    "secondary": "Вторичные материалы",
    "opinion": "Экспертные мнения",
}


def _headers_by_cell(report: Report) -> dict[str, BlockHeader]:
    return {h.cell: h for h in report.block_headers}


def _sorted_blocks(report: Report):
    headers = _headers_by_cell(report)

    def _key(block):
        h = headers.get(block.cell)
        prio = PRIORITY_ORDER.get(h.priority if h else "", 3)
        score = (
            -(h.score_novelty + h.score_concreteness + h.score_applicability) if h else 0
        )
        return (prio, score, block.cell)

    return sorted(report.blocks, key=_key)


# ---------- markdown ----------


def to_markdown(report: Report) -> str:
    lines: list[str] = ["# Аналитический отчёт\n", f"**Цель:** {report.goal}\n"]

    es = report.exec_summary
    if es is not None:
        lines.append("\n## Executive Summary\n")
        lines.append(f"**Цель:** {es.goal_restate}\n")
        lines.append("\n### Матрица доменов\n")
        lines.append(es.matrix_table_md.strip() + "\n")
        lines.append("\n### Топ-5 находок\n")
        for f in es.top_findings:
            lines.append(f"- **[{f.block_cell}]** {f.headline}")
        lines.append("\n### Топ-3 кросс-доменных связи\n")
        for c in es.top_connections:
            doms = " ↔ ".join(c.domains) if c.domains else ""
            lines.append(f"- **{doms}** — {c.headline}" if doms else f"- {c.headline}")
        lines.append("\n### Ключевые пробелы\n")
        for g in es.key_gaps:
            lines.append(f"- {g}")
        lines.append("\n---\n")

    lines.append("\n## Матрица доменов (развёрнуто)\n")
    for d in report.matrix.domains:
        lines.append(f"### {d.name}\n")
        lines.append(f"_{d.rationale}_\n")
        for layer in d.layers:
            lines.append(f"- **{layer.name}** — {layer.description}")
        lines.append("")

    lines.append("\n## Блоки (по приоритету)\n")
    headers = _headers_by_cell(report)
    for b in _sorted_blocks(report):
        h = headers.get(b.cell)
        emoji = PRIORITY_EMOJI.get(h.priority, "⚪") if h else "⚪"
        lines.append(f"### {emoji} {b.cell}\n")
        if h is not None:
            lines.append(f"- **Вывод:** {h.one_liner}")
            lines.append(f"- **Сильнейшая цифра:** {h.strongest_number}")
            lines.append(f"- **Главный пробел:** {h.main_gap}")
            lines.append(
                f"- **Оценка:** новизна {h.score_novelty}/3 · конкретность "
                f"{h.score_concreteness}/3 · применимость {h.score_applicability}/3 → "
                f"**{h.priority}**"
            )
            lines.append("")
        lines.append(b.summary.strip() + "\n")
        if b.findings:
            lines.append("**Источники:**\n")
            for f in b.findings:
                num = "📊 " if f.has_numbers else ""
                lines.append(f"- {num}{f.claim} — [{f.source_type}] {f.source}")
            lines.append("")
        if b.assumptions:
            lines.append("**Ключевые допущения (Key Assumptions Check):**\n")
            for a in b.assumptions:
                lines.append(f"- {a}")
            lines.append("")
        if getattr(b, "analogies", None):
            lines.append("**Структурные аналогии:**\n")
            for a in b.analogies:
                loc = f" ({a.location})" if a.location else ""
                lines.append(f"- **{a.situation}**{loc} — ожидалось: {a.expected}; по факту: {a.actual}. _Урок:_ {a.lesson}")
            lines.append("")
        if getattr(b, "indicators", None):
            lines.append("**Индикаторы конкурирующих гипотез:**\n")
            for iw in b.indicators:
                lines.append(f"- _{iw.hypothesis}_ → наблюдать: {iw.indicator} ({iw.timeframe})")
            lines.append("")
        if getattr(b, "decision_point", None):
            lines.append(f"**Точка решения:** {b.decision_point}\n")
        if b.gaps:
            lines.append("**Пробелы / куда копать дальше:**\n")
            for g in b.gaps:
                lines.append(f"- {g}")
            lines.append("")

    if report.connections:
        lines.append("\n## Кросс-доменные связи (бисоциация)\n")
        for c in report.connections:
            lines.append(
                f"### {' ↔ '.join(c.domains)} · {c.nature} · _{c.strength}_\n"
            )
            lines.append(f"**Общая переменная:** {c.shared_entity}\n")
            lines.append(c.description + "\n")
            if c.anchors:
                lines.append("**Опоры в блоках:**\n")
                for a in c.anchors:
                    lines.append(f"- {a}")
                lines.append("")
            if c.novelty:
                lines.append(f"**Что нового:** {c.novelty}\n")

    cone = getattr(report, "scenario_cone", None)
    if cone and cone.scenarios:
        lines.append(f"\n## Конус сценариев · горизонт {cone.question_horizon}\n")
        if cone.key_uncertainties:
            lines.append("**Ключевые неопределённости:**\n")
            for u in cone.key_uncertainties:
                lines.append(f"- {u}")
            lines.append("")
        for s in cone.scenarios:
            lines.append(f"### {s.name} · _{s.probability}_\n")
            lines.append(s.description + "\n")
            lines.append(f"**Драйвер:** {s.key_driver}\n")
            if s.implications:
                lines.append("**Следствия:**\n")
                for it in s.implications:
                    lines.append(f"- {it}")
                lines.append("")
            if s.indicators:
                lines.append("**Индикаторы:**\n")
                for ind in s.indicators:
                    lines.append(f"- {ind}")
                lines.append("")
        if cone.wild_card:
            lines.append(f"### Wild card · _{cone.wild_card.probability}_\n")
            lines.append(cone.wild_card.description + "\n")
            lines.append(f"**Эффект:** {cone.wild_card.impact}\n")
        if cone.conditional_verdict:
            lines.append(f"\n**Условный вердикт:** {cone.conditional_verdict}\n")

    inversions = getattr(report, "assumption_inversions", None) or []
    if inversions:
        lines.append("\n## Проверка допущений (Quadrant Crunching)\n")
        for bi in inversions:
            lines.append(f"### {bi.block_cell}\n")
            if bi.unfalsifiable_flag:
                lines.append("_⚠ Нет критических допущений — вывод может быть нефальсифицируем._\n")
            for inv in bi.inversions:
                dep = f" · **{inv.dependency}**"
                lines.append(
                    f"- _{inv.assumption}_ → **если ложно:** {inv.inversion} "
                    f"(P={inv.probability}{dep}). Следствие: {inv.consequence}. "
                    f"Ранний сигнал: {inv.early_signal}"
                )
            lines.append("")
    return "\n".join(lines)


# ---------- docx ----------

# Business color palette
_CLR_PRIMARY = RGBColor(0x1F, 0x38, 0x64)   # deep blue — title, H1
_CLR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)    # blue — H2
_CLR_ACCENT2 = RGBColor(0x44, 0x62, 0x8F)   # muted blue — H3
_CLR_TEXT = RGBColor(0x2B, 0x2B, 0x2B)      # near-black body
_CLR_SUBTLE = RGBColor(0x59, 0x59, 0x59)    # captions
_CLR_MUTED = RGBColor(0x80, 0x80, 0x80)     # very muted
_CLR_HIGH = RGBColor(0xC0, 0x00, 0x00)
_CLR_MEDIUM = RGBColor(0xC5, 0x5A, 0x11)
_CLR_LOW = RGBColor(0x54, 0x82, 0x35)
_CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

_PRIORITY_COLOR = {"high": _CLR_HIGH, "medium": _CLR_MEDIUM, "low": _CLR_LOW}
_PRIORITY_SHADE = {"high": "F7DADA", "medium": "FCE5CD", "low": "E2EFDA"}

_SHADE_CALLOUT = "F2F2F2"
_SHADE_CONCLUSION = "DEEBF7"
_SHADE_WARN = "FFF2CC"
_SHADE_TABLE_HEAD = "1F3864"
_SHADE_TABLE_ALT = "F2F5FA"


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _shade_paragraph(para, color_hex: str) -> None:
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def _set_para_border(para, color_hex: str = "BFBFBF", side: str = "left", size_pt: int = 24) -> None:
    """Add a colored border to a paragraph — used for left-bar callouts."""
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size_pt))
    el.set(qn("w:space"), "4")
    el.set(qn("w:color"), color_hex)
    borders.append(el)
    p_pr.append(borders)


_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)")


def _add_runs_from_markdown(para, text: str, base_size_pt: int = 11, color: RGBColor | None = None) -> None:
    """Render a single line of prose with **bold**, *italic*, `code` as runs.
    Keeps it simple — no nested emphasis, no links. Good enough for LLM-generated prose."""
    if not text:
        return
    for chunk in _INLINE_RE.split(text):
        if not chunk:
            continue
        run = para.add_run()
        if chunk.startswith("**") and chunk.endswith("**"):
            run.text = chunk[2:-2]
            run.bold = True
        elif chunk.startswith("__") and chunk.endswith("__"):
            run.text = chunk[2:-2]
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run.text = chunk[1:-1]
            run.italic = True
        elif chunk.startswith("_") and chunk.endswith("_"):
            run.text = chunk[1:-1]
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run.text = chunk[1:-1]
            run.font.name = "Consolas"
        else:
            run.text = chunk
        run.font.size = Pt(base_size_pt)
        if color is not None:
            run.font.color.rgb = color


def _add_field(paragraph, field_code: str) -> None:
    """Insert a Word field (e.g. PAGE, TOC) into the given paragraph."""
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ""
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(placeholder)
    run._r.append(fld_char3)


def _add_page_footer(doc, label: str) -> None:
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        left = para.add_run(label[:80] + (" … " if len(label) > 80 else " · "))
        left.font.size = Pt(9)
        left.font.color.rgb = _CLR_SUBTLE
        left.italic = True
        _add_field(para, "PAGE")
        for r in para.runs[-1:]:
            r.font.size = Pt(9)
            r.font.color.rgb = _CLR_SUBTLE
        sep = para.add_run(" / ")
        sep.font.size = Pt(9)
        sep.font.color.rgb = _CLR_SUBTLE
        _add_field(para, "NUMPAGES")


def _add_toc(doc) -> None:
    doc.add_heading("Содержание", level=1)
    intro = doc.add_paragraph()
    r = intro.add_run("После открытия документа нажмите F9 или правой кнопкой → «Обновить поле», чтобы заполнить оглавление.")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = _CLR_MUTED
    intro.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    _add_field(p, r'TOC \o "1-2" \h \z \u')
    doc.add_page_break()


def _set_doc_defaults(doc) -> None:
    """Set document-wide font, margins, line spacing."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = _CLR_TEXT
    rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
    if rFonts is None:
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:cs"), "Calibri")

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _style_heading(doc, level: int) -> None:
    """Adjust built-in Heading styles once per document."""
    if level == 0:
        s = doc.styles["Title"]
        s.font.name = "Calibri"
        s.font.size = Pt(28)
        s.font.bold = True
        s.font.color.rgb = _CLR_PRIMARY
    else:
        name = f"Heading {level}"
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = "Calibri"
            s.font.bold = True
            if level == 1:
                s.font.size = Pt(20)
                s.font.color.rgb = _CLR_PRIMARY
            elif level == 2:
                s.font.size = Pt(15)
                s.font.color.rgb = _CLR_ACCENT
            elif level == 3:
                s.font.size = Pt(12)
                s.font.color.rgb = _CLR_ACCENT2


def _add_kv_paragraph(doc, label: str, value: str, *, label_color: RGBColor | None = None) -> None:
    """Bold label + normal value in one paragraph."""
    p = doc.add_paragraph()
    r_label = p.add_run(f"{label}: ")
    r_label.bold = True
    r_label.font.size = Pt(11)
    if label_color is not None:
        r_label.font.color.rgb = label_color
    _add_runs_from_markdown(p, value, base_size_pt=11)


def _add_callout(doc, text: str, *, shade: str = _SHADE_CALLOUT, border_color: str = "BFBFBF", label: str | None = None, label_color: RGBColor | None = None) -> None:
    """Shaded paragraph with optional bold label prefix and left-border accent."""
    p = doc.add_paragraph()
    _shade_paragraph(p, shade)
    _set_para_border(p, color_hex=border_color, side="left", size_pt=24)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.3)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    if label:
        lr = p.add_run(f"{label} ")
        lr.bold = True
        lr.font.size = Pt(11)
        if label_color is not None:
            lr.font.color.rgb = label_color
    _add_runs_from_markdown(p, text, base_size_pt=11)


def _add_bullet(doc, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(0.6 * (level + 1))
    _add_runs_from_markdown(p, text, base_size_pt=11)


def _hex_from_rgb(clr: RGBColor) -> str:
    return f"{clr[0]:02X}{clr[1]:02X}{clr[2]:02X}"


def _style_table_header_row(row) -> None:
    for cell in row.cells:
        _set_cell_shading(cell, _SHADE_TABLE_HEAD)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            for r in para.runs:
                r.font.bold = True
                r.font.color.rgb = _CLR_WHITE
                r.font.size = Pt(10.5)


def _style_table_body_cell(cell, *, zebra: bool = False) -> None:
    if zebra:
        _set_cell_shading(cell, _SHADE_TABLE_ALT)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(2)
        for r in para.runs:
            r.font.size = Pt(10.5)


def _stats(report: Report) -> dict:
    n_findings = sum(len(b.findings) for b in report.blocks)
    n_numeric = sum(1 for b in report.blocks for f in b.findings if f.has_numbers)
    n_primary = sum(
        1 for b in report.blocks for f in b.findings
        if f.source_type in ("primary_academic", "primary_official", "primary_data")
    )
    domains = set()
    for d in report.matrix.domains:
        domains.add(d.name)
    return {
        "domains": len(domains),
        "blocks": len(report.blocks),
        "findings": n_findings,
        "with_numbers": n_numeric,
        "primary": n_primary,
        "connections": len(report.connections),
    }


def _synth_tldr(report: Report, limit: int = 5) -> list[str]:
    """Extract 3-5 plain-language insights. Fallback chain: top_findings → high/medium one_liners
    → any non-empty one_liner → top_connections → key_gaps. Degenerate reports still get a TL;DR."""
    out: list[str] = []
    seen: set[str] = set()

    def _push(text: str) -> bool:
        t = (text or "").strip()
        if not t or t in seen:
            return False
        seen.add(t)
        out.append(t)
        return len(out) >= limit

    if report.exec_summary and report.exec_summary.top_findings:
        for tf in report.exec_summary.top_findings[:limit]:
            if _push(tf.headline):
                return out

    headers = _headers_by_cell(report)
    for priority_tier in (("high", "medium"), ("low",)):
        if len(out) >= limit:
            break
        for b in _sorted_blocks(report):
            h = headers.get(b.cell)
            if h and h.priority in priority_tier and h.one_liner:
                if _push(h.one_liner):
                    return out

    if report.exec_summary:
        for c in report.exec_summary.top_connections or []:
            if _push(c.headline):
                return out
        for g in report.exec_summary.key_gaps or []:
            if _push(g):
                return out

    return out[:limit]


_DEGENERATE_NUM_RE = re.compile(r"^\s*0\s+[^\d]", flags=re.UNICODE)


def _is_meaningful_number(text: str) -> bool:
    """Filter out degenerate LLM outputs like '0 исследований…' that render as if they were numbers.
    Require: non-empty, contains at least one digit, not of pattern '0 <non-digit…>'."""
    t = (text or "").strip()
    if not t or len(t) < 2:
        return False
    if not any(ch.isdigit() for ch in t):
        return False
    if _DEGENERATE_NUM_RE.match(t):
        return False
    return True


def _collect_key_numbers(report: Report, limit: int = 5) -> list[tuple[str, str]]:
    """Return list of (number_phrase, source_label) from block headers, skipping empties and degenerate phrases."""
    out: list[tuple[str, str]] = []
    headers = _headers_by_cell(report)
    for b in _sorted_blocks(report):
        h = headers.get(b.cell)
        if not h or not _is_meaningful_number(h.strongest_number):
            continue
        label = b.cell.split(" / ")[0]
        out.append((h.strongest_number, label))
        if len(out) >= limit:
            break
    return out


def _collect_decisions(report: Report, limit: int = 5) -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    for b in _sorted_blocks(report):
        dp = getattr(b, "decision_point", None)
        if dp:
            out.append((dp, b.cell))
            if len(out) >= limit:
                break
    return out


def _collect_watchlist(report: Report, limit: int = 8) -> list[str]:
    """Indicators + key uncertainties — what to monitor."""
    out: list[str] = []
    cone = getattr(report, "scenario_cone", None)
    if cone and cone.key_uncertainties:
        out.extend(cone.key_uncertainties[:3])
    for b in _sorted_blocks(report):
        for iw in getattr(b, "indicators", []) or []:
            out.append(f"{iw.indicator} ({iw.timeframe})")
            if len(out) >= limit:
                return out[:limit]
    return out[:limit]


def _add_cover(doc, report: Report) -> None:
    stats = _stats(report)
    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.paragraph_format.space_before = Pt(48)
    r = t.add_run("Аналитический отчёт")
    r.font.name = "Calibri"
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = _CLR_PRIMARY

    # Goal subheading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Цель исследования")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = _CLR_SUBTLE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(report.goal)
    r.font.size = Pt(14)
    r.font.color.rgb = _CLR_TEXT

    # Stats panel
    p = doc.add_paragraph()
    _shade_paragraph(p, _SHADE_CONCLUSION)
    _set_para_border(p, color_hex=_hex_from_rgb(_CLR_ACCENT), side="left", size_pt=24)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    line = (
        f"{stats['domains']} доменов · {stats['blocks']} тематических блоков · "
        f"{stats['findings']} находок (из них {stats['with_numbers']} с числами, "
        f"{stats['primary']} из первичных источников) · "
        f"{stats['connections']} кросс-доменных связей"
    )
    r = p.add_run(line)
    r.font.size = Pt(11)

    # Date footer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run(f"Дата: {date.today().isoformat()}")
    r.font.size = Pt(10)
    r.font.color.rgb = _CLR_SUBTLE

    p = doc.add_paragraph()
    r = p.add_run("Smart Report — многоагентный аналитический конвейер")
    r.font.size = Pt(10)
    r.font.color.rgb = _CLR_MUTED
    r.italic = True

    doc.add_page_break()


def _add_tldr(doc, report: Report) -> None:
    tldr = _synth_tldr(report)
    if not tldr:
        return
    doc.add_heading("Коротко — что важно знать", level=1)

    intro = doc.add_paragraph()
    r = intro.add_run(
        "Главные выводы отчёта в одном абзаце: что обнаружено, какие числа держат вывод, "
        "и на что это влияет. Детали и доказательства — в следующих разделах."
    )
    r.italic = True
    r.font.color.rgb = _CLR_SUBTLE
    r.font.size = Pt(10.5)
    intro.paragraph_format.space_after = Pt(10)

    for i, headline in enumerate(tldr, 1):
        p = doc.add_paragraph()
        _shade_paragraph(p, _SHADE_CALLOUT)
        _set_para_border(p, color_hex=_hex_from_rgb(_CLR_ACCENT), side="left", size_pt=18)
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        num = p.add_run(f"{i}. ")
        num.bold = True
        num.font.color.rgb = _CLR_PRIMARY
        num.font.size = Pt(11)
        _add_runs_from_markdown(p, headline, base_size_pt=11)


def _add_exec_summary(doc, report: Report) -> None:
    es = report.exec_summary
    if es is None:
        return
    doc.add_heading("Executive Summary", level=1)

    _add_callout(
        doc,
        es.goal_restate,
        shade=_SHADE_CONCLUSION,
        border_color=_hex_from_rgb(_CLR_PRIMARY),
        label="Наш вопрос:",
        label_color=_CLR_PRIMARY,
    )

    # Key numbers panel
    nums = _collect_key_numbers(report, limit=5)
    if nums:
        doc.add_heading("Ключевые цифры", level=2)
        for number, label in nums:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{number} ")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _CLR_PRIMARY
            r2 = p.add_run(f"— {label}")
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = _CLR_SUBTLE

    # Top findings
    if es.top_findings:
        doc.add_heading("Главные находки", level=2)
        for f in es.top_findings:
            p = doc.add_paragraph(style="List Bullet")
            tag = p.add_run(f"[{f.block_cell}] ")
            tag.font.size = Pt(9.5)
            tag.font.color.rgb = _CLR_SUBTLE
            _add_runs_from_markdown(p, f.headline, base_size_pt=11)

    # Top connections
    if es.top_connections:
        doc.add_heading("Связи между доменами", level=2)
        for c in es.top_connections:
            p = doc.add_paragraph(style="List Bullet")
            if c.domains:
                tag = p.add_run(" ↔ ".join(c.domains) + " — ")
                tag.bold = True
                tag.font.size = Pt(10.5)
                tag.font.color.rgb = _CLR_ACCENT
            _add_runs_from_markdown(p, c.headline, base_size_pt=11)

    # Decisions
    decisions = _collect_decisions(report, limit=5)
    if decisions:
        doc.add_heading("Что решать сейчас", level=2)
        for dp, cell in decisions:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_from_markdown(p, dp, base_size_pt=11)
            tag = p.add_run(f"  ({cell})")
            tag.font.size = Pt(9.5)
            tag.font.color.rgb = _CLR_MUTED
            tag.italic = True

    # Watchlist
    watch = _collect_watchlist(report, limit=6)
    if watch:
        doc.add_heading("За чем следить в ближайшие 6–12 месяцев", level=2)
        for w in watch:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_from_markdown(p, w, base_size_pt=11)

    # Gaps
    if es.key_gaps:
        doc.add_heading("Что осталось непонятным", level=2)
        intro = doc.add_paragraph()
        r = intro.add_run(
            "Пробелы в данных, которые стоит закрыть до принятия решения на основе отчёта."
        )
        r.italic = True
        r.font.color.rgb = _CLR_SUBTLE
        r.font.size = Pt(10.5)
        for g in es.key_gaps:
            _add_bullet(doc, g)

    doc.add_page_break()


def _add_matrix_table(doc, report: Report) -> None:
    doc.add_heading("Матрица доменов", level=1)

    intro = doc.add_paragraph()
    r = intro.add_run(
        "Вопрос исследования разложен на домены (столбцы-темы) и слои внутри каждого домена "
        "(подвопросы, под которые собирались данные). Ниже — структура, на которой держится остальной отчёт."
    )
    r.italic = True
    r.font.color.rgb = _CLR_SUBTLE
    r.font.size = Pt(10.5)
    intro.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Cm(4.5), Cm(5.0), Cm(7.5)]
    hdr = table.rows[0].cells
    hdr[0].text = "Домен"
    hdr[1].text = "Слой"
    hdr[2].text = "Что ищем"
    _style_table_header_row(table.rows[0])
    for i, w in enumerate(widths):
        hdr[i].width = w

    zebra = False
    for d in report.matrix.domains:
        for j, layer in enumerate(d.layers):
            row = table.add_row()
            cells = row.cells
            cells[0].text = d.name if j == 0 else ""
            cells[1].text = layer.name
            cells[2].text = layer.description
            for i, w in enumerate(widths):
                cells[i].width = w
            for c in cells:
                _style_table_body_cell(c, zebra=zebra)
            # Make domain column bold
            for para in cells[0].paragraphs:
                for r in para.runs:
                    r.font.bold = True
                    r.font.color.rgb = _CLR_PRIMARY
        zebra = not zebra

    # Rationale section under the table
    doc.add_heading("Почему именно эти домены", level=2)
    for d in report.matrix.domains:
        p = doc.add_paragraph()
        r = p.add_run(f"{d.name}. ")
        r.bold = True
        r.font.color.rgb = _CLR_PRIMARY
        r.font.size = Pt(11)
        _add_runs_from_markdown(p, d.rationale, base_size_pt=11)

    doc.add_page_break()


def _add_block(doc, block: Block, header: BlockHeader | None) -> None:
    prio = header.priority if header else "low"
    prio_color = _PRIORITY_COLOR.get(prio, _CLR_SUBTLE)

    # Priority label + cell as H2
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    badge = p.add_run(f"  {PRIORITY_LABEL_RU.get(prio, prio).upper()}  ")
    badge.bold = True
    badge.font.size = Pt(9)
    badge.font.color.rgb = _CLR_WHITE
    # crude badge look via shading on the paragraph — but it would shade full line.
    # Instead, show priority label as coloured inline text + dot marker.
    badge.font.color.rgb = prio_color
    _shade_paragraph(p, _PRIORITY_SHADE.get(prio, "F2F2F2"))
    _set_para_border(p, color_hex=_hex_from_rgb(prio_color), side="left", size_pt=30)
    p.paragraph_format.left_indent = Cm(0.3)

    h = doc.add_heading(block.cell, level=2)
    h.paragraph_format.space_before = Pt(2)

    # Conclusion callout — the single most important line per block
    if header and header.one_liner:
        _add_callout(
            doc,
            header.one_liner,
            shade=_SHADE_CONCLUSION,
            border_color=_hex_from_rgb(_CLR_PRIMARY),
            label="Главный вывод:",
            label_color=_CLR_PRIMARY,
        )

    # Strongest number — skip degenerate "0 исследований" style phrases
    if header and _is_meaningful_number(header.strongest_number):
        _add_callout(
            doc,
            header.strongest_number,
            shade=_SHADE_CALLOUT,
            border_color=_hex_from_rgb(_CLR_ACCENT),
            label="Ключевая цифра:",
            label_color=_CLR_ACCENT,
        )

    # Main gap
    if header and header.main_gap:
        _add_kv_paragraph(doc, "Главный пробел", header.main_gap, label_color=_CLR_SUBTLE)

    # Confidence/priority scores — single line
    if header:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(
            f"Оценка: новизна {header.score_novelty}/3 · конкретность {header.score_concreteness}/3 · "
            f"применимость {header.score_applicability}/3 → "
        )
        r.font.size = Pt(9.5)
        r.font.color.rgb = _CLR_SUBTLE
        r2 = p.add_run(PRIORITY_LABEL_RU.get(prio, prio))
        r2.font.size = Pt(9.5)
        r2.bold = True
        r2.font.color.rgb = prio_color

    # Prose summary — paragraph per blank-line chunk, inline markdown parsed
    for para in block.summary.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_runs_from_markdown(p, para, base_size_pt=11)

    # Decision point
    if getattr(block, "decision_point", None):
        _add_callout(
            doc,
            block.decision_point,
            shade=_SHADE_WARN,
            border_color="BF8F00",
            label="Точка решения:",
            label_color=RGBColor(0x9C, 0x65, 0x00),
        )

    # Evidence
    doc.add_heading("Доказательная база", level=3)
    if not block.findings:
        _add_callout(
            doc,
            "По этому срезу данных собрать не удалось — см. «Пробелы» ниже, там расписано, куда копать дальше.",
            shade=_SHADE_WARN,
            border_color="BF8F00",
            label="Данных в блоке не найдено.",
            label_color=RGBColor(0x9C, 0x65, 0x00),
        )
    if block.findings:
        for f in block.findings:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            _add_runs_from_markdown(p, f.claim, base_size_pt=11)
            # Source line below the claim
            sub = doc.add_paragraph()
            sub.paragraph_format.left_indent = Cm(0.75)
            sub.paragraph_format.space_after = Pt(4)
            label_parts = []
            if f.source_label:
                label_parts.append(f.source_label)
            elif f.source_type:
                label_parts.append(SOURCE_TYPE_RU.get(f.source_type, f.source_type))
            if f.year:
                label_parts.append(str(f.year))
            label = " · ".join(label_parts)
            r_lbl = sub.add_run(label) if label else sub.add_run("")
            r_lbl.font.size = Pt(9.5)
            r_lbl.font.color.rgb = _CLR_SUBTLE
            r_lbl.italic = True
            if f.source:
                sep = sub.add_run(" — ") if label else sub.add_run("")
                sep.font.size = Pt(9.5)
                sep.font.color.rgb = _CLR_SUBTLE
                r_src = sub.add_run(f.source)
                r_src.font.size = Pt(9.5)
                r_src.font.color.rgb = _CLR_ACCENT
            if f.verbatim_quote:
                q = doc.add_paragraph()
                q.paragraph_format.left_indent = Cm(0.75)
                q.paragraph_format.space_after = Pt(4)
                r_q = q.add_run(f"«{f.verbatim_quote}»")
                r_q.italic = True
                r_q.font.size = Pt(10)
                r_q.font.color.rgb = _CLR_TEXT

    # Assumptions
    if block.assumptions:
        doc.add_heading("Ключевые допущения", level=3)
        intro = doc.add_paragraph()
        r = intro.add_run("Если хотя бы одно из этих утверждений неверно — вывод блока надо пересматривать.")
        r.italic = True
        r.font.color.rgb = _CLR_SUBTLE
        r.font.size = Pt(10)
        for a in block.assumptions:
            _add_bullet(doc, a)

    # Analogies
    if getattr(block, "analogies", None):
        doc.add_heading("Структурные аналогии", level=3)
        for a in block.analogies:
            loc = f" · {a.location}" if a.location else ""
            p = doc.add_paragraph()
            r = p.add_run(f"{a.situation}{loc}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _CLR_ACCENT
            _add_kv_paragraph(doc, "Ожидалось", a.expected)
            _add_kv_paragraph(doc, "По факту", a.actual)
            if getattr(a, "why_diverged", ""):
                _add_kv_paragraph(doc, "Почему разошлось", a.why_diverged)
            _add_kv_paragraph(doc, "Урок", a.lesson, label_color=_CLR_PRIMARY)

    # Indicators
    if getattr(block, "indicators", None):
        doc.add_heading("Индикаторы конкурирующих гипотез", level=3)
        for iw in block.indicators:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{iw.hypothesis}. ")
            r.italic = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = _CLR_SUBTLE
            r2 = p.add_run(f"Наблюдать: {iw.indicator}")
            r2.font.size = Pt(11)
            if iw.timeframe:
                r3 = p.add_run(f" ({iw.timeframe})")
                r3.font.size = Pt(10)
                r3.font.color.rgb = _CLR_MUTED

    # Gaps
    if block.gaps:
        doc.add_heading("Пробелы и что исследовать дальше", level=3)
        for g in block.gaps:
            _add_bullet(doc, g)


def _add_connections(doc, report: Report) -> None:
    if not report.connections:
        return
    doc.add_page_break()
    doc.add_heading("Кросс-доменные связи: подробный разбор", level=1)
    intro = doc.add_paragraph()
    r = intro.add_run(
        "Кросс-доменные связи показывают, где выводы одного блока подкрепляют или противоречат другому. "
        "Это главный источник неочевидных инсайтов — то, что видно только при сопоставлении тем."
    )
    r.italic = True
    r.font.color.rgb = _CLR_SUBTLE
    r.font.size = Pt(10.5)
    intro.paragraph_format.space_after = Pt(10)

    for c in report.connections:
        doc.add_heading(" ↔ ".join(c.domains), level=2)
        p = doc.add_paragraph()
        r = p.add_run(f"Тип связи: {c.nature} · Сила: {c.strength}")
        r.font.size = Pt(10)
        r.font.color.rgb = _CLR_SUBTLE
        r.italic = True

        _add_kv_paragraph(doc, "Общая переменная", c.shared_entity, label_color=_CLR_ACCENT)
        p = doc.add_paragraph()
        _add_runs_from_markdown(p, c.description, base_size_pt=11)

        if c.anchors:
            doc.add_heading("Опоры в блоках", level=3)
            for a in c.anchors:
                _add_bullet(doc, a)

        if c.novelty:
            _add_callout(
                doc,
                c.novelty,
                shade=_SHADE_CONCLUSION,
                border_color=_hex_from_rgb(_CLR_PRIMARY),
                label="Что нового даёт связь:",
                label_color=_CLR_PRIMARY,
            )


def _add_scenarios(doc, report: Report) -> None:
    cone = getattr(report, "scenario_cone", None)
    if not cone or not cone.scenarios:
        return
    doc.add_page_break()
    doc.add_heading(f"Сценарии развития · горизонт {cone.question_horizon}", level=1)
    intro = doc.add_paragraph()
    r = intro.add_run(
        "Три базовых сценария + wild card. Читать так: базовый — если всё идёт как сейчас; оптимистичный/"
        "пессимистичный — если ключевой драйвер уходит в крайность; wild card — редкое событие с большим эффектом."
    )
    r.italic = True
    r.font.color.rgb = _CLR_SUBTLE
    r.font.size = Pt(10.5)
    intro.paragraph_format.space_after = Pt(10)

    if cone.key_uncertainties:
        doc.add_heading("Ключевые неопределённости", level=2)
        for u in cone.key_uncertainties:
            _add_bullet(doc, u)

    for s in cone.scenarios:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        r = p.add_run(s.name)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = _CLR_PRIMARY
        r_prob = p.add_run(f"  · вероятность {s.probability}")
        r_prob.font.size = Pt(11)
        r_prob.italic = True
        r_prob.font.color.rgb = _CLR_SUBTLE

        desc = doc.add_paragraph()
        _add_runs_from_markdown(desc, s.description, base_size_pt=11)

        _add_kv_paragraph(doc, "Главный драйвер", s.key_driver, label_color=_CLR_ACCENT)

        if s.implications:
            doc.add_heading("Что это значит на практике", level=3)
            for it in s.implications:
                _add_bullet(doc, it)
        if s.indicators:
            doc.add_heading("Сигналы, что сценарий реализуется", level=3)
            for ind in s.indicators:
                _add_bullet(doc, ind)

    if cone.wild_card:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        r = p.add_run("Wild card")
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = _CLR_HIGH
        r_prob = p.add_run(f"  · {cone.wild_card.probability}")
        r_prob.font.size = Pt(11)
        r_prob.italic = True
        r_prob.font.color.rgb = _CLR_SUBTLE
        desc = doc.add_paragraph()
        _add_runs_from_markdown(desc, cone.wild_card.description, base_size_pt=11)
        _add_kv_paragraph(doc, "Эффект", cone.wild_card.impact, label_color=_CLR_HIGH)

    if cone.conditional_verdict:
        _add_callout(
            doc,
            cone.conditional_verdict,
            shade=_SHADE_CONCLUSION,
            border_color=_hex_from_rgb(_CLR_PRIMARY),
            label="Условный вердикт:",
            label_color=_CLR_PRIMARY,
        )


def _add_inversions(doc, report: Report) -> None:
    inversions = getattr(report, "assumption_inversions", None) or []
    if not inversions:
        return
    doc.add_page_break()
    doc.add_heading("Проверка допущений: что если мы ошибаемся", level=1)
    intro = doc.add_paragraph()
    r = intro.add_run(
        "Для каждого ключевого допущения — что будет, если оно окажется ложным. "
        "Если ни одно не помечено critical, вывод блока нефальсифицируем и ему нельзя безоговорочно доверять."
    )
    r.italic = True
    r.font.color.rgb = _CLR_SUBTLE
    r.font.size = Pt(10.5)
    intro.paragraph_format.space_after = Pt(10)

    for bi in inversions:
        doc.add_heading(bi.block_cell, level=2)
        if bi.unfalsifiable_flag:
            _add_callout(
                doc,
                "Ни одно допущение не критично — вывод блока нельзя надёжно опровергнуть.",
                shade=_SHADE_WARN,
                border_color="BF8F00",
                label="Внимание:",
                label_color=RGBColor(0x9C, 0x65, 0x00),
            )
        for inv in bi.inversions:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(f"Допущение: ")
            r.bold = True
            r.font.color.rgb = _CLR_ACCENT
            r.font.size = Pt(10.5)
            _add_runs_from_markdown(p, inv.assumption, base_size_pt=10.5)

            _add_kv_paragraph(doc, "Если ложно", inv.inversion, label_color=_CLR_HIGH)
            _add_kv_paragraph(doc, "Следствие для вывода", inv.consequence)
            p = doc.add_paragraph()
            r = p.add_run(f"Вероятность: {inv.probability} · Критичность: ")
            r.font.size = Pt(10)
            r.font.color.rgb = _CLR_SUBTLE
            r2 = p.add_run(inv.dependency)
            r2.font.size = Pt(10)
            r2.bold = True
            crit_color = _CLR_HIGH if inv.dependency == "critical" else (
                _CLR_MEDIUM if inv.dependency == "important" else _CLR_LOW
            )
            r2.font.color.rgb = crit_color
            _add_kv_paragraph(doc, "Ранний сигнал", inv.early_signal, label_color=_CLR_ACCENT)


def _add_sources_appendix(doc, report: Report) -> None:
    # Collect unique findings by source URL
    seen: set[str] = set()
    by_type: dict[str, list[Finding]] = defaultdict(list)
    for b in report.blocks:
        for f in b.findings:
            key = (f.source or f.claim)[:200]
            if key in seen:
                continue
            seen.add(key)
            by_type[f.source_type or "secondary"].append(f)
    if not by_type:
        return
    doc.add_page_break()
    doc.add_heading("Источники", level=1)
    intro = doc.add_paragraph()
    r = intro.add_run(
        "Все источники, использованные в отчёте, сгруппированы по типу. Академические и официальные — "
        "первичные данные; вторичные и экспертные — для контекста, но требуют перепроверки."
    )
    r.italic = True
    r.font.color.rgb = _CLR_SUBTLE
    r.font.size = Pt(10.5)
    intro.paragraph_format.space_after = Pt(10)

    type_order = [
        "primary_academic",
        "primary_official",
        "primary_data",
        "secondary",
        "opinion",
    ]
    for t in type_order:
        items = by_type.get(t) or []
        if not items:
            continue
        doc.add_heading(SOURCE_TYPE_RU.get(t, t), level=2)
        for f in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            label_parts = []
            if f.source_label:
                label_parts.append(f.source_label)
            if f.year:
                label_parts.append(str(f.year))
            label = " · ".join(label_parts)
            if label:
                r_lbl = p.add_run(f"{label} — ")
                r_lbl.bold = True
                r_lbl.font.size = Pt(10.5)
            if f.source:
                r_src = p.add_run(f.source)
                r_src.font.size = Pt(10.5)
                r_src.font.color.rgb = _CLR_ACCENT


def _add_methodology(doc, report: Report) -> None:
    stats = _stats(report)
    doc.add_page_break()
    doc.add_heading("Методология", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Отчёт собран многоагентной системой. Планировщик разложил цель на домены и слои; "
        "скауты собирали evidence по каждому слою через web-поиск и академические API; "
        "аналитик собрал блоки с проверкой допущений и кросс-доменными связями; критик оценил качество и приоритет. "
        "Каждая находка связана с первоисточником; числа извлечены из текста и подтверждены цитатами там, где это возможно."
    )
    r.font.size = Pt(10.5)
    r.font.color.rgb = _CLR_TEXT

    doc.add_heading("Ограничения", level=2)
    for line in [
        "LLM-синтез может упускать нюансы, отсутствующие в найденных источниках.",
        "Приоритизация блоков автоматическая — по оценкам новизны, конкретности, применимости.",
        "Сценарии и допущения — вероятностные оценки, а не прогнозы.",
        "Числа проверены на наличие, но не на репрезентативность выборки.",
    ]:
        _add_bullet(doc, line)

    doc.add_heading("Охват", level=2)
    cov = (
        f"{stats['domains']} доменов, {stats['blocks']} блоков, {stats['findings']} находок, "
        f"из них {stats['with_numbers']} с числовыми данными и {stats['primary']} из первичных источников. "
        f"Построено {stats['connections']} кросс-доменных связей."
    )
    p = doc.add_paragraph()
    r = p.add_run(cov)
    r.font.size = Pt(10.5)


def to_docx(report: Report, path: Path) -> None:
    doc = Document()
    _set_doc_defaults(doc)
    for lvl in (0, 1, 2, 3):
        _style_heading(doc, lvl)

    _add_page_footer(doc, report.goal)
    _add_cover(doc, report)
    _add_toc(doc)
    _add_tldr(doc, report)
    _add_exec_summary(doc, report)
    _add_matrix_table(doc, report)

    doc.add_heading("Анализ по блокам", level=1)
    intro = doc.add_paragraph()
    r = intro.add_run(
        "Блоки отсортированы по приоритету (высокий → низкий). В каждом блоке — главный вывод, "
        "ключевая цифра, прозовый разбор, доказательная база, допущения и пробелы."
    )
    r.italic = True
    r.font.color.rgb = _CLR_SUBTLE
    r.font.size = Pt(10.5)
    intro.paragraph_format.space_after = Pt(10)

    headers = _headers_by_cell(report)
    for b in _sorted_blocks(report):
        _add_block(doc, b, headers.get(b.cell))

    _add_connections(doc, report)
    _add_scenarios(doc, report)
    _add_inversions(doc, report)
    _add_sources_appendix(doc, report)
    _add_methodology(doc, report)

    doc.save(str(path))


# ---------- json + save_all ----------


def to_json(report: Report) -> str:
    return json.dumps(report.model_dump(), ensure_ascii=False, indent=2)


def save_all(report: Report, out_dir: Path, stem: str = "report") -> dict[str, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = {
        "md": out_dir / f"{stem}.md",
        "json": out_dir / f"{stem}.json",
        "docx": out_dir / f"{stem}.docx",
    }
    paths["md"].write_text(to_markdown(report), encoding="utf-8")
    paths["json"].write_text(to_json(report), encoding="utf-8")
    to_docx(report, paths["docx"])
    return paths
