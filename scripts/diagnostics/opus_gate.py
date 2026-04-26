"""Opus gate-keeper for milestone verdicts (two-week brief §2).

Reads a milestone artefact bundle (DOCX text-extracted, audit_summary.json,
trace.jsonl head, M<n>_OPUS_REQUEST.md) and asks Opus 4 for PASS/PARTIAL/FAIL
verdict on the brief's numeric acceptance criteria.

Saves response to milestones/M<n>_VERDICT.md verbatim. The agent then PARSES
the verdict line and gates the next milestone.

Usage:
    python -m scripts.diagnostics.opus_gate --milestone 1 \
        --request milestones/M1_OPUS_REQUEST.md \
        --artefact-dir docs/run2_baseline/q1_ev/

Cost: ~$0.50 per gate call (Opus 4 with ~10k input + ~2k output).
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path

if sys.platform == "win32":
    os.environ.setdefault("PYTHONIOENCODING", "utf-8")
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

from dotenv import load_dotenv

REPO_ROOT = Path(__file__).parent.parent.parent
load_dotenv(dotenv_path=REPO_ROOT / ".env")

import httpx

OPUS_MODEL = "anthropic/claude-opus-4"
USD_RUB_RATE = 75.4

SYSTEM_PROMPT = (
    "You are the gate-keeper for Smart Report milestones. Read artefacts honestly. "
    "PASS only if all numeric criteria met. FAIL if any critical criterion missed. "
    "PARTIAL if some met some not.\n\n"
    "Format your response EXACTLY like this (these section headers are required):\n\n"
    "## Numeric criteria check\n"
    "1. <criterion 1>: yes/no — <evidence>\n"
    "2. <criterion 2>: yes/no — <evidence>\n"
    "...\n\n"
    "## Substance grade (1-10)\n"
    "<single integer> — <one-line rationale>\n\n"
    "## Missing sources (if any)\n"
    "- <bullet list of expected-but-absent sources, or 'None noted'>\n\n"
    "## VERDICT\n"
    "<PASS | PARTIAL | FAIL>\n\n"
    "## Rationale (3-5 lines)\n"
    "<concise rationale>\n\n"
    "## Fix-before-next-milestone (only if PARTIAL or FAIL)\n"
    "- <actionable item 1>\n"
    "- <actionable item 2>\n"
)


def _extract_docx_text(docx_path: Path) -> str:
    """Pull text content from a .docx file via python-docx."""
    from docx import Document  # local import — keeps script importable without optional dep

    doc = Document(str(docx_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip("|").strip():
                parts.append(row_text)
    return "\n".join(parts)


def _build_user_message(
    request_md: str,
    docx_text: str,
    audit_summary: dict,
    trace_head: list[dict],
) -> str:
    return (
        f"# Milestone verdict request\n\n{request_md}\n\n"
        f"---\n\n"
        f"## DOCX contents (text-extracted)\n\n"
        f"```\n{docx_text[:80000]}\n```\n\n"  # cap at ~20k tokens
        f"---\n\n"
        f"## audit_summary.json\n\n"
        f"```json\n{json.dumps(audit_summary, ensure_ascii=False, indent=2)[:30000]}\n```\n\n"
        f"---\n\n"
        f"## trace.jsonl (first {len(trace_head)} events)\n\n"
        f"```\n"
        + "\n".join(json.dumps(e, ensure_ascii=False) for e in trace_head)
        + "\n```\n"
    )


def call_opus(system: str, user: str) -> tuple[str, float, dict]:
    """Single Opus call via OpenRouter. Returns (text, cost_rub, usage_dict)."""
    key = os.environ.get("OPENROUTER_API_KEY")
    if not key:
        raise RuntimeError("OPENROUTER_API_KEY missing in .env")
    payload = {
        "model": OPUS_MODEL,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.1,
        "max_tokens": 4000,
    }
    r = httpx.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {key}",
            "Content-Type": "application/json",
        },
        json=payload,
        timeout=180.0,
    )
    r.raise_for_status()
    data = r.json()
    text = data["choices"][0]["message"]["content"]
    usage = data.get("usage", {}) or {}
    cost_usd = usage.get("cost", 0.0) or 0.0
    cost_rub = round(cost_usd * USD_RUB_RATE, 4)
    return text, cost_rub, usage


def parse_verdict(opus_text: str) -> str:
    """Pull PASS / PARTIAL / FAIL out of the structured response.

    Looks for the line under '## VERDICT' header. Returns 'UNKNOWN' if
    the response didn't follow format (unusual for Opus at temp=0.1).
    """
    lines = opus_text.splitlines()
    for i, line in enumerate(lines):
        if line.strip().lower().startswith("## verdict"):
            for j in range(i + 1, min(i + 5, len(lines))):
                stripped = lines[j].strip().upper()
                for v in ("PASS", "PARTIAL", "FAIL"):
                    if stripped == v or stripped.startswith(v + " ") or stripped.startswith(v + "."):
                        return v
    return "UNKNOWN"


def run_gate(
    milestone: int,
    request_path: Path,
    artefact_dir: Path,
    out_path: Path,
) -> int:
    if not request_path.exists():
        print(f"ERROR: request file missing: {request_path}")
        return 2
    request_md = request_path.read_text(encoding="utf-8")

    docx_path = artefact_dir / "report.docx"
    if not docx_path.exists():
        print(f"ERROR: DOCX missing: {docx_path}")
        return 2
    docx_text = _extract_docx_text(docx_path)

    audit_path = artefact_dir / "audit_summary.json"
    audit = {}
    if audit_path.exists():
        audit = json.loads(audit_path.read_text(encoding="utf-8"))

    trace_path = artefact_dir / "trace.jsonl"
    trace_head: list[dict] = []
    if trace_path.exists():
        with trace_path.open("r", encoding="utf-8") as f:
            for i, line in enumerate(f):
                if i >= 200:
                    break
                try:
                    trace_head.append(json.loads(line))
                except Exception:
                    continue

    user_msg = _build_user_message(request_md, docx_text, audit, trace_head)
    print(f"=== Opus gate M{milestone} ===")
    print(f"  request: {request_path}")
    print(f"  artefact: {artefact_dir}")
    print(f"  docx text len: {len(docx_text)} chars")
    print(f"  audit keys: {list(audit.keys())[:10]}")
    print(f"  trace head: {len(trace_head)} events")
    print(f"  user msg total: {len(user_msg)} chars (~{len(user_msg)//4} tokens)")
    print(f"\n  calling {OPUS_MODEL} ...")

    text, cost_rub, usage = call_opus(SYSTEM_PROMPT, user_msg)
    cost_usd = round(cost_rub / USD_RUB_RATE, 4)
    verdict = parse_verdict(text)

    print(f"\n  done. tokens_in={usage.get('prompt_tokens')} tokens_out={usage.get('completion_tokens')}")
    print(f"  cost: ${cost_usd:.4f} ({cost_rub:.2f} RUB)")
    print(f"  parsed verdict: {verdict}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    header = (
        f"# M{milestone} verdict\n\n"
        f"**Ran:** {datetime.now(timezone.utc).isoformat()}\n"
        f"**Model:** {OPUS_MODEL}\n"
        f"**Cost:** ${cost_usd:.4f}\n"
        f"**Parsed verdict:** {verdict}\n\n"
        f"---\n\n"
    )
    out_path.write_text(header + text, encoding="utf-8")
    print(f"  saved: {out_path}")

    return 0 if verdict == "PASS" else 1


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--milestone", type=int, required=True)
    ap.add_argument("--request", type=Path, required=True)
    ap.add_argument("--artefact-dir", type=Path, required=True)
    ap.add_argument("--out", type=Path, default=None)
    args = ap.parse_args()

    out_path = args.out or REPO_ROOT / "milestones" / f"M{args.milestone}_VERDICT.md"
    return run_gate(args.milestone, args.request, args.artefact_dir, out_path)


if __name__ == "__main__":
    raise SystemExit(main())
